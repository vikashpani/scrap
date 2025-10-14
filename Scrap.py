import pandas as pd
import win32com.client as win32
import os

# --- Step 1: Read Excel files ---
combined_df = pd.read_excel("combined_output.xlsx")
claims_df = pd.read_excel("claims_data.xlsx")

# --- Step 2: Filter severity and numeric ranking ---
if "FIELD_RANKING" in claims_df.columns:
    claims_df = claims_df[
        (claims_df["SEVERITY"].str.upper() == "CRITICAL") &
        (claims_df["FIELD_RANKING"].astype(str).str.isdigit())
    ]
else:
    claims_df = claims_df[claims_df["SEVERITY"].str.upper() == "CRITICAL"]

# --- Step 3: Helper functions ---
def get_field_status(group):
    matched = set(group.loc[group["MATCH"] == True, "FIELD_NAME"])
    unmatched = set(group.loc[group["MATCH"] == False, "FIELD_NAME"])
    return matched, unmatched


def get_field_value(group, field_keyword):
    row = group[group["FIELD_NAME"].str.lower().str.contains(field_keyword)]
    if not row.empty:
        return (
            str(row["PST_VALUE"].iloc[0]).strip(),
            str(row["HRP_VALUE"].iloc[0]).strip(),
            bool(row["MATCH"].iloc[0]),
        )
    return None, None, None


def process_claim(claim_df):
    claim_no = claim_df["HRP_CLAIM_NO"].iloc[0]
    matched_fields, unmatched_fields = get_field_status(claim_df)

    pst_status, hrp_status, _ = get_field_value(claim_df, "claimstatus")
    pst_allowed, hrp_allowed, allowed_match = get_field_value(claim_df, "allowedamount")
    pst_paid, hrp_paid, _ = get_field_value(claim_df, "paidamount")
    pst_denial, hrp_denial, _ = get_field_value(claim_df, "denialreasoncode")

    reason_codes = list(claim_df.loc[claim_df["MATCH"] == False, "HRP_VALUE"].unique())

    result = {
        "HRP_CLAIM_NO": claim_no,
        "pst_status": pst_status,
        "hrp_status": hrp_status,
        "pst_allowed_amount": pst_allowed,
        "hrp_allowed_amount": hrp_allowed,
        "pst_paid_amount": pst_paid,
        "hrp_paid_amount": hrp_paid,
        "reason_codes": reason_codes,
        "matched_fields": list(matched_fields),
        "unmatched_fields": list(unmatched_fields),
        "status_summary": "",
        "sheet": ""
    }

    pst_status = str(pst_status).lower()
    hrp_status = str(hrp_status).lower()

    # --- Logic ---
    if pst_status == hrp_status:
        if pst_status == "paid":
            result["sheet"] = "pst_paid-hrp_paid"
            result["status_summary"] = (
                "Allowed amount matched" if allowed_match else "Allowed amount not matched"
            )
        elif pst_status == "denied":
            result["sheet"] = "pst_denied-hrp_denied"
            result["status_summary"] = "Denied"
        elif pst_status == "pended":
            result["sheet"] = "pst_pended-hrp_pended"
            result["status_summary"] = (
                "Allowed amount matched" if allowed_match else "Allowed amount not matched"
            )
    else:
        result["sheet"] = f"pst_{pst_status}-hrp_{hrp_status}"
        result["status_summary"] = f"{pst_status.title()} vs {hrp_status.title()}"

    return result


# --- Step 4: Apply to all claims ---
results = []
for claim_no, group in claims_df.groupby("HRP_CLAIM_NO"):
    processed = process_claim(group)
    if processed:
        results.append(processed)

results_df = pd.DataFrame(results)

# --- Step 5: Confusion Matrix + Summary ---
confusion = (
    results_df.groupby(["pst_status", "hrp_status"])
    .size()
    .reset_index(name="Count")
    .pivot(index="pst_status", columns="hrp_status", values="Count")
    .fillna(0)
)

pst_summary = results_df["pst_status"].value_counts().reset_index()
pst_summary.columns = ["PST_STATUS", "Count"]

hrp_summary = results_df["hrp_status"].value_counts().reset_index()
hrp_summary.columns = ["HRP_STATUS", "Count"]

def is_good_claim(row):
    pst = str(row["pst_status"]).lower()
    hrp = str(row["hrp_status"]).lower()
    unmatched = {x.lower() for x in row["unmatched_fields"]}
    summary = str(row["status_summary"]).lower()
    
    if pst == "paid" and hrp == "paid":
        if (not unmatched or unmatched == {"allowedamountclaimlevel"} or "allowed amount matched" in summary):
            return True
    if pst == "pended" and hrp == "paid":
        if unmatched.issubset({"claimstatus", "claimlevelstatus"}):
            return True
    if pst == "denied" and hrp == "paid":
        return True
    return False

results_df["GOOD_CLAIM"] = results_df.apply(is_good_claim, axis=1)
good_claims = results_df["GOOD_CLAIM"].sum()
total_claims = len(results_df)
good_claim_pct = round((good_claims / total_claims) * 100, 2)

summary_rows = [
    {"Metric": "Total Claims", "Count": total_claims, "Percentage": 100},
    {"Metric": "Good Claims", "Count": good_claims, "Percentage": good_claim_pct},
    {"Metric": "Bad Claims", "Count": total_claims - good_claims, "Percentage": 100 - good_claim_pct},
]

summary_df = pd.DataFrame(summary_rows)

conf_flat = confusion.reset_index()
conf_flat.columns = ["PST_STATUS"] + [f"HRP_{c}" for c in conf_flat.columns[1:]]
summary_combined = pd.concat([summary_df, conf_flat], ignore_index=True)

# --- Step 6: Write to Excel ---
temp_file = "claim_analysis_output.xlsx"
final_file = "claim_analysis_pivot.xlsx"

with pd.ExcelWriter(temp_file) as writer:
    results_df.to_excel(writer, sheet_name="All_Claims", index=False)
    summary_combined.to_excel(writer, sheet_name="Result_Summary", index=False)
    pst_summary.to_excel(writer, sheet_name="PST_Summary", index=False)
    hrp_summary.to_excel(writer, sheet_name="HRP_Summary", index=False)

print("✅ Excel file created. Now adding pivot tables...")

# --- Step 7: Convert summaries into real Pivot Tables ---
excel = win32.Dispatch("Excel.Application")
excel.Visible = False

wb = excel.Workbooks.Open(os.path.abspath(temp_file))

# Keep only summary sheets
for sheet in wb.Sheets:
    if sheet.Name not in ["Result_Summary", "PST_Summary", "HRP_Summary"]:
        sheet.Delete()

for sheet_name in ["Result_Summary", "PST_Summary", "HRP_Summary"]:
    ws = wb.Sheets(sheet_name)
    data_range = ws.Range("A1").CurrentRegion
    pivot_sheet = wb.Worksheets.Add()
    pivot_sheet.Name = f"{sheet_name}_Pivot"

    pivot_cache = wb.PivotCaches().Create(
        SourceType=1, SourceData=data_range
    )
    pivot_table = pivot_cache.CreatePivotTable(
        TableDestination=pivot_sheet.Cells(3, 1),
        TableName=f"{sheet_name}_PivotTable"
    )

    try:
        if sheet_name == "Result_Summary":
            pivot_table.PivotFields("Metric").Orientation = 1
            pivot_table.PivotFields("Count").Orientation = 4
            pivot_table.PivotFields("Percentage").Orientation = 4
        elif sheet_name == "PST_Summary":
            pivot_table.PivotFields("PST_STATUS").Orientation = 1
            pivot_table.PivotFields("Count").Orientation = 4
        elif sheet_name == "HRP_Summary":
            pivot_table.PivotFields("HRP_STATUS").Orientation = 1
            pivot_table.PivotFields("Count").Orientation = 4
    except Exception as e:
        print(f"⚠️ Could not create pivot for {sheet_name}: {e}")

wb.SaveAs(os.path.abspath(final_file))
wb.Close(SaveChanges=True)
excel.Quit()

print(f"✅ Final pivot-enabled Excel saved as: {final_file}")










import pandas as pd

# --- Step 1: Read Excel files ---
combined_df = pd.read_excel("combined_output.xlsx")
claims_df = pd.read_excel("claims_data.xlsx")

# --- Step 2: Filter severity and numeric ranking ---
if "FIELD_RANKING" in claims_df.columns:
    claims_df = claims_df[
        (claims_df["SEVERITY"].str.upper() == "CRITICAL") &
        (claims_df["FIELD_RANKING"].astype(str).str.isdigit())
    ]
else:
    claims_df = claims_df[claims_df["SEVERITY"].str.upper() == "CRITICAL"]

# --- Step 3: Helper functions ---
def get_field_status(group):
    matched = set(group.loc[group["MATCH"] == True, "FIELD_NAME"])
    unmatched = set(group.loc[group["MATCH"] == False, "FIELD_NAME"])
    return matched, unmatched


def get_field_value(group, field_keyword):
    row = group[group["FIELD_NAME"].str.lower().str.contains(field_keyword)]
    if not row.empty:
        return (
            str(row["PST_VALUE"].iloc[0]).strip(),
            str(row["HRP_VALUE"].iloc[0]).strip(),
            bool(row["MATCH"].iloc[0]),
        )
    return None, None, None


def process_claim(claim_df):
    claim_no = claim_df["HRP_CLAIM_NO"].iloc[0]
    matched_fields, unmatched_fields = get_field_status(claim_df)

    pst_status, hrp_status, _ = get_field_value(claim_df, "claimstatus")
    pst_allowed, hrp_allowed, allowed_match = get_field_value(claim_df, "allowedamount")
    pst_paid, hrp_paid, _ = get_field_value(claim_df, "paidamount")
    pst_denial, hrp_denial, _ = get_field_value(claim_df, "denialreasoncode")

    reason_codes = list(claim_df.loc[claim_df["MATCH"] == False, "HRP_VALUE"].unique())

    result = {
        "HRP_CLAIM_NO": claim_no,
        "pst_status": pst_status,
        "hrp_status": hrp_status,
        "pst_allowed_amount": pst_allowed,
        "hrp_allowed_amount": hrp_allowed,
        "pst_paid_amount": pst_paid,
        "hrp_paid_amount": hrp_paid,
        "reason_codes": reason_codes,
        "matched_fields": list(matched_fields),
        "unmatched_fields": list(unmatched_fields),
        "status_summary": "",
        "sheet": ""
    }

    pst_status = str(pst_status).lower()
    hrp_status = str(hrp_status).lower()

    # --- Logic ---
    if pst_status == hrp_status:
        if pst_status == "paid":
            result["sheet"] = "pst_paid-hrp_paid"
            result["status_summary"] = (
                "Allowed amount matched" if allowed_match else "Allowed amount not matched"
            )
        elif pst_status == "denied":
            result["sheet"] = "pst_denied-hrp_denied"
            result["status_summary"] = "Denied"
        elif pst_status == "pended":
            result["sheet"] = "pst_pended-hrp_pended"
            result["status_summary"] = (
                "Allowed amount matched" if allowed_match else "Allowed amount not matched"
            )
    else:
        result["sheet"] = f"pst_{pst_status}-hrp_{hrp_status}"
        result["status_summary"] = f"{pst_status.title()} vs {hrp_status.title()}"

    return result


# --- Step 4: Apply to all claims ---
results = []
for claim_no, group in claims_df.groupby("HRP_CLAIM_NO"):
    processed = process_claim(group)
    if processed:
        results.append(processed)

results_df = pd.DataFrame(results)

# --- Step 5: Confusion Matrix + Summary ---
confusion = (
    results_df.groupby(["pst_status", "hrp_status"])
    .size()
    .reset_index(name="Count")
    .pivot(index="pst_status", columns="hrp_status", values="Count")
    .fillna(0)
)

# Calculate PST and HRP distributions
pst_summary = results_df["pst_status"].value_counts().reset_index()
pst_summary.columns = ["PST_STATUS", "Count"]

hrp_summary = results_df["hrp_status"].value_counts().reset_index()
hrp_summary.columns = ["HRP_STATUS", "Count"]

# Calculate good claim logic
def is_good_claim(row):
    pst = str(row["pst_status"]).lower()
    hrp = str(row["hrp_status"]).lower()
    unmatched = {x.lower() for x in row["unmatched_fields"]}
    summary = str(row["status_summary"]).lower()
    
    # Paid-paid
    if pst == "paid" and hrp == "paid":
        if (not unmatched or unmatched == {"allowedamountclaimlevel"} or "allowed amount matched" in summary):
            return True
    # Pended-paid
    if pst == "pended" and hrp == "paid":
        if unmatched.issubset({"claimstatus", "claimlevelstatus"}):
            return True
    # Denied-paid
    if pst == "denied" and hrp == "paid":
        return True
    return False

results_df["GOOD_CLAIM"] = results_df.apply(is_good_claim, axis=1)
good_claims = results_df["GOOD_CLAIM"].sum()
total_claims = len(results_df)
good_claim_pct = round((good_claims / total_claims) * 100, 2)

# --- Step 6: Build Enhanced Summary Sheet ---
summary_rows = [
    {"Metric": "Total Claims", "Count": total_claims, "Percentage": 100},
    {"Metric": "Good Claims", "Count": good_claims, "Percentage": good_claim_pct},
    {"Metric": "Bad Claims", "Count": total_claims - good_claims, "Percentage": 100 - good_claim_pct},
    {},
]

summary_df = pd.DataFrame(summary_rows)

# Add Confusion Matrix Below in Same Sheet (Flattened)
conf_flat = confusion.reset_index()
conf_flat.columns = ["PST_STATUS"] + [f"HRP_{c}" for c in conf_flat.columns[1:]]
summary_combined = pd.concat([summary_df, conf_flat], ignore_index=True)

# --- Step 7: Write to Excel ---
with pd.ExcelWriter("claim_analysis_output.xlsx") as writer:
    for sheet, sub_df in results_df.groupby("sheet"):
        sub_df.to_excel(writer, sheet_name=sheet[:30], index=False)

    # Confusion and summaries
    summary_combined.to_excel(writer, sheet_name="Result_Summary", index=False)
    pst_summary.to_excel(writer, sheet_name="PST_Summary", index=False)
    hrp_summary.to_excel(writer, sheet_name="HRP_Summary", index=False)
    confusion.to_excel(writer, sheet_name="Confusion_Matrix")

print("✅ Processing completed. Output saved as 'claim_analysis_output.xlsx'")










import pandas as pd

# Step 1: Combine all sheet results into a single DataFrame
results_df = pd.DataFrame(results)

# Step 2: Save all results by sheet
with pd.ExcelWriter("claim_analysis_output_v1.xlsx", engine="xlsxwriter") as writer:
    for sheet, sub_df in results_df.groupby("sheet"):
        sub_df.to_excel(writer, sheet_name=sheet[:30], index=False)

    # Step 3: Create summary by sheet
    summary = results_df.groupby("sheet").size().reset_index(name="Count")
    summary["Percentage"] = round((summary["Count"] / summary["Count"].sum()) * 100, 2)
    summary.loc[len(summary.index)] = ["Total", summary["Count"].sum(), 100.0]

    # Step 4: Create pivot tables for PST and HRP
    # Pivot 1: PST
    pst_pivot = pd.pivot_table(
        results_df,
        values="sheet",
        index=["PST_STATUS"],
        columns=["Final_Status"],
        aggfunc="count",
        fill_value=0,
    ).reset_index()

    # Pivot 2: HRP
    hrp_pivot = pd.pivot_table(
        results_df,
        values="sheet",
        index=["HRP_STATUS"],
        columns=["Final_Status"],
        aggfunc="count",
        fill_value=0,
    ).reset_index()

    # Step 5: Write everything into one sheet (side-by-side layout)
    startrow = 0
    summary.to_excel(writer, sheet_name="Result_Summary", index=False, startrow=startrow)
    startrow += len(summary) + 3

    worksheet = writer.sheets["Result_Summary"]
    worksheet.write(startrow - 2, 0, "PST Pivot Summary")
    pst_pivot.to_excel(writer, sheet_name="Result_Summary", index=False, startrow=startrow)
    startrow += len(pst_pivot) + 4

    worksheet.write(startrow - 2, 0, "HRP Pivot Summary")
    hrp_pivot.to_excel(writer, sheet_name="Result_Summary", index=False, startrow=startrow)

print("✅ Processing completed. Output saved as 'claim_analysis_output_v1.xlsx'")






fig = px.imshow(
        confusion_percent,
        text_auto=".1f",
        color_continuous_scale="Blues",
        labels=dict(x="HRP Status", y="PST Status", color="Percentage (%)"),
        title="Confusion Matrix Heatmap (PST vs HRP)",
    )
    fig.update_xaxes(side="top")
    st.plotly_chart(fig, use_container_width=True)





import pandas as pd
import streamlit as st

# -------------------------------
# STEP 1: Load Excel
# -------------------------------
uploaded_file = st.file_uploader("Upload Claims Excel", type=["xlsx", "xls", "csv"])
if uploaded_file:
    df = pd.read_excel(uploaded_file)

    # Normalize column names
    df.columns = df.columns.str.strip().str.upper()

    # Expected columns (adjust if your Excel uses slightly different names)
    required_cols = ["HRP_CLAIM_NO", "PST_STATUS", "HRP_STATUS", "UNMATCHED_FIELDS", "SUMMARY_MATCH"]
    missing = [c for c in required_cols if c not in df.columns]
    if missing:
        st.error(f"Missing columns: {missing}")
        st.stop()

    # -------------------------------
    # STEP 2: Compute Total Claims
    # -------------------------------
    total_claims = df["HRP_CLAIM_NO"].nunique()

    # -------------------------------
    # STEP 3: Define Good Claim Logic
    # -------------------------------
    def is_good_claim(row):
        pst_status = str(row["PST_STATUS"]).lower().strip()
        hrp_status = str(row["HRP_STATUS"]).lower().strip()
        unmatched = str(row["UNMATCHED_FIELDS"]).lower().strip()
        summary = str(row["SUMMARY_MATCH"]).lower().strip()

        # Condition 1: Both paid, unmatched fields empty, and allowed amount matched
        if pst_status == "paid" and hrp_status == "paid":
            if unmatched in ["", "{}", "set()", "[]"] and "allowed amount matched" in summary:
                return True

        # Condition 2: PST pended and HRP paid, only claimstatus or claimlevelstatus mismatched
        if pst_status == "pended" and hrp_status == "paid":
            if "claimstatus" in unmatched or "claimlevelstatus" in unmatched:
                return True

        # Condition 3: PST denied and HRP paid
        if pst_status == "denied" and hrp_status == "paid":
            return True

        return False

    df["GOOD_CLAIM"] = df.apply(is_good_claim, axis=1)

    # -------------------------------
    # STEP 4: Summary Counts
    # -------------------------------
    good_claims_count = df["GOOD_CLAIM"].sum()
    bad_claims_count = total_claims - good_claims_count
    good_claim_percentage = round((good_claims_count / total_claims) * 100, 2)

    # -------------------------------
    # STEP 5: Confusion Matrix
    # -------------------------------
    confusion = (
        df.groupby(["PST_STATUS", "HRP_STATUS"])
        .size()
        .reset_index(name="Count")
        .pivot(index="PST_STATUS", columns="HRP_STATUS", values="Count")
        .fillna(0)
    )

    # Compute percentages
    confusion_percent = confusion / confusion.values.sum() * 100
    confusion_display = confusion.copy()
    for c in confusion.columns:
        confusion_display[c] = confusion[c].astype(int).astype(str) + " (" + confusion_percent[c].round(1).astype(str) + "%)"

    # -------------------------------
    # STEP 6: Dashboard UI
    # -------------------------------
    st.title("Claims Validation Dashboard")

    # KPI Cards
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Claims", total_claims)
    col2.metric("Good Claims", good_claims_count)
    col3.metric("Bad Claims", bad_claims_count)
    col4.metric("Good Claims (%)", f"{good_claim_percentage}%")

    st.divider()

    # Confusion Matrix
    st.subheader("Confusion Matrix (PST vs HRP)")
    st.dataframe(confusion_display.style.set_caption("Counts (and %) by Status"))

    # Optional chart
    st.bar_chart(
        pd.DataFrame({
            "Type": ["Good Claims", "Bad Claims"],
            "Count": [good_claims_count, bad_claims_count]
        }).set_index("Type")
    )

    st.divider()

    # Detailed Good Claims
    st.subheader("✅ Good Claims Details")
    st.dataframe(
        df[df["GOOD_CLAIM"] == True][
            ["HRP_CLAIM_NO", "PST_STATUS", "HRP_STATUS", "UNMATCHED_FIELDS", "SUMMARY_MATCH"]
        ]
    )

    # Download option
    st.download_button(
        "Download Good Claims as Excel",
        data=df[df["GOOD_CLAIM"] == True].to_csv(index=False).encode("utf-8"),
        file_name="good_claims.csv",
        mime="text/csv",
    )
else:
    st.info("Please upload your claims Excel file to begin.")










summary = results_df.groupby("sheet").size().reset_index(name="Count")
summary["Percentage"] = round((summary["Count"] / summary["Count"].sum()) * 100, 2)

# Add total row
summary.loc[len(summary.index)] = ["Total", summary["Count"].sum(), 100.0]


DWH_df["MATCH"] = DWH_df["MATCH"].apply(lambda x: True if x == 1 or x == 1.0 else False)





mask = DWH_df["FIELD_RANKING"].astype(str).str.strip().str.isdigit()
digit_value = DWH_df.loc[mask, "FIELD_RANKING"].iloc[0] if mask.any() else None
print(digit_value)





def get_multivalue_field(group, field_keyword):
    """
    Return (set of PST_VALUEs, set of HRP_VALUEs, match_flag)
    for fields like denialReasonCode that can appear multiple times.
    """
    rows = group[group["FIELD_NAME"].str.lower().str.contains(field_keyword)]
    if not rows.empty:
        pst_values = set(str(v).strip() for v in rows["PST_VALUE"].dropna())
        hrp_values = set(str(v).strip() for v in rows["HRP_VALUE"].dropna())
        match_flag = pst_values == hrp_values
        return pst_values, hrp_values, match_flag
    return set(), set(), None
    










import pandas as pd

# --- Step 1: Read Excel files ---
combined_df = pd.read_excel("combined_output.xlsx")  # contains FIELD_NAME, FIELD_CATEGORY
claims_df = pd.read_excel("claims_data.xlsx")        # contains FIELD_NAME, HRP_CLAIM_NO, PST_VALUE, HRP_VALUE, MATCH, SEVERITY, FIELD_RANKING

# --- Step 2: Filter severity and ranking ---
if "FIELD_RANKING" in claims_df.columns:
    claims_df = claims_df[
        (claims_df["SEVERITY"].str.upper() == "CRITICAL") &
        (claims_df["FIELD_RANKING"].astype(str).str.isdigit())
    ]
else:
    claims_df = claims_df[claims_df["SEVERITY"].str.upper() == "CRITICAL"]

# --- Helper functions ---
def get_field_status(group):
    matched = set(group.loc[group["MATCH"] == True, "FIELD_NAME"])
    unmatched = set(group.loc[group["MATCH"] == False, "FIELD_NAME"])
    return matched, unmatched


def get_field_value(group, field_keyword):
    """Return (PST_VALUE, HRP_VALUE, MATCH) for a given field name keyword."""
    row = group[group["FIELD_NAME"].str.lower().str.contains(field_keyword)]
    if not row.empty:
        return (
            str(row["PST_VALUE"].iloc[0]).strip(),
            str(row["HRP_VALUE"].iloc[0]).strip(),
            bool(row["MATCH"].iloc[0]),
        )
    return None, None, None


def process_claim(claim_df):
    claim_no = claim_df["HRP_CLAIM_NO"].iloc[0]
    matched_fields, unmatched_fields = get_field_status(claim_df)

    # --- Extract required field values ---
    pst_status, hrp_status, _ = get_field_value(claim_df, "claimstatus")
    pst_allowed, hrp_allowed, allowed_match = get_field_value(claim_df, "allowedamount")
    pst_paid, hrp_paid, _ = get_field_value(claim_df, "paidamount")
    pst_denial, hrp_denial, _ = get_field_value(claim_df, "denialreasoncode")

    reason_codes = list(claim_df.loc[claim_df["MATCH"] == False, "HRP_VALUE"].unique())

    result = {
        "HRP_CLAIM_NO": claim_no,
        "pst_status": pst_status,
        "hrp_status": hrp_status,
        "pst_allowed_amount": pst_allowed,
        "hrp_allowed_amount": hrp_allowed,
        "pst_paid_amount": pst_paid,
        "hrp_paid_amount": hrp_paid,
        "reason_codes": reason_codes,
        "matched_fields": list(matched_fields),
        "unmatched_fields": list(unmatched_fields),
        "status_summary": "",
        "sheet": ""
    }

    pst_status = str(pst_status).lower()
    hrp_status = str(hrp_status).lower()

    # --- Step 5: Decision logic ---
    if pst_status == hrp_status:
        # Status matched
        if pst_status == "paid":
            if allowed_match:
                result["status_summary"] = "Allowed amount matched"
            else:
                result["status_summary"] = "Allowed amount not matched"
            result["sheet"] = "pst_paid-hrp_paid"

        elif pst_status == "denied":
            result["status_summary"] = "Denied"
            result["sheet"] = "pst_denied-hrp_denied"

        elif pst_status == "pended":
            if allowed_match:
                result["status_summary"] = "Allowed amount matched"
            else:
                result["status_summary"] = "Allowed amount not matched"
            result["sheet"] = "pst_pended-hrp_pended"

    else:
        # Status mismatched
        if pst_status == "paid" and hrp_status == "denied":
            result["status_summary"] = "Paid vs Denied"
            result["sheet"] = "pst_paid-hrp_denied"

        elif pst_status == "paid" and hrp_status == "pended":
            result["status_summary"] = "Paid vs Pended"
            result["sheet"] = "pst_paid-hrp_pended"

        elif pst_status == "denied" and hrp_status == "paid":
            result["status_summary"] = "Denied vs Paid"
            result["sheet"] = "pst_denied-hrp_paid"

        elif pst_status == "denied" and hrp_status == "pended":
            result["status_summary"] = "Denied vs Pended"
            result["sheet"] = "pst_denied-hrp_pended"

        elif pst_status == "pended" and hrp_status == "paid":
            result["status_summary"] = "Pended vs Paid"
            result["sheet"] = "pst_pended-hrp_paid"

        elif pst_status == "pended" and hrp_status == "denied":
            result["status_summary"] = "Pended vs Denied"
            result["sheet"] = "pst_pended-hrp_denied"

    return result


# --- Step 6: Apply to all claims ---
results = []
for claim_no, group in claims_df.groupby("HRP_CLAIM_NO"):
    processed = process_claim(group)
    if processed:
        results.append(processed)

results_df = pd.DataFrame(results)

# --- Step 7: Save to Excel with multiple sheets ---
with pd.ExcelWriter("claim_analysis_output.xlsx") as writer:
    for sheet, sub_df in results_df.groupby("sheet"):
        sub_df.to_excel(writer, sheet_name=sheet[:30], index=False)

print("✅ Processing completed. Output saved as 'claim_analysis_output.xlsx'")











import pandas as pd

# Assuming your DataFrames are: DWH_df, claims_df, messagecodes_df

# Step 1: Get unique claim numbers
claims_unique = claims_df['HRP_CLAIM_NO'].unique()
messagecodes_unique = messagecodes_df['hccclaimnumber'].unique()
dwh_claims = DWH_df['HRP_CLAIM_NO'].unique()

# Step 2: Find missing claims from DWH
missing_in_dwh_from_claims = [c for c in claims_unique if c not in dwh_claims]
missing_in_dwh_from_messages = [c for c in messagecodes_unique if c not in dwh_claims]

# Step 3: Convert to DataFrame for Excel
missing_summary = pd.DataFrame({
    'Source': ['claims_df']*len(missing_in_dwh_from_claims) + ['messagecodes_df']*len(missing_in_dwh_from_messages),
    'Missing_Claim': missing_in_dwh_from_claims + missing_in_dwh_from_messages
})

# Step 4: Save to Excel
missing_summary.to_excel("missing_claims_summary.xlsx", index=False)

print("✅ Missing claims summary created!")
print(missing_summary)







import pandas as pd
import numpy as np

# Assuming MATCH is already converted to boolean
# xl_file['MATCH'] = xl_file['MATCH'].apply(lambda x: bool(int(x)) if not pd.isna(x) else False)

# Step 1: Get invalid claims
unique_claims = xl_file.groupby('HRP_CLAIM_NO')['MATCH'].all()
invalid_claims = unique_claims[unique_claims == False].index.to_list()

# Step 2: Filter invalid rows
invalid_rows = xl_file[xl_file['HRP_CLAIM_NO'].isin(invalid_claims)]

# Step 3: Keep only critical + unmatched rows
critical_unmatched = invalid_rows[
    (invalid_rows['SEVERITY'].astype(str).str.lower() == 'critical') &
    (invalid_rows['MATCH'] == False)
]

# Step 4: Identify columns that are unmatched per row
# If you have multiple *_MATCH columns, dynamically find them
match_cols = [col for col in critical_unmatched.columns if col != 'HRP_CLAIM_NO' and col != 'SEVERITY']

def get_unmatched_fields(row):
    return [col for col in match_cols if not row[col]]

critical_unmatched['UNMATCHED_FIELDS'] = critical_unmatched.apply(get_unmatched_fields, axis=1)

# Step 5: Group by claim number and aggregate
summary = critical_unmatched.groupby('HRP_CLAIM_NO').agg({
    'UNMATCHED_FIELDS': lambda x: list(set([item for sublist in x for item in sublist])),  # unique fields
}).reset_index()

summary['UNMATCHED_COUNT'] = summary['UNMATCHED_FIELDS'].apply(len)

# Step 6: Save to Excel
summary.to_excel("critical_unmatched_summary.xlsx", index=False)

print("✅ Summary Excel created!")
print(summary)










import pandas as pd

# 📘 Step 1: Read your Excel file
excel_path = r"C:\path\to\combined_output.xlsx"
df = pd.read_excel(excel_path)

# Ensure column names match exactly
df.columns = df.columns.str.strip()  # remove any trailing spaces

# 📋 Step 2: Your list of numbers
num_list = [12345, 67890, 22222, 33333]  # 🔁 replace with your actual list

# 📘 Step 3: Filter the rows where Run2 HRP # matches any number in the list
matched_df = df[df["Run2 HRP #"].isin(num_list)][["Unique Key", "PST Model", "Run2 HRP #"]]

# 📘 Step 4: Find missing numbers
missing_numbers = [num for num in num_list if num not in df["Run2 HRP #"].values]

# 📘 Step 5: Output results
print("✅ Matching rows:")
print(matched_df)

print("\n❌ Missing numbers from the list:")
print(missing_numbers)

# Optional: Save the matching rows to a new Excel file
output_path = r"C:\path\to\Run2_HRP_Matched_Output.xlsx"
matched_df.to_excel(output_path, index=False)
print(f"\n✅ Matching data saved to: {output_path}")











import pandas as pd

# === Step 1: Load the two Excel files ===
excel1 = pd.read_excel("combined_output.xlsx")
excel2 = pd.read_excel("parallel_runbook.xlsx")

# Clean column names (strip spaces)
excel1.columns = excel1.columns.str.strip()
excel2.columns = excel2.columns.str.strip()

# === Step 2: Merge on the matching columns ===
merged = pd.merge(
    excel2,
    excel1,
    how='left',
    left_on=["Member ID", "Patient Control Number", "Total Charge Amount"],
    right_on=["MEMBER ID", "patient Account Number", "Total Claims Charged Amount"],
    suffixes=("_runbook", "_combined")
)

# === Step 3: Update Run2 HRP Claims # if match found ===
merged["Run2 HRP Claims #"] = merged["Legacy Claims Number"].combine_first(merged["Run2 HRP Claims #"])

# === Step 4: Collect matched rows (where Legacy Claims Number was found) ===
matched_rows = merged[merged["Legacy Claims Number"].notna()][[
    "unique_claim_keys_updated",
    "PST Model Region Claims #",
    "Run2 HRP Claims #"
]]

# === Step 5: Save updated parallel_runbook and output ===
# Save updated parallel runbook
updated_runbook = merged[excel2.columns]  # keep only original columns
updated_runbook.to_excel("parallel_runbook_updated.xlsx", index=False)

# Save matched info to output
matched_rows.to_excel("output.xlsx", index=False)

print("✅ Process complete!")
print("→ Updated file: parallel_runbook_updated.xlsx")
print("→ Output file: output.xlsx")












import pandas as pd
import os
from glob import glob

# 📂 Path to your folder
folder_path = r"C:\path\to\your\excel_folder"   # 🔁 change this to your folder

# 🔍 Get all Excel files in the folder
excel_files = glob(os.path.join(folder_path, "*.xlsx"))

# 🧩 List to collect DataFrames
dfs = []

for file in excel_files:
    try:
        # Read first sheet of each file
        df = pd.read_excel(file)
        df["Source_File"] = os.path.basename(file)  # Optional: track source
        dfs.append(df)
    except Exception as e:
        print(f"❌ Error reading {file}: {e}")

# 📊 Combine all DataFrames
if dfs:
    combined_df = pd.concat(dfs, ignore_index=True)
    print(f"✅ Combined {len(dfs)} files with total rows: {len(combined_df)}")

    # 💾 Write to single Excel file
    output_path = os.path.join(folder_path, "combined_output.xlsx")
    combined_df.to_excel(output_path, index=False)

    print(f"✅ Combined file saved to: {output_path}")
else:
    print("⚠️ No Excel files found or all failed to read.")













import pandas as pd

def filter_matching_claims(excel1_path, excel2_path, output_path):
    # Read both Excel files
    df1 = pd.read_excel(excel1_path)
    df2 = pd.read_excel(excel2_path)

    # Extract the comparison columns (strip spaces and convert to string for safety)
    col1 = df1.iloc[:, 20].astype(str).str.strip()   # Column U = index 20 (0-based)
    col2 = df2.iloc[:, 0].astype(str).str.strip()    # Column A = index 0

    # Filter rows in df1 where U value is present in df2 A column
    matched_df = df1[col1.isin(col2)]

    # Save the matched rows to a new Excel file
    matched_df.to_excel(output_path, index=False)

    print(f"✅ Matching claims saved to: {output_path}")
    print(f"Total matched rows: {len(matched_df)}")

# Example usage
excel1 = r"C:\path\to\excel1.xlsx"
excel2 = r"C:\path\to\excel2.xlsx"
output = r"C:\path\to\matched_claims.xlsx"

filter_matching_claims(excel1, excel2, output)









import os
import shutil

def load_all_files(folder, dest_folder):
    """
    Read all .dat files, check for unique supplier NPIs (NM1*85),
    and copy only files with new supplier NPIs to another folder.
    """
    file_contents = []
    unique_suppliers = []

    # Ensure destination folder exists
    os.makedirs(dest_folder, exist_ok=True)

    for fname in os.listdir(folder):
        if fname.lower().endswith(".dat"):
            file_path = os.path.join(folder, fname)
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    file_contents.append(content)

                # Check for NM1*85 segment
                if "NM1*85" in content:
                    lines = content.split("~")
                    for line in lines:
                        if line.startswith("NM1*85"):
                            parts = line.split("*")
                            supplier_npi = parts[-1].strip()

                            # If this supplier NPI is new, copy file and add to list
                            if supplier_npi not in unique_suppliers:
                                unique_suppliers.append(supplier_npi)
                                dest_path = os.path.join(dest_folder, fname)
                                shutil.copy(file_path, dest_path)
                                print(f"Copied {fname} for supplier {supplier_npi}")
                            break  # Stop after first NM1*85 found

            except Exception as e:
                print(f"Error reading {fname}: {e}")
                continue

    return unique_suppliers, file_contents










import os
import streamlit as st
import pandas as pd

from pandasai import SmartDataframe
from pandasai.llm import AzureOpenAI, OpenAI
import requests

# ------------------------------
# Streamlit UI
# ------------------------------
st.set_page_config(page_title="Excel Q&A with PandasAI", layout="wide")

st.title("📊 Excel Q&A using PandasAI + LLM")

# Sidebar for LLM selection
llm_choice = st.sidebar.selectbox(
    "Choose LLM Backend",
    ["Azure OpenAI", "Groq (LLaMA)"]
)

# Sidebar API keys
if llm_choice == "Azure OpenAI":
    azure_api_key = st.sidebar.text_input("Azure OpenAI API Key", type="password")
    azure_endpoint = st.sidebar.text_input("Azure Endpoint (https://xxx.openai.azure.com/)")
    azure_deployment = st.sidebar.text_input("Azure Deployment Name")
    azure_api_version = st.sidebar.text_input("Azure API Version", value="2024-06-01-preview")

elif llm_choice == "Groq (LLaMA)":
    groq_api_key = st.sidebar.text_input("Groq API Key", type="password")
    groq_model = st.sidebar.text_input("Groq Model", value="llama3-70b-8192")


uploaded_file = st.file_uploader("📥 Upload Excel File", type=["xlsx", "xls"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    st.write("### Preview of Data", df.head())

    query = st.text_input("🔎 Ask a question about your data:")

    if st.button("Run Query") and query:
        try:
            # ------------------------------
            # Setup LLM
            # ------------------------------
            if llm_choice == "Azure OpenAI":
                if not azure_api_key or not azure_endpoint or not azure_deployment:
                    st.error("⚠️ Please enter Azure API Key, Endpoint, and Deployment Name.")
                    st.stop()

                llm = AzureOpenAI(
                    api_token=azure_api_key,
                    azure_endpoint=azure_endpoint,
                    deployment_name=azure_deployment,
                    api_version=azure_api_version,
                )

            elif llm_choice == "Groq (LLaMA)":
                if not groq_api_key:
                    st.error("⚠️ Please enter your Groq API Key.")
                    st.stop()

                # Custom LLM wrapper for Groq (simple REST API call)
                class GroqLLM(OpenAI):
                    def __init__(self, api_key, model="llama3-70b-8192"):
                        super().__init__(api_token=api_key)
                        self.model = model
                        self.api_key = api_key

                    def call(self, prompt, context=None):
                        headers = {"Authorization": f"Bearer {self.api_key}"}
                        payload = {
                            "model": self.model,
                            "messages": [{"role": "user", "content": prompt}],
                        }
                        resp = requests.post(
                            "https://api.groq.com/openai/v1/chat/completions",
                            headers=headers,
                            json=payload,
                        )
                        resp.raise_for_status()
                        return resp.json()["choices"][0]["message"]["content"]

                llm = GroqLLM(api_key=groq_api_key, model=groq_model)

            # ------------------------------
            # Run PandasAI
            # ------------------------------
            sdf = SmartDataframe(df, config={"llm": llm})
            answer = sdf.chat(query)

            st.success("✅ Query Answered")
            st.write("### Result")
            st.write(answer)

            # ------------------------------
            # Export to Excel
            # ------------------------------
            output_path = "pandasai_output.xlsx"
            with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
                df.to_excel(writer, sheet_name="Original Data", index=False)

                # If result is DataFrame-like, also export
                if isinstance(answer, pd.DataFrame):
                    answer.to_excel(writer, sheet_name="Query Result", index=False)
                else:
                    pd.DataFrame({"Answer": [answer]}).to_excel(
                        writer, sheet_name="Query Result", index=False
                    )

            with open(output_path, "rb") as f:
                st.download_button(
                    label="⬇️ Download Result Excel",
                    data=f,
                    file_name="pandasai_result.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

















GROQ_API_KEY = "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"

11m = ChatGroq (api_key=GROQ_API_KEY, model="meta-llama/1lama-4-maverick-17b-128e-instruct"




"""
streamlit_pandasai_full_app.py

Full Streamlit app that:
- Loads an uploaded Excel file and lets the user pick a sheet
- Supports two LLM backends: Azure (LangChain's AzureChatOpenAI if available, else pandasai.AzureOpenAI) and Groq (via HTTP API)
- Uses PandasAI (PandasAI/PandasAI.PandasAI) to run natural-language queries over the DataFrame
- Forces the LLM to return tabular results as CSV (recommended prompt technique) and parses into a DataFrame
- Lets user preview results and download as Excel containing the original sheet and LLM result

Notes:
- Edit the env vars in your environment or paste API keys into the UI when prompted
- Make sure the packages in requirements_txt below are installed

Usage:
  pip install -r requirements.txt
  streamlit run streamlit_pandasai_full_app.py

"""

import streamlit as st
import pandas as pd
import os
import io
from datetime import datetime
from pandasai import PandasAI
from pandasai.llm import OpenAI as PandasAIOpenAI
import json
import textwrap
import requests
from io import StringIO

st.set_page_config(page_title="PandasAI Excel Chat — Azure or Groq", layout="wide")
st.title("PandasAI + Excel — AzureChatOpenAI or Groq (Full Code)")

st.markdown(
    """
    Upload an Excel file, pick a sheet, ask a question in natural language, and get tabular output that you can download as an Excel file.

    Implementation details:
    - We ask the LLM to **return results in CSV only** so we can reliably parse to a DataFrame.
    - Backends supported: Azure (LangChain AzureChatOpenAI adapter or PandasAI's AzureOpenAI) and Groq (HTTP inference API).
    """
)

# ------------------------- Sidebar: options & creds -------------------------
with st.sidebar:
    st.header("LLM backend & credentials")
    backend = st.selectbox("LLM backend", ["azure_langchain", "azure_pandasai", "groq"])
    conversational = st.checkbox("Conversational (PandasAI) mode", value=False)
    st.markdown("**Azure**: Either use LangChain AzureChatOpenAI (preferred) or PandasAI's adapter. Provide credentials below or use env vars.")
    if backend.startswith("azure"):
        az_key = st.text_input("Azure OpenAI API key (leave blank to use env AZURE_OPENAI_API_KEY)", type="password")
        az_endpoint = st.text_input("Azure endpoint (e.g. https://<resource>.openai.azure.com) or leave empty to use AZURE_OPENAI_ENDPOINT env")
        az_deployment = st.text_input("Azure deployment name (deployment for your model)", value=os.getenv("AZURE_DEPLOYMENT_NAME", ""))
        if az_key:
            os.environ["AZURE_OPENAI_API_KEY"] = az_key
        if az_endpoint:
            os.environ["AZURE_OPENAI_ENDPOINT"] = az_endpoint
        if az_deployment:
            os.environ["AZURE_DEPLOYMENT_NAME"] = az_deployment

    if backend == "groq":
        groq_key = st.text_input("Groq API key (leave blank to use env GROQ_API_KEY)", type="password")
        groq_model = st.text_input("Groq model name (e.g. 'gpt-3o-mini')", value=os.getenv("GROQ_MODEL", "gpt-3o-mini"))
        groq_endpoint = st.text_input("Groq endpoint (leave blank to use default)", value=os.getenv("GROQ_ENDPOINT", "https://api.groq.com/v1"))
        if groq_key:
            os.environ["GROQ_API_KEY"] = groq_key
        if groq_model:
            os.environ["GROQ_MODEL"] = groq_model
        if groq_endpoint:
            os.environ["GROQ_ENDPOINT"] = groq_endpoint

st.sidebar.markdown("---")
st.sidebar.markdown("**Prompt tips**: ask for `CSV` only. Example: `Return rows where Age > 50 with columns ID, Name, Age, Amount as CSV with header only, no commentary.`")

# ------------------------- Upload Excel -------------------------
uploaded_file = st.file_uploader("Upload Excel (.xlsx/.xls)", type=["xlsx", "xls"])
if uploaded_file is None:
    st.info("Upload an Excel file to continue")
    st.stop()

# Read sheets
try:
    xls = pd.read_excel(uploaded_file, sheet_name=None)
    sheet_names = list(xls.keys())
    sheet_choice = st.selectbox("Choose sheet", sheet_names)
    df = xls[sheet_choice].copy()
    st.subheader("Data preview — selected sheet")
    st.dataframe(df.head(200))
except Exception as e:
    st.error(f"Failed to read Excel: {e}")
    st.stop()

# ------------------------- Prompt input -------------------------
st.markdown("---")
prompt_user = st.text_area("Enter your question for the dataframe (be explicit: request CSV output)", height=180)
run_btn = st.button("Run")

# ------------------------- Helper: LLM wrappers -------------------------

class LangChainAzureAdapter:
    """Adapter that wraps LangChain's AzureChatOpenAI and exposes a simple .generate(prompt) method for PandasAI.
    We import LangChain lazily so the app doesn't immediately fail if LangChain isn't installed.
    """
    def __init__(self, deployment_name: str = None):
        try:
            from langchain.chat_models import AzureChatOpenAI
        except Exception as e:
            raise RuntimeError("LangChain or AzureChatOpenAI not installed. Install 'langchain' and retry.")
        self.deployment = deployment_name or os.getenv("AZURE_DEPLOYMENT_NAME")
        self.client = AzureChatOpenAI(deployment_name=self.deployment,
                                      openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
                                      openai_api_base=os.getenv("AZURE_OPENAI_ENDPOINT"),
                                      openai_api_version=os.getenv("OPENAI_API_VERSION", "2023-12-01-preview"))

    def generate(self, prompt: str):
        # LangChain Chat model has .predict_messages or .predict — wrap simply
        # We try several common methods depending on LangChain version.
        try:
            # newer versions: .predict
            out = self.client.predict(prompt)
            return out
        except Exception:
            try:
                # older: .predict_messages
                resp = self.client.predict_messages([{"role": "user", "content": prompt}])
                # resp might be a ChatMessage object or str
                return getattr(resp, "content", str(resp))
            except Exception as e:
                raise RuntimeError(f"Azure LangChain call failed: {e}")

class PandasAIAzureAdapter:
    """Use pandasai.llm.OpenAI (which can be configured for Azure via env vars)"""
    def __init__(self):
        # This adapter will be passed to PandasAI as llm param via a small wrapper exposing .run
        self.llm = PandasAIOpenAI(deployment_name=os.getenv("AZURE_DEPLOYMENT_NAME", None))

    def generate(self, prompt: str):
        # pandasai.OpenAI instances expect to be used by PandasAI directly; but for reliability
        # we'll call them via their text-generation interface if available.
        try:
            return self.llm(prompt)
        except Exception as e:
            # fallback: return prompt back to indicate failure
            raise RuntimeError(f"PandasAI Azure adapter failed call: {e}")

class GroqHTTPAdapter:
    """Simple HTTP adapter to call Groq's text-completion endpoint.
    This implementation assumes a POST JSON API at GROQ_ENDPOINT + /completions or similar.
    You may need to adapt it to your Groq account's exact API shape.
    """
    def __init__(self, model=None):
        self.api_key = os.getenv("GROQ_API_KEY")
        if not self.api_key:
            raise RuntimeError("GROQ_API_KEY not found in env; provide API key in sidebar")
        self.model = model or os.getenv("GROQ_MODEL", "gpt-3o-mini")
        self.endpoint = os.getenv("GROQ_ENDPOINT", "https://api.groq.com/v1")

    def generate(self, prompt: str):
        url = f"{self.endpoint}/completions"
        headers = {"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"}
        payload = {
            "model": self.model,
            "prompt": prompt,
            "max_tokens": 512,
            "temperature": 0.0,
        }
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
        if resp.status_code != 200:
            raise RuntimeError(f"Groq API error {resp.status_code}: {resp.text}")
        data = resp.json()
        # Extract text — shape depends on API; try common keys
        if isinstance(data, dict):
            # try multiple fallback paths
            for k in ("text", "output", "choices"):
                if k in data:
                    v = data[k]
                    if isinstance(v, list) and len(v) > 0:
                        # choices -> choices[0].text or choices[0].output_text
                        first = v[0]
                        if isinstance(first, dict):
                            return first.get("text") or first.get("output_text") or json.dumps(first)
                        return str(first)
                    else:
                        return str(v)
        return json.dumps(data)

# ------------------------- Run flow -------------------------
if run_btn:
    if not prompt_user or not prompt_user.strip():
        st.error("Please enter a question prompt (ask the model to return CSV).")
        st.stop()

    # Ensure the prompt asks for CSV-only output to maximize parseability
    forced_prompt = textwrap.dedent(f"""
    You are given a dataframe described as CSV input. Answer with CSV only (header row then rows) with no extra commentary or markdown. Columns must match those requested in the question.

    QUESTION: {prompt_user}

    IMPORTANT: Return **CSV only**. If there are no matching rows, return a CSV with only the header row.
    """)

    st.info("Calling LLM — this may take a few seconds depending on the backend.")

    # Instantiate the chosen adapter
    llm_adapter = None
    try:
        if backend == "azure_langchain":
            llm_adapter = LangChainAzureAdapter(deployment_name=os.getenv("AZURE_DEPLOYMENT_NAME"))
        elif backend == "azure_pandasai":
            llm_adapter = PandasAIAzureAdapter()
        elif backend == "groq":
            llm_adapter = GroqHTTPAdapter(model=os.getenv("GROQ_MODEL"))
        else:
            st.error("Unknown backend selected")
            st.stop()
    except Exception as e:
        st.error(f"Failed to prepare LLM adapter: {e}")
        st.stop()

    # Build a PandasAI instance with a small wrapper LLM object exposing .generate
    class WrapperLLMForPandasAI:
        def __init__(self, adapter):
            self.adapter = adapter

        def __call__(self, prompt: str, **kwargs):
            # PandasAI may call llm(prompt)
            return self.adapter.generate(prompt)

        def generate(self, prompt: str, **kwargs):
            return self.adapter.generate(prompt)

    # Create PandasAI with the wrapper
    pandalai_llm = WrapperLLMForPandasAI(llm_adapter)
    p = PandasAI(llm=pandalai_llm, conversational=conversational)

    # Run with a careful prompt: we attach the original df's column names to the prompt optionally
    cols_sample = ", ".join(list(df.columns[:30]))
    full_prompt = textwrap.dedent(f"""
    You have access to a dataframe with columns: {cols_sample} (and more if present).

    {forced_prompt}
    """)

    try:
        # PandasAI typically expects p.run(df, prompt)
        response = p.run(df, prompt=full_prompt)
    except Exception as e:
        st.error(f"PandasAI run failed: {e}")
        st.stop()

    # response might be text (CSV) or a DataFrame object depending on PandasAI
    result_df = None
    if isinstance(response, pd.DataFrame):
        result_df = response
    else:
        # try to coerce text into dataframe — response may be a string CSV
        resp_text = str(response)
        # Attempt to find first line with comma and treat as CSV
        try:
            result_df = pd.read_csv(StringIO(resp_text))
        except Exception:
            # try to extract CSV block from response (in case the LLM added surrounding text)
            import re
            csv_match = re.search(r"(?sm)(^[\w\W]*?\n)?((?:[^
]*,.*\n)+[^
]*,.*)", resp_text)
            if csv_match:
                csv_text = csv_match.group(2)
                try:
                    result_df = pd.read_csv(StringIO(csv_text))
                except Exception:
                    # final fallback: try to parse by splitting lines and commas
                    lines = [ln for ln in resp_text.splitlines() if "," in ln]
                    if len(lines) >= 1:
                        csv_text = "\n".join(lines)
                        try:
                            result_df = pd.read_csv(StringIO(csv_text))
                        except Exception:
                            result_df = None

    # Show result
    if result_df is None:
        st.warning("Could not parse a CSV table from model output. Showing raw output below.")
        st.subheader("Raw model output")
        st.code(str(response))
    else:
        st.subheader("LLM result (parsed as DataFrame)")
        st.dataframe(result_df.head(200))

        # Prepare downloadable Excel with both sheets
        out = io.BytesIO()
        ts = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="original")
            result_df.to_excel(writer, index=False, sheet_name="result")
        out.seek(0)
        st.download_button("Download original + result as Excel", data=out, file_name=f"pandasai_result_{ts}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ------------------------- Requirements text -------------------------
requirements_txt = """
streamlit
pandas
pandasai
openpyxl
requests
langchain  # only if you want azure_langchain option
"""

st.sidebar.markdown("---")
st.sidebar.subheader("Requirements (pip)")
st.sidebar.code(requirements_txt)

st.sidebar.markdown("If you want, paste small Azure/Groq keys in the sidebar fields; otherwise set env vars: AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_DEPLOYMENT_NAME, GROQ_API_KEY, GROQ_MODEL, GROQ_ENDPOINT.")

# End of file















import pandas as pd
import xml.etree.ElementTree as ET
from collections import defaultdict
import os

# ---------------------------
# Utility Functions
# ---------------------------
def normalize_value(val):
    return str(val).strip() if val else ""

def normalize_date_str(date_str):
    if not date_str or pd.isna(date_str):
        return ""
    try:
        return pd.to_datetime(date_str).strftime("%Y%m%d")
    except Exception:
        return str(date_str)

def get_tracking_partner_id(root: ET.Element) -> str:
    """
    Extract ISA06 (Sender ID) from ISA segment
    """
    isa_seg = root.find(".//seg[@id='ISA']")
    if isa_seg is None:
        return ""
    for ele in isa_seg.findall("ele"):
        if ele.attrib.get("id") == "ISA06":
            return ele.text.strip()
    return ""

# ---------------------------
# Excel Loader
# ---------------------------
def load_excel_keys(EXCEL_FILE):
    print("📂 Loading Excel file:", EXCEL_FILE)
    df = pd.read_excel(EXCEL_FILE, dtype=str).fillna("")

    cols_map = {
        "clm01": "ClaimID",
        "member": "member_hcc_id",
        "startdate": "StartDate",
        "supplier_npi": "SUPPLIER_NPI",
        "supplier_taxid": "SI_SUPPLIER_TAX",
        "rendering_npi": "PRACTITIONER_HCC_ID",
    }

    # Validate required columns
    for col, header in cols_map.items():
        if header not in df.columns:
            raise SystemExit(f"❌ ERROR: Missing required column {header} in Excel")

    # Collect normalized unique keys
    excel_keys = set()
    for _, row in df.iterrows():
        nm109 = normalize_value(row[cols_map["member"]])
        sdate = normalize_date_str(row[cols_map["startdate"]])
        supplier_npi = normalize_value(row[cols_map["supplier_npi"]])
        supplier_taxid = normalize_value(row[cols_map["supplier_taxid"]])
        rendering_npi = normalize_value(row[cols_map["rendering_npi"]])

        key = (nm109, sdate, supplier_npi, supplier_taxid, rendering_npi)
        excel_keys.add(key)

    print(f"✅ Excel unique keys loaded: {len(excel_keys)}")
    return excel_keys

# ---------------------------
# EDI Parser
# ---------------------------
def get_edi_keys(root: ET.Element, tracking_partner_id: str):
    edi_keys = set()

    # Iterate through CLM segments (claims)
    for clm_seg in root.findall(".//seg[@id='CLM']"):
        nm109 = ""
        sdate = ""
        supplier_npi = ""
        supplier_taxid = ""
        rendering_npi = ""

        # Find Member ID (NM109)
        nm1_seg = clm_seg.find(".//seg[@id='NM1']")
        if nm1_seg is not None:
            for ele in nm1_seg.findall("ele"):
                if ele.attrib.get("id") == "NM109":
                    nm109 = normalize_value(ele.text)

        # Find Service Date (DTP03)
        dtp_seg = clm_seg.find(".//seg[@id='DTP']")
        if dtp_seg is not None:
            for ele in dtp_seg.findall("ele"):
                if ele.attrib.get("id") == "DTP03":
                    sdate = normalize_date_str(ele.text)

        # Find Supplier Info
        supplier_seg = clm_seg.find(".//seg[@id='SUPPLIER']")
        if supplier_seg is not None:
            supplier_npi = normalize_value(supplier_seg.findtext("ele[@id='NPI']"))
            supplier_taxid = normalize_value(supplier_seg.findtext("ele[@id='TAXID']"))

        # Find Rendering Practitioner Info
        rendering_seg = clm_seg.find(".//seg[@id='RENDERING']")
        if rendering_seg is not None:
            rendering_npi = normalize_value(rendering_seg.findtext("ele[@id='NPI']"))

        key = (nm109, sdate, supplier_npi, supplier_taxid, rendering_npi)
        edi_keys.add(key)

    print(f"📑 EDI keys extracted: {len(edi_keys)} for partner {tracking_partner_id}")
    return edi_keys

# ---------------------------
# Main Driver
# ---------------------------
if __name__ == "__main__":
    # Load Excel keys
    excel_keys = load_excel_keys("claims.xlsx")

    # Folder with EDI XML files
    edi_folder = "edi_files"
    all_edi_keys = set()

    for edi_file in os.listdir(edi_folder):
        if not edi_file.endswith(".xml"):
            continue
        edi_path = os.path.join(edi_folder, edi_file)
        print(f"\n🔍 Parsing EDI file: {edi_file}")

        tree = ET.parse(edi_path)
        root = tree.getroot()

        tracking_partner_id = get_tracking_partner_id(root)
        edi_keys = get_edi_keys(root, tracking_partner_id)
        all_edi_keys |= edi_keys  # merge all files' keys

    # Compare Excel vs EDI
    missing_in_edi = excel_keys - all_edi_keys
    extra_in_edi = all_edi_keys - excel_keys

    print("\n📊 Comparison Results")
    print("✅ Keys in both Excel & EDI:", len(excel_keys & all_edi_keys))
    print("❌ Missing in EDI:", len(missing_in_edi))
    print("⚠️ Extra in EDI:", len(extra_in_edi))

    # Debug: show some missing/extra
    if missing_in_edi:
        print("\n❌ Sample Missing in EDI:", list(missing_in_edi)[:5])
    if extra_in_edi:
        print("\n⚠️ Sample Extra in EDI:", list(extra_in_edi)[:5])










import pandas as pd
from collections import defaultdict
import xml.etree.ElementTree as ET

# ---------------------------
# Utility Normalization Functions
# ---------------------------
def normalize_value(val):
    return str(val).strip() if val else ""

def normalize_amount(val):
    try:
        return f"{float(val):.2f}"
    except Exception:
        return "0.00"

def normalize_filename(fname):
    return fname.strip().replace(" ", "_").lower() if fname else ""

def normalize_date_str(date_str):
    if not date_str or pd.isna(date_str):
        return ""
    try:
        return pd.to_datetime(date_str).strftime("%Y%m%d")
    except Exception:
        return str(date_str)

# ---------------------------
# Excel Loader
# ---------------------------
def load_excel(EXCEL_FILE):
    print("Loading Excel file:", EXCEL_FILE)
    df = pd.read_excel(EXCEL_FILE, dtype=str).fillna("")

    # Map columns automatically (adjust names to your Excel headers)
    cols_map = {
        "clm01": "Claim ID",
        "member": "Member ID",
        "startdate": "Service Dates",
        "supplier_npi": "Supplier NPI",
        "supplier_taxid": "Supplier Tax ID",
        "rendering_npi": "Rendering Practitioner NPI",
    }

    # Validate required columns
    for col in ["clm01", "member", "startdate"]:
        if cols_map[col] not in df.columns:
            raise SystemExit(f"ERROR: Missing column {cols_map[col]} in Excel")

    excel_keys_by_file = defaultdict(list)

    for _, row in df.iterrows():
        clm01 = normalize_value(row[cols_map["clm01"]])
        nm109 = normalize_value(row[cols_map["member"]])
        sdate = normalize_date_str(row[cols_map["startdate"]])

        supplier_npi = normalize_value(row[cols_map["supplier_npi"]])
        supplier_taxid = normalize_value(row[cols_map["supplier_taxid"]])
        rendering_npi = normalize_value(row[cols_map["rendering_npi"]])

        # Group by claim file (or fallback to 'default')
        clmfin = "default_claims"
        excel_keys_by_file[clmfin].append(
            (clm01, nm109, sdate, supplier_npi, supplier_taxid, rendering_npi)
        )

    print(f"Excel entries loaded: {sum(len(v) for v in excel_keys_by_file.values())}")
    return excel_keys_by_file


# ---------------------------
# EDI Segment Parser
# ---------------------------
def parse_segment(seg: ET.Element, tracking_partner_id: str) -> str:
    seg_id = seg.attrib["id"]
    elements = [seg_id]

    # dictionary to hold element values with proper position
    element_dict = {}

    # partner-specific composite separators
    comp_symbol = {
        "030240928": "*",
        "332151969": "*",
        "910842999": "*",
        "133052274": "*",
        "BTQBROADS": ":",
        "FWOCR": "",
        "474513700": ":",
    }

    symbol = comp_symbol.get(tracking_partner_id, ":")

    for child in seg:
        if child.tag == "ele":
            # extract numeric part (position index)
            idx_str = child.attrib["id"].replace(seg_id, "")
            if idx_str.isdigit():
                idx = int(idx_str)
                element_dict[idx] = child.text or ""

        elif child.tag == "comp":
            sub_values = []
            mainidx = None
            for sub in child:
                if sub.tag == "subele":
                    mainidx = int(sub.attrib["id"].split("-")[0].replace(seg_id, ""))
                    idx_str = sub.attrib["id"].split("-")[-1]
                    if idx_str.isdigit():
                        sub_values.append(sub.text or "")
            if mainidx is not None:
                element_dict[mainidx] = symbol.join(sub_values)

    # Build elements in proper order, filling missing with ""
    max_idx = max(element_dict.keys(), default=0)
    for i in range(1, max_idx + 1):
        elements.append(element_dict.get(i, ""))

    return "*".join(elements) + "~"


# ---------------------------
# Example Usage
# ---------------------------
if __name__ == "__main__":
    # Load Excel keys
    excel_data = load_excel("claims_data.xlsx")
    print(excel_data)

    # Example EDI parsing
    xml_str = """
    <seg id="CLM">
        <ele id="CLM01">12345</ele>
        <comp id="CLM02">
            <subele id="CLM02-1">100</subele>
            <subele id="CLM02-2">50</subele>
        </comp>
    </seg>
    """
    seg = ET.fromstring(xml_str)
    parsed = parse_segment(seg, "030240928")
    print("Parsed EDI Segment:", parsed)








import pandas as pd
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import pyarrow.parquet as pq

# Paths
EXCEL_PATH = r"C:\Users\VIKASPK\Documents\Prod_claims.xlsx"
PARQUET_PATH = "claims.parquet"
INDEX_PATH = "claims.index"
ID_MAP_PATH = "id_map.npy"
TEXT_COLUMNS = ["col1", "col2", "col3"]  # change to your text columns

# Step 1: Load Excel only once and save to parquet
df = pd.read_excel(EXCEL_PATH)
df = df.fillna("")
df.to_parquet(PARQUET_PATH)  # much faster to reload later
texts = df[TEXT_COLUMNS].astype(str).agg(" ".join, axis=1).tolist()
print(f"Loaded {len(texts)} rows")

# Step 2: Load embedding model
model = SentenceTransformer("BAAI/bge-small-en")
dim = 384

# Step 3: Build FAISS index
index = faiss.IndexFlatIP(dim)
id_map = []

# Use bigger batch size (faster)
BATCH_SIZE = 5000

for start in range(0, len(texts), BATCH_SIZE):
    end = min(start + BATCH_SIZE, len(texts))
    batch_texts = texts[start:end]
    embeddings = model.encode(batch_texts, 
                              batch_size=256,  # 👈 use larger batch
                              normalize_embeddings=True, 
                              show_progress_bar=True)
    index.add(np.array(embeddings, dtype="float32"))
    id_map.extend(range(start, end))
    print(f"✅ Processed {start} → {end}")

faiss.write_index(index, INDEX_PATH)
np.save(ID_MAP_PATH, np.array(id_map))
print("🎉 Index built and saved.")











import streamlit as st
import faiss
import numpy as np
import pandas as pd
from sentence_transformers import SentenceTransformer
import time

# === Load Model, Index, Data ===
st.title("Excel Hybrid Semantic Search (FAISS + Filters)")

model = SentenceTransformer("BAAI/bge-small-en")
index = faiss.read_index("claims.index")
id_map = np.load("id_map.npy")
df = pd.read_parquet("claims.parquet")

st.success(f"✅ Loaded {len(df)} rows with FAISS index")

# === User Inputs ===
query = st.text_input("Enter your search query:")
top_k = st.slider("Top K results", 1, 20, 10)

# Optional filters
st.subheader("Filters (optional)")
member_id = st.text_input("Filter by MemberID (leave empty for all)")
min_amount = st.number_input("Minimum Amount", min_value=0, value=0)
max_amount = st.number_input("Maximum Amount", min_value=0, value=1000000)

if query:
    start = time.time()

    # === Apply Filters First ===
    filtered_df = df.copy()
    if member_id:
        filtered_df = filtered_df[filtered_df["MemberID"].astype(str) == member_id]
    if "Amount" in filtered_df.columns:
        filtered_df = filtered_df[
            (filtered_df["Amount"].astype(float) >= min_amount) &
            (filtered_df["Amount"].astype(float) <= max_amount)
        ]

    if filtered_df.empty:
        st.warning("⚠️ No rows match the filters!")
    else:
        st.write(f"🔎 {len(filtered_df)} rows after filtering")

        # === Embed Query ===
        q_emb = model.encode([query], normalize_embeddings=True)
        q_emb = np.array(q_emb, dtype="float32")

        # === Get candidate indices from filtered subset ===
        candidate_ids = filtered_df.index.tolist()
        embeddings_subset = index.reconstruct_n(0, len(id_map))  # load all embeddings
        embeddings_subset = embeddings_subset[np.isin(id_map, candidate_ids)]

        # === Build temporary FAISS index for filtered set ===
        sub_index = faiss.IndexFlatIP(embeddings_subset.shape[1])
        sub_index.add(embeddings_subset)

        # === Run Search ===
        scores, idxs = sub_index.search(q_emb, top_k)

        # Map results back to original DataFrame
        results = []
        for i, score in zip(idxs[0], scores[0]):
            row = filtered_df.iloc[i].to_dict()
            row["score"] = float(score)
            results.append(row)

        end = time.time()

        # === Display ===
        st.write(f"⏱️ Search took {end - start:.2f} seconds")
        st.write(pd.DataFrame(results))










import pandas as pd

# Load Excel files
active_df = pd.read_excel('active_members.xlsx')
non_active_df = pd.read_excel('non_active_members.xlsx')

# Drop duplicates (based on all columns or subset)
active_df = active_df.drop_duplicates()
non_active_df = non_active_df.drop_duplicates()

# Create an empty list to store replacements
replacements = []

# Iterate over non-active members
for idx, non_active in non_active_df.iterrows():
    # Find matching active members based on Age and Gender
    matches = active_df[
        (active_df['Age'] == non_active['Age']) &
        (active_df['Gender'] == non_active['Gender'])
    ]
    
    if not matches.empty:
        # If more than one match, you can choose to take all or the first
        for _, match in matches.iterrows():
            replacements.append({
                'NonActiveID': non_active['MemberID'],  # or the column name in non-active
                'ReplacementID': match['MemberID'],    # or the column name in active
                'Age': non_active['Age'],
                'Gender': non_active['Gender']
            })

# Convert replacements to DataFrame
replacements_df = pd.DataFrame(replacements)

# Save to Excel
replacements_df.to_excel('replaced_members.xlsx', index=False)

print("Replacement process completed. Check 'replaced_members.xlsx'.")









import pandas as pd
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import os

# === Settings ===
EXCEL_PATH = r"C:\Users\VIKASPK\Documents\Prod_claims.xlsx"
INDEX_FILE = r"C:\Users\VIKASPK\Documents\faiss_index.bin"
ID_MAP_FILE = r"C:\Users\VIKASPK\Documents\id_map.npy"
DATA_FILE = r"C:\Users\VIKASPK\Documents\claims.parquet"
BATCH_SIZE = 10000

# === Load Excel ===
df = pd.read_excel(EXCEL_PATH)
print(f"Loaded {len(df)} rows, columns: {df.columns.tolist()}")

# === Auto-detect text + numeric columns for embeddings ===
text_cols = df.select_dtypes(include=["object"]).columns.tolist()
numeric_cols = df.select_dtypes(include=["int64", "float64"]).columns.tolist()

print("Text columns:", text_cols)
print("Numeric columns:", numeric_cols)

# Combine all columns into one string per row
df["combined_text"] = df[text_cols + numeric_cols].astype(str).agg(" | ".join, axis=1)
texts = df["combined_text"].tolist()
print(f"Prepared {len(texts)} rows for embeddings")

# === Load BGE model ===
model = SentenceTransformer("BAAI/bge-small-en")  # 384-dim
dim = 384

# === Create FAISS index ===
index = faiss.IndexFlatIP(dim)  # cosine similarity
id_map = []

# === Batch embedding and add to FAISS ===
for start in range(0, len(texts), BATCH_SIZE):
    end = min(start + BATCH_SIZE, len(texts))
    batch_texts = texts[start:end]
    embeddings = model.encode(batch_texts, normalize_embeddings=True, batch_size=64)
    index.add(np.array(embeddings, dtype="float32"))
    id_map.extend(range(start, end))
    print(f"Indexed rows {start} → {end}")

# === Save index, id_map, and data ===
faiss.write_index(index, INDEX_FILE)
np.save(ID_MAP_FILE, np.array(id_map))
df.to_parquet(DATA_FILE)
print("✅ FAISS index and data saved successfully")













import streamlit as st
import pandas as pd
import faiss
import numpy as np
import time
import os
import pickle
from sentence_transformers import SentenceTransformer

# ======================
# Config (cache files)
# ======================
INDEX_FILE = "faiss_index.bin"
DF_FILE = "data.pkl"

st.title("🔍 Excel Semantic Search with BGE + FAISS (Chunked)")

# Upload Excel option
uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx", "xls"])

# Enter Excel path option
uploaded_file1 = st.text_input("Or enter the Excel file path")

df = None
if uploaded_file or uploaded_file1:
    if uploaded_file:
        df = pd.read_excel(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file1)

if df is not None:
    st.success(f"✅ Loaded {len(df)} rows from Excel")

    # Pick text column
    text_column = st.selectbox("Select the text column to embed", df.columns)

    # Load model
    @st.cache_resource
    def load_model():
        return SentenceTransformer("BAAI/bge-small-en")

    model = load_model()
    dimension = model.get_sentence_embedding_dimension()

    texts = df[text_column].astype(str).tolist()

    # ======================
    # Load or Build Index
    # ======================
    if os.path.exists(INDEX_FILE) and os.path.exists(DF_FILE):
        st.info("📂 Found cached index, loading from disk...")

        index = faiss.read_index(INDEX_FILE)

        with open(DF_FILE, "rb") as f:
            df_cached = pickle.load(f)

        if len(df) != len(df_cached) or not df.equals(df_cached):
            st.warning("⚠️ Uploaded file differs from cached one → rebuilding index...")
            rebuild = True
        else:
            rebuild = False
    else:
        rebuild = True

    if rebuild:
        st.info("🔄 Building FAISS index in chunks... (can handle millions of rows)")

        index = faiss.IndexFlatIP(dimension)  # cosine similarity (normalize vectors)

        batch_size = 5000  # number of rows per chunk (tune this if memory issues)
        for start in range(0, len(texts), batch_size):
            end = min(start + batch_size, len(texts))
            batch_texts = texts[start:end]

            embeddings = model.encode(batch_texts, batch_size=128, show_progress_bar=True)
            embeddings = np.array(embeddings, dtype="float32")
            faiss.normalize_L2(embeddings)

            index.add(embeddings)
            st.write(f"✅ Indexed rows {start}–{end}")

        # Save to disk
        faiss.write_index(index, INDEX_FILE)
        with open(DF_FILE, "wb") as f:
            pickle.dump(df, f)

        st.success("✅ Index built and cached on disk!")

    # ======================
    # Query Search
    # ======================
    query = st.text_input("Enter your search query (e.g. 'claims with amount 1200')")

    if query:
        start = time.time()
        query_emb = model.encode([query])
        faiss.normalize_L2(query_emb)
        D, I = index.search(np.array(query_emb, dtype="float32"), 10)
        end = time.time()

        st.write(f"⚡ Search took **{end - start:.4f} seconds**")

        st.write("### Top 10 Matches")
        for idx, score in zip(I[0], D[0]):
            row_data = df.iloc[idx].to_dict()
            st.write(f"**Score:** {score:.4f}")
            st.json(row_data)















import streamlit as st
import pandas as pd
import faiss
import numpy as np
import time
import os
import pickle
from sentence_transformers import SentenceTransformer

# ======================
# Config (cache files)
# ======================
INDEX_FILE = "faiss_index.bin"
EMB_FILE = "embeddings.pkl"
DF_FILE = "data.pkl"

st.title("🔍 Excel Semantic Search with BGE + FAISS")

# Upload Excel
uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx", "xls"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)

    # Pick text column
    text_column = st.selectbox("Select the text column to embed", df.columns)

    # Load model
    @st.cache_resource
    def load_model():
        return SentenceTransformer("BAAI/bge-small-en")

    model = load_model()

    texts = df[text_column].astype(str).tolist()
    st.write(f"✅ Loaded {len(texts)} rows")

    # ======================
    # Load or Build Index
    # ======================
    if os.path.exists(INDEX_FILE) and os.path.exists(EMB_FILE) and os.path.exists(DF_FILE):
        st.info("📂 Found cached index, loading from disk...")

        index = faiss.read_index(INDEX_FILE)

        with open(EMB_FILE, "rb") as f:
            embeddings = pickle.load(f)

        with open(DF_FILE, "rb") as f:
            df_cached = pickle.load(f)

        # If uploaded data is different, rebuild index
        if len(df) != len(df_cached) or not df.equals(df_cached):
            st.warning("⚠️ Uploaded file differs from cached one → rebuilding index...")
            rebuild = True
        else:
            rebuild = False
    else:
        rebuild = True

    if rebuild:
        st.info("🔄 Building FAISS index... (this may take time for millions of rows)")

        embeddings = model.encode(texts, batch_size=128, show_progress_bar=True)
        embeddings = np.array(embeddings, dtype="float32")
        faiss.normalize_L2(embeddings)

        dimension = embeddings.shape[1]
        index = faiss.IndexFlatIP(dimension)
        index.add(embeddings)

        # Save everything
        faiss.write_index(index, INDEX_FILE)
        with open(EMB_FILE, "wb") as f:
            pickle.dump(embeddings, f)
        with open(DF_FILE, "wb") as f:
            pickle.dump(df, f)

        st.success("✅ Index built and cached on disk!")

    # ======================
    # Query Search
    # ======================
    query = st.text_input("Enter your search query (e.g. 'claims with amount 1200')")

    if query:
        start = time.time()
        query_emb = model.encode([query])
        faiss.normalize_L2(query_emb)
        D, I = index.search(np.array(query_emb, dtype="float32"), 10)
        end = time.time()

        st.write(f"⚡ Search took **{end - start:.4f} seconds**")

        st.write("### Top 10 Matches")
        for idx, score in zip(I[0], D[0]):
            row_data = df.iloc[idx].to_dict()
            st.write(f"**Score:** {score:.4f}")
            st.json(row_data)











import streamlit as st
import pandas as pd
import faiss
import numpy as np
import time
import os
import pickle
from sentence_transformers import SentenceTransformer

# ======================
# Config
# ======================
INDEX_FILE = "faiss_index.bin"
EMB_FILE = "embeddings.pkl"
DF_FILE = "data.pkl"

st.title("Excel Semantic Search with BGE Embeddings + FAISS (Cached)")

# Upload Excel
uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx", "xls"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)

    # Pick text column
    text_column = st.selectbox("Select the text column to embed", df.columns)

    # Load model
    @st.cache_resource
    def load_model():
        return SentenceTransformer("BAAI/bge-small-en")

    model = load_model()

    texts = df[text_column].astype(str).tolist()
    st.write(f"✅ Loaded {len(texts)} rows")

    # ======================
    # Load or Build Index
    # ======================
    if os.path.exists(INDEX_FILE) and os.path.exists(EMB_FILE) and os.path.exists(DF_FILE):
        st.info("📂 Found cached index, loading from disk...")

        index = faiss.read_index(INDEX_FILE)

        with open(EMB_FILE, "rb") as f:
            embeddings = pickle.load(f)

        with open(DF_FILE, "rb") as f:
            df_cached = pickle.load(f)

        # If uploaded data is different, rebuild index
        if len(df) != len(df_cached) or not df.equals(df_cached):
            st.warning("⚠️ Uploaded file differs from cached one → rebuilding index...")
            rebuild = True
        else:
            rebuild = False
    else:
        rebuild = True

    if rebuild:
        st.info("🔄 Building FAISS index... this may take time for millions of rows")

        embeddings = model.encode(texts, batch_size=128, show_progress_bar=True)
        embeddings = np.array(embeddings, dtype="float32")
        faiss.normalize_L2(embeddings)

        dimension = embeddings.shape[1]
        index = faiss.IndexFlatIP(dimension)
        index.add(embeddings)

        # Save everything
        faiss.write_index(index, INDEX_FILE)
        with open(EMB_FILE, "wb") as f:
            pickle.dump(embeddings, f)
        with open(DF_FILE, "wb") as f:
            pickle.dump(df, f)

        st.success("✅ Index built and cached on disk!")

    # ======================
    # Query Search
    # ======================
    query = st.text_input("Enter your search query:")

    if query:
        start = time.time()
        query_emb = model.encode([query])
        faiss.normalize_L2(query_emb)
        D, I = index.search(np.array(query_emb, dtype="float32"), 10)
        end = time.time()

        st.write(f"🔍 Search took **{end - start:.4f} seconds**")

        st.write("### Top 10 Matches")
        for idx, score in zip(I[0], D[0]):
            row_data = df.iloc[idx].to_dict()
            st.write(f"**Score:** {score:.4f}")
            st.json(row_data)













import streamlit as st
import pandas as pd
import numpy as np
import faiss
import tempfile
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.schema import Document
from pdf2image import convert_from_path
import pytesseract

st.set_page_config(page_title="PDF Service Matcher", layout="wide")

st.title("PDF Service Matcher with Excel")

# ------------------------------
# Step 1: Load Excel B and create FAISS index
# ------------------------------
@st.cache_data
def load_excel_and_build_index(excel_path):
    df = pd.read_excel(excel_path)
    
    embedding_model = HuggingFaceBgeEmbeddings(
        model_name="BAAI/bge-large-en",
        model_kwargs={"device": "cpu"},  # or "cuda" if GPU available
        encode_kwargs={"normalize_embeddings": True}  # cosine similarity
    )
    
    service_names = df["Service_Name"].astype(str).tolist()
    vectors = embedding_model.embed_documents(service_names)
    vectors = np.array(vectors).astype("float32")
    
    dim = vectors.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(vectors)
    
    # Mapping FAISS index to Excel row
    index_to_row = {i: df.iloc[i] for i in range(len(df))}
    
    return df, embedding_model, index, index_to_row

st.info("Load Excel B (large service dataset)")
excel_file = st.file_uploader("Upload Excel B", type=["xlsx"])
if excel_file:
    with st.spinner("Building FAISS index..."):
        df_b, embedding_model, faiss_index, index_to_row = load_excel_and_build_index(excel_file)
    st.success(f"Loaded {len(df_b)} services from Excel B and built FAISS index.")

# ------------------------------
# Step 2: PDF Upload and OCR
# ------------------------------
def fallback_ocr_loader(pdf_path):
    images = convert_from_path(pdf_path)
    docs = []
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img)
            page_text = f"\n[Page {i+1}]\n{text.strip()}"
            docs.append(Document(page_content=page_text))
        except Exception as e:
            st.warning(f"OCR failed on page {i+1}: {e}")
    return docs

def extract_text_from_pdf(uploaded_pdf):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_pdf.read())
        tmp_path = tmp.name
    
    # Try normal text extraction
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(tmp_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        if text.strip() == "":
            raise ValueError("No text layer")
        docs = [Document(page_content=text)]
    except:
        # Use OCR fallback
        docs = fallback_ocr_loader(tmp_path)
    return docs

uploaded_pdfs = st.file_uploader("Upload PDF Contracts", type=["pdf"], accept_multiple_files=True)

# ------------------------------
# Step 3: Similarity search against Excel B
# ------------------------------
def search_services_in_excel(pdf_services, embedding_model, faiss_index, index_to_row, top_k=5, threshold=0.75):
    query_vectors = embedding_model.embed_documents(pdf_services)
    query_vectors = np.array(query_vectors).astype("float32")
    
    scores, indices = faiss_index.search(query_vectors, top_k)
    
    results = []
    for q_idx, idx_list in enumerate(indices):
        matched_rows = []
        for score, idx in zip(scores[q_idx], idx_list):
            if score < threshold or idx == -1:
                continue
            row = index_to_row[idx]
            matched_rows.append({
                "PDF_Service": pdf_services[q_idx],
                "Matched_Service": row["Service_Name"],
                "CPT_Code": row.get("CPT_Code", None),
                "Similarity_Score": score
            })
        if not matched_rows:
            matched_rows.append({
                "PDF_Service": pdf_services[q_idx],
                "Matched_Service": None,
                "CPT_Code": None,
                "Similarity_Score": None
            })
        results.extend(matched_rows)
    return pd.DataFrame(results)

# ------------------------------
# Step 4: Process PDFs and display results
# ------------------------------
if uploaded_pdfs and excel_file:
    all_results = []
    for pdf in uploaded_pdfs:
        with st.spinner(f"Processing {pdf.name}..."):
            docs = extract_text_from_pdf(pdf)
            pdf_services = []
            for doc in docs:
                # Simple split by lines or customize your extraction logic
                lines = [line.strip() for line in doc.page_content.split("\n") if line.strip()]
                pdf_services.extend(lines)
            if not pdf_services:
                st.warning(f"No services found in {pdf.name}")
                continue
            
            df_results = search_services_in_excel(pdf_services, embedding_model, faiss_index, index_to_row)
            df_results["PDF_File"] = pdf.name
            all_results.append(df_results)
    
    if all_results:
        final_df = pd.concat(all_results, ignore_index=True)
        st.success("All PDFs processed successfully!")
        st.dataframe(final_df)
        
        # Download
        output_file = "matched_services.xlsx"
        final_df.to_excel(output_file, index=False)
        with open(output_file, "rb") as f:
            st.download_button(
                "Download Results as Excel",
                data=f,
                file_name=output_file,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )










..import streamlit as st
import pandas as pd
import numpy as np
import faiss
import re
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from sklearn.metrics.pairwise import cosine_similarity

# -----------------------
# Config
# -----------------------
MODEL_NAME = "BAAI/bge-small-en"  # or "BAAI/bge-m3"
INDEX_FILE = "claims_index.faiss"

model_kwargs = {"device": "cpu"}  # or "cuda"
encode_kwargs = {"normalize_embeddings": True}

embedding_model = HuggingFaceBgeEmbeddings(
    model_name=MODEL_NAME,
    model_kwargs=model_kwargs,
    encode_kwargs=encode_kwargs
)

# -----------------------
# Helpers
# -----------------------

def row_to_text(row):
    return ", ".join([f"{col}: {row[col]}" for col in row.index])

def build_index(df):
    global index, id_to_row

    df = df.astype(str)
    texts = df.apply(row_to_text, axis=1).tolist()
    st.info(f"Generating embeddings for {len(texts)} rows...")
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings, dtype="float32")

    d = embeddings.shape[1]
    index = faiss.IndexFlatL2(d)
    index.add(embeddings)

    id_to_row = {i: idx for i, idx in enumerate(df.index)}
    faiss.write_index(index, INDEX_FILE)
    st.success("Index built & saved.")

def load_index():
    global df, index, id_to_row
    df = df.astype(str)
    index = faiss.read_index(INDEX_FILE)
    id_to_row = {i: idx for i, idx in enumerate(df.index)}

def detect_columns(query, df):
    """Dynamically detect relevant columns using embeddings."""
    query_emb = embedding_model.embed_query(query)
    col_embeddings = [embedding_model.embed_query(col) for col in df.columns]
    sims = cosine_similarity([query_emb], col_embeddings)[0]
    top_cols = [df.columns[i] for i in np.argsort(sims)[::-1] if sims[i] > 0.5]
    return top_cols

def hybrid_search(query, top_k=10):
    """Generic search combining semantic + structured filtering"""
    values = re.findall(r"[A-Za-z0-9\.\-]+", query)
    candidate_cols = detect_columns(query, df)
    matches = {}

    # Structured filtering
    for col in candidate_cols:
        mask = df[col].astype(str).str.contains("|".join(values), case=False, na=False)
        sub_df = df[mask]
        if not sub_df.empty:
            matches[col] = sub_df

    # Fall back to semantic FAISS search
    if not matches:
        q_embed = np.array([embedding_model.embed_query(query)], dtype="float32")
        distances, indices = index.search(q_embed, top_k)
        sem_results = [df.iloc[id_to_row[idx]] for idx in indices[0]]
        return pd.DataFrame(sem_results)

    # Return structured matches
    return matches

# -----------------------
# Streamlit UI
# -----------------------

st.title("Claims Vector Search with HuggingFace BGE + FAISS")

uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx", "csv"])

if uploaded_file:
    if uploaded_file.name.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)
    else:
        df = pd.read_csv(uploaded_file)

    st.write("Preview of your data:")
    st.dataframe(df.head())

    if st.button("Build / Rebuild Index"):
        build_index(df)

    st.write("---")
    query = st.text_input("Enter your search query (natural language)")

    if query:
        results = hybrid_search(query)
        st.write("### Search Results:")

        if isinstance(results, dict):
            for col, sub_df in results.items():
                st.write(f"**Matches in column `{col}`:**")
                st.dataframe(sub_df)
        else:
            st.dataframe(results)
















import pandas as pd
import faiss
import numpy as np
from langchain_community.embeddings import HuggingFaceBgeEmbeddings

# -----------------------
# 🔹 Config
# -----------------------
MODEL_NAME = "BAAI/bge-small-en"   # or "BAAI/bge-m3" for multilingual
INDEX_FILE = "claims_index.faiss"

# -----------------------
# 🔹 Initialize embedding model
# -----------------------
model_kwargs = {"device": "cpu"}   # change to "cuda" if GPU available
encode_kwargs = {"normalize_embeddings": True}

embedding_model = HuggingFaceBgeEmbeddings(
    model_name=MODEL_NAME,
    model_kwargs=model_kwargs,
    encode_kwargs=encode_kwargs
)

# -----------------------
# 🔹 Convert row to text
# -----------------------
def row_to_text(row):
    return f"Claim Number: {row['Claims number']}, POS: {row['pos']}, CPT Codes: {row['cpt codes']}"

# -----------------------
# 🔹 Build FAISS index
# -----------------------
def build_index(excel_file):
    global df, index, id_to_row

    # Load Excel
    df = pd.read_excel(excel_file)
    df = df[["Claims number", "pos", "cpt codes"]].astype(str)

    # Convert rows into text
    texts = df.apply(row_to_text, axis=1).tolist()

    # Generate embeddings
    print(f"Generating embeddings for {len(texts)} rows...")
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings, dtype="float32")

    # Create FAISS index
    d = embeddings.shape[1]
    index = faiss.IndexFlatL2(d)
    index.add(embeddings)

    # Mapping FAISS IDs -> DataFrame index
    id_to_row = {i: idx for i, idx in enumerate(df.index)}

    # Save FAISS index
    faiss.write_index(index, INDEX_FILE)
    print(f"✅ Index built & saved to {INDEX_FILE}")

# -----------------------
# 🔹 Load FAISS index
# -----------------------
def load_index(excel_file):
    global df, index, id_to_row

    df = pd.read_excel(excel_file)
    df = df[["Claims number", "pos", "cpt codes"]].astype(str)

    index = faiss.read_index(INDEX_FILE)
    id_to_row = {i: idx for i, idx in enumerate(df.index)}
    print(f"✅ Index loaded from {INDEX_FILE}")

# -----------------------
# 🔹 Search Query
# -----------------------
def search_query(query, top_k=5):
    q_embed = np.array([embedding_model.embed_query(query)], dtype="float32")
    distances, indices = index.search(q_embed, top_k)

    results = []
    for idx in indices[0]:
        row = df.iloc[id_to_row[idx]]
        results.append(row)
    return pd.DataFrame(results)

# -----------------------
# 🔹 Update FAISS with new rows
# -----------------------
def update_faiss(new_excel_file):
    global df, index, id_to_row

    # Load new data
    new_df = pd.read_excel(new_excel_file)
    new_df = new_df[["Claims number", "pos", "cpt codes"]].astype(str)

    # Convert to text + embeddings
    new_texts = new_df.apply(row_to_text, axis=1).tolist()
    new_embeddings = embedding_model.embed_documents(new_texts)
    new_embeddings = np.array(new_embeddings, dtype="float32")

    # Add to FAISS
    index.add(new_embeddings)

    # Update mapping
    start_id = len(id_to_row)
    for i, idx in enumerate(new_df.index):
        id_to_row[start_id + i] = idx

    # Merge DataFrames
    df = pd.concat([df, new_df], ignore_index=True)

    # Save updated index
    faiss.write_index(index, INDEX_FILE)
    print(f"✅ Index updated & saved to {INDEX_FILE}")

# -----------------------
# 🔹 Example Usage
# -----------------------
if __name__ == "__main__":
    excel_file = "claims_data.xlsx"

    # Step 1: Build index from scratch
    build_index(excel_file)

    # Step 2: Search
    query = "Find claims with CPT code 99213 in POS 11"
    results = search_query(query, top_k=5)
    print("\n🔍 Search Results:")
    print(results)

    # Step 3: Update index with new data (optional)
    # update_faiss("new_claims.xlsx")












import pandas as pd
import faiss
import torch
from transformers import AutoTokenizer, AutoModel

# -----------------------
# 🔹 Config
# -----------------------
MODEL_NAME = "BAAI/bge-small-en"  # or "BAAI/bge-m3" for multilingual
INDEX_FILE = "claims_index.faiss"

# -----------------------
# 🔹 Load HuggingFace BGE model
# -----------------------
print("Loading embedding model...")
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModel.from_pretrained(MODEL_NAME)

# -----------------------
# 🔹 Embedding function
# -----------------------
def get_embeddings(sentences, batch_size=32):
    all_embeddings = []
    for i in range(0, len(sentences), batch_size):
        batch = sentences[i:i+batch_size]
        encoded_input = tokenizer(batch, padding=True, truncation=True, return_tensors="pt")
        with torch.no_grad():
            model_output = model(**encoded_input)
        # Mean pooling
        embeddings = model_output.last_hidden_state.mean(dim=1)
        all_embeddings.append(embeddings)
    return torch.cat(all_embeddings, dim=0).cpu().numpy()

# -----------------------
# 🔹 Convert row -> text
# -----------------------
def row_to_text(row):
    return f"Claim Number: {row['Claims number']}, POS: {row['pos']}, CPT Codes: {row['cpt codes']}"

# -----------------------
# 🔹 Build FAISS index
# -----------------------
def build_index(excel_file):
    global df, index, id_to_row

    df = pd.read_excel(excel_file)
    df = df[["Claims number", "pos", "cpt codes"]].astype(str)

    texts = df.apply(row_to_text, axis=1).tolist()
    print(f"Generating embeddings for {len(texts)} rows...")
    embeddings = get_embeddings(texts, batch_size=64)

    d = embeddings.shape[1]
    index = faiss.IndexFlatL2(d)
    index.add(embeddings)

    id_to_row = {i: idx for i, idx in enumerate(df.index)}

    faiss.write_index(index, INDEX_FILE)
    print(f"✅ Index built & saved to {INDEX_FILE}")

# -----------------------
# 🔹 Load FAISS index
# -----------------------
def load_index(excel_file):
    global df, index, id_to_row

    df = pd.read_excel(excel_file)
    df = df[["Claims number", "pos", "cpt codes"]].astype(str)

    index = faiss.read_index(INDEX_FILE)
    id_to_row = {i: idx for i, idx in enumerate(df.index)}
    print(f"✅ Index loaded from {INDEX_FILE}")

# -----------------------
# 🔹 Search Query
# -----------------------
def search_query(query, top_k=5):
    q_embed = get_embeddings([query])
    distances, indices = index.search(q_embed, top_k)

    results = []
    for idx in indices[0]:
        row = df.iloc[id_to_row[idx]]
        results.append(row)
    return pd.DataFrame(results)

# -----------------------
# 🔹 Update Index
# -----------------------
def update_faiss(new_excel_file):
    global df, index, id_to_row

    new_df = pd.read_excel(new_excel_file)
    new_df = new_df[["Claims number", "pos", "cpt codes"]].astype(str)

    new_texts = new_df.apply(row_to_text, axis=1).tolist()
    new_embeds = get_embeddings(new_texts, batch_size=64)

    index.add(new_embeds)

    start_id = len(id_to_row)
    for i, idx in enumerate(new_df.index):
        id_to_row[start_id + i] = idx

    df = pd.concat([df, new_df], ignore_index=True)
    faiss.write_index(index, INDEX_FILE)
    print(f"✅ Index updated & saved to {INDEX_FILE}")

# -----------------------
# 🔹 Example Usage
# -----------------------
if __name__ == "__main__":
    excel_file = "claims_data.xlsx"

    # 1. Build index from scratch
    build_index(excel_file)

    # 2. Search a query
    query = "Find claims with CPT code 99213 in POS 11"
    results = search_query(query, top_k=10)
    print("\n🔍 Search Results:")
    print(results)

    # 3. Update with new file (optional)
    # update_faiss("new_claims.xlsx")















import streamlit as st
import pandas as pd
import os, json, zipfile, math
from io import StringIO, BytesIO
from pyx12.x12n_document import x12n_document
from pyx12.params import Params
from pyx12.error_handler import ErrorHandler
from langchain_groq import ChatGroq
from langchain.prompts import ChatPromptTemplate

# --- Helper: validate edi ---
def validate_x12(edi_content):
    errors = []
    params = Params()
    eh = ErrorHandler()
    edi_buffer = StringIO(edi_content)
    try:
        x12n_document(params, edi_buffer,
                      fd_997=None, fd_html=None,
                      fd_xmldoc=None, map_path=None, errh=eh)
        if eh.has_errors():
            return False, [str(e) for e in eh.get_errors()]
        return True, []
    except Exception as e:
        return False, [str(e)]

# --- Streamlit UI ---
st.title("EDI Bulk Modifier with Groq + Auto-chunking")

excel_file = st.file_uploader("Upload Excel with modification queries", type=["xlsx"])
edi_file = st.file_uploader("Upload Base EDI File", type=["edi","txt"])

if excel_file and edi_file:
    df = pd.read_excel(excel_file)
    queries = df["ChangeQuery"].tolist()
    edi_text = edi_file.read().decode("utf-8")

    chunk_size = st.number_input("Rows per LLM call (chunk size)", min_value=5, max_value=50, value=20)
    
    if st.button("Generate EDIs"):
        st.info("Processing... this may take a while ⏳")
        os.makedirs("output/valid", exist_ok=True)
        os.makedirs("output/invalid", exist_ok=True)
        
        llm = ChatGroq(model="llama3-70b-8192", temperature=0, max_tokens=4000)
        all_results = []

        total_rows = len(queries)
        num_chunks = math.ceil(total_rows / chunk_size)

        for i in range(num_chunks):
            start = i * chunk_size
            end = min((i + 1) * chunk_size, total_rows)
            chunk_queries = [{"row": j+1, "query": queries[j]} for j in range(start, end)]

            prompt = ChatPromptTemplate.from_template("""
            You are an EDI expert.
            Given an EDI file and modification instructions, update the EDI for each row.
            Return JSON like:
            {
             "results": [
               {"row": <row>, "edi_updated": "<edi>", "summary": "<what changed>"}
             ]
            }
            EDI File:
            {edi_file}

            Instructions:
            {instructions}
            """)

            response = llm.invoke(
                prompt.format_messages(
                    edi_file=edi_text,
                    instructions=json.dumps(chunk_queries)
                )
            )

            # Parse LLM JSON
            try:
                data_str = response.content.strip().replace("```json","").replace("```","")
                data_chunk = json.loads(data_str)
                all_results.extend(data_chunk["results"])
            except Exception as e:
                st.error(f"❌ Failed to parse LLM response for rows {start+1}-{end}: {e}")
                st.text(response.content)
                continue

        # --- Validate and save ---
        results_log = []
        for result in all_results:
            row = result["row"]
            edi_updated = result["edi_updated"]
            summary = result.get("summary", "")
            valid, errors = validate_x12(edi_updated)

            if valid:
                filepath = f"output/valid/row_{row}.edi"
                with open(filepath, "w") as f:
                    f.write(edi_updated)
                results_log.append({"Row": row, "Query": queries[row-1], "Status": "Valid", "Summary": summary})
            else:
                filepath = f"output/invalid/row_{row}.edi"
                with open(filepath, "w") as f:
                    f.write(edi_updated)
                with open(f"output/invalid/row_{row}_reason.txt","w") as f:
                    f.write("\n".join(errors))
                results_log.append({
                    "Row": row,
                    "Query": queries[row-1],
                    "Status": "Invalid",
                    "Summary": summary,
                    "Errors": "; ".join(errors)
                })

        # --- Show summary ---
        st.success("✅ Processing complete")
        st.dataframe(pd.DataFrame(results_log))

        # --- Download ZIP ---
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for folder in ["valid","invalid"]:
                folder_path = os.path.join("output", folder)
                for fname in os.listdir(folder_path):
                    zf.write(os.path.join(folder_path,fname), arcname=f"{folder}/{fname}")
        zip_buffer.seek(0)

        st.download_button(
            label="Download All Results (ZIP)",
            data=zip_buffer,
            file_name="edi_output.zip",
            mime="application/zip"
        )











import streamlit as st
import pandas as pd
import os, json, zipfile
from io import StringIO, BytesIO
from pyx12.x12n_document import x12n_document
from pyx12.params import Params
from pyx12.error_handler import ErrorHandler
from langchain_openai import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate

# --- Helper: validate edi ---
def validate_x12(edi_content):
    """Validate using pyx12 and return (is_valid, errors)."""
    errors = []
    params = Params()
    eh = ErrorHandler()
    edi_buffer = StringIO(edi_content)
    try:
        x12n_document(params, edi_buffer, fd_997=None, fd_html=None, fd_xmldoc=None, map_path=None, errh=eh)
        if eh.has_errors():
            return False, [str(e)]
        return True, []
    except Exception as e:
        return False, [str(e)]

# --- Streamlit UI ---
st.title("EDI Bulk Modifier with LLM + Validation")

excel_file = st.file_uploader("Upload Excel with modification queries", type=["xlsx"])
edi_file = st.file_uploader("Upload Base EDI File", type=["edi","txt"])

if excel_file and edi_file:
    df = pd.read_excel(excel_file)
    queries = df["ChangeQuery"].tolist()
    edi_text = edi_file.read().decode("utf-8")

    if st.button("Generate EDIs"):
        st.info("Processing... this may take some time ⏳")

        # Build input for LLM
        instructions = [{"row": i+1, "query": q} for i, q in enumerate(queries)]
        llm_input = {
            "edi_file": edi_text,
            "instructions": instructions
        }

        # Azure LLM
        llm = AzureChatOpenAI(
            deployment_name="your-deployment",
            temperature=0,
            max_tokens=4000
        )

        prompt = ChatPromptTemplate.from_template("""
        You are an EDI expert.
        Given an EDI file and modification instructions, update the EDI for each row.
        Return JSON like:
        {
         "results": [
           {"row": <row>, "edi_updated": "<edi>", "summary": "<what changed>"}
         ]
        }
        EDI File:
        {edi_file}

        Instructions:
        {instructions}
        """)

        response = llm.invoke(
            prompt.format_messages(
                edi_file=edi_text,
                instructions=json.dumps(instructions)
            )
        )

        # Parse response
        try:
            data_str = response.content.strip().replace("```json","").replace("```","")
            data = json.loads(data_str)
        except Exception as e:
            st.error(f"❌ Failed to parse LLM response: {e}")
            st.text(response.content)
            st.stop()

        # Create output folders
        os.makedirs("output/valid", exist_ok=True)
        os.makedirs("output/invalid", exist_ok=True)

        results_log = []

        for result in data["results"]:
            row = result["row"]
            edi_updated = result["edi_updated"]
            summary = result.get("summary","")

            # Validate
            valid, errors = validate_x12(edi_updated)

            if valid:
                filepath = f"output/valid/row_{row}.edi"
                with open(filepath, "w") as f:
                    f.write(edi_updated)
                results_log.append({"Row": row, "Query": queries[row-1], "Status": "Valid", "Summary": summary})
            else:
                filepath = f"output/invalid/row_{row}.edi"
                with open(filepath, "w") as f:
                    f.write(edi_updated)
                with open(f"output/invalid/row_{row}_reason.txt","w") as f:
                    f.write("\n".join(errors))
                results_log.append({"Row": row, "Query": queries[row-1], "Status": "Invalid", "Summary": summary, "Errors": "; ".join(errors)})

        # Show results table
        st.success("✅ Processing complete")
        st.dataframe(pd.DataFrame(results_log))

        # Create ZIP to download
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for folder in ["valid","invalid"]:
                folder_path = os.path.join("output", folder)
                for fname in os.listdir(folder_path):
                    zf.write(os.path.join(folder_path,fname), arcname=f"{folder}/{fname}")
        zip_buffer.seek(0)

        st.download_button(
            label="Download All Results (ZIP)",
            data=zip_buffer,
            file_name="edi_output.zip",
            mime="application/zip"
        )












import pandas as pd
import os, json
from io import StringIO
from pyx12.x12n_document import x12n_document
from pyx12.params import Params
from pyx12.error_handler import ErrorHandler

# --- Step 1: Read Excel ---
df = pd.read_excel("instructions.xlsx")
queries = df["ChangeQuery"].tolist()

# --- Step 2: Build LLM input ---
edi_text = open("input.edi").read()

instructions = [{"row": i+1, "query": q} for i, q in enumerate(queries)]
llm_input = {
    "edi_file": edi_text,
    "instructions": instructions
}

# --- Step 3: Call LLM ---
from langchain_openai import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate

llm = AzureChatOpenAI(
    deployment_name="your-deployment",
    temperature=0,
    max_tokens=4000
)

prompt = ChatPromptTemplate.from_template("""
You are an EDI expert.
Given an EDI file and modification instructions, update the EDI for each row.
Return JSON like:
{
 "results": [
   {"row": <row>, "edi_updated": "<edi>", "summary": "<what changed>"}
 ]
}
EDI File:
{edi_file}

Instructions:
{instructions}
""")

response = llm.invoke(prompt.format_messages(edi_file=edi_text, instructions=json.dumps(instructions)))

# --- Step 4: Parse response ---
data_str = response.content.strip().replace("```json","").replace("```","")
data = json.loads(data_str)

os.makedirs("output/valid", exist_ok=True)
os.makedirs("output/invalid", exist_ok=True)

# --- Step 5: Validate and Save ---
for result in data["results"]:
    row = result["row"]
    edi_updated = result["edi_updated"]
    summary = result["summary"]

    # Save file
    edi_file_path = f"output/temp_row_{row}.edi"
    with open(edi_file_path, "w") as f:
        f.write(edi_updated)

    # Validate
    try:
        valid, errors = validate_x12(edi_updated)
        if valid:
            with open(f"output/valid/row_{row}.edi", "w") as f:
                f.write(edi_updated)
        else:
            with open(f"output/invalid/row_{row}.edi", "w") as f:
                f.write(edi_updated)
            with open(f"output/invalid/row_{row}_reason.txt", "w") as f:
                f.write("\n".join(errors))
    except Exception as e:
        with open(f"output/invalid/row_{row}_reason.txt", "w") as f:
            f.write(str(e))

# --- Helper: validation function ---
def validate_x12(edi_content):
    """Validate using pyx12 and return (is_valid, errors)."""
    errors = []
    params = Params()
    eh = ErrorHandler()
    edi_buffer = StringIO(edi_content)
    try:
        x12n_document(params, edi_buffer, fd_997=None, fd_html=None, fd_xmldoc=None, map_path=None, errh=eh)
        if eh.has_errors():
            return False, eh.get_errors()
        return True, []
    except Exception as e:
        return False, [str(e)]











import streamlit as st
from langchain.chat_models import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate
import os, json

# ---- Azure OpenAI setup ----
llm = AzureChatOpenAI(
    deployment_name="your-deployment",
    openai_api_key=os.getenv("AZURE_OPENAI_KEY"),
    openai_api_base=os.getenv("AZURE_OPENAI_ENDPOINT"),
    openai_api_version="2023-05-15",
    temperature=0
)

st.title("EDI AI Editor")

uploaded_file = st.file_uploader("Upload EDI file", type=["edi", "txt"])
if uploaded_file:
    edi_content = uploaded_file.read().decode("utf-8")
    st.text_area("EDI Preview", edi_content, height=200)

    # Initialize session state
    if "edi_data" not in st.session_state:
        st.session_state.edi_data = edi_content
    if "change_log" not in st.session_state:
        st.session_state.change_log = []

    query = st.text_input("Enter your edit request (e.g. 'Change service line date')")

    if query and st.button("Apply Change"):
        prompt = ChatPromptTemplate.from_template("""
        You are an EDI expert. 
        Current EDI:
        {edi}

        User request: {query}

        Return your response in JSON with two fields:
        {{
            "edi_updated": "<the updated EDI only>",
            "summary": "<summary of what changed>"
        }}
        """)

        chain = prompt | llm
        result = chain.invoke({"edi": st.session_state.edi_data, "query": query})

        try:
            data = json.loads(result.content)
            st.session_state.edi_data = data["edi_updated"]
            st.session_state.change_log.append(data["summary"])
        except Exception:
            st.error("⚠️ Failed to parse response. Please retry.")
            st.text(result.content)

    # Show change log
    if st.session_state.change_log:
        st.subheader("✅ Change Log")
        for i, change in enumerate(st.session_state.change_log, 1):
            st.markdown(f"**{i}.** {change}")

    # Final download buttons
    if st.button("Generate Final Files"):
        # Download EDI
        st.download_button(
            "⬇️ Download Final EDI",
            st.session_state.edi_data,
            file_name="updated.edi"
        )

        # Download Summary log
        summary_text = "\n".join(
            [f"{i+1}. {c}" for i, c in enumerate(st.session_state.change_log)]
        )
        st.download_button(
            "⬇️ Download Change Summary",
            summary_text,
            file_name="change_summary.txt"
        )














"""
edi_agent.py

Requirements:
- pip install langchain openai pyx12  (pyx12 optional but recommended)
- Set env vars: AZURE_OPENAI_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_DEPLOYMENT
"""

import os
import json
import io
import logging
from typing import List, Tuple, Dict, Any, Optional
from copy import deepcopy

# LangChain / Azure OpenAI
from langchain.chat_models import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate

# --- Utility helpers ---


def find_nth_occurrence(string: str, substring: str, n: int) -> int:
    """
    Return index (0-based) of nth occurrence of substring in string. -1 if not found.
    """
    if n <= 0:
        return -1
    start = 0
    for i in range(n):
        idx = string.find(substring, start)
        if idx == -1:
            return -1
        start = idx + 1
    return idx


def detect_separators(edi_text: str) -> Tuple[str, str]:
    """
    Return (segment_terminator, element_separator). Common defaults are '~' and '*'.
    Try to detect from ISA/GS if possible.
    """
    # Very simple heuristics:
    if '~' in edi_text:
        seg = '~'
    elif '\n' in edi_text:
        seg = '\n'
    else:
        seg = '~'
    if '*' in edi_text:
        elem = '*'
    elif '|' in edi_text:
        elem = '|'
    else:
        elem = '*'
    return seg, elem


def split_segments(edi_text: str, seg_term: str) -> List[str]:
    # Keep trailing separators trimmed
    segments = [s.strip() for s in edi_text.split(seg_term) if s.strip()]
    return segments


def join_segments(segments: List[str], seg_term: str) -> str:
    return seg_term.join(segments) + seg_term


def apply_edits(edi_text: str, edits: List[Dict[str, Any]]) -> Tuple[str, List[str]]:
    """
    Apply a list of edits to the edi_text.

    Each edit item should be structured:
    {
      "segment": "DTP",
      "occurrence": 1,      # optional (1-based). If omitted or null -> apply to all occurrences
      "field_pos": 3,       # 1-based (1 = first element after segment tag)
      "subfield_pos": null, # optional, 1-based for component/subfield
      "match": { "old_value": "20240101" },  # optional condition: only replace where match holds
      "new_value": "20250101"
    }

    Returns (new_text, applied_descriptions) where applied_descriptions are human readable lines of what changed.
    """
    seg_term, elem = detect_separators(edi_text)
    segments = split_segments(edi_text, seg_term)
    applied = []
    new_segments = deepcopy(segments)

    for edit in edits:
        seg_tag = edit.get("segment")
        occ = edit.get("occurrence")  # 1-based
        field_pos = edit.get("field_pos")  # 1-based
        subfield_pos = edit.get("subfield_pos")
        new_value = edit.get("new_value")
        match = edit.get("match")

        # iterate each segment and apply edit where appropriate
        count = 0
        for idx, seg in enumerate(new_segments):
            if not seg:
                continue
            parts = seg.split(elem)
            if parts[0].upper() != seg_tag.upper():
                continue
            count += 1
            if occ and occurrence_mismatch := (count != int(occ)):
                continue

            # field index in parts: parts[1] is first field after tag
            if not field_pos or field_pos < 1:
                # if field_pos missing, we cannot apply reliably
                continue
            arr_idx = int(field_pos)  # 1-based
            parts_len = len(parts) - 1
            if arr_idx > parts_len:
                # extend with empty fields if necessary
                parts.extend([''] * (arr_idx - parts_len))

            # check match condition if exists
            apply_here = True
            if match:
                if match.get("old_value") is not None:
                    existing_val = parts[arr_idx]
                    if existing_val != match.get("old_value"):
                        apply_here = False

            if apply_here:
                old = parts[arr_idx]
                # handle subfield/component splitting if requested (component sep usually ':')
                if subfield_pos:
                    comp_sep = ':'
                    comps = old.split(comp_sep) if old is not None else ['']
                    sf_idx = int(subfield_pos) - 1
                    if sf_idx >= len(comps):
                        comps.extend([''] * (sf_idx - len(comps) + 1))
                    comps[sf_idx] = new_value
                    parts[arr_idx] = comp_sep.join(comps)
                else:
                    parts[arr_idx] = new_value

                new_seg = elem.join(parts)
                new_segments[idx] = new_seg
                applied.append(
                    f"Segment {seg_tag} occurrence {count}: field {field_pos} changed from '{old}' to '{parts[arr_idx]}'"
                )
                # if occurrence provided, stop after applying to that occurrence
                if occ:
                    break

    new_text = join_segments(new_segments, seg_term)
    return new_text, applied


# --- pyx12 validator wrapper ---
def pyx12_validate(edi_text: str, seg_term: Optional[str] = None) -> Tuple[bool, List[str]]:
    """
    Validate edi_text using pyx12 if available.

    Returns (is_valid, errors_list). If pyx12 is not installed or fails, 
    returns a lightweight syntax check result and messages.

    Note: pyx12 APIs differ across versions; this wrapper attempts common calls.
    Please change implementation to match your pyx12 version if needed.
    """
    try:
        # Attempt to import typical pyx12 modules
        import pyx12
        # Many pyx12 installs expose x12n_document in pyx12.x12n_document or pyx12.xml
        try:
            from pyx12.x12n_document import x12n_document
            from pyx12.params import ParamsBase
        except Exception:
            # try alternative import paths
            try:
                from pyx12.x12n_document import x12n_document
                from pyx12.params import ParamsBase
            except Exception:
                # give up pyx12 detailed call and fallback
                raise

        # prepare buffers
        seg_term_detected, elem = detect_separators(edi_text)
        edibuffer = io.StringIO(edi_text)
        xmlbuffer = io.StringIO()
        params = ParamsBase()

        # x12n_document signature varies; common invocation used widely:
        # result = x12n_document(params, edibuffer, fd_997=None, fd_html=None, fd_xmldoc=xmlbuffer, map_path=None)
        result = x12n_document(params, edibuffer, fd_997=None, fd_html=None, fd_xmldoc=xmlbuffer, map_path=None)

        # `result` sometimes is an object with attribute 'errors' or 'error_list'
        errors = []
        if hasattr(result, "errors"):
            # result.errors may be a list of tuples or strings
            for e in result.errors:
                errors.append(str(e))
        elif hasattr(result, "error_list"):
            for e in result.error_list:
                errors.append(str(e))
        else:
            # try to parse xmlbuffer for errors or rely on result truthiness
            xml_value = xmlbuffer.getvalue()
            if "error" in xml_value.lower():
                errors.append(xml_value)
        is_valid = len(errors) == 0
        return is_valid, errors
    except Exception as e:
        # pyx12 not present or failed — fallback simple checks
        fallback_errors = []
        text = edi_text.strip()
        if not text.startswith("ISA"):
            fallback_errors.append("Missing or malformed ISA segment at beginning.")
        if "IEA" not in text:
            fallback_errors.append("Missing IEA segment (interchange trailer).")
        # other basic checks
        if len(fallback_errors) == 0:
            return True, []
        return False, fallback_errors


def is_valid_x12(edi_file_name: str, segments: List[str], log_path: Optional[str] = None) -> Tuple[bool, List[str]]:
    """
    Example wrapper requested in the prompt: given a filename and segments list validate the joined EDI.
    This function writes logs if log_path provided.
    """
    logFormatter = logging.Formatter("%(asctime)s [%(levelname)-5.5s] (%(filename)s) %(message)s")
    logger = logging.getLogger("validation_x12")
    logger.setLevel(logging.INFO)
    # Prevent duplicate handlers
    if not logger.handlers:
        consoleHandler = logging.StreamHandler()
        consoleHandler.setFormatter(logFormatter)
        logger.addHandler(consoleHandler)
        if log_path is not None:
            fh = logging.FileHandler(os.path.join(log_path, "validation_X12.log"))
            fh.setFormatter(logFormatter)
            logger.addHandler(fh)

    edi_text = "\n".join(segments)
    logger.info(f"Validating file {edi_file_name} (length={len(edi_text)} chars)")
    is_valid, errors = pyx12_validate(edi_text)
    if is_valid:
        logger.info("Validation successful.")
    else:
        logger.warning("Validation failed: " + "; ".join(errors))
    return is_valid, errors


# --- LLM prompting ---


def create_llm():
    """
    Create AzureChatOpenAI LLM via LangChain. Expects env vars:
    AZURE_OPENAI_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_DEPLOYMENT
    """
    key = os.getenv("AZURE_OPENAI_KEY")
    base = os.getenv("AZURE_OPENAI_ENDPOINT")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")
    if not (key and base and deployment):
        raise EnvironmentError("Set AZURE_OPENAI_KEY, AZURE_OPENAI_ENDPOINT, and AZURE_OPENAI_DEPLOYMENT env variables.")
    llm = AzureChatOpenAI(
        deployment_name=deployment,
        openai_api_key=key,
        openai_api_base=base,
        openai_api_version="2023-05-15",
        temperature=0.0
    )
    return llm


# Prompt templates
EDIT_PLAN_PROMPT = ChatPromptTemplate.from_messages([
    SystemMessagePromptTemplate.from_template(
        "You are an EDI X12 expert. The user uploaded an EDI file and asks for a specific change."
    ),
    HumanMessagePromptTemplate.from_template(
        "EDI (first 2000 chars):\n{edi_preview}\n\n"
        "User request: {user_request}\n\n"
        "Produce a JSON array named 'edits' where each element describes a single deterministic edit.\n"
        "Each edit must follow this schema:\n"
        "{\n"
        "  \"segment\": \"<SEG_TAG>\",           # e.g. DTP\n"
        "  \"occurrence\": <N|null>,             # 1-based occurrence to change; null or omitted = all occurrences\n"
        "  \"field_pos\": <M>,                   # 1-based position of the element after the segment tag (1 = first element after tag)\n"
        "  \"subfield_pos\": <K|null>,           # optional for component sub-elements (1-based)\n"
        "  \"match\": {\"old_value\": \"...\"},  # optional: only apply when current value equals this\n"
        "  \"new_value\": \"...\"                # new value to place\n"
        "}\n\n"
        "Return ONLY valid JSON (no explanatory text) with a top-level object like: {\"edits\": [ ... ], \"explain\": \"short human summary\"}\n"
        "If you cannot identify segments/positions, return edits: [] and explain why."
    )
])


ERROR_FIX_PROMPT = ChatPromptTemplate.from_messages([
    SystemMessagePromptTemplate.from_template(
        "You are an EDI X12 expert. The LLM previously proposed edits, they were applied, and the EDI failed validation with errors. Please produce a set of additional edits or corrections to resolve the errors."
    ),
    HumanMessagePromptTemplate.from_template(
        "Current EDI (first 2000 chars):\n{edi_preview}\n\nValidation errors:\n{errors}\n\n"
        "Produce a JSON object: {\"edits\": [ ... ], \"explain\": \"short explanation\"}\n"
        "Edits should follow the same schema used earlier and should be minimal to resolve the validation errors."
    )
])


def llm_generate_edit_plan(llm, edi_text: str, user_request: str) -> Dict[str, Any]:
    edi_preview = edi_text[:2000]
    prompt_input = {"edi_preview": edi_preview, "user_request": user_request}
    res = EDIT_PLAN_PROMPT.format_prompt(**prompt_input).to_messages()
    llm_response = llm.generate(res)
    # Grab text
    text = llm_response.generations[0][0].text.strip()
    # Expect JSON only
    try:
        j = json.loads(text)
    except Exception as e:
        # Try to extract the first JSON object inside the text
        import re
        m = re.search(r'(\{.*\})', text, flags=re.S)
        if m:
            j = json.loads(m.group(1))
        else:
            raise ValueError("LLM did not return JSON edit plan. Raw response:\n" + text)
    return j


def llm_fix_errors(llm, edi_text: str, errors: List[str]) -> Dict[str, Any]:
    edi_preview = edi_text[:2000]
    prompt_input = {"edi_preview": edi_preview, "errors": "\n".join(errors)}
    res = ERROR_FIX_PROMPT.format_prompt(**prompt_input).to_messages()
    llm_response = llm.generate(res)
    text = llm_response.generations[0][0].text.strip()
    try:
        j = json.loads(text)
    except Exception:
        import re
        m = re.search(r'(\{.*\})', text, flags=re.S)
        if m:
            j = json.loads(m.group(1))
        else:
            raise ValueError("LLM did not return JSON on fix step. Raw response:\n" + text)
    return j


# --- Orchestration workflow ---


def run_edit_workflow(llm, original_edi: str, user_request: str, max_iterations: int = 3) -> Dict[str, Any]:
    """
    Main workflow:
    1. Ask LLM for edit plan (JSON)
    2. Apply edits deterministically
    3. Validate using pyx12_validate
    4. If errors, send errors back to LLM for corrections and repeat up to max_iterations

    Returns dictionary:
      {
        "final_edi": "...",
        "history": [
            {"edits": [...], "summary": "...", "applied": [...], "valid": True/False, "errors":[...]}
        ],
        "success": True/False
      }
    """
    current = original_edi
    history = []

    # Step 1: initial plan
    try:
        plan = llm_generate_edit_plan(llm, current, user_request)
    except Exception as e:
        return {"final_edi": current, "history": [{"error": str(e)}], "success": False}

    edits = plan.get("edits", [])
    explain = plan.get("explain", "")
    new_edi, applied = apply_edits(current, edits)
    is_valid, errors = pyx12_validate(new_edi)

    history.append({"edits": edits, "explain": explain, "applied": applied, "valid": is_valid, "errors": errors})
    current = new_edi

    iter_count = 0
    while not is_valid and iter_count < max_iterations:
        iter_count += 1
        # ask LLM to propose fixes based on errors
        try:
            fix_plan = llm_fix_errors(llm, current, errors)
        except Exception as e:
            history.append({"error_from_llm_fix": str(e)})
            break
        fix_edits = fix_plan.get("edits", [])
        fix_explain = fix_plan.get("explain", "")
        new_edi2, applied2 = apply_edits(current, fix_edits)
        is_valid2, errors2 = pyx12_validate(new_edi2)
        history.append({"edits": fix_edits, "explain": fix_explain, "applied": applied2, "valid": is_valid2, "errors": errors2})
        current = new_edi2
        is_valid = is_valid2
        errors = errors2

    return {"final_edi": current, "history": history, "success": is_valid}


# --- Example CLI / Streamlit usage snippet ---

if __name__ == "__main__":
    # quick demo in CLI, not a full streamlit app
    llm = create_llm()

    # Read an EDI from file for testing
    sample_path = "sample.edi"
    if os.path.exists(sample_path):
        with open(sample_path, "r", encoding="utf-8") as f:
            edi_text = f.read()
    else:
        edi_text = "ISA*00*          *00*          *ZZ*SOMEID       *ZZ*DESTID       *210101*1253*^*00501*000000905*0*T*:~GS*HC*SENDER*RECEIVER*20210101*1253*1*X*005010X222~"  # tiny sample

    print("Enter user request (example: change service line DTP to 20250101):")
    ur = input().strip()
    result = run_edit_workflow(llm, edi_text, ur, max_iterations=3)

    print("Success:", result["success"])
    for i, h in enumerate(result["history"], 1):
        print(f"Step {i}: valid={h.get('valid')} errors={h.get('errors')}")
        print("Applied:", h.get("applied"))
    print("---- Final EDI ----")
    print(result["final_edi"][:2000])











import streamlit as st
from langchain.chat_models import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.memory import ConversationBufferMemory
import os

# ---- Azure OpenAI setup ----
llm = AzureChatOpenAI(
    deployment_name="your-deployment",
    openai_api_key=os.getenv("AZURE_OPENAI_KEY"),
    openai_api_base=os.getenv("AZURE_OPENAI_ENDPOINT"),
    openai_api_version="2023-05-15",
    temperature=0
)

memory = ConversationBufferMemory(return_messages=True)

st.title("EDI AI Editor")

uploaded_file = st.file_uploader("Upload EDI file", type=["edi","txt"])
if uploaded_file:
    edi_content = uploaded_file.read().decode("utf-8")
    st.text_area("EDI Preview", edi_content, height=200)

    if "edi_data" not in st.session_state:
        st.session_state.edi_data = edi_content

    query = st.text_input("Enter your edit request (e.g. 'Change service line date')")
    if query:
        prompt = ChatPromptTemplate.from_template("""
        You are an EDI expert. 
        Current EDI:
        {edi}

        User request: {query}

        - Identify correct segment(s).
        - Modify only what is necessary.
        - Return the updated EDI.
        - Also summarize what was changed.
        """)

        chain = prompt | llm
        result = chain.invoke({"edi": st.session_state.edi_data, "query": query})

        st.session_state.edi_data = result.content  # update stored EDI
        st.subheader("✅ Change Summary")
        st.write(result.content[:500])  # show summary & partial result

    if st.button("Generate Final EDI"):
        st.download_button("Download Updated EDI", st.session_state.edi_data, file_name="updated.edi")










import os
import json
import streamlit as st
import pandas as pd
import xml.etree.ElementTree as ET

# Assume these helper functions exist from your imports
# from friday_helpers import change_file_into_segment_list, is_valid_x12

# -------------------------------
# Helpers
# -------------------------------
def get_column_name(seg_id, pos, rule_dict):
    """Get human-friendly column name from rule_dict, fallback to seg+pos"""
    if seg_id in rule_dict and pos in rule_dict[seg_id]:
        return f"{seg_id}{pos} ({rule_dict[seg_id][pos].get('description','')})"
    return f"{seg_id}{pos}"

def xml_to_flat_row(xml_obj, rule_dict):
    """Convert XML into one flat row dictionary"""
    row = {}
    for seg in xml_obj.findall(".//seg"):
        seg_id = seg.attrib.get("id", "")
        for idx, ele in enumerate(seg.findall("ele"), start=1):
            pos = f"{idx:02d}"
            col_name = get_column_name(seg_id, pos, rule_dict)
            row[col_name] = ele.text or ""

            # Handle subfields
            if seg_id in rule_dict and pos in rule_dict[seg_id]:
                subfields = rule_dict[seg_id][pos].get("subfields", [])
                if subfields and ele.text:
                    parts = ele.text.split(":")
                    for i, sub in enumerate(subfields):
                        sub_col = sub.get("description", f"{seg_id}{pos}-{i+1}")
                        key = f"{seg_id}{pos}-{i+1} ({sub_col})"
                        row[key] = parts[i] if i < len(parts) else ""
    return row

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="EDI Validator & Converter", layout="wide")
st.title("📑 EDI Validator & Converter")

# Upload rule_dict.json
rule_file = st.file_uploader("Upload rule_dict.json", type=["json"])
if not rule_file:
    st.stop()

rule_dict = json.load(rule_file)
st.success("✅ Rule dictionary loaded")

# Upload multiple EDI XML files
uploaded_files = st.file_uploader("Upload EDI XML/DAT files", type=["xml", "DAT"], accept_multiple_files=True)
if not uploaded_files:
    st.stop()

st.write("# 📋 Validation Results")

# Temporary folder to store uploaded files
os.makedirs("temp", exist_ok=True)

results = []

for uploaded_file in uploaded_files:
    filename = uploaded_file.name
    filepath = os.path.join("temp", filename)

    # Save uploaded file locally
    with open(filepath, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Convert file to segment list
    try:
        edi_segments = change_file_into_segment_list(filepath)
    except Exception as e:
        st.error(f"❌ Failed to parse {filename}: {e}")
        results.append({"file": filename, "valid": False})
        continue

    # Validate
    try:
        is_valid = is_valid_x12(filename, edi_segments, "temp")
        if is_valid:
            st.success(f"✅ {filename} is valid")
            results.append({"file": filename, "valid": True})
        else:
            st.error(f"❌ {filename} is invalid")
            results.append({"file": filename, "valid": False})
    except Exception as e:
        st.error(f"❌ {filename} Exception during validation: {e}")
        results.append({"file": filename, "valid": False})
        continue

# Display validation summary
st.subheader("Validation Summary")
st.dataframe(pd.DataFrame(results))

# -------------------------------
# Convert valid files to Excel
# -------------------------------
for uploaded_file in uploaded_files:
    filename = uploaded_file.name
    filepath = os.path.join("temp", filename)

    # Only process valid files
    if not any(r["file"] == filename and r["valid"] for r in results):
        continue

    try:
        xml_obj = ET.parse(filepath)
        row_dict = xml_to_flat_row(xml_obj, rule_dict)
        df = pd.DataFrame([row_dict])

        st.subheader(f"📄 Extracted Data from {filename}")
        st.dataframe(df)

        # Save Excel for download
        excel_out = f"{os.path.splitext(filename)[0]}_flat.xlsx"
        df.to_excel(excel_out, index=False)

        with open(excel_out, "rb") as f:
            st.download_button(
                label=f"📥 Download Excel ({filename})",
                data=f,
                file_name=excel_out,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

    except Exception as e:
        st.error(f"❌ Error converting {filename} to Excel: {e}")









import os
import json
import streamlit as st
import pandas as pd
import xml.etree.ElementTree as ET

# Assume these helper functions exist from your imports
# from friday_helpers import change_file_into_segment_list, is_valid_x12

# -------------------------------
# Helpers
# -------------------------------
def get_column_name(seg_id, pos, rule_dict):
    """Get human-friendly column name from rule_dict, fallback to seg+pos"""
    if seg_id in rule_dict and pos in rule_dict[seg_id]:
        return f"{seg_id}{pos} ({rule_dict[seg_id][pos].get('description','')})"
    return f"{seg_id}{pos}"

def xml_to_flat_row(xml_obj, rule_dict):
    """Convert XML into one flat row dictionary"""
    row = {}
    for seg in xml_obj.findall(".//seg"):
        seg_id = seg.attrib.get("id", "")
        for idx, ele in enumerate(seg.findall("ele"), start=1):
            pos = f"{idx:02d}"
            col_name = get_column_name(seg_id, pos, rule_dict)
            row[col_name] = ele.text or ""

            # Handle subfields
            if seg_id in rule_dict and pos in rule_dict[seg_id]:
                subfields = rule_dict[seg_id][pos].get("subfields", [])
                if subfields and ele.text:
                    parts = ele.text.split(":")
                    for i, sub in enumerate(subfields):
                        sub_col = sub.get("description", f"{seg_id}{pos}-{i+1}")
                        key = f"{seg_id}{pos}-{i+1} ({sub_col})"
                        row[key] = parts[i] if i < len(parts) else ""
    return row

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="EDI Validator & Converter", layout="wide")
st.title("📑 EDI Validator & Converter")

# Upload rule_dict.json
rule_file = st.file_uploader("Upload rule_dict.json", type=["json"])
if not rule_file:
    st.stop()

rule_dict = json.load(rule_file)
st.success("✅ Rule dictionary loaded")

# Upload multiple EDI XML files
uploaded_files = st.file_uploader("Upload EDI XML/DAT files", type=["xml", "DAT"], accept_multiple_files=True)
if not uploaded_files:
    st.stop()

st.write("# 📋 Validation Results")

# Temporary folder to store uploaded files
os.makedirs("temp", exist_ok=True)

results = []

for uploaded_file in uploaded_files:
    filename = uploaded_file.name
    filepath = os.path.join("temp", filename)

    # Save uploaded file locally
    with open(filepath, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Convert file to segment list
    try:
        edi_segments = change_file_into_segment_list(filepath)
    except Exception as e:
        st.error(f"❌ Failed to parse {filename}: {e}")
        results.append({"file": filename, "valid": False})
        continue

    # Validate
    try:
        is_valid = is_valid_x12(filename, edi_segments, "temp")
        if is_valid:
            st.success(f"✅ {filename} is valid")
            results.append({"file": filename, "valid": True})
        else:
            st.error(f"❌ {filename} is invalid")
            results.append({"file": filename, "valid": False})
    except Exception as e:
        st.error(f"❌ {filename} Exception during validation: {e}")
        results.append({"file": filename, "valid": False})
        continue

# Display validation summary
st.subheader("Validation Summary")
st.dataframe(pd.DataFrame(results))

# -------------------------------
# Convert valid files to Excel
# -------------------------------
for uploaded_file in uploaded_files:
    filename = uploaded_file.name
    filepath = os.path.join("temp", filename)

    # Only process valid files
    if not any(r["file"] == filename and r["valid"] for r in results):
        continue

    try:
        xml_obj = ET.parse(filepath)
        row_dict = xml_to_flat_row(xml_obj, rule_dict)
        df = pd.DataFrame([row_dict])

        st.subheader(f"📄 Extracted Data from {filename}")
        st.dataframe(df)

        # Save Excel for download
        excel_out = f"{os.path.splitext(filename)[0]}_flat.xlsx"
        df.to_excel(excel_out, index=False)

        with open(excel_out, "rb") as f:
            st.download_button(
                label=f"📥 Download Excel ({filename})",
                data=f,
                file_name=excel_out,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

    except Exception as e:
        st.error(f"❌ Error converting {filename} to Excel: {e}")









import os
import json
import streamlit as st
import pandas as pd
import xml.etree.ElementTree as ET

# -------------------------------
# Helpers
# -------------------------------
def get_column_name(seg_id, pos, rule_dict):
    """Get human-friendly column name from rule_dict, fallback to seg+pos"""
    if seg_id in rule_dict and pos in rule_dict[seg_id]:
        return f"{seg_id}{pos} ({rule_dict[seg_id][pos].get('description','')})"
    return f"{seg_id}{pos}"

def xml_to_flat_row(xml_obj, rule_dict):
    """Convert XML into one flat row dictionary"""
    row = {}
    for seg in xml_obj.findall(".//seg"):
        seg_id = seg.attrib.get("id", "")
        for idx, ele in enumerate(seg.findall("ele"), start=1):
            pos = f"{idx:02d}"
            col_name = get_column_name(seg_id, pos, rule_dict)
            row[col_name] = ele.text or ""

            # Handle subfields
            if seg_id in rule_dict and pos in rule_dict[seg_id]:
                subfields = rule_dict[seg_id][pos].get("subfields", [])
                if subfields and ele.text:
                    parts = ele.text.split(":")
                    for i, sub in enumerate(subfields):
                        sub_col = sub.get("description", f"{seg_id}{pos}-{i+1}")
                        key = f"{seg_id}{pos}-{i+1} ({sub_col})"
                        row[key] = parts[i] if i < len(parts) else ""
    return row

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="EDI Validator & Converter", layout="wide")
st.title("📑 EDI Validator & Converter")

# Upload JSON rules
rule_file = st.file_uploader("Upload rule_dict.json", type=["json"])
if rule_file:
    rule_dict = json.load(rule_file)
    st.success("✅ Rule dictionary loaded")
else:
    st.stop()

# Upload EDI XML file
uploaded_file = st.file_uploader("Upload EDI XML file", type=["xml"])
if uploaded_file:
    try:
        xml_obj = ET.parse(uploaded_file)
        st.success("✅ XML loaded")
    except Exception as e:
        st.error(f"❌ Error parsing XML: {e}")
        st.stop()

    # Convert XML → flat row
    row_dict = xml_to_flat_row(xml_obj, rule_dict)
    df = pd.DataFrame([row_dict])  # one row

    st.subheader("📋 Extracted Data (one row per transaction)")
    st.dataframe(df)

    # Save Excel
    out_path = "edi_output_flat.xlsx"
    df.to_excel(out_path, index=False)

    with open(out_path, "rb") as f:
        st.download_button(
            label="📥 Download Excel",
            data=f,
            file_name="edi_output_flat.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )











import os
import sys
import json
import streamlit as st
import pandas as pd

# -------------------------------
# Import your helper functions
# -------------------------------
# Assume your project structure is:
#   /friday_helpers/
#       __init__.py
#       helpers.py
#   /app.py   <-- this file

# Adjust Python path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from friday_helpers import (
    load_file_as_xml_obj,
    is_valid_x12,
    change_file_into_segment_list,
)

# -------------------------------
# Helper: Load rule_dict.json
# -------------------------------
def load_rule_dict(json_file):
    with open(json_file, "r", encoding="utf-8") as f:
        return json.load(f)

# -------------------------------
# Helper: Get description for segment/element
# -------------------------------
def get_column_name(seg_id, ele_pos, rule_dict):
    """
    Return friendly column name if found in rule_dict, else fallback SEGxx
    """
    if seg_id in rule_dict:
        if ele_pos in rule_dict[seg_id]:
            desc = rule_dict[seg_id][ele_pos].get("description")
            if desc:
                return f"{desc} ({seg_id}{ele_pos})"
    return f"{seg_id}{ele_pos}"

# -------------------------------
# Convert XML object -> DataFrame
# -------------------------------
def xml_to_dataframe(xml_obj, rule_dict):
    rows = []

    for seg in xml_obj.findall(".//seg"):
        seg_id = seg.attrib.get("id", "")
        row = {"SegmentID": seg_id}

        for idx, ele in enumerate(seg.findall("ele"), start=1):
            pos = f"{idx:02d}"  # "01", "02", ...
            col_name = get_column_name(seg_id, pos, rule_dict)
            row[col_name] = ele.text or ""

            # If subfields exist in rule_dict, split and map them
            if seg_id in rule_dict and pos in rule_dict[seg_id]:
                subfields = rule_dict[seg_id][pos].get("subfields", [])
                if subfields and ele.text:
                    parts = ele.text.split(":")
                    for i, sub in enumerate(subfields):
                        sub_col = sub.get("description", f"{seg_id}{pos}-{i+1}")
                        key = f"{sub_col} ({seg_id}{pos}-{i+1})"
                        row[key] = parts[i] if i < len(parts) else ""

        rows.append(row)

    return pd.DataFrame(rows)

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="EDI Validator & XML→Excel", layout="wide")
st.title("📑 EDI Validator & Converter")

uploaded_file = st.file_uploader("Upload EDI .DAT file", type=["dat", "DAT"])
rule_file = st.file_uploader("Upload rule_dict.json", type=["json"])

if uploaded_file and rule_file:
    # Save temp files
    tmp_path = os.path.join("tmp_input.dat")
    with open(tmp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Convert file → segments
    edi_segments = change_file_into_segment_list(tmp_path)

    # Run validation
    is_valid = is_valid_x12(uploaded_file.name, edi_segments, os.getcwd())

    if not is_valid:
        st.error("❌ File is not valid X12 EDI")
    else:
        st.success("✅ File is valid X12 EDI")

        # Load XML object
        try:
            xml_obj = load_file_as_xml_obj(tmp_path)
        except Exception as e:
            st.error(f"XML Parse error: {e}")
            st.stop()

        # Load rule dict
        rule_dict = json.load(rule_file)

        # Convert XML → DataFrame
        df = xml_to_dataframe(xml_obj, rule_dict)

        st.dataframe(df.head(50))

        # Download Excel
        out_path = "edi_output.xlsx"
        df.to_excel(out_path, index=False)

        with open(out_path, "rb") as f:
            st.download_button(
                label="📥 Download Excel",
                data=f,
                file_name="edi_output.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )





# app.py
import os
import streamlit as st

# Import your helper functions
from friday_helpers import (
    load_file_as_xml_obj,
    is_valid_x12,
    change_file_into_segment_list
)

st.set_page_config(page_title="EDI File Validator", layout="wide")

st.title("📑 EDI File Validator")

# -------------------
# File uploader
# -------------------
uploaded_files = st.file_uploader(
    "Upload one or more EDI (.DAT) files", 
    type=["dat"], 
    accept_multiple_files=True
)

# -------------------
# Run validation
# -------------------
if uploaded_files:
    st.write("### Validation Results")

    for uploaded_file in uploaded_files:
        filename = uploaded_file.name
        filepath = os.path.join("temp", filename)

        # Ensure temp folder exists
        os.makedirs("temp", exist_ok=True)

        # Save uploaded file locally
        with open(filepath, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Convert to segment list
        try:
            edi_segments = change_file_into_segment_list(filepath)
        except Exception as e:
            st.error(f"❌ {filename} - Failed to parse: {e}")
            continue

        # Validate
        try:
            result = is_valid_x12(filename, edi_segments, "temp")
        except Exception as e:
            st.error(f"❌ {filename} - Exception during validation: {e}")
            continue

        # Show results
        if result is True:
            st.success(f"✅ {filename} - Valid EDI")
        else:
            st.error(f"❌ {filename} - Invalid EDI")
            if isinstance(result, (str, list, dict)):
                st.json(result) if isinstance(result, (dict, list)) else st.text(result)

    st.info("Validation Check Done ✅")








import os
import xml.etree.ElementTree as ET
from datetime import datetime
from collections import defaultdict

BATCH_SIZE = 50  # claims per file

# -------------------------
# Utility: turn <seg> into EDI text
# -------------------------
def seg_to_edi_text(seg):
    """Convert <seg> element into raw EDI line"""
    if seg is None:
        return ""
    tag_id = seg.attrib.get("id", "")
    elems = []
    for e in seg.findall("elem"):
        elems.append(e.text if e.text else "")
    return tag_id + "*" + "*".join(elems) + "~"

# -------------------------
# Claim collection
# -------------------------
def collect_claims(xml_file):
    """Parse XML, collect all 2000B claims with partner and parent ST_LOOP"""
    tree = ET.parse(xml_file)
    root = tree.getroot()

    claims_by_partner = defaultdict(list)

    # each ST_LOOP holds an ISA->GS->ST grouping
    for st_loop in root.findall(".//loop[@id='ST_LOOP']"):
        # Get envelope info
        st03 = st_loop.find("./seg[@id='ST']/elem[3]").text.strip()
        isa06 = st_loop.find(".//seg[@id='ISA']/elem[6]").text.strip()

        # find all subscriber loops (2000A)
        for sub_loop in st_loop.findall("./loop[@id='2000A']"):
            # find patient loops (2000B)
            for pat_loop in sub_loop.findall("./loop[@id='2000B']"):
                # store patient loop + subscriber loop + parent ST_LOOP reference
                claims_by_partner[(st03, isa06)].append(
                    (sub_loop, pat_loop, st_loop)
                )

    return claims_by_partner

# -------------------------
# File writer
# -------------------------
def write_dat_files(claims_by_partner_map, out_dir):
    """Write .DAT batches preserving ISA/GS/ST + header loop + SE/GE/IEA"""
    today = datetime.now().strftime("%m%d%Y")

    for (claim_type, isa06), claims in claims_by_partner_map.items():
        if not claims:
            continue

        # Process in chunks
        for batch_idx in range(0, len(claims), BATCH_SIZE):
            batch = claims[batch_idx: batch_idx + BATCH_SIZE]
            first_sub, first_pat, st_loop = batch[0]  # use first claim’s ST_LOOP

            # Envelope segments
            isa_seg = st_loop.find(".//seg[@id='ISA']")
            gs_seg = st_loop.find(".//seg[@id='GS']")
            st_seg = st_loop.find("./seg[@id='ST']")
            se_seg = st_loop.find("./seg[@id='SE']")
            ge_seg = st_loop.find(".//seg[@id='GE']")
            iea_seg = st_loop.find(".//seg[@id='IEA']")

            # Collect HEADER (between ST and first 2000A)
            header_parts = []
            for child in st_loop:
                if child.tag == "loop" and child.attrib.get("id") == "2000A":
                    break
                if child.tag == "seg":
                    header_parts.append(seg_to_edi_text(child))
                elif child.tag == "loop":
                    for seg in child.findall(".//seg"):
                        header_parts.append(seg_to_edi_text(seg))

            # Build lines
            lines = [
                seg_to_edi_text(isa_seg),
                seg_to_edi_text(gs_seg),
                seg_to_edi_text(st_seg),
            ]
            lines.extend(header_parts)

            # Add claims
            for sub_loop, pat_loop, _ in batch:
                for seg in sub_loop.findall("./seg"):
                    lines.append(seg_to_edi_text(seg))
                for seg in pat_loop.findall("./seg"):
                    lines.append(seg_to_edi_text(seg))

            # Close out
            lines += [
                seg_to_edi_text(se_seg),
                seg_to_edi_text(ge_seg),
                seg_to_edi_text(iea_seg),
            ]

            # File name
            fname = f"{claim_type}_{isa06}_Batch{batch_idx//BATCH_SIZE + 1}_{today}.DAT"
            outpath = os.path.join(out_dir, fname)

            # Write
            with open(outpath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))

            print(f"Wrote {len(batch)} claims -> {outpath}")

# -------------------------
# Main
# -------------------------
if __name__ == "__main__":
    input_xml = "input.xml"   # your source XML
    out_dir = "out"
    os.makedirs(out_dir, exist_ok=True)

    claims_map = collect_claims(input_xml)
    write_dat_files(claims_map, out_dir)

















import os
import pandas as pd

# ---------- CONFIG ----------
EXCEL_PATH = "input.xlsx"            # input excel
INPUT_SOURCE_FOLDER = "source_files" # folder with .txt files
OUTPUT_EXCEL = "output_with_status.xlsx"
# -----------------------------


def load_all_files(folder):
    """Read all .txt files and return combined text content."""
    file_contents = []
    for fname in os.listdir(folder):
        if fname.lower().endswith(".txt"):
            try:
                with open(os.path.join(folder, fname), "r", encoding="utf-8", errors="ignore") as f:
                    file_contents.append(f.read())
            except Exception as e:
                print(f"⚠️ Could not read {fname}: {e}")
    return " ".join(file_contents)  # merge all into one big string


def main():
    # Load Excel
    df = pd.read_excel(EXCEL_PATH, dtype=str).fillna("")

    # Build claim keys
    df["clm_key"] = df.apply(
        lambda r: f"CLM*{r['patientcontrolnumber'].strip()}*{r['amount'].strip()}",
        axis=1
    )

    # Read all files once
    all_content = load_all_files(INPUT_SOURCE_FOLDER)

    # Check presence
    df["status"] = df["clm_key"].apply(lambda key: "Present" if key in all_content else "Missing")

    # Save output
    df.to_excel(OUTPUT_EXCEL, index=False)
    print(f"✅ Results written to {OUTPUT_EXCEL}")


if __name__ == "__main__":
    main()









import os
import pandas as pd

# ---------- CONFIG ----------
EXCEL_PATH = "input.xlsx"           # path to your Excel file
INPUT_SOURCE_FOLDER = "source_files"  # folder where .txt files are stored
OUTPUT_EXCEL = "output_with_status.xlsx"
# -----------------------------


def normalize_filename(fname: str) -> str:
    """Ensure filenames always end with .txt"""
    fname = str(fname).strip()
    if not fname.lower().endswith(".txt"):
        fname = fname + ".txt"
    return fname


def main():
    # Load Excel
    df = pd.read_excel(EXCEL_PATH, dtype=str).fillna("")

    # Normalize filenames
    df["claimfilename"] = df["claimfilename"].apply(normalize_filename)

    # Cache file contents for faster lookup
    file_cache = {}
    for fname in set(df["claimfilename"]):
        file_path = os.path.join(INPUT_SOURCE_FOLDER, fname)
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                file_cache[fname] = f.read()
        else:
            file_cache[fname] = None  # file missing

    # Check each row
    statuses = []
    for _, row in df.iterrows():
        claim_file = row["claimfilename"]
        patient_id = row["patientcontrolnumber"].strip()
        amount = row["amount"].strip()

        clm_key = f"CLM*{patient_id}*{amount}"

        if file_cache[claim_file] is None:
            statuses.append("File Not Found")
        elif clm_key in file_cache[claim_file]:
            statuses.append("Present")
        else:
            statuses.append("Missing")

    # Add status column
    df["status"] = statuses

    # Save updated Excel
    df.to_excel(OUTPUT_EXCEL, index=False)
    print(f"✅ Status written to {OUTPUT_EXCEL}")


if __name__ == "__main__":
    main()






import os
import pandas as pd

# ---------- CONFIG ----------
EXCEL_PATH = "input.xlsx"   # path to your Excel file
INPUT_SOURCE_FOLDER = "source_files"  # folder where .txt files are stored
OUTPUT_MISSED = "missed_files.txt"    # output report

# Define column names as in Excel
PATIENT_COL = "patientcontrolnumber"
AMOUNT_COL = "amount"
FILE_COL = "claimfilename"  # filename column in Excel (AS101 or AS101.txt)
# -----------------------------


def normalize_filename(fname: str) -> str:
    """Ensure filenames always end with .txt"""
    fname = str(fname).strip()
    if not fname.lower().endswith(".txt"):
        fname = fname + ".txt"
    return fname


def main():
    # Load Excel
    df = pd.read_excel(EXCEL_PATH, dtype=str).fillna("")

    # Group by filename → list of CLM keys
    claims_by_file = {}
    for _, row in df.iterrows():
        claim_file = normalize_filename(row[FILE_COL])
        patient_id = row[PATIENT_COL].strip()
        amount = row[AMOUNT_COL].strip()

        clm_key = f"CLM*{patient_id}*{amount}"
        claims_by_file.setdefault(claim_file, []).append(clm_key)

    missed_report = []

    for claim_file, claim_keys in claims_by_file.items():
        file_path = os.path.join(INPUT_SOURCE_FOLDER, claim_file)

        if not os.path.exists(file_path):
            missed_report.append(f"{claim_file} → File not found")
            continue

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            # Check all claim keys
            missing = [key for key in claim_keys if key not in content]

            if missing:
                missed_report.append(
                    f"{claim_file} → Missing {len(missing)}/{len(claim_keys)} claims\n  "
                    + "\n  ".join(missing)
                )

        except Exception as e:
            missed_report.append(f"{claim_file} → Error: {e}")

    # Write results
    with open(OUTPUT_MISSED, "w") as f:
        if missed_report:
            f.write("\n".join(missed_report))
        else:
            f.write("All claims found ✅")

    print(f"Check complete. Results saved in {OUTPUT_MISSED}")


if __name__ == "__main__":
    main()







import os
import pandas as pd

# ---------- CONFIG ----------
EXCEL_PATH = "input.xlsx"   # path to your Excel file
INPUT_SOURCE_FOLDER = "source_files"  # folder where .txt files are stored
OUTPUT_MISSED = "missed_files.txt"    # output report

# Define column names as in Excel
PATIENT_COL = "patientcontrolnumber"
AMOUNT_COL = "amount"
FILE_COL = "claimfilename"  # filename column in Excel (AS101 / AS101.txt)
# -----------------------------


def normalize_filename(fname: str) -> str:
    """Ensure filenames always end with .txt"""
    fname = str(fname).strip()
    if not fname.lower().endswith(".txt"):
        fname = fname + ".txt"
    return fname


def main():
    # Load Excel
    df = pd.read_excel(EXCEL_PATH, dtype=str)
    df = df.fillna("")

    missed_files = []

    for _, row in df.iterrows():
        patient_id = row[PATIENT_COL].strip()
        amount = row[AMOUNT_COL].strip()
        claim_file = normalize_filename(row[FILE_COL])

        file_path = os.path.join(INPUT_SOURCE_FOLDER, claim_file)

        if not os.path.exists(file_path):
            missed_files.append(f"{claim_file} (File not found)")
            continue

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            search_string = f"CLM*{patient_id}*{amount}"

            if search_string not in content:
                missed_files.append(f"{claim_file} (Missing {search_string})")

        except Exception as e:
            missed_files.append(f"{claim_file} (Error: {e})")

    # Write results
    with open(OUTPUT_MISSED, "w") as f:
        if missed_files:
            f.write("\n".join(missed_files))
        else:
            f.write("All records found. ✅")

    print(f"Check complete. Results saved in {OUTPUT_MISSED}")


if __name__ == "__main__":
    main()




import os
import shutil

def split_files_into_folders(source_dir, target_dir, batch_size=50):
    os.makedirs(target_dir, exist_ok=True)

    # Get all files from source directory
    files = [f for f in os.listdir(source_dir) if os.path.isfile(os.path.join(source_dir, f))]
    files.sort()  # optional: keeps them in order

    folder_count = 1
    for i in range(0, len(files), batch_size):
        # Create subfolder
        folder_name = os.path.join(target_dir, f"batch_{folder_count}")
        os.makedirs(folder_name, exist_ok=True)

        # Move files into subfolder
        for f in files[i:i+batch_size]:
            shutil.copy(os.path.join(source_dir, f), os.path.join(folder_name, f))

        print(f"Created {folder_name} with {len(files[i:i+batch_size])} files")
        folder_count += 1


# Example usage
source_dir = r"P:\CORE-Axiom Health EdgeSFTPTestFiles\Deloitte Parallel Testing\TRAVIS\Source_EDI"
target_dir = r"P:\CORE-Axiom Health EdgeSFTPTestFiles\Deloitte Parallel Testing\TRAVIS\Split_EDI"
split_files_into_folders(source_dir, target_dir, batch_size=50)






# -------------------
# Download Template
# -------------------
import io

def create_template_excel():
    template_df = pd.DataFrame([
        {
            "SourceEDI": "837",
            "SourceSegment": "CLM",
            "SourceFieldPos": 1,
            "SourceSubFieldPos": "",
            "SourceFieldName": "ClaimID",
            "TargetEDI": "820",
            "TargetSegment": "RMR",
            "TargetFieldPos": 2,
            "TargetSubFieldPos": "",
            "TargetFieldName": "ClaimID"
        },
        {
            "SourceEDI": "837",
            "SourceSegment": "NM1",
            "SourceFieldPos": 2,
            "SourceSubFieldPos": "",
            "SourceFieldName": "PatientName",
            "TargetEDI": "820",
            "TargetSegment": "NM1",
            "TargetFieldPos": 2,
            "TargetSubFieldPos": "",
            "TargetFieldName": "PatientName"
        },
        {
            "SourceEDI": "837",
            "SourceSegment": "AMT",
            "SourceFieldPos": 2,
            "SourceSubFieldPos": "",
            "SourceFieldName": "ClaimAmount",
            "TargetEDI": "820",
            "TargetSegment": "RMR",
            "TargetFieldPos": 4,
            "TargetSubFieldPos": "",
            "TargetFieldName": "ClaimAmount"
        }
    ])

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        template_df.to_excel(writer, index=False, sheet_name="MappingTemplate")
    output.seek(0)
    return output


with st.sidebar:
    st.header("Configuration")
    # Button to download Excel template
    template_excel = create_template_excel()
    st.download_button(
        label="📥 Download Mapping Template",
        data=template_excel,
        file_name="EDI_Mapping_Template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )







# app.py
import streamlit as st
import pandas as pd
import os

st.set_page_config(page_title="Generic EDI Comparison", layout="wide")

# -------------------
# Utilities
# -------------------

def parse_edi(content: str, edi_type: str, mapping: pd.DataFrame) -> list[dict]:
    """
    Parse EDI content into records based on mapping Excel rows.
    """
    records = []
    segments = content.split("~")

    for idx, row in mapping[mapping["SourceEDI"] == edi_type].iterrows():
        seg_name = row["SourceSegment"]
        field_pos = int(row["SourceFieldPos"])
        sub_pos = row.get("SourceSubFieldPos", None)

        for seg in segments:
            parts = seg.strip().split("*")
            if not parts or parts[0] != seg_name:
                continue

            # get field
            value = parts[field_pos] if field_pos < len(parts) else ""

            # subfield extraction
            if pd.notna(sub_pos) and ":" in value:
                subs = value.split(":")
                if int(sub_pos) < len(subs):
                    value = subs[int(sub_pos)]

            records.append({
                "FieldName": row["SourceFieldName"],
                "Value": value,
                "RowID": idx  # link back to Excel row
            })
    return records


def compare_files(source_txt, target_txt, mapping, source_edi, target_edi, filename):
    """
    Compare source vs target EDI based on mapping file.
    Returns dataframe of differences and match status.
    """
    source_records = parse_edi(source_txt, source_edi, mapping)
    target_records = parse_edi(target_txt, target_edi, mapping)

    results = []
    all_matched = True

    for idx, row in mapping.iterrows():
        # Source value
        src_val = next((r["Value"] for r in source_records if r["RowID"] == idx), "")
        tgt_val = next((r["Value"] for r in target_records if r["RowID"] == idx), "")

        status = "Match" if src_val == tgt_val else "Mismatch"
        if status == "Mismatch":
            all_matched = False

        results.append({
            "File": filename,
            "FieldName": row["SourceFieldName"],
            "SourceValue": src_val,
            "TargetValue": tgt_val,
            "Status": status
        })

    return pd.DataFrame(results), all_matched


def save_report(all_dfs: dict, output_file: str):
    """Save comparison results to Excel"""
    with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
        for name, df in all_dfs.items():
            df.to_excel(writer, sheet_name=name, index=False)
    return output_file


# -------------------
# Streamlit UI
# -------------------

st.title("📑 Generic EDI Comparison Tool")

with st.sidebar:
    st.header("Configuration")
    mapping_file = st.file_uploader("Upload Mapping Excel", type=["xlsx"])
    run = st.button("Run Comparison")

if mapping_file:
    mapping = pd.read_excel(mapping_file)

    # Get unique EDI types
    source_edi = mapping["SourceEDI"].iloc[0]
    target_edi = mapping["TargetEDI"].iloc[0]

    st.write(f"🔹 Source EDI type: **{source_edi}**")
    st.write(f"🔹 Target EDI type: **{target_edi}**")

    files_source = st.file_uploader(f"Upload {source_edi} Files", type=["txt", "edi"], accept_multiple_files=True)
    files_target = st.file_uploader(f"Upload {target_edi} Files", type=["txt", "edi"], accept_multiple_files=True)

    if run and files_source and files_target:
        all_results = {}
        matched_files = []

        for fs in files_source:
            src_txt = fs.read().decode("utf-8", errors="ignore")

            for ft in files_target:
                tgt_txt = ft.read().decode("utf-8", errors="ignore")
                df, all_matched = compare_files(src_txt, tgt_txt, mapping, source_edi, target_edi, fs.name)

                all_results[f"{fs.name}_vs_{ft.name}"] = df

                if all_matched:
                    matched_files.append(fs.name)

        # Show results
        for name, df in all_results.items():
            st.subheader(f"Results for {name}")
            st.dataframe(df)

        if matched_files:
            st.success(f"✅ Perfect match found in: {', '.join(matched_files)}")

        # Save report
        output_file = "EDI_Comparison_Report.xlsx"
        save_report(all_results, output_file)
        with open(output_file, "rb") as f:
            st.download_button("⬇️ Download Report", f, file_name=output_file)












mport os
import pandas as pd
import shutil

# -------- Settings --------
excel_file = "file_list.xlsx"       # Excel file containing file names
column_name = "FileName"            # Column in Excel with file names
search_folder = "D:/Data"           # Root folder to search in (can have subfolders)
output_folder = "D:/CollectedFiles" # Where to copy collected files

# Create output folder if not exists
os.makedirs(output_folder, exist_ok=True)

# Read Excel file
df = pd.read_excel(excel_file)
file_list = df[column_name].dropna().astype(str).tolist()

# Walk through folder and find files
found_files = []
for root, _, files in os.walk(search_folder):
    for f in files:
        if f in file_list:  # exact match
            source_path = os.path.join(root, f)
            target_path = os.path.join(output_folder, f)
            shutil.copy2(source_path, target_path)  # copy file
            found_files.append(f)

# Report
missing_files = set(file_list) - set(found_files)

print(f"✅ Collected {len(found_files)} files into {output_folder}")
if missing_files:
    print("⚠️ Missing files (not found):")
    for mf in missing_files:
        print(" -", mf)














# app.py
import streamlit as st
import pandas as pd
import os
from io import BytesIO
from typing import List, Dict

st.set_page_config(page_title="EDI Comparison (Config-Driven)", layout="wide")

# -------------------
# Utilities
# -------------------

def parse_edi(content: str, edi_type: str, mapping_df: pd.DataFrame) -> List[dict]:
    """Generic EDI parser using mapping configuration."""
    records = []
    segments = content.split("~")
    for seg in segments:
        parts = seg.strip().split("*")
        if not parts:
            continue
        # filter mapping rows for this EDI type + segment
        for _, row in mapping_df[(mapping_df["SourceEDI"] == edi_type) & (mapping_df["Segment"] == parts[0])].iterrows():
            val = ""
            try:
                if int(row["FieldPos"]) < len(parts):
                    val = parts[int(row["FieldPos"])]
                    if pd.notna(row["SubFieldPos"]):
                        subs = val.split(":")
                        if int(row["SubFieldPos"]) < len(subs):
                            val = subs[int(row["SubFieldPos"])]
            except Exception:
                val = ""
            records.append({
                "FieldName": row["FieldName"],
                "Value": val,
                "Segment": row["Segment"]
            })
    return records


def compare_records(source_records, target_records, mapping_df):
    """Compare extracted records based on mapping config."""
    results = []
    for _, row in mapping_df.iterrows():
        src_vals = [r["Value"] for r in source_records if r["FieldName"] == row["FieldName"]]
        tgt_vals = [r["Value"] for r in target_records if r["FieldName"] == row["TargetFieldName"]]

        src_val = ", ".join(src_vals) if src_vals else ""
        tgt_val = ", ".join(tgt_vals) if tgt_vals else ""

        status = "Match" if src_val == tgt_val else "Mismatch"

        results.append({
            "SourceField": row["FieldName"],
            "SourceValue": src_val,
            "TargetField": row["TargetFieldName"],
            "TargetValue": tgt_val,
            "Status": status
        })
    return results


def save_excel(comparison_results, output_path="comparison_results.xlsx"):
    """Save comparison results into Excel."""
    df = pd.DataFrame(comparison_results)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Comparison", index=False)
    return output_path


def get_template_excel() -> BytesIO:
    """Generate a sample mapping Excel template."""
    data = {
        "SourceEDI": ["837I", "837I", "837I"],
        "Segment": ["NM1", "NM1", "CLM"],
        "FieldPos": [3, 4, 1],
        "SubFieldPos": [None, None, None],
        "FieldName": ["LastName", "FirstName", "ClaimID"],
        "TargetEDI": ["820", "820", "820"],
        "TargetSegment": ["NM1", "NM1", "RMR"],
        "TargetFieldPos": [3, 4, 2],
        "TargetSubPos": [None, None, None],
        "TargetFieldName": ["LastName", "FirstName", "ClaimID"]
    }
    df = pd.DataFrame(data)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Mapping", index=False)
    buffer.seek(0)
    return buffer


# -------------------
# Streamlit UI
# -------------------

st.title("📊 Config-Driven EDI Comparison Tool")

# Sidebar: Download template
st.sidebar.header("📥 Download Template")
template_buffer = get_template_excel()
st.sidebar.download_button(
    label="Download Mapping Template",
    data=template_buffer,
    file_name="edi_mapping_template.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# Step 1: Upload mapping Excel
mapping_file = st.file_uploader("Upload Mapping Excel", type=["xlsx"])

if mapping_file:
    mapping_df = pd.read_excel(mapping_file)
    st.success("Mapping file loaded!")

    required_cols = ["SourceEDI", "Segment", "FieldPos", "SubFieldPos", "FieldName",
                     "TargetEDI", "TargetSegment", "TargetFieldPos", "TargetSubPos", "TargetFieldName"]
    if not all(col in mapping_df.columns for col in required_cols):
        st.error(f"Mapping file must contain columns: {required_cols}")
    else:
        # Step 2: Identify unique EDIs
        source_edi = mapping_df["SourceEDI"].iloc[0]
        target_edi = mapping_df["TargetEDI"].iloc[0]

        st.subheader("Step 2: Upload EDI Files")

        source_file = st.file_uploader(f"Upload {source_edi} File", type=["txt", "edi"])
        target_file = st.file_uploader(f"Upload {target_edi} File", type=["txt", "edi"])

        if source_file and target_file:
            if st.button("Run Comparison"):
                source_content = source_file.read().decode("utf-8")
                target_content = target_file.read().decode("utf-8")

                source_records = parse_edi(source_content, source_edi, mapping_df)
                target_records = parse_edi(target_content, target_edi, mapping_df)

                comparison_results = compare_records(source_records, target_records, mapping_df)

                # Show in UI
                df_results = pd.DataFrame(comparison_results)
                st.dataframe(df_results)

                # Save to Excel
                output_path = save_excel(comparison_results)
                with open(output_path, "rb") as f:
                    st.download_button(
                        label="📥 Download Comparison Report",
                        data=f,
                        file_name=os.path.basename(output_path),
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )









# app.py
import streamlit as st
import pandas as pd
import os
from typing import List, Dict

st.set_page_config(page_title="EDI Comparison (Config-Driven)", layout="wide")

# -------------------
# Utilities
# -------------------

def parse_edi(content: str, edi_type: str, mapping_df: pd.DataFrame) -> List[dict]:
    """Generic EDI parser using mapping configuration."""
    records = []
    segments = content.split("~")
    for seg in segments:
        parts = seg.strip().split("*")
        if not parts:
            continue
        # filter mapping rows for this EDI type + segment
        for _, row in mapping_df[(mapping_df["SourceEDI"] == edi_type) & (mapping_df["Segment"] == parts[0])].iterrows():
            val = ""
            try:
                if int(row["FieldPos"]) < len(parts):
                    val = parts[int(row["FieldPos"])]
                    if pd.notna(row["SubFieldPos"]):
                        subs = val.split(":")
                        if int(row["SubFieldPos"]) < len(subs):
                            val = subs[int(row["SubFieldPos"])]
            except Exception:
                val = ""
            records.append({
                "FieldName": row["FieldName"],
                "Value": val,
                "Segment": row["Segment"]
            })
    return records


def compare_records(source_records, target_records, mapping_df):
    """Compare extracted records based on mapping config."""
    results = []
    for _, row in mapping_df.iterrows():
        src_vals = [r["Value"] for r in source_records if r["FieldName"] == row["FieldName"]]
        tgt_vals = [r["Value"] for r in target_records if r["FieldName"] == row["TargetFieldName"]]

        src_val = ", ".join(src_vals) if src_vals else ""
        tgt_val = ", ".join(tgt_vals) if tgt_vals else ""

        status = "Match" if src_val == tgt_val else "Mismatch"

        results.append({
            "SourceField": row["FieldName"],
            "SourceValue": src_val,
            "TargetField": row["TargetFieldName"],
            "TargetValue": tgt_val,
            "Status": status
        })
    return results


def save_excel(comparison_results, output_path="comparison_results.xlsx"):
    """Save comparison results into Excel."""
    df = pd.DataFrame(comparison_results)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Comparison", index=False)
    return output_path


# -------------------
# Streamlit UI
# -------------------

st.title("📊 Config-Driven EDI Comparison Tool")

# Step 1: Upload mapping Excel
mapping_file = st.file_uploader("Upload Mapping Excel", type=["xlsx"])

if mapping_file:
    mapping_df = pd.read_excel(mapping_file)
    st.success("Mapping file loaded!")

    required_cols = ["SourceEDI", "Segment", "FieldPos", "SubFieldPos", "FieldName",
                     "TargetEDI", "TargetSegment", "TargetFieldPos", "TargetSubPos", "TargetFieldName"]
    if not all(col in mapping_df.columns for col in required_cols):
        st.error(f"Mapping file must contain columns: {required_cols}")
    else:
        # Step 2: Identify unique EDIs
        source_edi = mapping_df["SourceEDI"].iloc[0]
        target_edi = mapping_df["TargetEDI"].iloc[0]

        st.subheader("Step 2: Upload EDI Files")

        source_file = st.file_uploader(f"Upload {source_edi} File", type=["txt", "edi"])
        target_file = st.file_uploader(f"Upload {target_edi} File", type=["txt", "edi"])

        if source_file and target_file:
            if st.button("Run Comparison"):
                source_content = source_file.read().decode("utf-8")
                target_content = target_file.read().decode("utf-8")

                source_records = parse_edi(source_content, source_edi, mapping_df)
                target_records = parse_edi(target_content, target_edi, mapping_df)

                comparison_results = compare_records(source_records, target_records, mapping_df)

                # Show in UI
                df_results = pd.DataFrame(comparison_results)
                st.dataframe(df_results)

                # Save to Excel
                output_path = save_excel(comparison_results)
                with open(output_path, "rb") as f:
                    st.download_button(
                        label="📥 Download Comparison Report",
                        data=f,
                        file_name=os.path.basename(output_path),
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )









import os
import xml.etree.ElementTree as ET
from collections import defaultdict

def get_claim_type(st_impl):
    if "X222" in st_impl or "005010X222" in st_impl:
        return "institutional"
    elif "X223" in st_impl or "005010X223" in st_impl:
        return "professional"
    else:
        return "unknown"

def merge_edi_roots(xml_roots, output_folder):
    grouped = defaultdict(list)

    # Group by trading partner + claim type
    for root in xml_roots:
        isa06 = root.find(".//loop[@id='ISA_LOOP']/seg[@id='ISA']/ele[@id='ISA06']")
        partner_id = isa06.text if isa06 is not None else "UNKNOWN"

        st03 = root.find(".//loop[@id='ST_LOOP']/seg[@id='ST']/ele[@id='ST03']")
        claim_type = get_claim_type(st03.text if st03 is not None else "")

        grouped[(partner_id, claim_type)].append(root)

    for (partner_id, claim_type), roots in grouped.items():
        if claim_type == "unknown":
            continue

        # Create new merged root
        first_root = roots[0]
        isa_loop = ET.Element("loop", {"id": "ISA_LOOP"})

        # Copy ISA segment from first
        isa_seg = first_root.find(".//loop[@id='ISA_LOOP']/seg[@id='ISA']")
        isa_loop.append(isa_seg)

        # Create GS_LOOP
        gs_loop = ET.SubElement(isa_loop, "loop", {"id": "GS_LOOP"})

        # Copy GS and ST (with HEADER) from first
        gs_seg = first_root.find(".//loop[@id='GS_LOOP']/seg[@id='GS']")
        gs_loop.append(gs_seg)

        st_loop = ET.SubElement(gs_loop, "loop", {"id": "ST_LOOP"})
        st_seg = first_root.find(".//loop[@id='ST_LOOP']/seg[@id='ST']")
        st_loop.append(st_seg)

        header = first_root.find(".//loop[@id='ST_LOOP']/loop[@id='HEADER']")
        if header is not None:
            st_loop.append(header)

        # Merge DETAIL loops from all roots
        for root in roots:
            for detail in root.findall(".//loop[@id='ST_LOOP']/loop[@id='DETAIL']"):
                st_loop.append(detail)

        # Take SE from last root
        last_root = roots[-1]
        se_seg = last_root.find(".//loop[@id='ST_LOOP']/seg[@id='SE']")
        st_loop.append(se_seg)

        # Add GE and IEA from last root
        ge_seg = last_root.find(".//loop[@id='GS_LOOP']/seg[@id='GE']")
        isa_loop.append(ge_seg)

        iea_seg = last_root.find(".//seg[@id='IEA']")
        isa_loop.append(iea_seg)

        # --- Fix Counts ---
        # Update SE01 (# of included segments in ST loop)
        se01 = se_seg.find("./ele[@id='SE01']")
        if se01 is not None:
            total_segments = len(st_loop.findall(".//seg"))
            se01.text = str(total_segments)

        # Update GE01 (# of transaction sets included)
        ge01 = ge_seg.find("./ele[@id='GE01']")
        if ge01 is not None:
            ge01.text = "1"

        # Update IEA01 (# of functional groups)
        iea01 = iea_seg.find("./ele[@id='IEA01']")
        if iea01 is not None:
            iea01.text = "1"

        # --- Write merged file ---
        output_filename = f"{claim_type}_{partner_id}_merged.DAT"
        output_path = os.path.join(output_folder, output_filename)

        with open(output_path, "wb") as f:
            f.write(ET.tostring(isa_loop, encoding="utf-8"))

        print(f"Created merged file: {output_filename}")




import os
import xml.etree.ElementTree as ET
from collections import defaultdict

def get_claim_type(st_impl):
    """
    Determine claim type from ST03 value
    """
    if "X222" in st_impl or "005010X222" in st_impl:
        return "institutional"
    elif "X223" in st_impl or "005010X223" in st_impl:
        return "professional"
    else:
        return "unknown"

def extract_segments(root):
    """Extract loops: ISA, GS, ST, Header, Body, Footer"""
    isa_loop = [elem.text for elem in root.findall(".//ISA_loop//Segment")]
    gs_loop = [elem.text for elem in root.findall(".//GS_loop//Segment")]
    st_loop = [elem.text for elem in root.findall(".//ST_loop//Segment")]
    header_loop = [elem.text for elem in root.findall(".//Header_loop//Segment")]
    body = [elem.text for elem in root.findall(".//Header_loop//Body//Segment")]
    footer = [elem.text for elem in root.findall(".//Footer_loop//Segment")]
    return isa_loop, gs_loop, st_loop, header_loop, body, footer

def merge_edi_roots(xml_roots, output_folder):
    # Group by (ISA06, claim_type)
    grouped = defaultdict(list)

    for root in xml_roots:
        # Trading partner id
        isa06 = root.find(".//ISA06").text if root.find(".//ISA06") is not None else "UNKNOWN"
        # Claim type
        st_elem = root.find(".//ST03")
        st_impl = st_elem.text if st_elem is not None else ""
        claim_type = get_claim_type(st_impl)

        grouped[(isa06, claim_type)].append(root)

    # Process each group
    for (partner_id, claim_type), roots in grouped.items():
        if claim_type == "unknown":
            continue

        merged_segments = []
        all_body_segments = []
        footer_segments = []

        for idx, root in enumerate(roots):
            isa, gs, st, header, body, footer = extract_segments(root)

            if idx == 0:
                # Add header sections only once
                merged_segments.extend(isa)
                merged_segments.extend(gs)
                merged_segments.extend(st)
                merged_segments.extend(header)

            # Always collect body
            all_body_segments.extend(body)

            if idx == len(roots) - 1:
                # Footer only from last file
                footer_segments = footer

        # Final merged file content
        merged_segments.extend(all_body_segments)
        merged_segments.extend(footer_segments)

        # Write out to file
        output_filename = f"{claim_type}_{partner_id}_merged.DAT"
        output_path = os.path.join(output_folder, output_filename)
        with open(output_path, "w") as f:
            f.write("~".join(merged_segments) + "~")

        print(f"Created merged file: {output_filename}")


# Example usage
if __name__ == "__main__":
    input_folder = "source_xmls"
    output_folder = "merged_output"
    os.makedirs(output_folder, exist_ok=True)

    xml_roots = []
    for filename in os.listdir(input_folder):
        if filename.endswith(".xml"):
            tree = ET.parse(os.path.join(input_folder, filename))
            xml_roots.append(tree.getroot())

    merge_edi_roots(xml_roots, output_folder)











import os
import xml.etree.ElementTree as ET
from collections import defaultdict

def get_claim_type(st_impl):
    """
    Determine claim type from ST03 value
    Institutional: contains 'X222' or '005010X222'
    Professional: contains 'X223' or '005010X223'
    """
    if "X222" in st_impl or "005010X222" in st_impl:
        return "institutional"
    elif "X223" in st_impl or "005010X223" in st_impl:
        return "professional"
    else:
        return "unknown"

def merge_edi_roots(xml_roots, output_folder):
    # Group roots by (trading_partner_id, claim_type)
    grouped = defaultdict(list)

    for root in xml_roots:
        # Extract trading partner id (ISA06)
        isa06 = root.find(".//ISA06").text if root.find(".//ISA06") is not None else "UNKNOWN"

        # Extract ST03 to get claim type
        st_elem = root.find(".//ST03")
        st_impl = st_elem.text if st_elem is not None else ""
        claim_type = get_claim_type(st_impl)

        grouped[(isa06, claim_type)].append(root)

    # Merge per group
    for (partner_id, claim_type), roots in grouped.items():
        if claim_type == "unknown":
            continue  # skip unknowns

        # Build merged EDI segments
        merged_segments = []

        # Take header from the first root
        header_segments = []
        footer_segments = []
        body_segments = []

        for idx, root in enumerate(roots):
            segments = [elem.text for elem in root.findall(".//Segment")]
            if not segments:
                continue

            # ISA..BHT header (keep only from first file)
            if idx == 0:
                header_segments = segments[:segments.index("BHT")+1]

            # Footer (SE..IEA, keep only from last file)
            if idx == len(roots) - 1:
                footer_segments = segments[segments.index("SE"):]

            # Body (everything between BHT and SE)
            body_start = segments.index("BHT")+1
            body_end = segments.index("SE")
            body_segments.extend(segments[body_start:body_end])

        merged_segments.extend(header_segments)
        merged_segments.extend(body_segments)
        merged_segments.extend(footer_segments)

        # Write merged EDI file
        output_filename = f"{claim_type}_{partner_id}_merged.DAT"
        output_path = os.path.join(output_folder, output_filename)

        with open(output_path, "w") as f:
            f.write("~".join(merged_segments) + "~")

        print(f"Created merged file: {output_filename}")


# Example usage:
if __name__ == "__main__":
    # Example: parse all XML files from folder into roots
    input_folder = "source_xmls"
    output_folder = "merged_output"
    os.makedirs(output_folder, exist_ok=True)

    xml_roots = []
    for filename in os.listdir(input_folder):
        if filename.endswith(".xml"):
            tree = ET.parse(os.path.join(input_folder, filename))
            xml_roots.append(tree.getroot())

    # Merge based on trading partner & claim type
    merge_edi_roots(xml_roots, output_folder)











import os
import hashlib
import xml.etree.ElementTree as ET
from copy import deepcopy

# -------------------
# Fingerprint Logic
# -------------------
def fingerprint_element(elem: ET.Element) -> str:
    """Create deterministic fingerprint for claim loops to deduplicate."""
    parts = []
    for node in elem.iter():
        if node.tag == 'seg':
            segid = node.attrib.get('id', '')
            ele_texts = [(ele.text or '').strip() for ele in node.findall('ele')]
            parts.append(segid + "/" + "|".join(ele_texts))
    return hashlib.sha256("||".join(parts).encode("utf-8")).hexdigest()

# -------------------
# Convert XML seg → EDI line
# -------------------
def parse_segment(seg: ET.Element) -> str:
    seg_id = seg.attrib["id"]
    element_dict = {}

    for child in seg:
        if child.tag == "ele":
            idx_str = child.attrib.get("id", "").replace(seg_id, "")
            if idx_str.isdigit():
                element_dict[int(idx_str)] = child.text or ""

        elif child.tag == "comp":
            sub_values = []
            for sub in child:
                if sub.tag == "subele":
                    sub_values.append(sub.text or "")
            idx_str = sub.attrib.get("id", "").replace(seg_id, "")
            if idx_str.isdigit():
                element_dict[int(idx_str)] = ":".join(sub_values)

    # preserve gaps
    max_idx = max(element_dict.keys() or [0])
    elements = [seg_id] + [element_dict.get(i, "") for i in range(1, max_idx + 1)]
    return "*".join(elements) + "~"

# -------------------
# Fix Control Counts
# -------------------
def fix_control_counts(root: ET.Element):
    """Fix SE, GE, IEA counts after merge."""
    # SE01: number of segments between ST and SE (including both)
    segments = [seg for seg in root.findall(".//seg")]
    st_index = None
    for i, seg in enumerate(segments):
        if seg.attrib.get("id") == "ST":
            st_index = i
        elif seg.attrib.get("id") == "SE" and st_index is not None:
            count = i - st_index + 1
            se01 = seg.find("ele[@id='SE01']")
            if se01 is not None:
                se01.text = str(count)

    # GE01 = number of STs
    st_count = len(root.findall(".//seg[@id='ST']"))
    for ge in root.findall(".//seg[@id='GE']"):
        ge01 = ge.find("ele[@id='GE01']")
        if ge01 is not None:
            ge01.text = str(st_count)

    # IEA01 = number of GSs
    gs_count = len(root.findall(".//seg[@id='GS']"))
    for iea in root.findall(".//seg[@id='IEA']"):
        iea01 = iea.find("ele[@id='IEA01']")
        if iea01 is not None:
            iea01.text = str(gs_count)

# -------------------
# Extract Trading Partner + Claim Type
# -------------------
def extract_ids(root: ET.Element):
    isa06 = None
    st01 = None
    for seg in root.findall(".//seg[@id='ISA']"):
        elems = seg.findall("ele")
        if len(elems) >= 6:
            isa06 = (elems[5].text or "").strip()
    for seg in root.findall(".//seg[@id='ST']"):
        elems = seg.findall("ele")
        if elems:
            st01 = (elems[0].text or "").strip()
    return isa06, st01

# -------------------
# Merge EDI XMLs
# -------------------
def merge_edis(xml_roots, output_file):
    if not xml_roots:
        raise ValueError("No EDI roots provided")

    isa06_list, st01_list = [], []
    for root in xml_roots:
        isa06, st01 = extract_ids(root)
        isa06_list.append(isa06)
        st01_list.append(st01)

    if len(set(isa06_list)) != 1 or len(set(st01_list)) != 1:
        raise ValueError("Trading partner ISA06 or Claim type ST01 mismatch – cannot merge.")

    merged_root = ET.Element("EDI837")
    seen = set()

    # Headers (take only from first)
    first_root = xml_roots[0]
    for seg in first_root.findall(".//seg[@id='ISA']") + \
                first_root.findall(".//seg[@id='GS']") + \
                first_root.findall(".//seg[@id='ST']") + \
                first_root.findall(".//seg[@id='BHT']"):
        merged_root.append(deepcopy(seg))

    # Add body loops (deduplicated)
    for root in xml_roots:
        for loop in root.findall(".//loop[@id='2000A']") + \
                    root.findall(".//loop[@id='2000B']") + \
                    root.findall(".//loop[@id='2300']"):
            fp = fingerprint_element(loop)
            if fp in seen:
                continue
            seen.add(fp)
            merged_root.append(deepcopy(loop))

    # Footer (from first)
    for seg in first_root.findall(".//seg[@id='SE']") + \
                first_root.findall(".//seg[@id='GE']") + \
                first_root.findall(".//seg[@id='IEA']"):
        merged_root.append(deepcopy(seg))

    # Fix control counters
    fix_control_counts(merged_root)

    # Convert back to EDI text
    edi_lines = [parse_segment(seg) for seg in merged_root.findall(".//seg")]
    edi_text = "\n".join(edi_lines)

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(edi_text)

    print(f"✅ Merged EDI written: {output_file}")
    return edi_text











import os

def read_edi_file(filepath):
    """Read an EDI file and return list of segments."""
    with open(filepath, "r") as f:
        content = f.read().strip()
    segments = content.split("~")
    segments = [seg.strip() for seg in segments if seg.strip()]
    return segments

def split_edi_parts(segments):
    """Split EDI into header, claims, footer."""
    # ISA/GS header
    header = []
    # GE/IEA footer
    footer = []
    # claim body (loops)
    body = []

    in_header = True
    in_footer = False

    for seg in segments:
        seg_id = seg[:3]

        if seg_id in ("GE", "IEA"):
            in_footer = True

        if in_header:
            header.append(seg)
            if seg_id == "ST":  # ST marks start of claim set
                in_header = False
        elif in_footer:
            footer.append(seg)
        else:
            body.append(seg)

    return header, body, footer


def merge_edis(file1, file2, outputfile):
    # Read both files
    segs1 = read_edi_file(file1)
    segs2 = read_edi_file(file2)

    # Split header/body/footer
    header1, body1, footer1 = split_edi_parts(segs1)
    header2, body2, footer2 = split_edi_parts(segs2)

    # For merging → keep header1 + combine claim bodies + footer1
    merged_segments = header1 + body1 + body2 + footer1

    # Update segment counts (GE, IEA, SE)
    merged_segments = update_control_segments(merged_segments)

    # Write output
    with open(outputfile, "w") as f:
        f.write("~".join(merged_segments) + "~")

    print(f"✅ Merged EDI written to {outputfile}")


def update_control_segments(segments):
    """Fix counts in SE, GE, IEA segments after merging."""
    new_segments = []
    seg_count = 0
    for seg in segments:
        seg_id = seg[:3]
        if seg_id == "SE":
            parts = seg.split("*")
            parts[1] = str(seg_count + 1)  # SE segment count
            seg = "*".join(parts)
        elif seg_id == "GE":
            parts = seg.split("*")
            parts[1] = "1"  # only one functional group
            seg = "*".join(parts)
        elif seg_id == "IEA":
            parts = seg.split("*")
            parts[1] = "1"  # only one interchange
            seg = "*".join(parts)

        new_segments.append(seg)
        seg_count += 1

    return new_segments


if __name__ == "__main__":
    file1 = "837_file1.edi"
    file2 = "837_file2.edi"
    outputfile = "837_merged.edi"

    merge_edis(file1, file2, outputfile)








import os
import hashlib
import xml.etree.ElementTree as ET
from copy import deepcopy

# -------------------
# Fingerprint Logic
# -------------------
def fingerprint_element(elem: ET.Element) -> str:
    """Create deterministic fingerprint for claim loops to deduplicate."""
    parts = []
    for node in elem.iter():
        if node.tag == 'seg':
            segid = node.attrib.get('id', '')
            ele_texts = []
            for ele in node.findall('ele'):
                ele_texts.append((ele.text or '').strip())
            parts.append(segid + "/" + "|".join(ele_texts))
    joined = "||".join(parts)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()

# -------------------
# Convert XML seg → EDI line
# -------------------
def parse_segment(seg: ET.Element) -> str:
    seg_id = seg.attrib["id"]
    element_dict = {}

    # capture ele children
    for child in seg:
        if child.tag == "ele":
            idx_str = child.attrib.get("id", "").replace(seg_id, "")
            if idx_str.isdigit():
                idx = int(idx_str)
                element_dict[idx] = child.text or ""

        elif child.tag == "comp":
            sub_values = []
            for sub in child:
                if sub.tag == "subele":
                    sub_values.append(sub.text or "")
            idx_str = sub.attrib.get("id", "").replace(seg_id, "")
            if idx_str.isdigit():
                idx = int(idx_str)
                element_dict[idx] = ":".join(sub_values)

    # build with gaps ("***")
    max_idx = max(element_dict.keys() or [0])
    elements = [seg_id] + [element_dict.get(i, "") for i in range(1, max_idx+1)]
    return "*".join(elements) + "~"

# -------------------
# Fix Control Counts
# -------------------
def fix_control_counts(root: ET.Element):
    """Fix SE, GE, IEA counts."""
    # Count segments between ST and SE
    seg_count = 0
    for seg in root.findall(".//seg"):
        if seg.attrib.get("id") == "ST":
            seg_count = 1  # count ST itself
        elif seg.attrib.get("id") == "SE":
            se01 = seg.find("ele[@id='SE01']")
            if se01 is not None:
                se01.text = str(seg_count + 1)  # include SE
        else:
            seg_count += 1

    # GE01 = number of ST segments
    st_count = len(root.findall(".//seg[@id='ST']"))
    for ge in root.findall(".//seg[@id='GE']"):
        ge01 = ge.find("ele[@id='GE01']")
        if ge01 is not None:
            ge01.text = str(st_count)

    # IEA01 = number of GS segments
    gs_count = len(root.findall(".//seg[@id='GS']"))
    for iea in root.findall(".//seg[@id='IEA']"):
        iea01 = iea.find("ele[@id='IEA01']")
        if iea01 is not None:
            iea01.text = str(gs_count)

# -------------------
# Merge EDI XMLs
# -------------------
def merge_edis(xml_roots, output_file):
    merged_root = ET.Element("EDI837")
    seen = set()

    # Header from first file
    first_root = xml_roots[0]
    for seg in first_root.findall(".//seg[@id='ISA']") + \
                first_root.findall(".//seg[@id='GS']") + \
                first_root.findall(".//seg[@id='ST']") + \
                first_root.findall(".//seg[@id='HEADER']"):
        merged_root.append(deepcopy(seg))

    # Claims (dedup)
    for root in xml_roots:
        for claim in root.findall(".//loop[@id='2000A']") + \
                     root.findall(".//loop[@id='2000B']") + \
                     root.findall(".//loop[@id='2300']"):
            fp = fingerprint_element(claim)
            if fp in seen:
                continue
            seen.add(fp)
            merged_root.append(deepcopy(claim))

    # Footer from first file
    for seg in first_root.findall(".//seg[@id='SE']") + \
                first_root.findall(".//seg[@id='GE']") + \
                first_root.findall(".//seg[@id='IEA']"):
        merged_root.append(deepcopy(seg))

    # Fix counters
    fix_control_counts(merged_root)

    # Convert back to EDI text
    edi_lines = [parse_segment(seg) for seg in merged_root.findall(".//seg")]
    edi_text = "\n".join(edi_lines)

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(edi_text)

    print(f"✅ Merged EDI written: {output_file}")
    return edi_text

# -------------------
# Example usage
# -------------------
if __name__ == "__main__":
    # Suppose you already built two XML rootnodes:
    # root1 = load_edi_segments_as_xml_obj(segments1)
    # root2 = load_edi_segments_as_xml_obj(segments2)
    # For demo I assume you already have them
    xml_roots = [root1, root2]

    out = merge_edis(xml_roots, "merged_837.dat")
    print(out[:500])  # print first 500 chars














if heal_file:
    valid_edi_text = heal_file.read().decode("utf-8")

    if st.button("Generate Healing Suggestions"):
        proposals = propose_fixes(
            invalid_results,
            valid_edi_text,
            st.session_state["rules_dict"],
            llm
        )
        st.session_state["proposals"] = proposals  # ✅ save proposals in session
        st.session_state["approvals"] = []         # reset approvals when regenerating

    # Show proposals if present
    if "proposals" in st.session_state and st.session_state["proposals"]:
        st.write("### Proposed Fixes (approve to apply):")

        # Load or create approvals state
        if "approvals" not in st.session_state:
            st.session_state["approvals"] = []

        approvals = []
        for i, p in enumerate(st.session_state["proposals"]):
            seg, pos = p.get("segment"), p.get("position")
            label = f"{seg}:{pos}"
            if "sub_pos" in p:
                label += f"-{p['sub_pos']}"

            with st.expander(label):
                st.json(p)

            approve = st.checkbox(
                f"Approve {label}",
                key=f"approve_{i}",
                value=(p in st.session_state["approvals"])  # ✅ persist across reruns
            )
            if approve:
                approvals.append(p)

        # Update session approvals
        st.session_state["approvals"] = approvals

        # Apply fixes
        if st.button("Apply Approved Fixes"):
            modified_rules = st.session_state["rules_dict"]

            for p in st.session_state["approvals"]:
                seg, pos = p["segment"], p["position"]
                sub_pos = p.get("sub_pos")
                change = p["proposed_change"]

                if sub_pos is None:  # ✅ field-level
                    if "update_usage" in change:
                        modified_rules[seg][pos]["usage"] = change["update_usage"]

                    if change.get("empty_codes"):
                        modified_rules[seg][pos]["accepted_codes"] = []

                    if "update_codes" in change:
                        modified_rules[seg][pos]["accepted_codes"] = change["update_codes"]

                    if "add_code" in change:
                        modified_rules[seg][pos].setdefault("accepted_codes", [])
                        modified_rules[seg][pos]["accepted_codes"].append(change["add_code"])

                else:  # ✅ subfield-level
                    subfields = modified_rules[seg][pos].get("subfields", [])
                    for sf in subfields:
                        if sf["sub_pos"] == sub_pos:
                            if "update_usage" in change:
                                sf["usage"] = change["update_usage"]

                            if change.get("empty_codes"):
                                sf["accepted_codes"] = []

                            if "update_codes" in change:
                                sf["accepted_codes"] = change["update_codes"]

                            if "add_code" in change:
                                sf.setdefault("accepted_codes", [])
                                sf["accepted_codes"].append(change["add_code"])

            # Save new version
            version_name = save_version(modified_rules)
            st.success(f"✅ Saved healed rules as version {version_name}")
            st.session_state["rules_dict"] = modified_rules
            

if heal_file:
    valid_edi_text = heal_file.read().decode("utf-8")

    if st.button("Generate Healing Suggestions"):
        proposals = propose_fixes(
            invalid_results,
            valid_edi_text,
            st.session_state["rules_dict"],
            llm
        )

        if proposals:
            st.write("### Proposed Fixes (approve to apply):")
            approvals = []

            for i, p in enumerate(proposals):
                seg, pos = p.get("segment"), p.get("position")
                label = f"{seg}:{pos}"
                if "sub_pos" in p:
                    label += f"-{p['sub_pos']}"

                with st.expander(label):
                    st.json(p)

                approve = st.checkbox(f"Approve {label}", key=f"approve_{i}")
                if approve:
                    approvals.append(p)

            if st.button("Apply Approved Fixes"):
                modified_rules = st.session_state["rules_dict"]

                for p in approvals:
                    seg, pos = p["segment"], p["position"]
                    sub_pos = p.get("sub_pos")
                    change = p["proposed_change"]

                    # Field-level
                    if sub_pos is None:
                        if "update_usage" in change:
                            modified_rules[seg][pos]["usage"] = change["update_usage"]

                        if "empty_codes" in change and change["empty_codes"]:
                            modified_rules[seg][pos]["accepted_codes"] = []

                        if "update_codes" in change:
                            modified_rules[seg][pos]["accepted_codes"] = change["update_codes"]

                        if "add_code" in change:
                            if "accepted_codes" not in modified_rules[seg][pos]:
                                modified_rules[seg][pos]["accepted_codes"] = []
                            modified_rules[seg][pos]["accepted_codes"].append(change["add_code"])

                    # Subfield-level
                    else:
                        subfields = modified_rules[seg][pos].get("subfields", [])
                        for sf in subfields:
                            if sf["sub_pos"] == sub_pos:
                                if "update_usage" in change:
                                    sf["usage"] = change["update_usage"]

                                if "empty_codes" in change and change["empty_codes"]:
                                    sf["accepted_codes"] = []

                                if "update_codes" in change:
                                    sf["accepted_codes"] = change["update_codes"]

                                if "add_code" in change:
                                    if "accepted_codes" not in sf:
                                        sf["accepted_codes"] = []
                                    sf["accepted_codes"].append(change["add_code"])

                # Save new version
                version_name = save_version(modified_rules)
                st.success(f"✅ Saved healed rules as version {version_name}")
                st.session_state["rules_dict"] = modified_rules
        else:
            st.info("No suggestions generated by LLM.")
else:
    st.info("No invalid results found. Nothing to heal.")







def propose_fixes(invalid_results, valid_edi_text, rules_dict, llm):
    """
    Use LLM to propose minimal rule fixes based on invalid validation results
    and a valid EDI reference. Proposals may include both usage updates and
    accepted code additions inside `proposed_change`. Works for both fields
    and subfields.
    """
    import json
    from langchain.prompts import ChatPromptTemplate

    prompt = ChatPromptTemplate.from_template("""
You are an EDI rules repair assistant.

We have a base rules JSON, invalid validation results, and a valid EDI sample.

Base Rules (do not change structure, only update usage or accepted codes where necessary):
{rules_dict}

Invalid Results (these fields failed validation):
{invalid_results}

Valid EDI (reference):
{valid_edi_text}

Task:
- Suggest minimal modifications to the rules.  
- Only update fields, subfields, or segments that are invalid.  
- Preserve the JSON structure (do not delete or add unrelated keys).  

Rules:
- If a REQUIRED field or subfield is empty in EDI, propose BOTH:
  1. update_usage → "SITUATIONAL"
  2. add_code with "Code": "" and a Definition explaining blank is allowed.  
- If a code is present in EDI but missing from accepted_codes, propose to add it.  
- If accepted_codes exist but are wrong, propose either:
  - replacement (with `update_codes`)  
  - or clearing them (`empty_codes: true`).  
- If a missing subfield is detected, include "sub_pos" in your proposal so we can locate it.  

Return ONLY a JSON list of suggestions in this format:

[
  {
    "segment": "ISA",
    "position": "02",
    "proposed_change": {
      "update_usage": "SITUATIONAL",
      "add_code": {
        "Code": "",
        "Definition": "Authorization Information may be blank"
      }
    },
    "reason": "Field was REQUIRED in rules but EDI shows no value. Suggest changing to SITUATIONAL so blanks are allowed."
  },
  {
    "segment": "SVC",
    "position": "01",
    "sub_pos": "2",
    "proposed_change": {
      "add_code": {
        "Code": "AD",
        "Definition": "American Dental Association Codes"
      }
    },
    "reason": "Subfield SVC-01-2 value 'AD' is present in EDI but missing from accepted codes."
  }
]
""")

    chain = prompt | llm
    response = chain.invoke({
        "rules_dict": json.dumps(rules_dict, indent=2),
        "invalid_results": json.dumps(invalid_results, indent=2),
        "valid_edi_text": valid_edi_text
    })

    try:
        parsed_rules = response.content.strip().replace("```json", "").replace("```", "")
        return json.loads(parsed_rules)
    except Exception:
        return []
       






from langchain.prompts import ChatPromptTemplate

prompt = ChatPromptTemplate.from_template("""
You are an EDI rules repair assistant.

We have a base rules JSON, invalid validation results, and a valid EDI sample.

Base Rules (do not change structure, only update usage or accepted codes where necessary):
{rules_dict}

Invalid Results (these fields failed validation):
{invalid_results}

Valid EDI (reference):
{valid_edi_text}

---

### Task:
Suggest **minimal modifications** to the rules so they become valid against the EDI.

1. Only update the fields/segments/subfields that are invalid.
2. Preserve the JSON structure (do not delete or add unrelated keys).
3. If you suggest new accepted codes, use the following schema:
   ```json
   {
     "segment": "...",
     "position": "...",
     "sub_position": "..." (optional, if the issue is in a subfield),
     "proposed_change": {
       "add_code": {"Code": "...", "Definition": "..."}
     },
     "reason": "..."
   }






import streamlit as st
import json
import os
import datetime
from typing import Dict, List

from langchain_openai import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate


# ----------------------
# Azure OpenAI Setup
# ----------------------
def get_llm():
    return AzureChatOpenAI(
        openai_api_base=st.session_state.get("AZURE_OPENAI_ENDPOINT"),
        openai_api_version="2023-07-01-preview",
        deployment_name=st.session_state.get("AZURE_DEPLOYMENT_NAME"),
        openai_api_key=st.session_state.get("AZURE_OPENAI_KEY"),
        temperature=0
    )


# ----------------------
# Helpers
# ----------------------

def load_json(path):
    with open(path, "r") as f:
        return json.load(f)

def save_json(obj, path):
    with open(path, "w") as f:
        json.dump(obj, f, indent=2)

def save_version(rules_dict, version_dir="rule_versions"):
    os.makedirs(version_dir, exist_ok=True)
    version_name = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(version_dir, f"rules_{version_name}.json")
    save_json(rules_dict, path)
    return version_name, path

def list_versions(version_dir="rule_versions"):
    if not os.path.exists(version_dir):
        return []
    return sorted(os.listdir(version_dir))

def load_version(filename, version_dir="rule_versions"):
    path = os.path.join(version_dir, filename)
    return load_json(path)


# ----------------------
# Healing with LLM
# ----------------------

def propose_fixes(invalid_results, valid_edi_text, rules_dict):
    llm = get_llm()
    prompt = ChatPromptTemplate.from_template("""
You are an EDI rules repair assistant.
We have a base rules JSON, invalid validation results, and a valid EDI sample.

Base Rules (do not change structure, only update usage or accepted codes where necessary):
{rules_dict}

Invalid Results (these fields failed validation):
{invalid_results}

Valid EDI (reference):
{valid_edi_text}

Task:
- Suggest minimal modifications to the rules.
- Only update fields and segments that are invalid.
- Preserve the JSON structure (do not delete or add unrelated keys).
- If you suggest new accepted codes, return them in the form:
  {{"segment": "...", "position": "...", "proposed_add": "...", "reason": "..."}}

Return ONLY a JSON list of suggestions.
""")

    chain = prompt | llm
    response = chain.invoke({
        "rules_dict": json.dumps(rules_dict, indent=2),
        "invalid_results": json.dumps(invalid_results, indent=2),
        "valid_edi_text": valid_edi_text
    })

    try:
        return json.loads(response.content)
    except:
        return []


# ----------------------
# Healing UI
# ----------------------

st.subheader("Heal Rules")

# Azure credentials in sidebar
with st.sidebar:
    st.text_input("Azure OpenAI Endpoint", key="AZURE_OPENAI_ENDPOINT")
    st.text_input("Deployment Name", key="AZURE_DEPLOYMENT_NAME")
    st.text_input("API Key", type="password", key="AZURE_OPENAI_KEY")

# Load validation results
if os.path.exists("validation_results.json") and "rules_dict" in st.session_state:
    results = load_json("validation_results.json")
    invalid_results = [r for r in results if r["status"] != "Valid"]

    if invalid_results:
        heal_file = st.file_uploader("Upload valid EDI for healing", type=["txt", "edi"], key="heal")

        if heal_file:
            valid_edi_text = heal_file.read().decode("utf-8")

            if st.button("Generate Healing Suggestions"):
                proposals = propose_fixes(invalid_results, valid_edi_text, st.session_state["rules_dict"])

                if proposals:
                    st.write("🔍 Proposed Fixes (approve to apply):")

                    approvals = []
                    for i, p in enumerate(proposals):
                        with st.expander(f"{p['segment']}:{p['position']}"):
                            st.json(p)
                            approve = st.checkbox(f"Approve {p['segment']}:{p['position']}", key=f"approve_{i}")
                            if approve:
                                approvals.append(p)

                    if st.button("Apply Approved Fixes"):
                        modified_rules = st.session_state["rules_dict"]

                        for p in approvals:
                            seg, pos = p["segment"], p["position"]
                            if "accepted_codes" not in modified_rules[seg][pos]:
                                modified_rules[seg][pos]["accepted_codes"] = []
                            modified_rules[seg][pos]["accepted_codes"].append(
                                {"Code": p["proposed_add"], "Definition": "Healed by LLM + user approval"}
                            )

                        version_name, _ = save_version(modified_rules)
                        st.success(f"Saved healed rules as version {version_name}")
                        st.session_state["rules_dict"] = modified_rules
                else:
                    st.info("No suggestions generated by LLM.")
    else:
        st.info("✅ No invalid results found. Nothing to heal.")













import streamlit as st
import json
import os
import datetime
from typing import Dict, List

# ----------------------
# Helper functions
# ----------------------

def load_rules(file) -> Dict:
    return json.load(file)

def convert_rules(json_rules: List[Dict]) -> Dict:
    """Convert flat list into nested rules_dict (like before)."""
    rules_dict = {}
    for rule in json_rules:
        seg = rule.get("SegmentName", "").strip()
        pos = str(rule.get("FieldPosition", "")).zfill(2)
        subpos = str(rule.get("SubPosition", "")).strip()

        if not seg or not pos:
            continue

        if seg not in rules_dict:
            rules_dict[seg] = {}

        if pos not in rules_dict[seg]:
            rules_dict[seg][pos] = {
                "usage": rule.get("Usage", ""),
                "description": rule.get("ShortDescription", ""),
                "accepted_codes": rule.get("AcceptedCodes", []),
                "subfields": {}
            }

        # Handle subfields
        if subpos:
            rules_dict[seg][pos]["subfields"][subpos] = {
                "usage": rule.get("Usage", ""),
                "description": rule.get("ShortDescription", ""),
                "accepted_codes": rule.get("AcceptedCodes", [])
            }

    return rules_dict

def validate_edi(edi_text: str, rules: Dict) -> List[Dict]:
    """Fake validator stub - plug your validate_segment here"""
    results = [{"edi_line": line, "status": "Valid"} for line in edi_text.split("~") if line.strip()]
    return results

def save_version(rules_dict, version_dir="rule_versions"):
    """Save a new version of rules_dict with timestamp."""
    os.makedirs(version_dir, exist_ok=True)
    version_name = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(version_dir, f"rules_{version_name}.json")
    with open(path, "w") as f:
        json.dump(rules_dict, f, indent=2)
    return version_name, path

def list_versions(version_dir="rule_versions"):
    if not os.path.exists(version_dir):
        return []
    return sorted(os.listdir(version_dir))

def load_version(filename, version_dir="rule_versions"):
    path = os.path.join(version_dir, filename)
    with open(path, "r") as f:
        return json.load(f)


# ----------------------
# Streamlit UI
# ----------------------

st.sidebar.title("EDI Rule Validator")

# Upload base rules
st.sidebar.subheader("Base Rules")
base_rules_file = st.sidebar.file_uploader("Upload base rules.json", type="json")

llm_details = st.sidebar.text_area("LLM Details (API keys, endpoint, etc.)", "")

if base_rules_file:
    base_rules = load_rules(base_rules_file)
    rules_dict = convert_rules(base_rules)
    st.session_state["base_rules"] = base_rules
    st.session_state["rules_dict"] = rules_dict
    st.sidebar.success("Base rules loaded and converted ✅")

    if st.sidebar.button("Save Initial Rules Version"):
        version_name, _ = save_version(rules_dict)
        st.sidebar.success(f"Saved as version {version_name}")


# Main page
st.title("EDI Validator")

# Upload EDI
edi_file = st.file_uploader("Upload EDI file", type=["txt", "edi"])

if edi_file and "rules_dict" in st.session_state:
    edi_text = edi_file.read().decode("utf-8")

    if st.button("Validate"):
        results = validate_edi(edi_text, st.session_state["rules_dict"])
        with open("validation_results.json", "w") as f:
            json.dump(results, f, indent=2)
        st.json(results)
        st.success("Validation results saved to validation_results.json ✅")


# Healing Section
st.subheader("Heal Rules")
if edi_file and "rules_dict" in st.session_state:
    heal_file = st.file_uploader("Upload valid EDI for healing", type=["txt", "edi"], key="heal")
    if heal_file and st.button("Modify Rules"):
        edi_text = heal_file.read().decode("utf-8")
        # TODO: Add logic: analyze edi_text + failed validations → update rules
        # For now, simulate a modification
        modified_rules = st.session_state["rules_dict"]
        modified_rules["ISA"]["01"]["accepted_codes"].append({"Code": "99", "Definition": "Test Override"})

        version_name, _ = save_version(modified_rules)
        st.success(f"Rules healed and saved as version {version_name}")


# Version Management
st.subheader("Rule Versions")
versions = list_versions()
if versions:
    selected_version = st.selectbox("Select version", versions)
    if st.button("Load Version"):
        st.session_state["rules_dict"] = load_version(selected_version)
        st.success(f"Loaded {selected_version}")










import re
from typing import Dict, List

# Example: rules dict after your JSON conversion
# rules = {
#   "ISA": {
#     "01": {
#       "usage": "REQUIRED",
#       "description": "Authorization Information Qualifier",
#       "accepted_codes": [
#           {"Code": "00", "Definition": "No Authorization Information Present"},
#           {"Code": "03", "Definition": "Additional Data Identification"}
#       ]
#     },
#     "02": {...}
#   }
# }

REQUIRED_SEGMENTS = {"ISA", "GS", "ST"}  # adjust as per your guide


def validate_segment(segment: str, rules: Dict) -> Dict:
    """
    Validate one EDI segment against rules.
    Handles subfields (e.g., SVC*01-1:01-2*...).
    """
    elements = segment.split("*")
    seg_id = elements[0].strip()

    field_reasons = {}
    overall_status = "Matched"

    # Syntax check
    if not seg_id.isalpha():
        return {
            "edi_line": segment,
            "status": "Invalid",
            "rule_line": seg_id,
            "reason": {"Syntax": "Invalid segment ID (not alphabetic)"}
        }

    # Remove from required segments if present
    if seg_id in REQUIRED_SEGMENTS:
        REQUIRED_SEGMENTS.remove(seg_id)

    # Field-level validation
    if seg_id in rules:
        seg_rules = rules[seg_id]

        for pos, rule in seg_rules.items():
            idx = int(pos)  # 01 -> 1, etc.

            try:
                value = elements[idx]
            except IndexError:
                field_reasons[f"{seg_id}-{pos}"] = f"Invalid: Missing required field ({rule['description']})"
                overall_status = "Invalid"
                continue

            # Subfield check (e.g., 01-1:01-2)
            subfields = value.split(":") if ":" in value else [value]

            # Loop through subfields
            for sub_idx, sub_val in enumerate(subfields, start=1):
                sub_key = f"{seg_id}-{pos}-{sub_idx}" if len(subfields) > 1 else f"{seg_id}-{pos}"

                # Usage check
                if rule.get("usage", "").upper() == "REQUIRED" and not sub_val.strip():
                    field_reasons[sub_key] = f"Invalid: {rule['description']} is required"
                    overall_status = "Invalid"
                    continue

                # Accepted codes check
                valid_codes = []
                valid_defs = {}
                if rule.get("accepted_codes"):
                    if isinstance(rule["accepted_codes"], list):
                        for c in rule["accepted_codes"]:
                            if isinstance(c, dict) and "Code" in c:
                                valid_codes.append(c["Code"])
                                valid_defs[c["Code"]] = c.get("Definition", "")
                            elif isinstance(c, str):
                                valid_codes.append(c)
                                valid_defs[c] = ""

                if valid_codes:
                    if sub_val not in valid_codes:
                        defs_str = ", ".join([f"{c}: {valid_defs[c]}" if valid_defs[c] else c for c in valid_codes])
                        field_reasons[sub_key] = f"Invalid: Invalid code '{sub_val}'. Valid codes: {defs_str}"
                        overall_status = "Invalid"
                        continue

                # If passed all checks
                field_reasons[sub_key] = f"Valid: {rule['description']} ({sub_val})"

    else:
        field_reasons[seg_id] = "No field-level rules; syntax OK"

    return {
        "edi_line": segment,
        "status": overall_status,
        "rule_line": seg_id,
        "reason": field_reasons
    }


def validate_edi_file(edi_text: str, rules: Dict) -> List[Dict]:
    """
    Validate full EDI file line by line.
    """
    segments = [seg.strip() for seg in edi_text.split("~") if seg.strip()]
    all_results = [validate_segment(seg, rules) for seg in segments]

    # Final mandatory check – for missing segments
    if REQUIRED_SEGMENTS:
        for seg in list(REQUIRED_SEGMENTS):
            all_results.append({
                "edi_line": "N/A",
                "status": "Invalid",
                "rule_line": seg,
                "reason": {seg: f"Mandatory segment {seg} missing in file"}
            })

    return all_results
    






from typing import List, Dict

# Example REQUIRED_SEGMENTS (you’ll define based on Companion Guide)
REQUIRED_SEGMENTS = {"ISA", "GS", "ST", "SE", "GE", "IEA"}  

def validate_edi_file(edi_text: str, rules: Dict) -> List[Dict]:
    """
    Validate full EDI file line by line against Companion Guide rules.
    Also performs a final check for missing mandatory segments.
    """
    # Split EDI text into segments (~ is the segment terminator)
    segments = [seg.strip() for seg in edi_text.split("~") if seg.strip()]

    all_results = []

    # Copy REQUIRED_SEGMENTS so we don’t mutate global set directly
    remaining_required = set(REQUIRED_SEGMENTS)

    for seg in segments:
        result = validate_segment(seg, rules, remaining_required)
        all_results.append(result)

    # Final mandatory check — see if any required segments were never found
    if remaining_required:
        for seg in remaining_required:
            all_results.append({
                "edi_line": "N/A",
                "status": "Invalid",
                "rule_line": seg,
                "reason": {seg: f"Mandatory segment {seg} missing in file"}
            })

    return all_results
    





def validate_segment(segment, rules, REQUIRED_SEGMENTS, llm=None):
    field_reasons = {}
    overall_status = "Matched"

    # Split into fields
    elements = segment.split("*")

    # Normalize segment ID
    if len(elements[0]) > 3:
        elements[0] = elements[0][:3]

    seg_id = elements[0].strip()

    # -----------------
    # Syntax check
    # -----------------
    if not seg_id.isalpha():
        return {
            "edi line": segment,
            "status": "Invalid",
            "rule line": seg_id,
            "reason": {"Syntax": "Invalid segment ID (not alphabetic)"}
        }

    # -----------------
    # Mandatory check
    # -----------------
    if seg_id in REQUIRED_SEGMENTS:
        REQUIRED_SEGMENTS.remove(seg_id)

    # -----------------
    # Field-level validation (including subfields)
    # -----------------
    if seg_id in rules:
        seg_rules = rules[seg_id]

        for pos, rule in seg_rules.items():
            idx = int(pos)
            desc = rule["description"]

            try:
                value = elements[idx]
            except IndexError:
                field_reasons[f"{seg_id}-{pos}"] = f"Invalid: Missing required field ({desc})"
                overall_status = "Invalid"
                continue

            # --- Subfield handling ---
            if ":" in value:
                subfields = value.split(":")
                for sub_idx, sub_val in enumerate(subfields, start=1):
                    sub_key = f"{seg_id}-{pos}-{sub_idx}"

                    if rule.get("usage") == "Required" and not sub_val.strip():
                        field_reasons[sub_key] = f"Invalid: Subfield {sub_idx} of {desc} is required"
                        overall_status = "Invalid"
                        continue

                    if rule.get("accepted_codes") and sub_val not in rule["accepted_codes"]:
                        llm_reason = None
                        if llm:
                            llm_reason = llm([
                                HumanMessage(content=f"Field {seg_id}-{pos} subfield {sub_idx} "
                                                     f"has invalid code '{sub_val}'. "
                                                     f"Valid codes are {rule['accepted_codes']}. "
                                                     f"Describe the error in one sentence.")
                            ])

                        reason = llm_reason.content if llm_reason else f"Invalid code '{sub_val}'"
                        field_reasons[sub_key] = f"Invalid: {reason}"
                        overall_status = "Invalid"
                        continue

                    field_reasons[sub_key] = f"Valid: {desc} (subfield {sub_idx}: {sub_val})"

            else:
                # --- Normal field handling ---
                if rule.get("usage") == "Required" and not value.strip():
                    field_reasons[f"{seg_id}-{pos}"] = f"Invalid: {desc} is required"
                    overall_status = "Invalid"
                    continue

                if rule.get("accepted_codes") and value not in rule["accepted_codes"]:
                    llm_reason = None
                    if llm:
                        llm_reason = llm([
                            HumanMessage(content=f"Field {seg_id}-{pos} has invalid code '{value}'. "
                                                 f"Valid codes are {rule['accepted_codes']}. "
                                                 f"Describe the error in one sentence.")
                        ])
                    reason = llm_reason.content if llm_reason else f"Invalid code '{value}'"
                    field_reasons[f"{seg_id}-{pos}"] = f"Invalid: {reason}"
                    overall_status = "Invalid"
                    continue

                field_reasons[f"{seg_id}-{pos}"] = f"Valid: {desc} ({value})"
    else:
        field_reasons[seg_id] = "No field-level rules; syntax OK"

    # -----------------
    # Final result
    # -----------------
    return {
        "edi line": segment,
        "status": overall_status,
        "rule line": seg_id,
        "reason": field_reasons
    }
    

import json
from typing import List, Dict, Any

def convert_rules(json_rules: List[Dict[str, Any]]) -> Dict:
    """
    Convert flat rules list into nested dict:
    {Segment: {FieldPos: {details + subfields}}}
    Dedupes codes, merges subfields into a list.
    """

    rules_dict: Dict[str, Dict[str, Any]] = {}

    for rule in json_rules:
        seg = rule.get("SegmentName", "").strip()
        field_pos = str(rule.get("FieldPosition", "")).zfill(2)  # "01"
        sub_pos = str(rule.get("SubPosition", "")).strip()

        usage = rule.get("Usage", "").strip()
        desc = rule.get("ShortDescription", "").strip()
        codes = rule.get("AcceptedCodes", [])

        # Normalize codes
        norm_codes = []
        for c in codes:
            if isinstance(c, dict):
                code_val = str(c.get("Code", "")).strip()
                def_val = str(c.get("Definition", "")).strip()
                if code_val:
                    norm_codes.append({"Code": code_val, "Definition": def_val})
            elif isinstance(c, str):
                norm_codes.append({"Code": c, "Definition": ""})

        # Init segment + field
        if seg not in rules_dict:
            rules_dict[seg] = {}
        if field_pos not in rules_dict[seg]:
            rules_dict[seg][field_pos] = {
                "usage": usage,
                "description": desc,
                "accepted_codes": [],
                "subfields": []
            }

        if sub_pos and sub_pos != "":  # It's a subfield
            # Merge into subfields list
            sub_entry = {
                "sub_pos": sub_pos,
                "usage": usage,
                "description": desc,
                "accepted_codes": []
            }
            existing_codes = {(c["Code"], c["Definition"]) for c in sub_entry["accepted_codes"]}
            for c in norm_codes:
                key = (c["Code"], c["Definition"])
                if key not in existing_codes:
                    sub_entry["accepted_codes"].append(c)
                    existing_codes.add(key)

            # Avoid duplicate subfields
            existing_subs = rules_dict[seg][field_pos]["subfields"]
            if not any(sf["sub_pos"] == sub_pos for sf in existing_subs):
                existing_subs.append(sub_entry)

        else:  # It's the main field
            existing_codes = {(c["Code"], c["Definition"]) for c in rules_dict[seg][field_pos]["accepted_codes"]}
            for c in norm_codes:
                key = (c["Code"], c["Definition"])
                if key not in existing_codes:
                    rules_dict[seg][field_pos]["accepted_codes"].append(c)
                    existing_codes.add(key)

    # Save to JSON
    with open("rules_dict.json", "w") as f:
        json.dump(rules_dict, f, indent=2)

    return rules_dict











import pandas as pd

# Load Excel
df = pd.read_excel("input.xlsx")

# Suppose the column name is "date_column"
df["date_column"] = pd.to_datetime(df["date_column"]).dt.strftime("%Y%m%d")

# Save back to Excel
df.to_excel("output.xlsx", index=False)


import json
import re
from typing import List, Dict, Any

def convert_rules(json_rules: List[Dict[str, Any]]) -> Dict:
    """
    Convert flat list of rules into nested dict:
    {
      SegmentName: {
        FieldPosition: {
          SubPosition: {
            "usage": str,
            "description": str,
            "accepted_codes": [{"Code": str, "Definition": str}]
          }
        }
      }
    }
    Handles duplicates by merging accepted codes.
    """

    rules_dict: Dict[str, Dict[str, Dict[str, Any]]] = {}

    for rule in json_rules:
        seg = rule.get("SegmentName", "").strip()
        raw_pos = str(rule.get("FieldPosition", "")).strip()
        sub_pos = str(rule.get("SubPosition", "")).strip()

        # Normalize field position (digits only, padded to 2)
        digits = re.findall(r"\d+", raw_pos)
        if not digits:
            continue
        pos = digits[0].zfill(2)

        usage = rule.get("Usage", "").strip()
        desc = rule.get("ShortDescription", "").strip()
        codes = rule.get("AcceptedCodes", [])

        # Normalize codes into list of dicts
        norm_codes = []
        for c in codes:
            if isinstance(c, dict):
                code_val = c.get("Code", "").strip()
                def_val = c.get("Definition", "").strip()
                if code_val:
                    norm_codes.append({"Code": code_val, "Definition": def_val})
            elif isinstance(c, str):
                norm_codes.append({"Code": c, "Definition": ""})

        # Ensure nesting exists
        if seg not in rules_dict:
            rules_dict[seg] = {}
        if pos not in rules_dict[seg]:
            rules_dict[seg][pos] = {}
        if sub_pos not in rules_dict[seg][pos]:
            rules_dict[seg][pos][sub_pos] = {
                "usage": usage,
                "description": desc,
                "accepted_codes": []
            }

        # Merge accepted codes without duplication
        existing_codes = rules_dict[seg][pos][sub_pos]["accepted_codes"]
        existing_set = {(c["Code"], c["Definition"]) for c in existing_codes}
        for c in norm_codes:
            key = (c["Code"], c["Definition"])
            if key not in existing_set:
                existing_codes.append(c)
                existing_set.add(key)

    # Save to JSON
    with open("rules_dict.json", "w") as f:
        json.dump(rules_dict, f, indent=2)

    return rules_dict
    






prompt = f"""
You are given part of an EDI implementation guideline.

Extract validation rules for segments and fields.

Rules must include these keys:

- "SegmentName": The segment name (example: ISA, GS, ST, SVC).
- "FieldPosition": The field position number (example: 01, 02, 03).
- "SubPosition": The sub-position if available (example: 8), otherwise empty string.
- "Usage": One of [Required, Situational, Not Used].
- "ShortDescription": A short description of the field.
- "AcceptedCodes": A list of objects with keys:
    - "Code": the code value
    - "Definition": explanation of the code

Input text:
{chunk}

Return only a JSON list of rules in this exact format:

[
  {{
    "SegmentName": "ST",
    "FieldPosition": "01",
    "SubPosition": "",
    "Usage": "Required",
    "ShortDescription": "Transaction Set Identifier Code",
    "AcceptedCodes": [
      {{
        "Code": "835",
        "Definition": "Health Care Claim Payment/Advice"
      }}
    ]
  }}
]
"""






import re
from typing import List, Dict, Optional
from langchain_core.documents import Document

# Regex: usage keyword + segment + position + optional sub-position
HEADER_PATTERN = re.compile(
    r"\b(REQUIRED|SITUATIONAL|NOT USED)\b\s+([A-Z]+)(\d{2})(?:\s*-\s*(\d+))?",
    re.IGNORECASE
)

def element_based_parsing(docs: List[Document]) -> List[Dict[str, Optional[str]]]:
    """
    Parse PDF docs into element-based chunks.
    Ensures multi-page elements are merged, and only valid headers are chunk splitters.
    """

    chunks = []
    current_chunk = []
    current_meta = None

    # Iterate through all pages in sequence
    for doc in docs:
        for line in doc.page_content.splitlines():
            line = line.strip()
            if not line:
                continue

            match = HEADER_PATTERN.search(line)
            if match:
                # Found a valid header -> close previous chunk
                if current_chunk and current_meta:
                    chunks.append({
                        "Usage": current_meta[0].upper(),
                        "Element": f"{current_meta[1]}{current_meta[2]}",
                        "SubPos": current_meta[3],
                        "Text": " ".join(current_chunk).strip()
                    })

                # Start new chunk
                current_meta = match.groups()
                current_chunk = [line]
            else:
                # Continuation of current element
                if current_meta:
                    current_chunk.append(line)

    # Save last chunk
    if current_chunk and current_meta:
        chunks.append({
            "Usage": current_meta[0].upper(),
            "Element": f"{current_meta[1]}{current_meta[2]}",
            "SubPos": current_meta[3],
            "Text": " ".join(current_chunk).strip()
        })

    # 🔑 Merge chunks across pages if same element/subpos
    merged = []
    for chunk in chunks:
        if merged and (
            merged[-1]["Element"] == chunk["Element"]
            and merged[-1]["SubPos"] == chunk["SubPos"]
            and merged[-1]["Usage"] == chunk["Usage"]
        ):
            # extend previous chunk text
            merged[-1]["Text"] += " " + chunk["Text"]
        else:
            merged.append(chunk)

    return merged
    

import re
from typing import List, Dict, Optional

def extract_chunks(text: str) -> List[Dict[str, Optional[str]]]:
    """
    Extract chunks from text based on usage keywords followed by segment/position.
    Example valid header: 'REQUIRED GS01 - 1' or 'SITUATIONAL ISA14'
    """

    # Regex pattern: usage + segment + position + optional sub-position
    pattern = re.compile(
        r"\b(REQUIRED|SITUATIONAL|NOT USED)\b\s+([A-Z]+)(\d{2})(?:\s*-\s*(\d+))?",
        re.IGNORECASE
    )

    chunks = []
    current_chunk = []
    current_meta = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        match = pattern.search(line)
        if match:
            # save old chunk
            if current_chunk and current_meta:
                chunks.append({
                    "Usage": current_meta[0].upper(),  # REQUIRED, etc.
                    "Element": f"{current_meta[1]}{current_meta[2]}",  # GS01, ISA14
                    "SubPos": current_meta[3],  # sub-position if exists
                    "Text": " ".join(current_chunk).strip()
                })

            # start new chunk
            current_meta = match.groups()
            current_chunk = [line]
        else:
            if current_meta:
                current_chunk.append(line)

    # save last chunk
    if current_chunk and current_meta:
        chunks.append({
            "Usage": current_meta[0].upper(),
            "Element": f"{current_meta[1]}{current_meta[2]}",
            "SubPos": current_meta[3],
            "Text": " ".join(current_chunk).strip()
        })

    return chunks
    




import re
from typing import List, Dict
from langchain_core.documents import Document

def parse_elements_from_page(text: str) -> List[Dict[str, str]]:
    """
    Parse one page of EDI spec text into chunks by REQUIRED/SITUATIONAL/NOT USED.
    Extracts element IDs (e.g., GS01, GS01-1).
    """
    usage_pattern = re.compile(r"\b(REQUIRED|SITUATIONAL|NOT USED)\b", re.IGNORECASE)

    parts = usage_pattern.split(text)
    chunks = []

    for i in range(1, len(parts), 2):  # usage marker, then block of text
        usage = parts[i].strip().upper()
        block = parts[i+1].strip()

        elem_match = re.search(r"\b([A-Z]{2,3}\d{2}(?:-\d+)?)\b", block)
        if not elem_match:
            continue

        element_id = elem_match.group(1)

        # break down GS01-1 → segment=GS, position=01, subfield=1
        seg_match = re.match(r"([A-Z]+)(\d{2})(?:-(\d+))?", element_id)
        segment, position, subfield = None, None, None
        if seg_match:
            segment, position, subfield = seg_match.groups()

        chunks.append({
            "Usage": usage,
            "Element": element_id,
            "Segment": segment,
            "Position": position,
            "Subfield": subfield,
            "Text": block
        })

    return chunks


def parse_documents(docs: List[Document]) -> List[Dict[str, str]]:
    """
    Process a list of Documents (page by page) into structured element chunks.
    """
    all_chunks = []
    for doc in docs:
        page_chunks = parse_elements_from_page(doc.page_content)
        all_chunks.extend(page_chunks)
    return all_chunks


# ---------------- Example Usage ----------------
docs = [
    Document(page_content="""
REQUIRED GS01 479 Functional Identifier Code
Code identifying a group of application related transaction sets

REQUIRED GS02 142 Application Sender’s Code
Use this code to identify the unit sending the information.

REQUIRED GS03-1 124 Application Receiver’s Code
Use this code to identify the unit receiving the information.
""")
]

parsed_chunks = parse_documents(docs)

for ch in parsed_chunks:
    print(ch)
    






transformer = LLMGraphTransformer(
    llm=patched_llm,
    prompt=(
        "Extract nodes and relationships as JSON. "
        "Each relationship must include: source_node_id, source_node_type, "
        "target_node_id, target_node_type, type, properties (object). "
        "Do not use 'source' or 'target'."
    )
)






from typing import List, Dict, Any
from langchain_experimental.graph_transformers import LLMGraphTransformer
from langchain_core.documents import Document

class SafeLLMGraphTransformer:
    """
    A wrapper around LLMGraphTransformer that sanitizes relationships and nodes
    before returning, so validation errors like 'target_node_id missing' are avoided.
    """

    def __init__(self, llm, **kwargs):
        self._transformer = LLMGraphTransformer(llm=llm, **kwargs)

    def _sanitize_relationship(self, rel: Dict[str, Any]) -> Dict[str, Any]:
        """Ensure required keys exist for each relationship."""
        rel.setdefault("source_node_id", rel.get("source", "unknown_source"))
        rel.setdefault("target_node_id", rel.get("target", "unknown_target"))
        rel.setdefault("source_node_type", rel.get("source_type", "UnknownType"))
        rel.setdefault("target_node_type", rel.get("target_type", "UnknownType"))
        rel.setdefault("type", rel.get("type", "UNKNOWN_REL"))
        return rel

    def _sanitize_node(self, node: Dict[str, Any]) -> Dict[str, Any]:
        """Ensure required keys exist for each node."""
        node.setdefault("id", node.get("id", node.get("name", "unknown_id")))
        node.setdefault("type", node.get("type", "UnknownType"))
        node.setdefault("properties", node.get("properties", {}))
        return node

    def convert_to_graph_documents(self, documents: List[Document], **kwargs):
        graph_docs = self._transformer.convert_to_graph_documents(documents, **kwargs)

        # Sanitize graph docs
        for gdoc in graph_docs:
            gdoc.nodes = [self._sanitize_node(node) for node in gdoc.nodes]
            gdoc.relationships = [self._sanitize_relationship(rel) for rel in gdoc.relationships]

        return graph_docs


# ------------------ Example Usage ------------------

from langchain_openai import AzureChatOpenAI

llm = AzureChatOpenAI(
    deployment_name="your-deployment",
    temperature=0,
    model="gpt-4o-mini"
)

safe_transformer = SafeLLMGraphTransformer(llm=llm)

# Suppose your parsed PDF docs are already wrapped in langchain Documents:
docs = [
    Document(page_content="Example text about SVC01 element and its mapping")
]

graph_docs = safe_transformer.convert_to_graph_documents(docs)

print(graph_docs[0].nodes)
print(graph_docs[0].relationships)





from langchain_openai import AzureChatOpenAI
from langchain_core.output_parsers import PydanticOutputParser
from typing import Type
from pydantic import BaseModel


class LLMWithStructuredOutput:
    """Wrapper to patch missing .with_structured_output in LangChain 0.3.x"""

    def __init__(self, llm):
        self.llm = llm

    def __getattr__(self, name):
        # Forward everything else to the underlying llm
        return getattr(self.llm, name)

    def with_structured_output(self, schema: Type[BaseModel], **kwargs):
        parser = PydanticOutputParser(pydantic_object=schema)
        return self.llm | parser


# ---- Example usage ----

# Your AzureChatOpenAI instance
azure_llm = AzureChatOpenAI(
    deployment_name="your-deployment",
    model="gpt-4o-mini",
    temperature=0,
)

# Wrap it
patched_llm = LLMWithStructuredOutput(azure_llm)

# Now pass this into LLMGraphTransformer
from langchain_experimental.graph_transformers import LLMGraphTransformer

llm_transformer = LLMGraphTransformer(
    llm=patched_llm,
    # optional schema arguments...
)






import re
from collections import defaultdict

def build_segment_dicts(elements: list) -> dict:
    """
    Convert element-level docs into segment-wise dictionary.
    Handles subfields like SVC01-1 by parsing from Text.
    """
    segments = defaultdict(dict)

    for elem in elements:
        element_name = elem["Element"]   # e.g., "SVC01"
        text = elem["Text"]

        # Base segment name (ISA, GS, SVC, etc.)
        if len(element_name) >= 3:
            segment = element_name[:3]
        else:
            segment = element_name[:2]

        # --- Look for subfield markers like "SVC01 1", "SVC01 2" inside text ---
        matches = re.findall(rf"{element_name}\s*(\d+)", text)

        if matches:
            # For each subfield found in the text, build key SVC01-<n>
            for sub in matches:
                field_position = f"{element_name}-{sub}"
                segments[segment][field_position] = text
        else:
            # No subfields, store as normal
            field_position = element_name[-2:]
            segments[segment][element_name] = text

    return segments







from collections import defaultdict

def build_segment_dicts(elements: list) -> dict:
    """
    Convert element-level docs into segment-wise dictionary.
    Handles subfields like SVC01-1.
    """
    segments = defaultdict(list)

    for elem in elements:
        element_name = elem["Element"]   # e.g., "ISA01", "SVC01-1", "GS02"
        text = elem.get("Text", "").strip()

        # Handle composite/subfield case (e.g., "SVC01-1")
        if "-" in element_name:
            segment = element_name[:3]                 # "SVC"
            field_position = element_name[3:]          # "01-1"
        else:
            if len(element_name) > 3:                  # e.g., "ISA01", "GS02"
                segment = element_name[:3]
                field_position = element_name[3:]
            else:                                      # fallback (very short tags)
                segment = element_name[:2]
                field_position = element_name[2:]

        segments[segment].append({
            "field_position": field_position,
            "text": text
        })

    return dict(segments)
    




def build_rules(elements, llm):
    """
    Main pipeline: process all elements → build segment dicts → query LLM → merge into rules.json
    """
    segments = build_segment_dicts(elements)
    rules = {"segments": []}

    for segment_name, fields in segments.items():
        enriched = query_llm_for_segment(segment_name, fields, llm)
        rules["segments"].append(json.loads(enriched))

    return rules

from collections import defaultdict

def build_segment_dicts(elements: list) -> dict:
    """
    Convert element-level docs into segment-wise dictionary.
    Handles subfields like SVC01-1.
    """
    segments = defaultdict(list)

    for elem in elements:
        element_name = elem["Element"]  # e.g., "ISA01", "SVC01-1"
        
        if "-" in element_name:
            # subfield (composite)
            segment = element_name[:3]   # e.g., SVC
            field_position = element_name
        elif len(element_name) == 5:
            segment = element_name[:3]   # e.g., ISA
            field_position = element_name[3:]
        else:
            segment = element_name[:2]   # e.g., GS
            field_position = element_name[2:]

        text = elem["Text"]

        segments[segment].append({
            "field_position": field_position,
            "text": text
        })

    return segments


import re

def element_based_chunking(text: str):
    """
    Splits EDI implementation guide text into chunks per element (ISA01, REF01, SVC01-1, etc.)
    """
    element_pattern = re.compile(r"\b(REQUIRED|SITUATIONAL|NOT USED)\s+([A-Z]{2,3}\d{2}(?:-\d{1,2})?)", re.IGNORECASE)
    
    chunks = []
    current_chunk = []
    current_element = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        match = element_pattern.search(line)
        if match:
            # save previous element chunk
            if current_chunk and current_element:
                chunks.append({
                    "Element": current_element,
                    "Text": " ".join(current_chunk).strip()
                })
            # start new element
            current_element = match.group(2)
            current_chunk = [line]
        else:
            # continuation line
            if current_element:
                current_chunk.append(line)

    # save last element
    if current_chunk and current_element:
        chunks.append({
            "Element": current_element,
            "Text": " ".join(current_chunk).strip()
        })

    return chunks


def query_llm_for_segment(segment_name: str, fields: list, llm) -> dict:
    """
    Send one segment (with all its fields, including composites) to the LLM 
    and get enriched rules back.
    """
    prompt = f"""
You are analyzing an EDI X12 segment called {segment_name}.

Each field has a position and some extracted text.
Some fields are composites (e.g., SVC01) and may contain subfields (e.g., SVC01-1 … SVC01-8).
If subfields are provided, treat them as children of the parent field.

For each field, return JSON with:
- field_position (e.g., "01", "SVC01-1")
- usage ("Required", "Optional", "Situational")
- short_description
- accepted_codes (as an object, {"CODE": "Definition"}), empty if none

Input fields:
{json.dumps(fields, indent=2)}

Return JSON only, structured like:
{{
  "segment": "{segment_name}",
  "fields": [
    {{
      "field_position": "01",
      "usage": "Required",
      "short_description": "Authorization Information Qualifier",
      "accepted_codes": {{
        "00": "No Authorization",
        "01": "Authorization Present"
      }}
    }},
    {{
      "field_position": "SVC01-1",
      "usage": "Required",
      "short_description": "Product/Service ID Qualifier",
      "accepted_codes": {{
        "AD": "American Dental Association Codes",
        "HC": "Healthcare Common Procedural Coding System"
      }}
    }}
  ]
}}
No markdown, no explanation.
"""
    return llm.invoke(prompt).content











import json
import re
from typing import List, Dict

def convert_rules(json_rules: List[Dict]) -> Dict:
    """
    Convert flat list of rules from companion guide JSON
    into nested dictionary: {Segment: {Position: {rule details}}}
    Handles duplicates by merging accepted codes.
    """
    rules_dict = {}

    for rule in json_rules:
        key_list = list(rule.values())

        seg = key_list[0]  # Segment name
        raw_pos = str(key_list[1])

        # Extract digits only (normalize field position)
        digits = re.findall(r"\d+", raw_pos)
        if not digits:
            continue
        pos = digits[0].zfill(2)  # e.g., "1" → "01"

        usage = key_list[2]
        desc = key_list[3]

        # Collect accepted codes
        codes = []
        try:
            if isinstance(key_list[4], list):
                if key_list[4] and isinstance(key_list[4][0], str):
                    codes = key_list[4]
                else:
                    codes = [c["Code"] for c in key_list[4] if "Code" in c]
        except Exception:
            codes = []

        # Initialize nested dict
        if seg not in rules_dict:
            rules_dict[seg] = {}

        if pos not in rules_dict[seg]:
            # Create new entry
            rules_dict[seg][pos] = {
                "usage": usage,
                "description": desc,
                "accepted_codes": set(codes)
            }
        else:
            # Merge with existing entry
            existing = rules_dict[seg][pos]
            existing["accepted_codes"].update(codes)

    # Convert sets back to lists for JSON serialization
    for seg in rules_dict:
        for pos in rules_dict[seg]:
            rules_dict[seg][pos]["accepted_codes"] = list(rules_dict[seg][pos]["accepted_codes"])

    with open("rules_dict.json", "w") as f:
        json.dump(rules_dict, f, indent=2)

    return rules_dict








import json
from typing import List, Dict
from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage

# --------------------------
# 1. Load and convert rules
# --------------------------

def convert_rules(json_rules: List[Dict]) -> Dict:
    """
    Convert flat list of rules from companion guide JSON
    into a nested dictionary: {Segment: {Position: {rule details}}}
    """
    rules_dict = {}
    for rule in json_rules:
        seg = rule["Segment Name"]
        pos = rule["Field Position"]
        usage = rule.get("Usage", "")
        desc = rule.get("Short Description", "")
        codes = [c["Code"] for c in rule.get("Accepted Codes", [])]

        if seg not in rules_dict:
            rules_dict[seg] = {}

        rules_dict[seg][pos] = {
            "usage": usage,
            "description": desc,
            "accepted_codes": codes
        }
    return rules_dict


# --------------------------
# 2. Setup LLM (Azure OpenAI)
# --------------------------

llm = AzureChatOpenAI(
    deployment_name="gpt-4o",  # replace with your deployment
    model="gpt-4o",
    temperature=0
)


# --------------------------
# 3. Validation Logic
# --------------------------

REQUIRED_SEGMENTS = ["ISA", "GS", "ST", "BPR", "TRN", "N1", "CLP", "SE", "GE", "IEA"]

def validate_segment(segment: str, rules: Dict) -> Dict:
    """
    Validate one EDI segment and return a single dictionary
    with overall status and per-field reasons.
    """
    elements = segment.split("*")
    seg_id = elements[0].strip()

    field_reasons = {}
    overall_status = "Matched"

    # ---------- Syntax check ----------
    if not seg_id.isalpha():
        return {
            "edi_line": segment,
            "status": "Invalid",
            "rule_line": seg_id,
            "reason": {"Syntax": "Invalid segment ID (not alphabetic)"}
        }

    # ---------- Mandatory check ----------
    if seg_id in REQUIRED_SEGMENTS:
        REQUIRED_SEGMENTS.remove(seg_id)

    # ---------- Field-level validation ----------
    if seg_id in rules:
        seg_rules = rules[seg_id]
        for pos, rule in seg_rules.items():
            idx = int(pos) - 1
            desc = rule["description"]

            try:
                value = elements[idx]
            except IndexError:
                field_reasons[f"{seg_id}{pos}"] = f"Invalid - Missing required field {desc}"
                overall_status = "Invalid"
                continue

            # Required field check
            if rule.get("usage") == "Required" and not value.strip():
                field_reasons[f"{seg_id}{pos}"] = f"Invalid - {desc} is required"
                overall_status = "Invalid"
                continue

            # Accepted codes check
            if rule.get("accepted_codes"):
                if value not in rule["accepted_codes"]:
                    llm_reason = llm([HumanMessage(
                        content=f"Field {seg_id}{pos} has invalid code '{value}'. "
                                f"Valid codes are {rule['accepted_codes']}. "
                                f"Describe the error in one sentence."
                    )])
                    reason = llm_reason.content if llm_reason else f"Invalid code {value}"
                    field_reasons[f"{seg_id}{pos}"] = f"Invalid - {reason}"
                    overall_status = "Invalid"
                    continue

            # If passed all checks
            field_reasons[f"{seg_id}{pos}"] = f"Valid - {desc} = {value}"
    else:
        field_reasons[seg_id] = "No field-level rules; syntax OK"

    return {
        "edi_line": segment,
        "status": overall_status,
        "rule_line": seg_id,
        "reason": field_reasons
    }


def validate_edi_file(edi_text: str, rules: Dict) -> List[Dict]:
    """
    Validate full EDI file line by line
    """
    segments = [seg.strip() for seg in edi_text.split("~") if seg.strip()]
    all_results = [validate_segment(seg, rules) for seg in segments]

    # Final mandatory check for missing segments
    if REQUIRED_SEGMENTS:
        for seg in REQUIRED_SEGMENTS:
            all_results.append({
                "edi_line": "N/A",
                "status": "Invalid",
                "rule_line": seg,
                "reason": {seg: f"Mandatory segment {seg} missing in file"}
            })
    return all_results


# --------------------------
# 4. Run Example
# --------------------------
if __name__ == "__main__":
    # Load rules.json (list of dicts)
    with open("rules.json") as f:
        json_rules = json.load(f)

    RULES = convert_rules(json_rules)

    # Load 835 EDI file
    with open("835.edi") as f:
        edi_text = f.read()

    results = validate_edi_file(edi_text, RULES)

    # Save output to JSON
    with open("validation_results.json", "w") as f:
        json.dump(results, f, indent=2)

    print(json.dumps(results[:3], indent=2))  # show first 3 results






import re

def parse_segments(text):
    segments = []
    # Split on segment terminator (~) or newline if PDF extracted
    raw_segments = re.split(r'~|\n', text)
    
    for raw in raw_segments:
        raw = raw.strip()
        if not raw:
            continue
        # Segment name = first 2-3 chars
        match = re.match(r'^([A-Z0-9]{2,3})(.*)', raw)
        if not match:
            continue
        seg_name, rest = match.groups()
        elements = rest.strip().split('*')[1:] if '*' in rest else []
        
        seg_dict = {
            "segment": seg_name,
            "elements": [
                {
                    "position": str(i+1).zfill(2),
                    "text": elem.strip(),
                    "usage": "",
                    "short_description": "",
                    "accepted_codes": []
                }
                for i, elem in enumerate(elements)
            ]
        }
        segments.append(seg_dict)
    return segments





def parse_segments(docs):
    """
    Input: docs (list of LangChain Document objects with .page_content)
    Output: structured segments dict for LLM
    """
    segments = []

    for doc in docs:
        text = doc.page_content

        # Split into lines for processing
        lines = [line.strip() for line in text.split("\n") if line.strip()]

        current_segment = None
        current_fields = []

        for line in lines:
            # If line looks like start of a new segment (e.g., ISA, GS, ST, SE, etc.)
            if len(line) >= 2 and line[:2].isalpha() and line[:2].isupper():
                # save previous segment if exists
                if current_segment:
                    segments.append({
                        "segment": current_segment,
                        "fields": current_fields
                    })
                # start new segment
                parts = line.split()
                current_segment = parts[0]  # e.g., ISA
                current_fields = [{"field_position": "01", "text": " ".join(parts[1:])}]
            
            else:
                # looks like continuation or element description
                if current_fields:
                    field_pos = str(len(current_fields) + 1).zfill(2)
                    current_fields.append({
                        "field_position": field_pos,
                        "text": line
                    })

        # last segment flush
        if current_segment:
            segments.append({
                "segment": current_segment,
                "fields": current_fields
            })

    return segments







import json
from collections import defaultdict
from openai import OpenAI

client = OpenAI()

def build_segment_dicts(elements):
    """
    Convert element-level docs into segment-wise dictionary
    """
    segments = defaultdict(list)
    for elem in elements:
        element_name = elem["Element"]  # e.g., ISA01
        segment = element_name[:3]      # ISA
        field_position = element_name[3:]  # 01
        text = elem["text"]

        segments[segment].append({
            "field_position": field_position,
            "text": text
        })
    return segments


def query_llm_for_segment(segment_name, fields):
    """
    Send one segment (with all its fields) to the LLM
    and get enriched rules back.
    """
    prompt = f"""
    You are analyzing an EDI X12 segment called {segment_name}.
    Each field has a position and some extracted text.
    For each field, return:
    - field_position
    - usage (required/optional/conditional if identifiable)
    - short_description
    - accepted_codes with description (if applicable)

    Input fields:
    {json.dumps(fields, indent=2)}

    Return JSON only, structured like:
    {{
      "segment": "{segment_name}",
      "fields": [
        {{
          "field_position": "01",
          "usage": "Required",
          "short_description": "Authorization Information Qualifier",
          "accepted_codes": {{"00": "No Authorization", "01": "Authorization Present"}}
        }}
      ]
    }}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )

    return json.loads(response.choices[0].message.content)


def build_rules(elements):
    """
    Main pipeline: process all elements → segment dict → query LLM → merge into rules.json
    """
    segments = build_segment_dicts(elements)
    rules = {"segments": []}

    for segment_name, fields in segments.items():
        enriched = query_llm_for_segment(segment_name, fields)
        rules["segments"].append(enriched)

    return rules


# Example usage
elements = [
    {"Element": "ISA01", "text": "Authorization Information Qualifier"},
    {"Element": "ISA02", "text": "Authorization Information"},
    {"Element": "ISA03", "text": "Security Information Qualifier"},
    {"Element": "ISA04", "text": "Security Information"},
    {"Element": "GS01", "text": "Functional Identifier Code"},
    {"Element": "GS02", "text": "Application Sender’s Code"}
]

rules_json = build_rules(elements)

with open("rules.json", "w") as f:
    json.dump(rules_json, f, indent=2)










import re

def element_based_chunking(text: str):
    """
    Splits EDI implementation guide text into chunks per element (ISA01, ISA02...).
    """
    # Pattern: Usage + ISAxx
    element_pattern = re.compile(r"\b(REQUIRED|SITUATIONAL|NOT USED)\s+(ISA\d{2})", re.IGNORECASE)

    chunks = []
    current_chunk = []
    current_element = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        match = element_pattern.search(line)
        if match:
            # save previous element chunk
            if current_chunk and current_element:
                chunks.append({
                    "Element": current_element,
                    "Text": " ".join(current_chunk).strip()
                })
                current_chunk = []

            # start new element
            current_element = match.group(2)
            current_chunk.append(line)
        else:
            # continuation line
            if current_element:
                current_chunk.append(line)

    # save last
    if current_chunk and current_element:
        chunks.append({
            "Element": current_element,
            "Text": " ".join(current_chunk).strip()
        })

    return chunks











prompt = f"""
You are given part of an EDI implementation guideline (example: ISA segment fields 01–16).

Your task: Extract *all* validation rules for every field in every segment.

Output format: A JSON array of objects. No markdown, no comments, no explanations.

Each object must contain exactly these keys:
- "Segment Name": string
- "Field Position": integer (1–N, no leading zeros, numeric only)
- "Usage": one of ["Required","Situational","Not Used"]
- "Short Description": string
- "Accepted Codes": list of {{"code": string, "description": string}} (empty list if none)

STRICT INSTRUCTIONS:
1. If the segment has 16 fields, you must return **all 16 entries** (Field Position 1 through 16). Never skip a field, even if it has no accepted codes.
2. Do not merge different fields. Each field must have its own entry.
3. Capture every accepted code that appears, from tables, inline text, bullets, or parentheses. Do not drop or shorten them. If description is not available, use "" (empty string).
4. Field Position must always increase sequentially (1, 2, 3…) without missing numbers.
5. Usage must exactly match the guideline text (e.g., "Required" if marked as M, "Situational" if marked as S, "Not Used" if marked as X).
6. Short Description should be the full label or definition (not abbreviated unless that is exactly how it appears).
7. Accepted Codes list must contain all codes for that field, no duplicates, and with full descriptions.

Validation before output:
- Ensure Segment Name + Field Position combination is unique.
- Ensure numeric field positions are sequential (no skips).
- Ensure Accepted Codes are captured completely.

Return only the JSON array.

Text:
{chunk.page_content}
"""







prompt = f"""
You are given part of an EDI implementation guideline.

Goal: Extract validation rules for segments and fields with NO duplicates and NO missing accepted codes.

Output: Return ONLY a JSON array of objects. No markdown, no comments, no extra text.

For each rule include exactly these keys:
- "Segment Name": string
- "Field Position": integer (numeric only)
- "Usage": one of ["Required","Situational","Not Used"]
- "Short Description": string
- "Accepted Codes": list of {{"code": string, "description": string}} (empty list if none specified)

STRICT PROCESS (follow exactly; do not output these steps):
1) SILENTLY build a dictionary keyed by K = "<Segment Name>|<Field Position>" while reading the text.
2) When you see the same K again, MERGE into the existing entry instead of creating a new one:
   - Short Description: keep the longest/most informative version.
   - Usage: if conflicting, choose the strictest (Required > Situational > Not Used).
   - Accepted Codes: UNION of all codes found for K (no duplicates).
3) EXHAUSTIVE CODE CAPTURE:
   - Capture codes from any phrasing like "Accepted Codes", "Valid Values", "Values", "must be one of", "qualifier values", "(code – description)", tables, bullets, comma-separated lists, or text in parentheses after the field name.
   - Include codes even if descriptions are missing (use empty string for description).
   - Normalize codes by trimming spaces; keep original case as shown (usually uppercase).
4) ONLY AFTER finishing the entire text, convert the dictionary values to a JSON list.
5) Sort final list by "Segment Name" (A→Z) then by "Field Position" (ascending).
6) Ensure there is exactly ONE object per unique ("Segment Name","Field Position").

Validation before output:
- No duplicate ("Segment Name","Field Position") pairs.
- "Field Position" is an integer.
- "Accepted Codes" lists contain unique codes only.

Text:
{chunk.page_content}
"""












def parse_segment(seg: ET.Element) -> str:
    seg_id = seg.attrib["id"]
    elements = [seg_id]

    # dictionary to hold element values with proper position
    element_dict = {}

    for child in seg:
        if child.tag == "ele":
            # extract numeric part (position index)
            idx_str = child.attrib["id"].replace(seg_id, "")
            if idx_str.isdigit():
                idx = int(idx_str)
                element_dict[idx] = child.text or ""
        elif child.tag == "comp":
            idx_str = child.attrib["id"].replace(seg_id, "")
            if idx_str.isdigit():
                idx = int(idx_str)
                sub_values = [sub.text or "" for sub in child if sub.tag == "subele"]
                element_dict[idx] = ":".join(sub_values)

    # Build elements in proper order, filling missing with ""
    max_idx = max(element_dict.keys(), default=0)
    for i in range(1, max_idx + 1):
        elements.append(element_dict.get(i, ""))

    return "*".join(elements) + "~"









import xml.etree.ElementTree as ET

def parse_segment(seg: ET.Element) -> str:
    seg_id = seg.attrib["id"]

    # dictionary for elements by position
    positions = {}

    for child in seg:
        if child.tag == "ele":
            idx = int(child.attrib["id"].replace(seg_id, ""))  # e.g. CLM02 → 2
            positions[idx] = child.text or ""
        elif child.tag == "comp":
            idx = int(child.attrib["id"].replace(seg_id, ""))  # e.g. CLM05 → 5
            sub_values = [sub.text or "" for sub in child if sub.tag == "subele"]
            positions[idx] = ":".join(sub_values)

    # fill gaps with empty "" so *** appear
    max_idx = max(positions.keys())
    elements = [seg_id] + [positions.get(i, "") for i in range(1, max_idx + 1)]

    return "*".join(elements) + "~"
















#!/usr/bin/env python3
"""
Extract matched claims from large 837I/837P XMLs and build EDI .DAT files grouped by claim type + ISA06 (trading partner).

- Deduplicates identical ISA_LOOP (full ISA->IEA interchange).
- Matches claims using Excel lookup: (CLM01, NM109, CLM02, StartDate).
- Groups matched claims by (ClaimType, ISA06) and by subscriber (2000A).
- Writes proper EDI .DAT files reusing original ISA/GS/ST/.../SE/GE/IEA segments.

CONFIG: set INPUT_XML_FOLDER, EXCEL_FILE, OUTPUT_FOLDER
"""
import os
import hashlib
import xml.etree.ElementTree as ET
from collections import defaultdict
from datetime import datetime
import pandas as pd

# -----------------------
# CONFIG - change these
# -----------------------
INPUT_XML_FOLDER = "edi_claim_extractor_source"   # folder with 837I/837P XMLs
EXCEL_FILE = "edi_claim_extractor_input_file.xlsx"
OUTPUT_FOLDER = "edi_claim_outputs"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Excel column names tolerance (try to auto-detect common variations)
EXCEL_CANDIDATES = {
    "clm01": ["Patient Control Number", "CLM01", "patient_control_number", "patient_control_no"],
    "member": ["Member ID", "NM109", "MemberID", "member_id"],
    "amount": ["Total Bill Amount", "CLM02", "TotalAmount", "total_amount"],
    "startdate": ["Start Date", "Start Date of Service", "StartDate", "start_date"]
}


# -----------------------
# Helpers
# -----------------------
def pick_excel_columns(df):
    """Return column names for clm01, member, amount, startdate by checking candidate headers."""
    cols = {k: None for k in EXCEL_CANDIDATES}
    lower_map = {c.lower().strip(): c for c in df.columns}
    for key, candidates in EXCEL_CANDIDATES.items():
        for cand in candidates:
            if cand in df.columns:
                cols[key] = cand
                break
            if cand.lower() in lower_map:
                cols[key] = lower_map[cand.lower()]
                break
    # If any missing, try fuzzy fallback: find columns containing substrings
    for key, val in cols.items():
        if val is None:
            for col in df.columns:
                lc = col.lower()
                if key == "clm01" and ("clm" in lc and "01" in lc or "control" in lc):
                    cols[key] = col; break
                if key == "member" and ("member" in lc or "nm1" in lc):
                    cols[key] = col; break
                if key == "amount" and ("amount" in lc or "clm02" in lc or "total" in lc):
                    cols[key] = col; break
                if key == "startdate" and ("date" in lc):
                    cols[key] = col; break
    return cols


def normalize_date_str(s: str) -> str:
    """Return digits-only normalized date for comparisons (e.g., 20250103 or 01032025 etc)."""
    if s is None:
        return ""
    s = str(s).strip()
    if not s:
        return ""
    # remove non-digits
    digits = "".join(ch for ch in s if ch.isdigit())
    return digits


def normalize_value(s: str) -> str:
    if s is None:
        return ""
    return str(s).strip()


def fingerprint_element(elem: ET.Element) -> str:
    """Fingerprint an element (and its child seg/ele structure) deterministically."""
    parts = []
    for node in elem.iter():
        if node.tag == 'seg':
            segid = node.attrib.get('id', '')
            ele_texts = []
            for ele in node.findall('ele'):
                ele_texts.append((ele.text or '').strip())
            parts.append(segid + "|" + "|".join(ele_texts))
    joined = "||".join(parts)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()


def seg_to_edi_text(seg: ET.Element) -> str:
    """Convert <seg id='ISA'> <ele id='ISA01'>..</ele> ... </seg> -> 'ISA*ele1*ele2*...~'"""
    if seg is None:
        return ""
    segid = seg.attrib.get('id', '').strip()
    # preserve element order
    ele_texts = []
    for ele in seg.findall('ele'):
        # keep empty elements as empty strings to preserve field positions
        ele_texts.append((ele.text or '').strip())
    if ele_texts:
        return segid + "*" + "*".join(ele_texts) + "~"
    else:
        return segid + "~"


def get_seg_ele(scope: ET.Element, seg_id: str, ele_id: str) -> str:
    """Find first <seg id='seg_id'>/<ele id='ele_id'> descendant under scope and return text."""
    node = scope.find(f".//seg[@id='{seg_id}']/ele[@id='{ele_id}']")
    return (node.text or "").strip() if node is not None and node.text else ""


def find_all_isa_loops(root: ET.Element):
    """Return list of ISA_LOOP elements. Try common XPaths/fallbacks."""
    loops = root.findall(".//loop[@id='ISA_LOOP']")
    if loops:
        return loops
    # try just seg grouping
    return root.findall(".//seg[@id='ISA']/..")  # parent of ISA seg (likely a loop element)


def get_isa06_from_isa_loop(isa_loop: ET.Element) -> str:
    return get_seg_ele(isa_loop, "ISA", "ISA06")


def get_gs08_from_isa_loop(isa_loop: ET.Element) -> str:
    return get_seg_ele(isa_loop, "GS", "GS08")


def get_st_implref_from_st(st_seg: ET.Element) -> str:
    # ST3 (implementation reference) sometimes stored as ST03
    if st_seg is None:
        return ""
    node = st_seg.find("ele[@id='ST03']")
    if node is not None and (node.text or "").strip():
        return (node.text or "").strip()
    return ""


def get_dtp_start_date(loop2000b: ET.Element, target_tag: str) -> str:
    """Return DTP03 where DTP01 matches target_tag (e.g., '472' or '434')."""
    for dtp in loop2000b.findall(".//seg[@id='DTP']"):
        dtp01 = dtp.find("ele[@id='DTP01']")
        if dtp01 is not None and (dtp01.text or "").strip() == target_tag:
            dtp03 = dtp.find("ele[@id='DTP03']")
            return (dtp03.text or "").strip() if dtp03 is not None and dtp03.text else ""
    return ""


# -----------------------
# Load Excel lookup
# -----------------------
print("Loading Excel file:", EXCEL_FILE)
df = pd.read_excel(EXCEL_FILE, dtype=str).fillna("")
cols_map = pick_excel_columns(df)
if not cols_map['clm01'] or not cols_map['member'] or not cols_map['amount']:
    raise SystemExit("ERROR: Could not auto-detect necessary Excel columns. Please ensure Excel has Patient Control Number, Member ID, Total Bill Amount columns.")

excel_keys = set()
for _, row in df.iterrows():
    clm01 = normalize_value(row[cols_map['clm01']])
    nm109 = normalize_value(row[cols_map['member']])
    clm02 = normalize_value(row[cols_map['amount']])
    # start date may be optional in Excel; if present normalize
    sdate = normalize_date_str(row[cols_map['startdate']]) if cols_map['startdate'] else ""
    excel_keys.add((clm01, nm109, clm02, sdate))
print(f"Excel entries loaded: {len(excel_keys)}")

# -----------------------
# Process XML files: dedupe ISA_LOOP, match claims
# -----------------------
seen_isa_hashes = set()
# structure: claims_by_partner[(claim_type, isa06)] -> { subscriber_fp: (subscriber_2000a_elem, set_of_2000b_elems) }
claims_by_partner = defaultdict(lambda: defaultdict(lambda: {"2000a": None, "2000b_set": [], "seen_b_fp": set()}))

xml_files = [f for f in os.listdir(INPUT_XML_FOLDER) if f.lower().endswith(".xml")]
print("Found XML files:", xml_files)

for xmlf in xml_files:
    path = os.path.join(INPUT_XML_FOLDER, xmlf)
    print("Parsing:", path)
    tree = ET.parse(path)
    root = tree.getroot()

    isa_loops = find_all_isa_loops(root)
    if not isa_loops:
        print("  No ISA_LOOP found in", xmlf)
        continue

    for isa_loop in isa_loops:
        isa_fp = fingerprint_element(isa_loop)
        if isa_fp in seen_isa_hashes:
            # skip identical full interchange
            # print("Skipping duplicate ISA_LOOP fingerprint:", isa_fp[:8])
            continue
        seen_isa_hashes.add(isa_fp)

        isa06 = get_isa06_from_isa_loop(isa_loop)
        if not isa06:
            # cannot group without trading partner; skip
            # print("  No ISA06; skipping interchange")
            continue

        # find header/footer segs
        isa_seg = isa_loop.find(".//seg[@id='ISA']")
        iea_seg = isa_loop.find(".//seg[@id='IEA']")

        # find GS groups
        gs_loops = isa_loop.findall(".//loop[@id='GS_LOOP']")
        if not gs_loops:
            # fallback: try to build from GS segments under this ISA
            gs_segs = isa_loop.findall(".//seg[@id='GS']")
            for seg in gs_segs:
                l = ET.Element("loop", attrib={"id": "GS_LOOP"})
                l.append(seg)
                # try to attach GE if present under ISA
                ge = isa_loop.find(".//seg[@id='GE']")
                if ge is not None:
                    l.append(ge)
                gs_loops.append(l)

        for gs_loop in gs_loops:
            gs_seg = gs_loop.find(".//seg[@id='GS']")
            ge_seg = gs_loop.find(".//seg[@id='GE']")

            # find ST loops under this GS
            st_loops = gs_loop.findall(".//loop[@id='ST_LOOP']")
            if not st_loops:
                st_segs = gs_loop.findall(".//seg[@id='ST']")
                for sseg in st_segs:
                    st_loop = ET.Element("loop", attrib={"id": "ST_LOOP"})
                    st_loop.append(sseg)
                    # attach matching SE if exists
                    se_candidate = gs_loop.find(".//seg[@id='SE']")
                    if se_candidate is not None:
                        st_loop.append(se_candidate)
                    st_loops.append(st_loop)

            for st_loop in st_loops:
                st_seg = st_loop.find(".//seg[@id='ST']")
                se_seg = st_loop.find(".//seg[@id='SE']")
                if st_seg is None:
                    continue

                # determine claim type: prefer GS08 then ST03
                gs08 = (gs_seg.find("ele[@id='GS08']").text or "").strip() if gs_seg is not None and gs_seg.find("ele[@id='GS08']") is not None else ""
                st_impl = get_st_implref_from_st(st_seg)
                claim_type = None
                if "X222" in gs08 or "005010X222" in gs08 or "X222" in st_impl or "005010X222" in st_impl:
                    claim_type = "Professional"
                elif "X223" in gs08 or "005010X223" in gs08 or "X223" in st_impl or "005010X223" in st_impl:
                    claim_type = "Institutional"
                else:
                    # fallback: if filename contains 'p' or 'i'
                    if "p" in xmlf.lower():
                        claim_type = "Professional"
                    elif "i" in xmlf.lower():
                        claim_type = "Institutional"
                    else:
                        # unknown, skip
                        continue

                # iterate 2000A loops inside this ST
                for loop2000a in st_loop.findall(".//loop[@id='2000A']"):
                    # fingerprint subscriber-level 2000A (to dedupe later)
                    a_fp = fingerprint_element(loop2000a)

                    # we'll collect any matched 2000B under this 2000A
                    matched_b_list = []

                    for loop2000b in loop2000a.findall(".//loop[@id='2000B']"):
                        # extract CLM01, NM109, CLM02
                        clm01 = get_seg_ele(loop2000b, "CLM", "CLM01")
                        nm109 = get_seg_ele(loop2000b, "NM1", "NM109")
                        if not nm109:
                            # sometimes member id is stored under different ele id; try known variants
                            nm109 = get_seg_ele(loop2000b, "NM1", "NM1_09") or get_seg_ele(loop2000b, "NM1", "NM1_08")
                        clm02 = get_seg_ele(loop2000b, "CLM", "CLM02")

                        # get start dates for both tags
                        dtp472 = get_dtp_start_date(loop2000b, "472")  # professional
                        dtp434 = get_dtp_start_date(loop2000b, "434")  # institutional

                        # normalize strings for matching
                        n_clm01 = normalize_value(clm01)
                        n_nm109 = normalize_value(nm109)
                        n_clm02 = normalize_value(clm02)
                        n_dtp472 = normalize_date_str(dtp472)
                        n_dtp434 = normalize_date_str(dtp434)

                        # Match rules:
                        # - if ST type is Professional try match using dtp472; else dtp434
                        matched_as = None
                        if claim_type == "Professional" and n_dtp472:
                            if (n_clm01, n_nm109, n_clm02, n_dtp472) in excel_keys:
                                matched_as = "Professional"
                        elif claim_type == "Institutional" and n_dtp434:
                            if (n_clm01, n_nm109, n_clm02, n_dtp434) in excel_keys:
                                matched_as = "Institutional"
                        # fallback: try matching ignoring dtp if Excel startdate empty or not provided
                        if not matched_as:
                            # try match where Excel maybe has empty start date
                            if (n_clm01, n_nm109, n_clm02, "") in excel_keys:
                                matched_as = claim_type

                        if matched_as:
                            # we will keep this 2000B
                            matched_b_list.append(loop2000b)

                    # if we found any matched 2000B under this 2000A, record them grouped by (claim_type, isa06)
                    if matched_b_list:
                        key = (claim_type, isa06)
                        # ensure this subscriber (by a_fp) is recorded
                        entry = claims_by_partner[key][a_fp]
                        if entry["2000a"] is None:
                            # store a shallow copy of the 2000A loop without inner 2000B children (we will append matched ones)
                            # Create new element and copy seg children except nested 2000B loops
                            a_copy = ET.Element("loop", attrib={"id": "2000A"})
                            for seg in loop2000a:
                                # if seg is a loop 2000B skip; if seg is seg element, append
                                if seg.tag == "loop" and seg.attrib.get("id","").startswith("2000B"):
                                    continue
                                a_copy.append(seg)
                            entry["2000a"] = a_copy
                        # For each matched 2000B, append if not duplicate
                        for b in matched_b_list:
                            b_fp = fingerprint_element(b)
                            if b_fp not in entry["seen_b_fp"]:
                                entry["seen_b_fp"].add(b_fp)
                                # append a deep copy of b (to avoid references)
                                entry["2000b_set"].append(b)

# -----------------------
# Build and write .DAT files
# -----------------------
def write_dat_files(claims_by_partner_map, out_dir):
    today = datetime.now().strftime("%m%d%Y")
    for (claim_type, isa06), subs_dict in claims_by_partner_map.items():
        # collect all non-empty subscribers
        subs = [v for v in subs_dict.values() if v["2000b_set"]]
        if not subs:
            continue

        # Use header/footer from first matched subscriber's stored tuple? We stored only 2000a/2000b; need header from original scanning
        # But in our earlier collection we didn't keep isa/gs/st/se/ge/iea per key to avoid repetition.
        # For this script we will find the first occurrence in XML files again to get header/footer segments.
        # Simpler: reconstruct header/footer from any matched 2000b's ancestor segs.

        # For header reuse, extract the ISA/GS/ST/SE/GE/IEA segments from the original 2000a/2000b elements' ancestors.
        # We'll find ancestor segs by searching upwards — but ElementTree does not have parent.
        # Instead, to keep it simple and robust: search the entire XML folder again for the first ISA_LOOP that contains any of these 2000B fingerprints and reuse its headers.
        header_isa_seg = None
        header_gs_seg = None
        header_st_seg = None
        footer_se_seg = None
        footer_ge_seg = None
        footer_iea_seg = None

        # find one original location for header/footer (search input XMLs)
        found_header = False
        # Build set of b fingerprints for search
        all_b_fps = set()
        for sub in subs:
            for b in sub["2000b_set"]:
                all_b_fps.add(fingerprint_element(b))

        # search xml files to find a matching containership
        for xmlf in xml_files:
            path = os.path.join(INPUT_XML_FOLDER, xmlf)
            tree = ET.parse(path)
            root = tree.getroot()
            isa_loops = find_all_isa_loops(root)
            for isa_loop in isa_loops:
                isa_loop_fps_b = set()
                for loop2000b in isa_loop.findall(".//loop[@id='2000B']"):
                    isa_loop_fps_b.add(fingerprint_element(loop2000b))
                if all_b_fps & isa_loop_fps_b:
                    # use header/footer from this isa_loop
                    header_isa_seg = isa_loop.find(".//seg[@id='ISA']")
                    header_gs_seg = isa_loop.find(".//seg[@id='GS']")
                    header_st_seg = isa_loop.find(".//seg[@id='ST']")
                    footer_se_seg = isa_loop.find(".//seg[@id='SE']")
                    footer_ge_seg = isa_loop.find(".//seg[@id='GE']")
                    footer_iea_seg = isa_loop.find(".//seg[@id='IEA']")
                    found_header = True
                    break
            if found_header:
                break

        # fallback: if not found, skip writing (shouldn't happen)
        if not found_header:
            print("Warning: could not locate header/footer for", claim_type, isa06)
            continue

        # Build EDI lines: ISA, GS, ST, optional HEADER segs, then 2000A blocks with their included 2000B children, then SE, GE, IEA
        lines = []
        lines.append(seg_to_edi_text(header_isa_seg))
        lines.append(seg_to_edi_text(header_gs_seg))
        lines.append(seg_to_edi_text(header_st_seg))
        # optional: if any 'HEADER' seg exists under header_st_seg's context, include - try find under st siblings
        header_seg = None
        # search for 'HEADER' seg near ST (under same ST_LOOP)
        # Try: find seg id HEADER under the root (first instance)
        header_seg_generic = None
        for xmlf in xml_files:
            path = os.path.join(INPUT_XML_FOLDER, xmlf)
            tree = ET.parse(path)
            root = tree.getroot()
            h = root.find(".//seg[@id='HEADER']")
            if h is not None:
                header_seg_generic = h
                break
        if header_seg_generic is not None:
            lines.append(seg_to_edi_text(header_seg_generic))

        # append unique 2000A blocks with their matched 2000B children
        for sub in subs:
            a_elem = sub["2000a"]
            if a_elem is None:
                continue
            # append segs inside a_elem (non-loop segs)
            for seg in a_elem.findall(".//seg"):
                lines.append(seg_to_edi_text(seg))
            # append each matched 2000B under this subscriber
            for b in sub["2000b_set"]:
                # b may be a <loop id='2000B'> containing many seg children
                for seg in b.findall(".//seg"):
                    lines.append(seg_to_edi_text(seg))

        # append trailers
        if footer_se_seg is not None:
            lines.append(seg_to_edi_text(footer_se_seg))
        if footer_ge_seg is not None:
            lines.append(seg_to_edi_text(footer_ge_seg))
        if footer_iea_seg is not None:
            lines.append(seg_to_edi_text(footer_iea_seg))

        # write file
        fname = f"{claim_type}_{isa06}_{today}.DAT"
        outpath = os.path.join(out_dir, fname)
        with open(outpath, "w", encoding="utf-8") as w:
            # join with newline for human readability; EDI segment terminator `~` is already included
            w.write("\n".join(lines))
        print(f"Wrote {len(subs)} subscribers -> {outpath}")

write_dat_files(claims_by_partner, OUTPUT_FOLDER)
print("Completed. Outputs in:", OUTPUT_FOLDER)







import os
import re
import pandas as pd
import xml.etree.ElementTree as ET
from collections import defaultdict
from datetime import datetime

# -------------------
# Config
# -------------------
input_xml_folder = "edi_claim_extractor_source"
excel_file = "edi_claim_extractor_input_file.xlsx"
output_folder = "edi_claim_outputs"
os.makedirs(output_folder, exist_ok=True)

# -------------------
# Load Excel claims of interest
# -------------------
df = pd.read_excel(excel_file, dtype=str).fillna("")

claims_lookup = set(
    (str(row["Patient Control Number"]).strip(),
     str(row["Member ID"]).strip(),
     str(row["Total Bill Amount"]).strip(),
     str(row["Start Date"]).strip())
    for _, row in df.iterrows()
)

# -------------------
# Utilities
# -------------------
def get_text(element, path):
    node = element.find(path)
    return node.text.strip() if node is not None and node.text else ""

def match_claim(clm01, nm109, clm02, start_date):
    return (clm01.strip(), nm109.strip(), clm02.strip(), start_date.strip()) in claims_lookup

# -------------------
# Process XML files
# -------------------
claims_by_partner = defaultdict(lambda: {"Professional": [], "Institutional": []})

for file in os.listdir(input_xml_folder):
    if not file.endswith(".xml"):
        continue

    tree = ET.parse(os.path.join(input_xml_folder, file))
    root = tree.getroot()

    # Extract trading partner from ISA06
    isa06 = get_text(root, ".//ISA06")
    if not isa06:
        continue

    # Determine claim type by looking at ST segment
    st_val = get_text(root, ".//ST01")
    if st_val == "837":
        # Professional vs Institutional indicator
        claim_type = "Professional" if "HC" in file else "Institutional"
    else:
        continue

    # Get headers and footers
    isa = root.find(".//ISA")
    gs = root.find(".//GS")
    st = root.find(".//ST")
    se = root.find(".//SE")
    ge = root.find(".//GE")
    iea = root.find(".//IEA")

    # Traverse 2000A → 2000B
    for loop2000a in root.findall(".//2000A"):
        a_copy = ET.Element("2000A", attrib=loop2000a.attrib)
        added_b = False

        for loop2000b in loop2000a.findall(".//2000B"):
            clm01 = get_text(loop2000b, ".//CLM01")
            nm109 = get_text(loop2000b, ".//NM109")
            clm02 = get_text(loop2000b, ".//CLM02")

            # DTP start date (472 for professional, 434 for institutional)
            dtp_tag = "472" if claim_type == "Professional" else "434"
            start_date = ""
            for dtp in loop2000b.findall(".//DTP"):
                if dtp_tag in (dtp.text or ""):
                    start_date = dtp.find("DTP03").text if dtp.find("DTP03") is not None else ""

            if match_claim(clm01, nm109, clm02, start_date):
                a_copy.append(loop2000b)
                added_b = True

        if added_b:
            claims_by_partner[isa06][claim_type].append((isa, gs, st, a_copy, se, ge, iea))

# -------------------
# Write outputs
# -------------------
today = datetime.now().strftime("%m%d%Y")

for partner, types in claims_by_partner.items():
    for claim_type, blocks in types.items():
        if not blocks:
            continue

        # Build new XML root
        new_root = ET.Element("EDI")
        # take header/footer from first block
        isa, gs, st, _, se, ge, iea = blocks[0]
        new_root.append(isa)
        new_root.append(gs)
        new_root.append(st)

        # add all unique 2000A
        seen_a = set()
        for _, _, _, a_block, _, _, _ in blocks:
            a_str = ET.tostring(a_block)
            if a_str not in seen_a:
                seen_a.add(a_str)
                new_root.append(a_block)

        new_root.append(se)
        new_root.append(ge)
        new_root.append(iea)

        out_file = f"{claim_type}_{partner}_{today}.DAT"
        ET.ElementTree(new_root).write(os.path.join(output_folder, out_file), encoding="utf-8", xml_declaration=True)

print(f"✅ Extraction completed. Output written to {output_folder}")












import os
import xml.etree.ElementTree as ET
from io import StringIO
from x12.core import x12n_document, ParamsBase

def change_file_into_segment_list(edi_file_path):
    with open(edi_file_path, 'r', encoding='utf-8') as file:
        segment_list = [line.strip() for line in file.readlines()]
    return segment_list

def load_edi_segments_as_xml_obj(edi_segments_list):
    edibuffer = StringIO('\n'.join(edi_segments_list))
    param = ParamsBase()
    xmlbuffer = StringIO()
    x12n_document(param, edibuffer, fd_997=None, fd_html=None, fd_xmldoc=xmlbuffer, map_path=None)
    xmlbuffer.seek(0)
    rootnode = ET.fromstring(xmlbuffer.getvalue())
    return rootnode

def identify_edi_type(segments):
    """
    Returns 'I' for Institutional, 'P' for Professional, or None if unsupported (like Dental).
    """
    for seg in segments:
        if seg.startswith("ST*837"):
            if "005010X223" in seg:  # Institutional
                return "I"
            elif "005010X222" in seg:  # Professional
                return "P"
            elif "005010X224" in seg:  # Dental (skip)
                return None
    return None

def build_combined_xml(input_folder, output_folder):
    institutional_root = ET.Element("Institutional837")
    professional_root = ET.Element("Professional837")

    for filename in os.listdir(input_folder):
        if not filename.endswith(".txt") and not filename.endswith(".edi"):
            continue

        file_path = os.path.join(input_folder, filename)
        segments = change_file_into_segment_list(file_path)
        edi_type = identify_edi_type(segments)

        if edi_type is None:  # Skip dental or unknown
            continue

        rootnode = load_edi_segments_as_xml_obj(segments)

        if edi_type == "I":
            institutional_root.append(rootnode)
        elif edi_type == "P":
            professional_root.append(rootnode)

    # Write Institutional
    if len(institutional_root):
        inst_tree = ET.ElementTree(institutional_root)
        inst_tree.write(os.path.join(output_folder, "837I_combined.xml"), encoding="utf-8", xml_declaration=True)

    # Write Professional
    if len(professional_root):
        prof_tree = ET.ElementTree(professional_root)
        prof_tree.write(os.path.join(output_folder, "837P_combined.xml"), encoding="utf-8", xml_declaration=True)

    print("✅ Combined XMLs generated successfully.")















"""
Combine multiple X12 EDI .dat files into TWO XML outputs (Institutional 837I vs Professional 837P)
using pyx12's x12n_document.

Features:
- Scans a source folder for EDI files (*.dat).
- Auto-detects transaction set version (ICN/ICVN) from ISA/GS/ST segments.
- Loads the correct pyx12 map folder automatically.
- Separates Institutional (837I, e.g. 005010X223A2) and Professional (837P, e.g. 005010X222A1).
- Writes two XML files: combined_institutional.xml and combined_professional.xml.

Usage:
------
python combine_edi_to_xml.py \
    --source "/path/to/source/folder" \
    --outdir "/path/to/output/folder"

Notes:
------
- The 2000A/2000B loops are preserved by pyx12.
- Institutional vs Professional are split into two separate XML files.
"""

from __future__ import annotations
import argparse
import io
import sys
from pathlib import Path
from typing import Iterable, Tuple
import xml.etree.ElementTree as ET

# --- pyx12 imports ---
try:
    from pyx12.params import params
    from pyx12.x12n_document import x12n_document
except Exception as e:
    print("[ERROR] pyx12 not available or unexpected import error:", e, file=sys.stderr)
    raise


def find_edi_files(src_dir: Path, patterns: Tuple[str, ...] = ("*.dat",)) -> Iterable[Path]:
    for pat in patterns:
        yield from src_dir.rglob(pat)


def detect_version_and_txset(edi_path: Path) -> Tuple[str, str, str]:
    """
    Detect version (ICVN), transaction set (ST01), and convention (ST03).
    Returns (icvn, txset, implref).
    """
    with open(edi_path, "rt", encoding="utf-8", errors="ignore") as f:
        content = f.read(5000)
    segments = content.split("~")
    icvn = ""
    txset = ""
    implref = ""
    for seg in segments:
        fields = seg.split("*")
        if fields[0] == "GS" and len(fields) > 8:
            icvn = fields[8]
        elif fields[0] == "ST" and len(fields) > 2:
            txset = fields[1]
            implref = fields[2]
        if icvn and txset:
            break
    return icvn, txset, implref


def get_default_map_dir(icvn: str) -> Path:
    import pyx12
    base = Path(pyx12.__file__).resolve().parent / "maps"
    return base / icvn[:6]


def convert_edi_to_xml_bytes(edi_path: Path, p: "params") -> bytes:
    buf = io.BytesIO()
    with open(edi_path, "rb") as fp_in:
        x12n_document(fp_in, buf, p, str(edi_path))
    return buf.getvalue()


def parse_xml_bytes(xml_bytes: bytes) -> ET.Element:
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError as e:
        snippet = xml_bytes[:300].decode("utf-8", errors="replace")
        raise RuntimeError(f"Failed to parse XML: {e}\nSnippet:\n{snippet}")
    return root


def build_combined_tree(doc_roots: Iterable[Tuple[Path, ET.Element]]) -> ET.ElementTree:
    combined_root = ET.Element("CombinedX12")
    for src_path, root in doc_roots:
        wrapper = ET.SubElement(combined_root, "Document", attrib={"file": str(src_path)})
        wrapper.append(root)
    return ET.ElementTree(combined_root)


def pretty_write(tree: ET.ElementTree, out_path: Path) -> None:
    try:
        ET.indent(tree, space="  ", level=0)
    except Exception:
        pass
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tree.write(out_path, encoding="utf-8", xml_declaration=True)


def combine_folder(src_dir: Path, out_dir: Path, file_patterns: Tuple[str, ...] = ("*.dat",)) -> Tuple[int, int]:
    edi_files = list(find_edi_files(src_dir, file_patterns))
    if not edi_files:
        raise FileNotFoundError(f"No EDI files found in {src_dir} matching {file_patterns}")

    institutional_roots = []
    professional_roots = []

    for edi_path in sorted(edi_files):
        icvn, txset, implref = detect_version_and_txset(edi_path)
        if txset != "837":
            continue  # skip non-837s

        map_dir = get_default_map_dir(icvn)
        p = params()
        p.set("map_path", str(map_dir))
        p.set("icvn", icvn)

        xml_bytes = convert_edi_to_xml_bytes(edi_path, p)
        root = parse_xml_bytes(xml_bytes)

        # Classify Professional vs Institutional based on implementation reference
        if "X222" in icvn or "X222" in implref:
            professional_roots.append((edi_path, root))
        elif "X223" in icvn or "X223" in implref:
            institutional_roots.append((edi_path, root))
        else:
            # Default to professional if not clearly institutional
            professional_roots.append((edi_path, root))

    if institutional_roots:
        inst_tree = build_combined_tree(institutional_roots)
        pretty_write(inst_tree, out_dir / "combined_institutional.xml")
    if professional_roots:
        prof_tree = build_combined_tree(professional_roots)
        pretty_write(prof_tree, out_dir / "combined_professional.xml")

    return len(institutional_roots), len(professional_roots)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Combine multiple EDI files into two XMLs: Institutional vs Professional")
    parser.add_argument("--source", required=True, type=Path, help="Folder containing EDI .dat files")
    parser.add_argument("--outdir", required=True, type=Path, help="Folder to write XML outputs")
    parser.add_argument("--glob", nargs="+", default=["*.dat"], help="Glob patterns to include (default: *.dat)")
    args = parser.parse_args(argv)

    inst_count, prof_count = combine_folder(src_dir=args.source, out_dir=args.outdir, file_patterns=tuple(args.glob))
    print(f"Wrote {inst_count} institutional and {prof_count} professional files into {args.outdir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())













import os
import pandas as pd
from pyx12.x12file import X12Reader
from pyx12.writer import X12Writer
from datetime import datetime

# Load claims of interest from Excel
claims_df = pd.read_excel("edi_claim_extractor_input_file.xlsx")

# Create dict for fast lookup
claims_lookup = {
    (str(row["PatientControlNumber"]),
     str(row["MemberID"]),
     str(row["TotalBillAmount"]),
     str(row["StartDateOfService"]))
    for _, row in claims_df.iterrows()
}

source_folder = "edi_claim_extractor_source_folder"
output_folder = "Extracted_EDIs"
os.makedirs(output_folder, exist_ok=True)

for file in os.listdir(source_folder):
    if not file.endswith(".edi"):
        continue
    
    with open(os.path.join(source_folder, file), 'r') as f:
        edi = X12Reader(f)

        current_claim = {}
        collected_claims = {}
        isa_sender = None
        st_type = None

        for seg in edi:
            seg_id = seg.get_seg_id()

            if seg_id == "ISA":
                isa_sender = seg.get_value('ISA06')

            if seg_id == "ST":
                st_type = seg.get_value('ST01')  # 837P or 837I

            if seg_id == "CLM":
                current_claim["PatientControlNumber"] = seg.get_value("CLM01")
                current_claim["TotalBillAmount"] = seg.get_value("CLM02")

            if seg_id == "NM1" and seg.get_value("NM101") == "IL":  # Subscriber
                current_claim["MemberID"] = seg.get_value("NM109")

            if seg_id == "DTP":
                if seg.get_value("DTP01") in ["472", "434"]:  # Service dates
                    dt_str = seg.get_value("DTP03")
                    current_claim["StartDateOfService"] = dt_str

            if seg_id == "SE":  # End of claim
                key = (current_claim.get("PatientControlNumber"),
                       current_claim.get("MemberID"),
                       current_claim.get("TotalBillAmount"),
                       current_claim.get("StartDateOfService"))

                if key in claims_lookup:
                    claim_type = "Professional" if st_type == "837P" else "Institutional"
                    collected_claims.setdefault((claim_type, isa_sender), []).append(edi.copy_current_loop())

                current_claim = {}

        # Write filtered EDI files
        for (claim_type, partner), claims in collected_claims.items():
            out_file = f"{claim_type}_{partner}_{datetime.now().strftime('%m%d













import os
import re
import json
from typing import List, Dict, Any, Tuple
from langchain_openai import AzureChatOpenAI

# =========================
# Azure OpenAI (LangChain)
# =========================
def get_azure_llm() -> AzureChatOpenAI:
    """
    Configure Azure OpenAI via LangChain. Reads creds from environment.
    Required env vars:
      - AZURE_OPENAI_ENDPOINT
      - AZURE_OPENAI_API_KEY
      - AZURE_OPENAI_API_VERSION (e.g. 2024-05-01-preview)
      - AZURE_OPENAI_DEPLOYMENT (e.g. "gpt-4o")
    """
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-05-01-preview")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")

    if not endpoint or not api_key:
        raise RuntimeError(
            "Missing Azure OpenAI configuration. "
            "Set AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY."
        )

    llm = AzureChatOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version=api_version,
        deployment_name=deployment,
        temperature=0,
    )
    return llm

# =========================
# EDI Parsing Utilities
# =========================
def parse_edi_file(edi_text: str) -> List[Dict[str, Any]]:
    """
    Split by '~' into segments. Each segment split by '*'.
    Returns list of dicts: {segment, line, fields}
    - fields[0] is the segment name (e.g., 'ISA')
    - positions 1..n map to fields[1..n]
    """
    segments = []
    for raw in edi_text.split("~"):
        line = raw.strip()
        if not line:
            continue
        parts = line.split("*")
        seg = parts[0].strip()
        segments.append({"segment": seg, "line": line, "fields": parts})
    return segments

def normalize_position(field_position: Any, segment_name: str) -> int:
    """
    Normalize a rule's FieldPosition to a 1-based integer:
    Accepts: 1, "1", "01", "001", "ISA01", "ISA1", "SegmentName01", "NM101", etc.
    Returns integer position (1..N). Raises ValueError if no digits found.
    """
    if field_position is None:
        raise ValueError("FieldPosition is None")

    s = str(field_position).strip()

    # If it's purely digits (e.g., "1", "01", "001")
    if s.isdigit():
        return int(s)

    # Common forms: "ISA01", "ISA1", "NM101", "REF02", "SegmentName15"
    # Extract trailing digits
    m = re.search(r"(\d+)$", s)
    if m:
        return int(m.group(1))

    # Edge case: "segmentname01" with varying case
    s_upper = s.upper()
    seg_upper = segment_name.upper()
    if s_upper.startswith(seg_upper):
        tail = s[len(segment_name):]
        m2 = re.search(r"(\d+)$", tail)
        if m2:
            return int(m2.group(1))

    raise ValueError(f"Cannot normalize FieldPosition='{field_position}' for {segment_name}")

def unique_rule_positions(rules_for_segment: List[Dict[str, Any]], segment_name: str) -> List[int]:
    """
    Deduplicate rule positions that are the same spot with different notations (01, 1, ISA01, etc.)
    Returns sorted unique integer positions.
    """
    pos_set = set()
    for r in rules_for_segment:
        try:
            p = normalize_position(r.get("FieldPosition"), segment_name)
            pos_set.add(p)
        except Exception:
            # Ignore malformed positions; we won't validate those
            pass
    return sorted(pos_set)

# =========================
# Validation Core
# =========================
def validate_required(value: str) -> Tuple[bool, str]:
    if value is None:
        return False, "Required field missing (no element at this position)."
    if str(value).strip() == "":
        return False, "Required field empty."
    return True, "Required field present."

def validate_not_used(value: str) -> Tuple[bool, str]:
    if value is None or str(value).strip() == "":
        return True, "Field correctly not used."
    return False, f"Field marked 'Not Used' but value '{value}' found."

def validate_situational_with_llm(
    llm: AzureChatOpenAI,
    segment_name: str,
    position: int,
    value: str,
    description: str,
    all_fields: List[str],
    edi_line: str
) -> Dict[str, Any]:
    """
    Ask LLM to judge situational validity following your rules:
      - If description has an explicit condition referencing other fields, obey it.
      - If no explicit condition is present in description, BOTH present or absent are valid.
      - 'Not Used' logic never applies here (this is situational branch).
      - Treat standard EDI codes like '00','01','ZZ','T','P' as legitimate values (not 'empty').
    """
    # Build field map for clarity (1-based)
    field_map = {f"{segment_name}{i}": (all_fields[i] if i < len(all_fields) else "")
                 for i in range(1, len(all_fields))}

    prompt = f"""
You are an EDI validator. Validate one situational field based ONLY on the given description and values.

Rules for interpretation:
- "Situational" = conditionally required. If the description specifies a condition (e.g., depends on another field), check it.
- If NO explicit condition is present in the description, BOTH present and absent are valid.
- Do not treat legitimate codes like "00", "01", "ZZ", "T", "P" as empty.
- Only use information from the provided inputs. Do not invent extra conditions.

Segment: {segment_name}
Field Position: {position}
Field Tag: {segment_name}{position}
Field Value: {value}
Description: {description}

All Fields (1-based map):
{json.dumps(field_map, indent=2)}

EDI Line:
{edi_line}

Respond ONLY JSON, no markdown, no extra text:
{{
  "status": "Matched" or "Invalid",
  "rule_line": "{segment_name}{position}",
  "reason": "one-line reason"
}}
""".strip()

    resp = llm.invoke(prompt)
    content = (resp.content or "").strip()
    try:
        data = json.loads(content)
        # Guardrails for schema
        st = data.get("status")
        rl = data.get("rule_line", f"{segment_name}{position}")
        rs = data.get("reason", "Situational check done.")
        if st not in ("Matched", "Invalid"):
            raise ValueError("Bad status from LLM.")
        return {"status": st, "rule_line": rl, "edi_line": edi_line, "reason": rs}
    except Exception:
        # Fallback: if LLM fails, default to permissive per your policy
        ok = (value is None) or (str(value).strip() == "") or (str(value).strip() != "")
        return {
            "status": "Matched" if ok else "Invalid",
            "rule_line": f"{segment_name}{position}",
            "edi_line": edi_line,
            "reason": "Situational check fallback: no explicit condition enforced."
        }

def validate_segment_against_rules(
    segment: Dict[str, Any],
    rules_for_segment: List[Dict[str, Any]],
    llm: AzureChatOpenAI
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    """
    Validate one parsed segment against its rules.
    Returns:
      - field_results: list of {status, rule_line, edi_line, reason}
      - segment_summary: {status, ediline, reason}
    """
    seg_name = segment["segment"]
    line = segment["line"]
    fields = segment["fields"]  # fields[0] == seg_name; positions start at 1

    # Deduplicate positions & pre-check for extra elements
    rule_positions = unique_rule_positions(rules_for_segment, seg_name)
    max_rule_pos = max(rule_positions) if rule_positions else 0

    field_results: List[Dict[str, Any]] = []

    # If there are more elements than rules cover (beyond max position), mark as extras
    if len(fields) - 1 > max_rule_pos:
        # For each extra position > max_rule_pos, mark as invalid (no rule defined)
        for i in range(max_rule_pos + 1, len(fields)):
            field_results.append({
                "status": "Invalid",
                "rule_line": f"{seg_name}{i}",
                "edi_line": line,
                "reason": f"Extra element at position {i} without a defined rule."
            })

    # Validate per rule
    for rule in rules_for_segment:
        usage = str(rule.get("Usage", "")).strip().lower()
        desc = str(rule.get("ShortDescription", "")).strip()
        try:
            pos = normalize_position(rule.get("FieldPosition"), seg_name)
        except Exception:
            # If we cannot normalize, skip but record an issue
            field_results.append({
                "status": "Invalid",
                "rule_line": f"{seg_name}{rule.get('FieldPosition')}",
                "edi_line": line,
                "reason": "Unrecognized FieldPosition format in rules."
            })
            continue

        # Fetch value at position (1-based)
        value = fields[pos] if pos < len(fields) else None
        rule_line = f"{seg_name}{pos}"

        # Required
        if usage == "required":
            ok, reason = validate_required(value)
            field_results.append({
                "status": "Matched" if ok else "Invalid",
                "rule_line": rule_line,
                "edi_line": line,
                "reason": reason
            })
            continue

        # Not Used
        if usage == "not used":
            ok, reason = validate_not_used(value)
            field_results.append({
                "status": "Matched" if ok else "Invalid",
                "rule_line": rule_line,
                "edi_line": line,
                "reason": reason
            })
            continue

        # Situational → LLM
        if usage == "situational":
            field_results.append(
                validate_situational_with_llm(
                    llm=llm,
                    segment_name=seg_name,
                    position=pos,
                    value=value if value is not None else "",
                    description=desc,
                    all_fields=fields,
                    edi_line=line
                )
            )
            continue

        # Any other usage → treat as optional/valid
        field_results.append({
            "status": "Matched",
            "rule_line": rule_line,
            "edi_line": line,
            "reason": "Optional field (no strict validation)."
        })

    # Build segment summary
    invalids = [r for r in field_results if r["status"] == "Invalid"]
    if not invalids:
        segment_summary = {
            "status": "Matched",
            "ediline": line,
            "reason": "All required, situational, and not-used fields validated successfully."
        }
    else:
        segment_summary = {
            "status": "Invalid",
            "ediline": line,
            "reason": "; ".join(dict.fromkeys([r["reason"] for r in invalids]))  # unique-preserving
        }

    return field_results, segment_summary

# =========================
# Segment Summary via LLM
# =========================
def summarize_segment_with_llm(
    llm: AzureChatOpenAI,
    edi_line: str,
    field_results: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Ask LLM to produce the final single-dict summary with:
      - status: "Matched" if all fields matched, else "Invalid"
      - ediline: the original segment line
      - reason: one-line reason (either all good, or concise failures)
    """
    prompt = f"""
Summarize these field-level EDI validation results into ONE JSON dictionary.
Rules:
- If ALL fields are "Matched" -> status = "Matched".
- If ANY field is "Invalid" -> status = "Invalid".
- 'ediline' must equal the original EDI line verbatim.
- 'reason' must be a single concise line. If invalid, mention the most important failure(s).

EDI line:
{edi_line}

Field results:
{json.dumps(field_results, indent=2)}

Respond ONLY JSON, no markdown:
{{
  "status": "Matched" or "Invalid",
  "ediline": "{edi_line}",
  "reason": "<one line>"
}}
""".strip()

    resp = llm.invoke(prompt)
    content = (resp.content or "").strip()
    try:
        data = json.loads(content)
        # guardrails
        if "status" not in data or "ediline" not in data or "reason" not in data:
            raise ValueError("Missing keys.")
        if data["status"] not in ("Matched", "Invalid"):
            raise ValueError("Bad status.")
        return data
    except Exception:
        # Local fallback mirroring the same rules
        any_invalid = any(r["status"] == "Invalid" for r in field_results)
        if any_invalid:
            return {
                "status": "Invalid",
                "ediline": edi_line,
                "reason": "; ".join(dict.fromkeys([r["reason"] for r in field_results if r["status"] == "Invalid"]))
            }
        return {
            "status": "Matched",
            "ediline": edi_line,
            "reason": "All required, situational, and not-used fields validated successfully."
        }

# =========================
# Example Driver
# =========================
def load_rules_from_json(path: str) -> List[Dict[str, Any]]:
    """
    Expects a list of dicts, each having:
      - SegmentName
      - FieldPosition (e.g., 1, "01", "ISA01", "NM101")
      - Usage ("Required" | "Situational" | "Not Used" | ...)
      - ShortDescription
    """
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def run_validation(edi_text: str, rules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Validates all segments in the EDI text with their respective rules.
    Returns list of per-segment summaries (each: {status, ediline, reason}).
    """
    llm = get_azure_llm()
    segments = parse_edi_file(edi_text)

    # Group rules by segment
    rules_by_segment: Dict[str, List[Dict[str, Any]]] = {}
    for r in rules:
        seg = str(r.get("SegmentName", "")).strip()
        if not seg:
            continue
        rules_by_segment.setdefault(seg, []).append(r)

    summaries: List[Dict[str, Any]] = []

    for seg in segments:
        seg_name = seg["segment"]
        if seg_name not in rules_by_segment:
            # No rules for this segment -> single Invalid record or treat as pass?
            # We'll mark as Invalid with reason for transparency.
            summaries.append({
                "status": "Invalid",
                "ediline": seg["line"],
                "reason": f"No rules defined for segment '{seg_name}'."
            })
            continue

        field_results, local_summary = validate_segment_against_rules(
            segment=seg,
            rules_for_segment=rules_by_segment[seg_name],
            llm=llm
        )

        # Option: Let LLM craft the final single-dict summary (with fallback)
        final_summary = summarize_segment_with_llm(llm, seg["line"], field_results)

        # Ensure we always return exactly the requested keys/types
        summaries.append({
            "status": final_summary["status"],
            "ediline": final_summary["ediline"],
            "reason": final_summary["reason"]
        })

    return summaries

# =========================
# __main__
# =========================
if __name__ == "__main__":
    # ---- Example usage ----
    # 1) Load rules from a local JSON file (produced from your PDF pipeline)
    #    e.g., rules.json = [
    #      {"SegmentName":"ISA","FieldPosition":"ISA01","Usage":"Required","ShortDescription":"Authorization Information Qualifier"},
    #      {"SegmentName":"ISA","FieldPosition":"ISA14","Usage":"Situational","ShortDescription":"Acknowledgment can be requested through data element ISA14."},
    #      {"SegmentName":"ISA","FieldPosition":"ISA15","Usage":"Required","ShortDescription":"Test Indicator"},
    #      {"SegmentName":"ISA","FieldPosition":"ISA16","Usage":"Required","ShortDescription":"Component Element Separator"}
    #    ]
    RULES_PATH = "rules.json"
    rules_list = load_rules_from_json(RULES_PATH)

    # 2) Read an EDI file (raw text)
    #    Example single line shown; in real use, read a whole file with multiple segments.
    edi_text = (
        "ISA*00*          *00*          *ZZ*SENDERID      *ZZ*RECEIVERID    *250101*1253*^*00501*000000905*1*T*:~"
        "GS*HC*SENDER*RECEIVER*20250101*1253*1*X*005010X222A1~"
        "ST*837*0001*005010X222A1~"
    )

    # 3) Run validation (returns list of per-segment summaries)
    summaries = run_validation(edi_text, rules_list)

    # 4) Print results
    print(json.dumps(summaries, indent=2))
















import json

isa_rules = [
    {'SegmentName': 'ISA', 'FieldPosition': 1, 'Usage': 'Required',
     'ShortDescription': 'Authorization Information Qualifier'},
    {'SegmentName': 'ISA', 'FieldPosition': 2, 'Usage': 'Required',
     'ShortDescription': 'Authorization Information'},
    {'SegmentName': 'ISA', 'FieldPosition': 3, 'Usage': 'Required',
     'ShortDescription': 'Security Information Qualifier'},
    {'SegmentName': 'ISA', 'FieldPosition': 4, 'Usage': 'Required',
     'ShortDescription': 'Security Information'},
    {'SegmentName': 'ISA', 'FieldPosition': 14, 'Usage': 'Situational',
     'ShortDescription': 'Acknowledgment requested'},
    {'SegmentName': 'ISA', 'FieldPosition': 15, 'Usage': 'Required',
     'ShortDescription': 'Test Indicator'},
    {'SegmentName': 'ISA', 'FieldPosition': 16, 'Usage': 'Required',
     'ShortDescription': 'Subelement Separator'}
]

edi_line = "ISA*00*          *01*SECRET    *ZZ*SUBMITTERS.ID*ZZ*RECEIVERS.ID*030101*1253*^*00501*000000905*1*T*:~"
fields = edi_line.split("*")

def validate_isa(fields, rules):
    results = []
    for rule in rules:
        pos = rule['FieldPosition']
        idx = pos  # our rules already give int position
        
        # Prepare rule line text
        rule_line = f"{rule['SegmentName']}{pos} - {rule['ShortDescription']} ({rule['Usage']})"
        
        if idx >= len(fields):
            results.append({
                "status": "Invalid",
                "rule_line": rule_line,
                "reason": "Field missing in EDI"
            })
            continue

        value = fields[idx].strip()
        if rule['Usage'] == "Required" and not value:
            results.append({
                "status": "Invalid",
                "rule_line": rule_line,
                "reason": "Required field empty"
            })
        else:
            results.append({
                "status": "Matched",
                "rule_line": rule_line,
                "reason": "Field present and valid"
            })
    return results

# Run validation
validation_output = validate_isa(fields, isa_rules)

# Pretty print JSON
print(json.dumps(validation_output, indent=2))









def validate_isa(fields, rules):
    errors = []
    for rule in rules:
        pos = rule['FieldPosition']
        # Handle if rule uses numbers
        if isinstance(pos, int):
            idx = pos  # field index in split
        elif pos.startswith("ISA"):
            idx = int(pos.replace("ISA", ""))  # e.g. ISA14 → 14
        else:
            continue

        if idx >= len(fields):
            errors.append(f"Missing field {pos}: {rule['ShortDescription']}")
            continue

        value = fields[idx]
        if rule['Usage'] == "Required" and not value.strip():
            errors.append(f"Field {pos} ({rule['ShortDescription']}) is required but empty")

    return errors















prompt = f"""
You are validating an EDI segment against rules.

Interpretation rules:
- "Required" means the field must exist in the segment. Even if the value is "00" or empty string symbols like "00", "ZZ", it is still valid as long as it is a standard code or allowed value.
- "Situational" means the field may be present or absent depending on conditions. If no condition is given in the rule, then both present and absent are acceptable.
- "Not Used" means the field must be blank or omitted.
- Never flag standard EDI codes like "00", "01", "ZZ" as empty values.
- Only mark as invalid if the field is missing entirely, in the wrong position, or violates a clearly defined situational dependency.

EDI line:
{line.strip()}

Fields:
{fields}

Rules:
{json.dumps(matching_rules, indent=2)}

Respond only in JSON (no explanation):

{{
  "status": "Matched" or "Invalid",
  "rule_line": "matching rule or None",
  "reason": "one-line reason"
}}
"""













prompt = f"""
You are an EDI 837/ISA segment validator.

Validate this EDI line against the provided rules and fields.

EDI line:
{line.strip()}

Fields extracted (with positions):
{fields}

Rules for this segment (from implementation guide):
{json.dumps(matching_rules, indent=2)}

Validation rules to apply:
1. If a field is marked as "Required", it must have a non-empty value in the EDI line. If missing → Invalid.
2. If a field is marked as "Not Used", it must always be empty. If it has a value → Invalid.
3. If a field is marked as "Situational":
   - If the rule description specifies a condition (e.g. "Required when field X = Y"), check that condition.
   - If condition is met → value must be present.
   - If condition is not met or no condition is described → the field may be empty or filled, both are valid.

Respond strictly in JSON with the following format:
{{
  "status": "Matched" or "Invalid",
  "rule_line": "The rule line from the guide that matched or failed",
  "reason": "One-line explanation of why it is valid or invalid"
}}
"""



















ISA*00*

*00*

*ZZ*PAYERSENDER

*ZZ*PROVIDERRECV

*240601*1200*^*00501*000000001*0*T*:~

GS*HP*PAYERSENDER*PROVIDERRECV*20240601*1200*1*X*005010X221A1~

ST*835*0001~

BPR*I*1500.00*C*CHK*123456789*DA*987654321*1234567890~

TRN*1*1234567890*9876543210~

N1*PR* INSURANCE COMPANY~

N3*123 MAIN STREET~

N4*ANYTOWN*CA*90210~

N1*PE*PROVIDER NAME~

N3*456 MEDICAL AVE~

N4*OTHERTOWN*CA*90001~

CLP 12345*1*200.00*150.00*50.00*MC*1234567890~

CAS CO*45*50.00~

NM1*QC*1*DOE JOHN****MI*987654321~

DTM*405*20240601~

REF EVA123456789~

SE*14*0001~

GE*1*1~

IEA*1*000000001~

















from langchain_openai import AzureChatOpenAI
from langchain.prompts import ChatPromptTemplate

# Initialize Azure OpenAI model
llm = AzureChatOpenAI(
    deployment_name="your-deployment-name",  # e.g., "gpt-4"
    model="gpt-4",
    temperature=0,
    openai_api_version="2023-05-15",
    azure_endpoint="https://YOUR-RESOURCE-NAME.openai.azure.com/",
    api_key="YOUR-AZURE-OPENAI-KEY"
)

# Prompt template
prompt_template = ChatPromptTemplate.from_template("""
You are validating EDI segments against rules.

### Segment:
{segment}

### Rules:
{rules}

### Task:
1. Check if this segment follows the rules.
2. If valid, return JSON:
   {{
     "segment": "{segment}",
     "matched_rule": "<exact rule line>",
     "reason": "Valid because ..."
   }}
3. If invalid, return JSON:
   {{
     "segment": "{segment}",
     "matched_rule": null,
     "reason": "Invalid because <explanation>"
   }}
""")

def validate_segment(segment: str, rules: str):
    prompt = prompt_template.format(segment=segment, rules=rules)
    response = llm.predict(prompt)
    return response



















import os
import re
import json
import pandas as pd
import pytesseract
from pdf2image import convert_from_path
from langchain_openai import AzureChatOpenAI
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# =============== STEP 1: PDF Extraction ==================
def extract_pdf(pdf_path, poppler_path=None):
    images = convert_from_path(pdf_path, poppler_path=poppler_path)
    docs = []

    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img)
            page_text = f"\n[Page {i+1}]\n{text.strip()}"
            docs.append(Document(page_content=page_text))
        except Exception as e:
            print(f"OCR failed on page {i+1}: {e}")
    return docs


# =============== STEP 2: Chunking ==================
def chunk_documents(docs, chunk_size=1500, overlap=200):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    chunks = []
    for doc in docs:
        chunks.extend(text_splitter.split_documents([doc]))
    return chunks


# =============== STEP 3: Build Rules with LLM ==================
def build_rules_from_chunks(chunks, llm):
    rules = []
    for i, chunk in enumerate(chunks):
        prompt = f"""
        You are given part of an EDI implementation guideline.
        Extract validation rules for segments and fields.
        Rules must include:
        - Segment name
        - Field position
        - Usage (Required / Situational / Not Used)
        - Short description

        Text:
        {chunk.page_content}

        Return JSON list of rules.
        """
        resp = llm.invoke(prompt)
        try:
            parsed = json.loads(resp.content)
            rules.extend(parsed)
        except Exception:
            print(f"Failed to parse chunk {i}")
    return rules


# =============== STEP 4: Store Rules ==================
def save_rules(rules, output_path="rules.json"):
    with open(output_path, "w") as f:
        json.dump(rules, f, indent=2)
    print(f"Rules stored at {output_path}")


# =============== STEP 5: Validate Uploaded EDI ==================
def validate_edi(edi_file, rules, llm):
    results = []
    with open(edi_file, "r") as f:
        edi_lines = f.readlines()

    for line in edi_lines:
        seg_id = line.split("*")[0]
        matching_rules = [r for r in rules if r.get("segment") == seg_id]

        if not matching_rules:
            results.append({
                "edi_line": line.strip(),
                "rule_line": None,
                "status": "Invalid",
                "reason": "No matching rule for this segment"
            })
            continue

        # Let LLM validate field-level usage
        prompt = f"""
        Validate this EDI line against the rules.

        EDI line:
        {line.strip()}

        Rules:
        {json.dumps(matching_rules, indent=2)}

        Respond JSON:
        {{
          "status": "Matched/Invalid",
          "rule_line": "...",
          "reason": "one-line reason"
        }}
        """
        resp = llm.invoke(prompt)
        try:
            parsed = json.loads(resp.content)
            parsed["edi_line"] = line.strip()
            results.append(parsed)
        except Exception:
            results.append({
                "edi_line": line.strip(),
                "rule_line": None,
                "status": "Error",
                "reason": "LLM parse failure"
            })

    return results


# =============== STEP 6: Dump to Excel ==================
def dump_to_excel(results, output_path="validation_results.xlsx"):
    df = pd.DataFrame(results)
    df.to_excel(output_path, index=False)
    print(f"Validation results saved to {output_path}")


# =============== STEP 7: HTML Highlight ==================
def generate_html(results, output_path="validation_report.html"):
    html_lines = []
    for res in results:
        if res["status"] == "Invalid":
            html_lines.append(
                f"<p style='color:red;' title='{res['reason']}'>{res['edi_line']}</p>"
            )
        else:
            html_lines.append(
                f"<p style='color:green;' title='{res['reason']}'>{res['edi_line']}</p>"
            )
    with open(output_path, "w") as f:
        f.write("<html><body>" + "\n".join(html_lines) + "</body></html>")
    print(f"HTML report saved to {output_path}")


# ================= RUN ==================
if __name__ == "__main__":
    # Azure OpenAI Config
    llm = AzureChatOpenAI(
        deployment_name="gpt-4o",  # change to your deployment
        temperature=0,
        api_version="2024-05-01-preview"
    )

    # Step 1 & 2: Extract + Chunk
    docs = extract_pdf("EDI_Guideline.pdf", poppler_path="C:/poppler/bin")
    chunks = chunk_documents(docs)

    # Step 3 & 4: Build + Save Rules
    rules = build_rules_from_chunks(chunks, llm)
    save_rules(rules)

    # Step 5: Validate
    results = validate_edi("sample.edi", rules, llm)

    # Step 6: Excel
    dump_to_excel(results)

    # Step 7: HTML
    generate_html(results)



# =============== STEP 4B: Load Rules ==================
def load_rules(input_path="rules.json"):
    if os.path.exists(input_path):
        with open(input_path, "r") as f:
            return json.load(f)
    return None


# ================= RUN ==================
if __name__ == "__main__":
    # Azure OpenAI Config
    llm = AzureChatOpenAI(
        deployment_name="gpt-4o",  # change to your deployment name
        temperature=0,
        api_version="2024-05-01-preview"
    )

    # Step 3 & 4: Build rules only if not found locally
    rules = load_rules("rules.json")
    if rules:
        print("✅ Loaded rules from local rules.json")
    else:
        print("⚡ No local rules found, extracting from PDF...")
        docs = extract_pdf("EDI_Guideline.pdf", poppler_path="C:/poppler/bin")
        chunks = chunk_documents(docs)
        rules = build_rules_from_chunks(chunks, llm)
        save_rules(rules, "rules.json")

    # Step 5: Validate EDI
    results = validate_edi("sample.edi", rules, llm)

    # Step 6: Excel
    dump_to_excel(results, "validation_results.xlsx")

    # Step 7: HTML
    generate_html(results, "validation_report.html")












import os

def update_service_date_in_folder(folder_path: str, new_service_date: str, output_folder: str):
    """
    Update service dates in EDI files inside a folder.
    
    - Institutional claims (837I) → update DTP*434
    - Professional claims (837P) → update DTP*472
    
    Args:
        folder_path (str): Path to the folder containing EDI files.
        new_service_date (str): New date of service (YYYYMMDD).
        output_folder (str): Path to save updated files.
    """

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(folder_path):
        if not filename.lower().endswith((".edi", ".txt")):
            continue  # Skip non-EDI files

        input_file = os.path.join(folder_path, filename)
        output_file = os.path.join(output_folder, filename)

        with open(input_file, "r") as f:
            edi_content = f.read()

        segments = edi_content.split("~")
        updated_segments = []
        claim_type = None

        for seg in segments:
            if seg.startswith("ST*837*"):
                if "005010X223A2" in edi_content:  
                    claim_type = "Institutional"
                elif "005010X222A1" in edi_content:  
                    claim_type = "Professional"

            if claim_type == "Institutional" and seg.startswith("DTP*434*"):
                parts = seg.split("*")
                if len(parts) >= 3:
                    parts[-1] = new_service_date
                    seg = "*".join(parts)

            elif claim_type == "Professional" and seg.startswith("DTP*472*"):
                parts = seg.split("*")
                if len(parts) >= 3:
                    parts[-1] = new_service_date
                    seg = "*".join(parts)

            updated_segments.append(seg)

        updated_edi = "~".join(updated_segments)

        with open(output_file, "w") as f:
            f.write(updated_edi)

        print(f"✅ Updated {filename} and saved to {output_file}")


# Example usage:
# folder with edi files = "input_edi"
# updated files will go to "output_edi"
# new service date = "20250818"

update_service_date_in_folder(
    folder_path="input_edi",
    new_service_date="20250818",
    output_folder="output_edi"
)













# app.py
import streamlit as st
import pandas as pd
import re
import os
from decimal import Decimal, InvalidOperation
from collections import defaultdict
from typing import Dict, List, Tuple

st.set_page_config(page_title="837I ↔ 820 Comparison (Strict match)", layout="wide")

# -------------------
# Utilities
# -------------------
def norm_str(x):
    return (x or "").strip()

def sanitize_sheet_name(name: str) -> str:
    if not name:
        return "UNKNOWN"
    s = re.sub(r'[\[\]\*\?/\\:]', "_", name)
    return s[:31]

def to_decimal(x):
    if x is None:
        return None
    s = str(x).replace(",", "").strip()
    if s == "":
        return None
    try:
        return Decimal(s)
    except InvalidOperation:
        return None

def equal_amounts(a, b):
    aa = to_decimal(a)
    bb = to_decimal(b)
    if aa is None or bb is None:
        return False
    return aa == bb

def name_match(first1, last1, first2, last2):
    return norm_str(first1).upper() == norm_str(first2).upper() and norm_str(last1).upper() == norm_str(last2).upper()

# -------------------
# ADX Code Definitions
# -------------------
ADX_CODES = {
    "52": "Credit for Overpayment",
    "53": "Remittance for Previous Underpayment",
    "80": "Overpayment",
    "81": "Credit as Agreed",
    "86": "Duplicate Payment",
    "BJ": "Insurance Charge",
    "H1": "Information Forthcoming",
    "H6": "Partial Payment Remitted",
    "RU": "Interest",
    "WO": "Overpayment Recovery",
    "WW": "Overpayment Credit",
}

# -------------------
# Parsers
# -------------------

def parse_837i(content: str) -> Dict[str, List[dict]]:
    providers = defaultdict(list)
    current_provider = None
    current_first = ""
    current_last = ""

    for raw in content.split("~"):
        seg = raw.strip()
        if not seg:
            continue
        parts = seg.split("*")
        if len(parts) < 1:
            continue

        if parts[0] == "NM1" and len(parts) >= 4 and parts[1] == "41":
            current_provider = norm_str(parts[3])
        elif parts[0] == "NM1" and len(parts) >= 5 and parts[1] == "IL":
            current_last = norm_str(parts[3])
            current_first = norm_str(parts[4])
        elif parts[0] == "CLM" and len(parts) >= 3:
            clm01 = parts[1]
            amt = norm_str(parts[2])
            claim_type = ""
            claim_id = ""
            member_id = ""

            if clm01.startswith("C-"):
                claim_type = "C"
                m = re.match(r"^C-([^\-\s]+)", clm01)
                if m:
                    claim_id = m.group(1)
            elif clm01.startswith("P-"):
                claim_type = "P"
                m = re.match(r"^P-([^\-\s]+)", clm01)
                if m:
                    member_id = m.group(1)
                    if len(member_id) > 1 and re.match(r".*[A-Z]$", member_id) and re.search(r"\d", member_id[:-1]):
                        member_id = member_id[:-1]

            if current_provider is None:
                current_provider = "UNKNOWN_PROVIDER"

            record = {
                "Provider": current_provider,
                "FirstName": current_first,
                "LastName": current_last,
                "Type": claim_type,
                "ClaimID": claim_id,
                "MemberID": member_id,
                "Amount": amt
            }
            providers[current_provider].append(record)

    return providers

def parse_820(content: str) -> Dict[str, List[dict]]:
    records = []
    payees = set()
    current_qe_first = ""
    current_qe_last = ""
    current_qe_memberid = ""
    last_record = None

    for raw in content.split("~"):
        seg = raw.strip()
        if not seg:
            continue
        parts = seg.split("*")
        if len(parts) < 1:
            continue

        if parts[0] == "N1" and len(parts) >= 3 and parts[1] == "PE":
            payee_name = norm_str(parts[2])
            if payee_name:
                payees.add(payee_name)
        if parts[0] == "NM1" and len(parts) >= 4 and parts[1] == "PE":
            payee_name = norm_str(parts[3])
            if payee_name:
                payees.add(payee_name)

        if parts[0] == "NM1" and len(parts) >= 3 and parts[1] == "QE":
            if len(parts) >= 5:
                current_qe_last = norm_str(parts[3])
                current_qe_first = norm_str(parts[4])
            else:
                name_field = norm_str(parts[3]) if len(parts) >= 4 else ""
                if name_field:
                    tokens = name_field.split()
                    if len(tokens) >= 2:
                        current_qe_last = tokens[0]
                        current_qe_first = " ".join(tokens[1:])
                    else:
                        current_qe_last = name_field
                        current_qe_first = ""
            current_qe_memberid = norm_str(parts[-1]) if len(parts) >= 1 else ""

        if parts[0] == "RMR" and len(parts) >= 3 and parts[1] == "IK":
            claim_id = norm_str(parts[2])
            amount = ""
            for elem in parts[3:]:
                if re.match(r'^[\+\-]?\d+(?:\.\d+)?$', elem.strip()):
                    amount = elem.strip()
                    break

            rec = {
                "ClaimID": claim_id,
                "Amount": amount,
                "MemberID": current_qe_memberid,
                "FirstName": current_qe_first,
                "LastName": current_qe_last,
                "ADX_Code": "",
                "ADX_Description": ""
            }
            records.append(rec)
            last_record = rec

        if parts[0] == "ADX" and last_record is not None:
            reason_code = norm_str(parts[-1])
            last_record["ADX_Code"] = reason_code
            last_record["ADX_Description"] = ADX_CODES.get(reason_code, "")

    return {"payees": payees, "records": records}

# -------------------
# Comparison logic
# -------------------

def compare_providers_vs_820(providers_data: Dict[str, List[dict]], parsed_820_files: List[Tuple[str, dict]]):
    results = defaultdict(list)
    per_file_indexes = []

    for fname, parsed in parsed_820_files:
        claim_index = defaultdict(list)
        member_index = defaultdict(list)
        for r in parsed["records"]:
            cid = norm_str(r.get("ClaimID",""))
            mid = norm_str(r.get("MemberID",""))
            if cid:
                claim_index[cid].append(r)
            if mid:
                member_index[mid].append(r)
        per_file_indexes.append((fname, parsed["payees"], claim_index, member_index))

    for provider, claims in providers_data.items():
        for c in claims:
            p_first = c.get("FirstName","")
            p_last = c.get("LastName","")
            p_type = c.get("Type","")
            p_claimid = norm_str(c.get("ClaimID",""))
            p_memberid = norm_str(c.get("MemberID",""))
            p_amount = norm_str(c.get("Amount",""))

            matched = False
            matched_file = ""
            matched_rec = None
            status = "Member not found in 820"

            for fname, payees, claim_index, member_index in per_file_indexes:
                if p_type == "C" and p_claimid:
                    candidates = claim_index.get(p_claimid, [])
                    for cand in candidates:
                        if name_match(p_first, p_last, cand.get("FirstName",""), cand.get("LastName","")) and equal_amounts(p_amount, cand.get("Amount","")):
                            matched = True
                            matched_file = fname
                            matched_rec = cand
                            break
                    if matched:
                        status = "Match found"
                        break

                if p_type == "P" and p_memberid:
                    candidates = member_index.get(p_memberid, [])
                    for cand in candidates:
                        if name_match(p_first, p_last, cand.get("FirstName",""), cand.get("LastName","")) and equal_amounts(p_amount, cand.get("Amount","")):
                            matched = True
                            matched_file = fname
                            matched_rec = cand
                            break
                    if matched:
                        status = "Match found"
                        break

            row = {
                "Provider": provider,
                "FirstName": p_first,
                "LastName": p_last,
                "Type": p_type,
                "ClaimID": p_claimid,
                "MemberID": p_memberid,
                "Amount_837I": p_amount,
                "Matched_820_File": matched_file if matched else "",
                "Amount_820": matched_rec.get("Amount","") if matched_rec else "",
                "820_FirstName": matched_rec.get("FirstName","") if matched_rec else "",
                "820_LastName": matched_rec.get("LastName","") if matched_rec else "",
                "ADX_Code": matched_rec.get("ADX_Code","") if matched_rec else "",
                "ADX_Description": matched_rec.get("ADX_Description","") if matched_rec else "",
                "Status": status
            }
            results[provider].append(row)

    return results

# -------------------
# Excel writer
# -------------------

def save_excel(provider_results: Dict[str, List[dict]], basename: str, output_dir="output_reports"):
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{os.path.splitext(basename)[0]}_comparison.xlsx")

    provider_summary = []
    users_rows = []

    with pd.ExcelWriter(out_path, engine="xlsxwriter") as writer:
        for provider, rows in provider_results.items():
            df = pd.DataFrame(rows)
            if not df.empty:
                df = df.drop_duplicates(subset=["FirstName","LastName","Type","ClaimID","MemberID","Amount_837I"])
            sheet_name = sanitize_sheet_name(provider)
            df.to_excel(writer, sheet_name=sheet_name, index=False)

            total_claims = df.shape[0]
            unique_users = sorted(set((r["LastName"], r["FirstName"]) for r in rows))
            provider_summary.append({"Provider": provider, "TotalClaims": total_claims, "UniqueUsersCount": len(unique_users)})

            for ln, fn in unique_users:
                cnt = df[(df["LastName"] == ln) & (df["FirstName"] == fn)].shape[0]
                users_rows.append({"Provider": provider, "UserLastName": ln, "UserFirstName": fn, "UserClaimCount": cnt})

        report_df = pd.DataFrame(provider_summary) if provider_summary else pd.DataFrame(columns=["Provider","TotalClaims","UniqueUsersCount"])
        users_df = pd.DataFrame(users_rows) if users_rows else pd.DataFrame(columns=["Provider","UserLastName","UserFirstName","UserClaimCount"])

        report_df.to_excel(writer, sheet_name="Report", index=False, startrow=0)
        startrow = len(report_df) + 3 if not report_df.empty else 3
        if not users_df.empty:
            ws = writer.sheets["Report"]
            ws.write(startrow - 1, 0, "Unique Users per Provider")
            users_df.to_excel(writer, sheet_name="Report", index=False, startrow=startrow)

    return out_path

# -------------------
# Streamlit UI
# -------------------

st.title("837I ↔ 820 Comparison (strict id/member + name + amount)")

with st.sidebar:
    st.header("Upload files")
    files_837 = st.file_uploader("Upload 837I files (multiple)", type=["txt","edi"], accept_multiple_files=True)
    files_820 = st.file_uploader("Upload 820 files (multiple)", type=["txt","edi"], accept_multiple_files=True)
    run = st.button("Run Comparison")

if run:
    if not files_837 or not files_820:
        st.error("Please upload at least one 837I file and one 820 file.")
    else:
        parsed_820 = []
        st.info("Parsing 820 files...")
        for f in files_820:
            txt = f.read().decode("utf-8", errors="ignore")
            parsed = parse_820(txt)
            parsed_820.append((f.name, parsed))
            st.write(f"- {f.name}: RMR rows = {len(parsed['records'])}; payees detected = {', '.join(sorted(parsed['payees'])) if parsed['payees'] else '(none)'}")

        saved_files = []
        st.info("Processing 837I files and comparing...")
        for f in files_837:
            txt = f.read().decode("utf-8", errors="ignore")
            providers = parse_837i(txt)

            with st.expander(f"837I preview: {f.name} (Providers & claims)"):
                for prov, rows in providers.items():
                    st.write(f"Provider: {prov} — claims: {len(rows)}")
                    if rows:
                        st.dataframe(pd.DataFrame(rows).head(30))

            results = compare_providers_vs_820(providers, parsed_820)
            outpath = save_excel(results, f.name, output_dir="output_reports")
            saved_files.append(outpath)

            provider_summary = []
            users_rows = []
            for prov, rows in results.items():
                df = pd.DataFrame(rows).drop_duplicates(subset=["FirstName","LastName","Type","ClaimID","MemberID","Amount_837I"]) if rows else pd.DataFrame()
                total_claims = df.shape[0]
                unique_users = sorted(set((r["LastName"], r["FirstName"]) for r in rows))
                provider_summary.append({"Provider": prov, "TotalClaims": total_claims, "UniqueUsersCount": len(unique_users)})
                for ln, fn in unique_users:
                    cnt = df[(df["LastName"] == ln) & (df["FirstName"] == fn)].shape[0]
                    users_rows.append({"Provider": prov, "UserLastName": ln, "UserFirstName": fn, "UserClaimCount": cnt})

            st.subheader(f"Report for {f.name}")
            if provider_summary:
                report_df = pd.DataFrame(provider_summary)
                st.dataframe(report_df, use_container_width=True)
            else:
                st.write("No providers/claims found in this 837I file.")

            if users_rows:
                st.write("Unique users per provider:")
                users_df = pd.DataFrame(users_rows)
                st.dataframe(users_df, use_container_width=True)

            with st.expander(f"Detailed comparisons for {f.name}"):
                for prov, rows in results.items():
                    st.write(f"Provider: {prov}")
                    st.dataframe(pd.DataFrame(rows), use_container_width=True)

        outdir = os.path.abspath("output_reports")
        st.success(f"Reports saved to folder: {outdir}")
        st.write("Saved files:")
        for p in saved_files:
            st.write(f"• {os.path.basename(p)}")














# app.py
import streamlit as st
import pandas as pd
import re
import os
from decimal import Decimal, InvalidOperation
from collections import defaultdict
from typing import Dict, List, Tuple

st.set_page_config(page_title="837I ↔ 820 Comparison (Strict match)", layout="wide")

# -------------------
# Utilities
# -------------------
def norm_str(x):
    return (x or "").strip()

def sanitize_sheet_name(name: str) -> str:
    if not name:
        return "UNKNOWN"
    s = re.sub(r'[\[\]\*\?\/\\:]', "_", name)
    return s[:31]

def to_decimal(x):
    if x is None:
        return None
    s = str(x).replace(",", "").strip()
    if s == "":
        return None
    try:
        return Decimal(s)
    except InvalidOperation:
        return None

def equal_amounts(a, b):
    aa = to_decimal(a)
    bb = to_decimal(b)
    if aa is None or bb is None:
        return False
    return aa == bb

def name_match(first1, last1, first2, last2):
    return norm_str(first1).upper() == norm_str(first2).upper() and norm_str(last1).upper() == norm_str(last2).upper()

# -------------------
# Parsers
# -------------------

def parse_837i(content: str) -> Dict[str, List[dict]]:
    """
    Parse 837I content and return providers dict:
      provider_name -> list of claims dicts:
        {
          "Provider", "FirstName", "LastName", "Type"("C"|"P"), "ClaimID", "MemberID", "Amount"
        }
    """
    providers = defaultdict(list)
    current_provider = None
    current_first = ""
    current_last = ""

    for raw in content.split("~"):
        seg = raw.strip()
        if not seg:
            continue
        parts = seg.split("*")
        if len(parts) < 1:
            continue

        # Provider: NM1*41*...
        if parts[0] == "NM1" and len(parts) >= 4 and parts[1] == "41":
            current_provider = norm_str(parts[3])

        # Subscriber: NM1*IL*1*LAST*FIRST...
        elif parts[0] == "NM1" and len(parts) >= 5 and parts[1] == "IL":
            current_last = norm_str(parts[3])
            current_first = norm_str(parts[4])

        # Claim: CLM*<token>*<amount>...
        elif parts[0] == "CLM" and len(parts) >= 3:
            clm01 = parts[1]
            amt = norm_str(parts[2])
            claim_type = ""
            claim_id = ""
            member_id = ""

            # Patterns in your examples: "C-<claimid>-..." or "P-<memberid>-..."
            if clm01.startswith("C-"):
                claim_type = "C"
                m = re.match(r"^C-([^-\s]+)", clm01)
                if m:
                    claim_id = m.group(1)
            elif clm01.startswith("P-"):
                claim_type = "P"
                m = re.match(r"^P-([^-\s]+)", clm01)
                if m:
                    member_id = m.group(1)
                    # optional normalization: drop trailing single uppercase letter (if present and desirable)
                    if len(member_id) > 1 and re.match(r".*[A-Z]$", member_id) and re.search(r"\d", member_id[:-1]):
                        # only strip if it ends with a letter and there is a digit earlier (safe heuristic)
                        member_id = member_id[:-1]

            if current_provider is None:
                current_provider = "UNKNOWN_PROVIDER"

            record = {
                "Provider": current_provider,
                "FirstName": current_first,
                "LastName": current_last,
                "Type": claim_type,
                "ClaimID": claim_id,
                "MemberID": member_id,
                "Amount": amt
            }
            providers[current_provider].append(record)

    return providers

def parse_820(content: str) -> Dict[str, List[dict]]:
    """
    Parse 820 content to extract RMR*IK rows and associate them with nearest prior NM1*QE (if any).
    Returns:
      {
        "payees": set(payee names),
        "records": [ {"ClaimID":..., "Amount":..., "MemberID":..., "FirstName":..., "LastName":...}, ... ]
      }
    """
    records = []
    payees = set()
    current_qe_first = ""
    current_qe_last = ""
    current_qe_memberid = ""

    for raw in content.split("~"):
        seg = raw.strip()
        if not seg:
            continue
        parts = seg.split("*")
        if len(parts) < 1:
            continue

        # N1*PE or NM1*PE can contain payee name (optional)
        if parts[0] == "N1" and len(parts) >= 3 and parts[1] == "PE":
            payee_name = norm_str(parts[2])
            if payee_name:
                payees.add(payee_name)
        if parts[0] == "NM1" and len(parts) >= 4 and parts[1] == "PE":
            # NM1*PE*2*PAYEENAME...
            payee_name = norm_str(parts[3])
            if payee_name:
                payees.add(payee_name)

        # NM1*QE sets context for subsequent RMR rows; last element is often member id
        if parts[0] == "NM1" and len(parts) >= 3 and parts[1] == "QE":
            # Typical: NM1*QE*1*LAST*FIRST*...*N*<memberid>
            # We'll try to set last and first with available fields
            if len(parts) >= 5:
                current_qe_last = norm_str(parts[3])
                current_qe_first = norm_str(parts[4])
            else:
                # fallback: try to split the single name field into last/first
                name_field = norm_str(parts[3]) if len(parts) >= 4 else ""
                if name_field:
                    # try splitting by space: last is first token, rest first name
                    tokens = name_field.split()
                    if len(tokens) >= 2:
                        current_qe_last = tokens[0]
                        current_qe_first = " ".join(tokens[1:])
                    else:
                        current_qe_last = name_field
                        current_qe_first = ""
            # The member id is often last element
            current_qe_memberid = norm_str(parts[-1]) if len(parts) >= 1 else ""

        # RMR*IK core row
        if parts[0] == "RMR" and len(parts) >= 3 and parts[1] == "IK":
            claim_id = norm_str(parts[2])  # may be blank for P-premium submissions
            amount = ""
            # find first numeric-like element after p[2]
            for elem in parts[3:]:
                if re.match(r'^[\+\-]?\d+(?:\.\d+)?$', elem.strip()):
                    amount = elem.strip()
                    break

            rec = {
                "ClaimID": claim_id,
                "Amount": amount,
                "MemberID": current_qe_memberid,
                "FirstName": current_qe_first,
                "LastName": current_qe_last
            }
            records.append(rec)

    return {"payees": payees, "records": records}

# -------------------
# Comparison logic (strict: id/member AND names AND amount)
# -------------------

def compare_providers_vs_820(providers_data: Dict[str, List[dict]], parsed_820_files: List[Tuple[str, dict]]):
    """
    providers_data: provider -> list of claim dicts (from parse_837i)
    parsed_820_files: list of (filename, parsed_dict) where parsed_dict = parse_820(...)
    Returns: provider -> list of result rows (dictionaries)
    """
    results = defaultdict(list)

    # Build per-file indexes (do not merge files; we need filename where match found)
    per_file_indexes = []
    for fname, parsed in parsed_820_files:
        claim_index = defaultdict(list)
        member_index = defaultdict(list)
        # Also store name-indexed records for strict checking if needed
        for r in parsed["records"]:
            cid = norm_str(r.get("ClaimID",""))
            mid = norm_str(r.get("MemberID",""))
            amt = norm_str(r.get("Amount",""))
            # index by claim id (if present)
            if cid:
                claim_index[cid].append(r)
            # index by member id (if present)
            if mid:
                member_index[mid].append(r)
        per_file_indexes.append((fname, parsed["payees"], claim_index, member_index))

    # Iterate every provider claim and attempt find first match across 820 files
    for provider, claims in providers_data.items():
        for c in claims:
            p_first = c.get("FirstName","")
            p_last = c.get("LastName","")
            p_type = c.get("Type","")
            p_claimid = norm_str(c.get("ClaimID",""))
            p_memberid = norm_str(c.get("MemberID",""))
            p_amount = norm_str(c.get("Amount",""))

            matched = False
            matched_file = ""
            matched_rec = None
            status = "Member not found in 820"

            for fname, payees, claim_index, member_index in per_file_indexes:
                # For C-type: match ClaimID presence AND names AND amount
                if p_type == "C" and p_claimid:
                    candidates = claim_index.get(p_claimid, [])
                    for cand in candidates:
                        # require names and amount to match
                        if name_match(p_first, p_last, cand.get("FirstName",""), cand.get("LastName","")) and equal_amounts(p_amount, cand.get("Amount","")):
                            matched = True
                            matched_file = fname
                            matched_rec = cand
                            break
                    if matched:
                        status = "Match found"
                        break

                # For P-type: match MemberID and names and amount
                if p_type == "P" and p_memberid:
                    candidates = member_index.get(p_memberid, [])
                    for cand in candidates:
                        if name_match(p_first, p_last, cand.get("FirstName",""), cand.get("LastName","")) and equal_amounts(p_amount, cand.get("Amount","")):
                            matched = True
                            matched_file = fname
                            matched_rec = cand
                            break
                    if matched:
                        status = "Match found"
                        break

                # If neither candidate or types missing, continue to next 820 file

            # Build output row
            row = {
                "Provider": provider,
                "FirstName": p_first,
                "LastName": p_last,
                "Type": p_type,
                "ClaimID": p_claimid,
                "MemberID": p_memberid,
                "Amount_837I": p_amount,
                "Matched_820_File": matched_file if matched else "",
                "Amount_820": matched_rec.get("Amount","") if matched_rec else "",
                "820_FirstName": matched_rec.get("FirstName","") if matched_rec else "",
                "820_LastName": matched_rec.get("LastName","") if matched_rec else "",
                "Status": status
            }
            results[provider].append(row)

    return results

# -------------------
# Excel writer (saves to output_reports/)
# -------------------

def save_excel(provider_results: Dict[str, List[dict]], basename: str, output_dir="output_reports"):
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{os.path.splitext(basename)[0]}_comparison.xlsx")

    provider_summary = []
    users_rows = []

    with pd.ExcelWriter(out_path, engine="xlsxwriter") as writer:
        for provider, rows in provider_results.items():
            df = pd.DataFrame(rows)
            # Dedupe subscriber duplicate rows by key columns
            if not df.empty:
                df = df.drop_duplicates(subset=["FirstName","LastName","Type","ClaimID","MemberID","Amount_837I"])
            sheet_name = sanitize_sheet_name(provider)
            df.to_excel(writer, sheet_name=sheet_name, index=False)

            total_claims = df.shape[0]
            unique_users = sorted(set((r["LastName"], r["FirstName"]) for r in rows))
            provider_summary.append({"Provider": provider, "TotalClaims": total_claims, "UniqueUsersCount": len(unique_users)})

            for ln, fn in unique_users:
                cnt = df[(df["LastName"] == ln) & (df["FirstName"] == fn)].shape[0]
                users_rows.append({"Provider": provider, "UserLastName": ln, "UserFirstName": fn, "UserClaimCount": cnt})

        # Write Report sheet
        report_df = pd.DataFrame(provider_summary) if provider_summary else pd.DataFrame(columns=["Provider","TotalClaims","UniqueUsersCount"])
        users_df = pd.DataFrame(users_rows) if users_rows else pd.DataFrame(columns=["Provider","UserLastName","UserFirstName","UserClaimCount"])

        report_df.to_excel(writer, sheet_name="Report", index=False, startrow=0)
        # Write users table below
        startrow = len(report_df) + 3 if not report_df.empty else 3
        if not users_df.empty:
            ws = writer.sheets["Report"]
            ws.write(startrow - 1, 0, "Unique Users per Provider")
            users_df.to_excel(writer, sheet_name="Report", index=False, startrow=startrow)

    return out_path

# -------------------
# Streamlit UI
# -------------------

st.title("837I ↔ 820 Comparison (strict id/member + name + amount)")

with st.sidebar:
    st.header("Upload files")
    files_837 = st.file_uploader("Upload 837I files (multiple)", type=["txt","edi"], accept_multiple_files=True)
    files_820 = st.file_uploader("Upload 820 files (multiple)", type=["txt","edi"], accept_multiple_files=True)
    run = st.button("Run Comparison")

if run:
    if not files_837 or not files_820:
        st.error("Please upload at least one 837I file and one 820 file.")
    else:
        # Read & parse all 820 files (preserve their order)
        parsed_820 = []
        st.info("Parsing 820 files...")
        for f in files_820:
            txt = f.read().decode("utf-8", errors="ignore")
            parsed = parse_820(txt)
            parsed_820.append((f.name, parsed))
            st.write(f"- {f.name}: RMR rows = {len(parsed['records'])}; payees detected = {', '.join(sorted(parsed['payees'])) if parsed['payees'] else '(none)'}")

        saved_files = []
        st.info("Processing 837I files and comparing...")
        for f in files_837:
            txt = f.read().decode("utf-8", errors="ignore")
            providers = parse_837i(txt)

            # preview
            with st.expander(f"837I preview: {f.name} (Providers & claims)"):
                for prov, rows in providers.items():
                    st.write(f"Provider: {prov} — claims: {len(rows)}")
                    if rows:
                        st.dataframe(pd.DataFrame(rows).head(30))

            # compare
            results = compare_providers_vs_820(providers, parsed_820)

            # Save excel
            outpath = save_excel(results, f.name, output_dir="output_reports")
            saved_files.append(outpath)

            # Show Report sheet in UI
            # Recreate Report and Users tables same as save_excel does
            provider_summary = []
            users_rows = []
            for prov, rows in results.items():
                df = pd.DataFrame(rows).drop_duplicates(subset=["FirstName","LastName","Type","ClaimID","MemberID","Amount_837I"]) if rows else pd.DataFrame()
                total_claims = df.shape[0]
                unique_users = sorted(set((r["LastName"], r["FirstName"]) for r in rows))
                provider_summary.append({"Provider": prov, "TotalClaims": total_claims, "UniqueUsersCount": len(unique_users)})
                for ln, fn in unique_users:
                    cnt = df[(df["LastName"] == ln) & (df["FirstName"] == fn)].shape[0]
                    users_rows.append({"Provider": prov, "UserLastName": ln, "UserFirstName": fn, "UserClaimCount": cnt})

            st.subheader(f"Report for {f.name}")
            if provider_summary:
                report_df = pd.DataFrame(provider_summary)
                st.dataframe(report_df, use_container_width=True)
            else:
                st.write("No providers/claims found in this 837I file.")

            if users_rows:
                st.write("Unique users per provider:")
                users_df = pd.DataFrame(users_rows)
                st.dataframe(users_df, use_container_width=True)

            # Optionally show detailed comparisons
            with st.expander(f"Detailed comparisons for {f.name}"):
                for prov, rows in results.items():
                    st.write(f"Provider: {prov}")
                    st.dataframe(pd.DataFrame(rows), use_container_width=True)

        # Final message
        outdir = os.path.abspath("output_reports")
        st.success(f"Reports saved to folder: {outdir}")
        st.write("Saved files:")
        for p in saved_files:
            st.write(f"• {os.path.basename(p)}")











import os
import pandas as pd

def create_excel(provider_results, filename):
    # Ensure output folder exists
    output_folder = "output_reports"
    os.makedirs(output_folder, exist_ok=True)

    # Build output file path
    file_path = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}_results.xlsx")

    total_claims = 0
    summary_rows = []

    with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
        for provider, claims in provider_results.items():
            df = pd.DataFrame(claims)
            df.to_excel(writer, sheet_name=provider[:31], index=False)
            total_claims += len(claims)

            unique_users = len(set([f"{c['firstname']} {c['lastname']}" for c in claims]))
            summary_rows.append({
                "Provider": provider,
                "Total Claims": len(claims),
                "Unique Users": unique_users
            })

        # Summary sheet
        pd.DataFrame(summary_rows).to_excel(writer, sheet_name="Report", index=False)

    return file_path













import streamlit as st
import pandas as pd
import re
import os
from io import BytesIO
from collections import defaultdict

# ---------- Parsing Functions ----------

def parse_837i(file_content):
    """Parse 837I EDI content into structured data grouped by provider."""
    providers_data = defaultdict(list)
    current_provider = None
    current_subscriber = None

    lines = file_content.split("~")
    for line in lines:
        segments = line.strip().split("*")

        if segments[0] == "NM1" and segments[1] == "41":
            current_provider = segments[2] if len(segments) > 2 else "UNKNOWN_PROVIDER"

        if segments[0] == "NM1" and segments[1] == "IL":
            current_subscriber = {
                "firstname": segments[4] if len(segments) > 4 else "",
                "lastname": segments[3] if len(segments) > 3 else "",
                "claim_id": "",
                "member_id": "",
                "claim_amount": "",
                "type": ""
            }

        if segments[0] == "CLM":
            claim_type = segments[1].split("-")[0][-1]  # 'C' or 'P'
            if claim_type.upper() == "C":
                claim_id = segments[1].split("-")[1]
                amount = segments[2]
                current_subscriber["claim_id"] = claim_id
                current_subscriber["type"] = "C"
                current_subscriber["claim_amount"] = amount
            elif claim_type.upper() == "P":
                member_id = segments[1].split("-")[1]
                amount = segments[2]
                current_subscriber["member_id"] = member_id
                current_subscriber["type"] = "P"
                current_subscriber["claim_amount"] = amount
            providers_data[current_provider].append(current_subscriber)

    return providers_data


def parse_820(file_content):
    """Parse 820 EDI content into structured data."""
    records = []
    lines = file_content.split("~")
    current_payee = None
    current_member_id = None
    current_claim_id = None
    current_amount = None

    for line in lines:
        segments = line.strip().split("*")

        if segments[0] == "NM1" and segments[1] == "QE":
            current_payee = segments[3] + " " + segments[4] if len(segments) > 4 else segments[3]
            current_member_id = segments[-1]

        if segments[0] == "RMR" and segments[1] == "IK":
            current_claim_id = segments[2]

        if segments[0] == "AMT" and segments[1] == "R8":
            current_amount = segments[2]
            records.append({
                "payee": current_payee.strip(),
                "member_id": current_member_id.strip(),
                "claim_id": current_claim_id.strip() if current_claim_id else "",
                "amount": current_amount.strip()
            })
            current_claim_id = None
            current_amount = None

    return records

# ---------- Comparison Logic ----------

def compare_837i_820(providers_data, all_820_data):
    """Compare 837I providers data with all 820 data."""
    result_per_provider = defaultdict(list)

    for provider, claims in providers_data.items():
        for claim in claims:
            match_found = False
            found_in_file = None

            for fname, records in all_820_data.items():
                for rec in records:
                    if claim["type"] == "C" and claim["claim_id"] and rec["claim_id"] == claim["claim_id"]:
                        match_found = True
                        found_in_file = fname
                        break
                    elif claim["type"] == "P" and claim["member_id"] and rec["member_id"] == claim["member_id"]:
                        match_found = True
                        found_in_file = fname
                        break
                if match_found:
                    break

            claim_result = claim.copy()
            claim_result["found_in"] = found_in_file if match_found else "Member not found in 820"
            result_per_provider[provider].append(claim_result)

    return result_per_provider

# ---------- Excel Export ----------

def create_excel(provider_results, filename):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        total_claims = 0
        summary_rows = []
        
        for provider, claims in provider_results.items():
            df = pd.DataFrame(claims)
            df.to_excel(writer, sheet_name=provider[:31], index=False)
            total_claims += len(claims)
            unique_users = len(set([f"{c['firstname']} {c['lastname']}" for c in claims]))
            summary_rows.append({
                "Provider": provider,
                "Total Claims": len(claims),
                "Unique Users": unique_users
            })
        
        # Summary sheet
        pd.DataFrame(summary_rows).to_excel(writer, sheet_name="Report", index=False)

    return output.getvalue()

# ---------- Streamlit App ----------

st.title("837I vs 820 Claims Comparison")

st.sidebar.header("Upload EDI Files")
uploaded_837i_files = st.sidebar.file_uploader("Upload 837I files", type=["txt"], accept_multiple_files=True)
uploaded_820_files = st.sidebar.file_uploader("Upload 820 files", type=["txt"], accept_multiple_files=True)

if uploaded_837i_files and uploaded_820_files:
    # Parse all 820 first
    all_820_data = {}
    for file in uploaded_820_files:
        content = file.read().decode("utf-8")
        all_820_data[file.name] = parse_820(content)

    # Process each 837I and compare
    for file in uploaded_837i_files:
        content = file.read().decode("utf-8")
        providers_data = parse_837i(content)
        provider_results = compare_837i_820(providers_data, all_820_data)
        excel_data = create_excel(provider_results, file.name)
        
        st.download_button(
            label=f"Download Excel for {file.name}",
            data=excel_data,
            file_name=f"{os.path.splitext(file.name)[0]}_comparison.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )









import streamlit as st
import pandas as pd
import re
import os
from io import BytesIO

st.set_page_config(page_title="837I vs 820 Claim Comparison", layout="wide")

# --- Helper functions to parse 837I ---
def parse_837i(content):
    claims = []
    segments = content.split("~")
    
    first_name = last_name = claim_id = claim_amount = None
    
    for seg in segments:
        if seg.startswith("NM1*IL"):
            parts = seg.split("*")
            last_name = parts[3] if len(parts) > 3 else ""
            first_name = parts[4] if len(parts) > 4 else ""
        
        elif seg.startswith("CLM"):
            parts = seg.split("*")
            claim_id = parts[1] if len(parts) > 1 else ""
            claim_amount = parts[2] if len(parts) > 2 else ""
            
            if first_name and last_name and claim_id and claim_amount:
                claims.append({
                    "FirstName": first_name,
                    "LastName": last_name,
                    "ClaimID": claim_id,
                    "ClaimAmount": claim_amount
                })
                claim_id = claim_amount = None  # reset for next claim
                
    return pd.DataFrame(claims)

# --- Helper functions to parse 820 ---
def parse_820(content):
    payments = []
    segments = content.split("~")
    
    first_name = last_name = claim_id = claim_amount = None
    
    for seg in segments:
        if seg.startswith("NM1*QE"):
            parts = seg.split("*")
            last_name = parts[3] if len(parts) > 3 else ""
            first_name = parts[4] if len(parts) > 4 else ""
        
        elif seg.startswith("RMR*IK"):
            parts = seg.split("*")
            claim_id = parts[2] if len(parts) > 2 else ""
        
        elif seg.startswith("AMT*"):
            parts = seg.split("*")
            if parts[0] == "AMT" or parts[0].startswith("AMT"):
                claim_amount = parts[2] if len(parts) > 2 else ""
            
            if first_name and last_name and claim_id and claim_amount:
                payments.append({
                    "FirstName": first_name,
                    "LastName": last_name,
                    "ClaimID": claim_id,
                    "ClaimAmount": claim_amount
                })
                claim_id = claim_amount = None
                
    return pd.DataFrame(payments)

# --- Comparison logic ---
def compare_claims(df837, df820):
    results = []
    for _, row in df837.iterrows():
        match = df820[
            (df820["FirstName"].str.lower() == row["FirstName"].lower()) &
            (df820["LastName"].str.lower() == row["LastName"].lower()) &
            (df820["ClaimID"].str.strip() == row["ClaimID"].strip()) &
            (df820["ClaimAmount"].astype(str).str.strip() == str(row["ClaimAmount"]).strip())
        ]
        if match.empty:
            status = "Member not found in 820"
        else:
            status = "Match found"
        
        result_row = row.to_dict()
        result_row["Status"] = status
        results.append(result_row)
    return pd.DataFrame(results)

# --- Streamlit UI ---
st.sidebar.header("Upload EDI Files")
file_837 = st.sidebar.file_uploader("Upload 837I EDI File", type=["txt", "edi"])
file_820 = st.sidebar.file_uploader("Upload 820 EDI File", type=["txt", "edi"])

if file_837 and file_820:
    content_837 = file_837.read().decode("utf-8", errors="ignore")
    content_820 = file_820.read().decode("utf-8", errors="ignore")
    
    df837 = parse_837i(content_837)
    df820 = parse_820(content_820)
    
    st.write("### Parsed 837I Claims")
    st.dataframe(df837)
    st.write("### Parsed 820 Claims")
    st.dataframe(df820)
    
    comparison_df = compare_claims(df837, df820)
    st.write("### Comparison Results")
    st.dataframe(comparison_df)
    
    # Create Excel with one sheet per subscriber
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        for name, group in comparison_df.groupby(["FirstName", "LastName"]):
            sheet_name = f"{name[0]}_{name[1]}"[:31]  # Excel sheet name limit
            group.to_excel(writer, index=False, sheet_name=sheet_name)
    
    st.download_button(
        label="Download Excel Report",
        data=output.getvalue(),
        file_name="837I_vs_820_Comparison.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )








# app.py
import os
import streamlit as st
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.docstore import InMemoryDocstore
from langchain.embeddings import AzureOpenAIEmbeddings
from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage
import tempfile
import json

# ---------------------------
# Streamlit sidebar inputs
# ---------------------------
st.set_page_config(page_title="EDI Companion Guide Rule Ingestion", layout="wide")
st.sidebar.header("Upload Companion Guide")

edi_type = st.sidebar.selectbox("EDI Type", ["837i", "820"])
uploaded_file = st.sidebar.file_uploader("Upload Companion Guide (PDF)", type=["pdf"])
test_query = st.sidebar.text_input("Test retrieval query")
process_button = st.sidebar.button("Extract & Store Rules")

# ---------------------------
# Azure settings from env/secrets
# ---------------------------
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY", st.secrets.get("AZURE_OPENAI_API_KEY"))
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", st.secrets.get("AZURE_OPENAI_ENDPOINT"))
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", st.secrets.get("AZURE_OPENAI_API_VERSION"))
EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", st.secrets.get("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"))
CHAT_DEPLOYMENT = os.getenv("AZURE_OPENAI_GPT_DEPLOYMENT", st.secrets.get("AZURE_OPENAI_GPT_DEPLOYMENT"))

# ---------------------------
# LangChain helpers
# ---------------------------
def get_embeddings():
    return AzureOpenAIEmbeddings(
        deployment=EMBEDDING_DEPLOYMENT,
        model="text-embedding-3-large",
        openai_api_key=AZURE_OPENAI_API_KEY,
        openai_api_base=AZURE_OPENAI_ENDPOINT,
        openai_api_type="azure",
        openai_api_version=AZURE_OPENAI_API_VERSION
    )

def get_llm():
    return AzureChatOpenAI(
        deployment_name=CHAT_DEPLOYMENT,
        temperature=0,
        openai_api_key=AZURE_OPENAI_API_KEY,
        openai_api_base=AZURE_OPENAI_ENDPOINT,
        openai_api_type="azure",
        openai_api_version=AZURE_OPENAI_API_VERSION
    )

def load_pdf(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name
    loader = PyPDFLoader(tmp_path)
    docs = loader.load()
    os.remove(tmp_path)
    return docs

def chunk_documents(documents, chunk_size=1200, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(documents)

def save_faiss_index(vectorstore, edi_type):
    save_path = f"vectorstores/{edi_type}"
    os.makedirs(save_path, exist_ok=True)
    vectorstore.save_local(save_path)

def load_faiss_index(edi_type):
    path = f"vectorstores/{edi_type}"
    if os.path.exists(path):
        return FAISS.load_local(path, get_embeddings(), allow_dangerous_deserialization=True)
    return None

def extract_rules_from_chunk(text, edi_type):
    llm = get_llm()
    prompt = f"""
You are an expert in EDI {edi_type} companion guides.
Given the following guide content, extract ALL segment rules in JSON format.
Each rule must have:
- segment (like "NM1*IL")
- description (1 short sentence)
- fields (list of objects with field_name and description)

Content:
{text}

Return ONLY JSON array, no extra text.
"""
    response = llm([HumanMessage(content=prompt)])
    try:
        rules = json.loads(response.content)
        return rules
    except Exception:
        return []

# ---------------------------
# Processing
# ---------------------------
if process_button and uploaded_file:
    st.write(f"Extracting rules from {edi_type} companion guide...")

    docs = load_pdf(uploaded_file)
    chunks = chunk_documents(docs)

    all_rules_texts = []
    for chunk in chunks:
        rules = extract_rules_from_chunk(chunk.page_content, edi_type)
        for r in rules:
            r_text = json.dumps(r, ensure_ascii=False)
            all_rules_texts.append(r_text)

    if not all_rules_texts:
        st.error("No rules extracted. Check guide format or LLM prompt.")
    else:
        embeddings = get_embeddings()
        vectorstore = FAISS.from_texts(all_rules_texts, embeddings, docstore=InMemoryDocstore({}))
        save_faiss_index(vectorstore, edi_type)
        st.success(f"Extracted {len(all_rules_texts)} rules and saved FAISS index for {edi_type}.")

# ---------------------------
# Test retrieval
# ---------------------------
if test_query:
    vs = load_faiss_index(edi_type)
    if vs:
        results = vs.similarity_search(test_query, k=3)
        st.write("### Top matching rules from FAISS:")
        for i, r in enumerate(results, 1):
            try:
                parsed = json.loads(r.page_content)
                st.json(parsed)
            except:
                st.text(r.page_content)
    else:
        st.error(f"No FAISS index found for {edi_type}. Please extract rules first.")
















# app.py
import os
import streamlit as st
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.docstore import InMemoryDocstore
from langchain.embeddings import AzureOpenAIEmbeddings
from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage
import tempfile
import shutil

# ---------------------------
# Streamlit sidebar inputs
# ---------------------------
st.set_page_config(page_title="EDI Companion Guide Ingestion", layout="wide")
st.sidebar.header("Upload Companion Guide")

edi_type = st.sidebar.selectbox("EDI Type", ["837i", "820"])
uploaded_file = st.sidebar.file_uploader("Upload Companion Guide (PDF)", type=["pdf"])
test_query = st.sidebar.text_input("Test retrieval query")
process_button = st.sidebar.button("Process Guide")

# ---------------------------
# Azure settings from env/secrets
# ---------------------------
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY", st.secrets.get("AZURE_OPENAI_API_KEY"))
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", st.secrets.get("AZURE_OPENAI_ENDPOINT"))
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", st.secrets.get("AZURE_OPENAI_API_VERSION"))
EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", st.secrets.get("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"))
CHAT_DEPLOYMENT = os.getenv("AZURE_OPENAI_GPT_DEPLOYMENT", st.secrets.get("AZURE_OPENAI_GPT_DEPLOYMENT"))

# ---------------------------
# Helper functions
# ---------------------------
def get_embeddings():
    return AzureOpenAIEmbeddings(
        deployment=EMBEDDING_DEPLOYMENT,
        model="text-embedding-3-large",
        openai_api_key=AZURE_OPENAI_API_KEY,
        openai_api_base=AZURE_OPENAI_ENDPOINT,
        openai_api_type="azure",
        openai_api_version=AZURE_OPENAI_API_VERSION
    )

def get_llm():
    return AzureChatOpenAI(
        deployment_name=CHAT_DEPLOYMENT,
        temperature=0,
        openai_api_key=AZURE_OPENAI_API_KEY,
        openai_api_base=AZURE_OPENAI_ENDPOINT,
        openai_api_type="azure",
        openai_api_version=AZURE_OPENAI_API_VERSION
    )

def load_pdf(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name
    loader = PyPDFLoader(tmp_path)
    docs = loader.load()
    os.remove(tmp_path)
    return docs

def chunk_documents(documents, chunk_size=1000, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(documents)

def save_faiss_index(vectorstore, edi_type):
    save_path = f"vectorstores/{edi_type}"
    os.makedirs(save_path, exist_ok=True)
    vectorstore.save_local(save_path)

def load_faiss_index(edi_type):
    path = f"vectorstores/{edi_type}"
    if os.path.exists(path):
        return FAISS.load_local(path, get_embeddings(), allow_dangerous_deserialization=True)
    return None

# ---------------------------
# Processing
# ---------------------------
if process_button and uploaded_file:
    st.write(f"Processing {edi_type} companion guide...")

    # Load & chunk
    docs = load_pdf(uploaded_file)
    chunks = chunk_documents(docs)

    # Create vectorstore
    embeddings = get_embeddings()
    vectorstore = FAISS.from_documents(chunks, embeddings, docstore=InMemoryDocstore({}))

    # Save index
    save_faiss_index(vectorstore, edi_type)
    st.success(f"FAISS index for {edi_type} saved successfully.")

# ---------------------------
# Test retrieval
# ---------------------------
if test_query:
    vs = load_faiss_index(edi_type)
    if vs:
        results = vs.similarity_search(test_query, k=3)
        st.write("### Top matches from FAISS index:")
        for i, r in enumerate(results, 1):
            st.markdown(f"**Result {i}:**\n```\n{r.page_content}\n```")
    else:
        st.error(f"No FAISS index found for {edi_type}. Please process a guide first.")
















import pandas as pd
import json

def extract_bold_field_name(description):
    """Extract field name from bold markdown text like '**ISA01 - Authorization Info Qualifier**'."""
    import re
    match = re.search(r"\*\*(.*?)\*\*", description)
    if match:
        return match.group(1).strip()
    return None

def export_output(data, base_name):
    # Group rows by segment
    segment_rows = {}

    for item in data:
        segment = item["segment"]
        fields = item["fields"]

        row = {}
        for field_desc, field_value in fields:
            field_name = extract_bold_field_name(field_desc)
            if field_name:
                row[field_name] = field_value

        # Append the row to the corresponding segment's list
        segment_rows.setdefault(segment, []).append(row)

    # Write to Excel with multiple sheets
    excel_file = f"{base_name}.xlsx"
    with pd.ExcelWriter(excel_file) as writer:
        for segment, rows in segment_rows.items():
            df = pd.DataFrame(rows)
            # Sheet names can't be more than 31 characters or contain some symbols
            safe_segment = segment[:31].replace('/', '_').replace('\\', '_')
            df.to_excel(writer, sheet_name=safe_segment, index=False)

    # Write full JSON as backup
    json_file = f"{base_name}.json"
    with open(json_file, "w") as jf:
        json.dump(data, jf, indent=4)

















def enrich_with_descriptions(segments_dict, chain):
    enriched_segments = {}

    for segment, data in segments_dict.items():
        # Join all field names together and send one LLM call per segment
        fields = data.get("fields", [])
        
        # Prepare input prompt: ask for description of all fields at once
        joined_fields = ", ".join(fields)
        prompt = f"Describe the following fields for the EDI segment '{segment}': {joined_fields}"

        # Trigger LLM ONCE per segment
        enriched_text = chain.run(prompt)

        # Split back the enriched descriptions intelligently (you can refine this logic)
        enriched_fields = enriched_text.split(",") if "," in enriched_text else [enriched_text]

        # Map each original field with its enriched description (fallback if count mismatch)
        descriptions = {}
        for i, field in enumerate(fields):
            desc = enriched_fields[i].strip() if i < len(enriched_fields) else ""
            descriptions[field] = desc

        # Store back in result
        enriched_segments[segment] = {
            "fields": fields,
            "descriptions": descriptions
        }

    return enriched_segments













def export_output(data_rows, output_path="out/output_segments.xlsx"):
    """
    Export rows to separate Excel sheets based on segment extracted by `extract_bold_field_name`.
    """
    from collections import defaultdict
    import pandas as pd
    import os

    segment_sheets = defaultdict(list)

    for row in data_rows:
        for field in row:
            segment = extract_bold_field_name(field)
            segment_sheets[segment].append(row)
            break  # only consider the first field for segment grouping

    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for segment, rows in segment_sheets.items():
            df = pd.DataFrame(rows)
            safe_segment = segment[:31]  # Excel sheet name limit
            df.to_excel(writer, sheet_name=safe_segment, index=False)

    print(f"Exported to {output_path}")


def extract_bold_field_name(text):
    """
    Extract bold field names enclosed in '**' and return the segment.
    If multiple '**' are found, use the last one.
    Fallback to extracting segment from text like 'ENT01 - something'
    """
    import re
    matches = re.findall(r'\*\*(.*?)\*\*', text)
    if matches:
        segment_name = matches[-1].strip().split()[0]  # Use last match and take first word
    else:
        segment_name = text.strip().split()[0]  # Fallback, take first word (like 'ENT01' → 'ENT')

    return segment_name.upper()














import re
import pandas as pd
import json

def extract_bold_field_name(text):
    """Extract the **bold** part of the field description."""
    match = re.search(r"\*\*(.*?)\*\*", text)
    return match.group(1).strip() if match else None

def export_output(data, base_name):
    rows = []
    for item in data:
        row = {"Segment": item["segment"]}
        for field_desc, field_value in item["fields"]:
            field_name = extract_bold_field_name(field_desc)
            if field_name:
                row[field_name] = field_value
        rows.append(row)

    df = pd.DataFrame(rows)
    excel_file = f"{base_name}.xlsx"
    json_file = f"{base_name}.json"

    df.to_excel(excel_file, index=False)
    
    # Full JSON with descriptions retained
    with open(json_file, "w") as f:
        json.dump(data, f, indent=2)

    return excel_file, json_file













import streamlit as st
from pdf2image import convert_from_path
import pytesseract
from langchain_core.documents import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import AzureChatOpenAI
from langchain.chains import RetrievalQA
import pandas as pd
import json
import os

# === OCR FUNCTION FOR PDF ===
def extract_text_from_pdf_with_ocr(pdf_path, poppler_path=None):
    images = convert_from_path(pdf_path, poppler_path=poppler_path)
    docs = []
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img)
            page_text = f"\n[Page {i+1}]\n{text.strip()}"
            docs.append(Document(page_content=page_text))
        except Exception as e:
            print(f"OCR failed on page {i+1}: {e}")
    return docs

# === VECTOR INDEX ===
def create_vector_index(docs, index_path):
    embedding_model = HuggingFaceBgeEmbeddings(
        model_name="BAAI/bge-large-en",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': False}
    )
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(docs)
    db = FAISS.from_documents(chunks, embedding_model)
    db.save_local(index_path)
    return db

# === LOAD VECTOR INDEX + QA CHAIN ===
def get_retrieval_qa_chain(index_path, azure_api_key, azure_endpoint):
    embedding_model = HuggingFaceBgeEmbeddings(
        model_name="BAAI/bge-large-en",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": False}
    )
    db = FAISS.load_local(index_path, embedding_model)
    llm = AzureChatOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=azure_api_key,
        deployment_name="gpt-4",
        api_version="2023-05-15",
        temperature=0
    )
    retriever = db.as_retriever()
    chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
    return chain

# === PARSE RAW EDI ===
def parse_edi_file(edi_text):
    segments = edi_text.split("~")
    data = []
    for seg in segments:
        parts = seg.strip().split("*")
        if len(parts) > 1:
            data.append({"segment": parts[0], "fields": parts[1:]})
    return data

# === LLM-ASSISTED SEGMENT + FIELD ANNOTATION ===
def enrich_with_descriptions(data, chain):
    enriched = []
    for entry in data:
        segment = entry["segment"]
        fields = entry["fields"]
        try:
            seg_question = f"In an EDI file, what does the segment '{segment}' represent?"
            seg_description = chain.run(seg_question)
        except Exception:
            seg_description = "Not available"
        
        field_names = []
        for idx, field in enumerate(fields):
            try:
                field_question = f"In the EDI segment '{segment}', what does field number {idx + 1} represent?"
                field_description = chain.run(field_question)
            except Exception:
                field_description = f"Field{idx+1}"
            field_names.append((field_description, field))
        
        enriched.append({
            "segment": segment,
            "description": seg_description,
            "fields": field_names
        })
    return enriched

# === EXPORT TO EXCEL + JSON ===
def export_output(data, base_name):
    rows = []
    for item in data:
        row = {
            "Segment": item["segment"],
            "Description": item["description"]
        }
        for desc, val in item["fields"]:
            row[desc] = val
        rows.append(row)

    df = pd.DataFrame(rows)
    excel_file = f"{base_name}.xlsx"
    json_file = f"{base_name}.json"
    df.to_excel(excel_file, index=False)

    with open(json_file, "w") as f:
        json.dump(data, f, indent=2)

    return excel_file, json_file


# ===================
# ==== STREAMLIT ====
# ===================

st.set_page_config(page_title="EDI Data Extractor", layout="wide")
st.title("📑 Intelligent EDI Data Extractor")

# === SIDEBAR CONFIG ===
st.sidebar.header("⚙️ Configuration")
edi_type = st.sidebar.selectbox("Select EDI Format", ["837I", "820"])
pdf_file = st.sidebar.file_uploader("Upload Companion Guide PDF", type=["pdf"], key="pdf")
generate_index = st.sidebar.button("🔁 Generate/Update Vector Index")

azure_key = st.sidebar.text_input("Azure OpenAI API Key", type="password")
azure_endpoint = st.sidebar.text_input("Azure OpenAI Endpoint")

# === INDEX CHECK ===
index_dir = "indices"
index_path = os.path.join(index_dir, f"{edi_type}_index")
index_exists = os.path.exists(os.path.join(index_path, "index.faiss"))

# === GENERATE INDEX IF NEEDED ===
if generate_index:
    if not pdf_file:
        st.sidebar.warning("Please upload a companion guide PDF.")
    else:
        os.makedirs(index_dir, exist_ok=True)
        with open("temp.pdf", "wb") as f:
            f.write(pdf_file.read())
        with st.spinner("Running OCR and creating vector index..."):
            docs = extract_text_from_pdf_with_ocr("temp.pdf")
            create_vector_index(docs, index_path)
        st.sidebar.success("Vector index created successfully.")
        index_exists = True

# === MAIN APP AREA ===
st.subheader(f"📂 Upload and Extract from {edi_type} EDI File")
edi_file = st.file_uploader("Upload EDI File (.txt)", type=["txt"], key="edi")
extract_button = st.button("📤 Extract Data to Excel")

if extract_button:
    if not edi_file:
        st.error("Please upload an EDI file.")
    elif not azure_key or not azure_endpoint:
        st.error("Please provide Azure API Key and Endpoint in the sidebar.")
    elif not index_exists:
        st.warning(f"No index found for {edi_type}. Please upload PDF and generate the index from sidebar.")
    else:
        with st.spinner("Processing..."):
            edi_text = edi_file.read().decode("utf-8")
            segments = parse_edi_file(edi_text)
            chain = get_retrieval_qa_chain(index_path, azure_key, azure_endpoint)
            enriched = enrich_with_descriptions(segments, chain)
            xls_file, json_file = export_output(enriched, edi_type)
        st.success("Extraction complete!")
        with open(xls_file, "rb") as f:
            st.download_button("📥 Download Excel", f, file_name=xls_file)
        with open(json_file, "rb") as f:
            st.download_button("📥 Download JSON", f, file_name=json_file)













def process_edi_files(edi_files):
    import json

    # Load Cached Guide Data
    with open("cached_guide_segments.json", "r") as f:
        guide_data = json.load(f)

    # --- FIX 1: Flatten Nested Lists in guide_data ---
    flattened_guide_data = []
    for item in guide_data:
        if isinstance(item, list):
            flattened_guide_data.extend(item)  # Unpack nested list
        elif isinstance(item, dict):
            flattened_guide_data.append(item)  # Keep dict entries
    guide_data = flattened_guide_data

    # --- Main Processing Loop ---
    extracted_claims = []

    for edi_file in edi_files:
        content = edi_file.read().decode('utf-8', errors='ignore')
        segments = content.split('~')

        current_claim = {}
        for segment in segments:
            segment = segment.strip()
            if not segment:
                continue  # Skip empty lines

            seg_id = segment.split('*')[0].strip()

            # --- FIX 2: Safe Segment Lookup ---
            guide_segment = next(
                (item for item in guide_data if isinstance(item, dict) and item.get('segment_id') == seg_id),
                None
            )

            # Debugging
            st.write(f"Processing Segment ID: {seg_id}")
            if guide_segment:
                st.write(f"Matched Guide Segment: {guide_segment['segment_id']}")
            else:
                st.warning(f"No mapping found for segment: {seg_id}")
                continue

            # Extract Data using LLM
            extracted = extract_data_from_segment(segment, guide_segment)

            # Update Current Claim Dict
            current_claim.update(extracted)

            # Detect End of Claim (For Example: CLM Segment starts a new claim)
            if seg_id == 'CLM' and current_claim:
                extracted_claims.append(current_claim.copy())
                current_claim = {}

    # In case there's any remaining data
    if current_claim:
        extracted_claims.append(current_claim)

    return extracted_claims








def extract_data_from_segment(segment_line, guide_segment):
    # Convert fields list to flat dictionary for prompt clarity
    field_mappings = {field['code']: field['name'] for field in guide_segment['fields']}

    prompt = f"""
    You are an EDI Data Extractor.

    EDI Segment Line:
    {segment_line}

    Companion Guide Mapping:
    Segment ID: {guide_segment['segment_id']}
    Description: {guide_segment['description']}
    Fields Mapping:
    {json.dumps(field_mappings, indent=2)}

    Extract only the data elements from the segment based on this mapping.
    Return as JSON key-value pairs.
    """
    response = client.chat.completions.create(
        model="gpt-35-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    try:
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.warning(f"Failed to parse extraction for segment: {segment_line[:30]}... Error: {e}")
        return {}










import streamlit as st
import os
import json
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from openai import AzureOpenAI

# Azure OpenAI Configuration
AZURE_OPENAI_KEY = "YOUR_AZURE_OPENAI_KEY"
AZURE_OPENAI_ENDPOINT = "YOUR_AZURE_OPENAI_ENDPOINT"
client = AzureOpenAI(api_key=AZURE_OPENAI_KEY, azure_endpoint=AZURE_OPENAI_ENDPOINT)

# --- UI Layout ---
st.title("📄 EDI Companion Guide & Data Extractor")

# Session State
if 'guide_processed' not in st.session_state:
    st.session_state['guide_processed'] = os.path.exists("cached_guide_segments.json")

# --- Companion Guide Upload ---
companion_guide_file = st.file_uploader("Upload Companion Guide (PDF/DOCX)", type=["pdf", "docx"])

if st.button("🔄 Re-Extract Companion Guide Data"):
    if companion_guide_file:
        with st.spinner("Processing Companion Guide using AI..."):
            guide_data = extract_segments_from_guide(companion_guide_file)

            with open("cached_guide_segments.json", "w") as f:
                json.dump(guide_data, f, indent=2)

            st.session_state['guide_processed'] = True
            st.success("Companion Guide processed and cached successfully!")
    else:
        st.error("Please upload a Companion Guide file first.")

# --- EDI Files Upload ---
uploaded_edi_files = st.file_uploader("Upload EDI Files", type=["txt"], accept_multiple_files=True)

if st.button("🚀 Extract Data from EDI Files"):
    if uploaded_edi_files and st.session_state['guide_processed']:
        with st.spinner("Extracting data from EDI files..."):
            extracted_data = process_edi_files(uploaded_edi_files)
            export_to_excel(extracted_data)
            st.success("✅ Data extracted and saved to Extracted_EDI_Data.xlsx!")
            st.download_button("📥 Download Excel", data=open("Extracted_EDI_Data.xlsx", "rb"), file_name="Extracted_EDI_Data.xlsx")
    else:
        st.error("Upload EDI files and ensure Companion Guide is processed.")


# ------------------ FUNCTION DEFINITIONS -------------------

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_segments_from_guide(file):
    # Extract raw text from file
    if file.name.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file)
    elif file.name.endswith(".docx"):
        raw_text = extract_text_from_docx(file)
    else:
        return []

    chunks = split_into_segments(raw_text)

    extracted_data = []
    for chunk in chunks:
        prompt = f"""
        You are an EDI Companion Guide Parser.

        Given Text:
        {chunk}

        Extract Segment ID, Description, and Field Mappings.
        Return JSON like:
        {{
            "segment_id": "SEGMENT_ID",
            "description": "Description",
            "fields": {{
                "ElementCode1": "Field Name 1",
                "ElementCode2": "Field Name 2"
            }}
        }}
        """
        response = client.chat.completions.create(
            model="gpt-35-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        try:
            extracted_data.append(json.loads(response.choices[0].message.content))
        except Exception as e:
            st.warning(f"Failed to parse segment: {chunk[:30]}... Error: {e}")
    return extracted_data

def split_into_segments(raw_text):
    # Example: split by 'Segment:' or segment heading patterns in your guide
    segments = []
    current_segment = ""
    for line in raw_text.split('\n'):
        if line.strip().startswith("Segment") or line.strip().startswith("ISA") or line.strip().startswith("CLM"):
            if current_segment:
                segments.append(current_segment.strip())
            current_segment = line + "\n"
        else:
            current_segment += line + "\n"
    if current_segment:
        segments.append(current_segment.strip())
    return segments

def process_edi_files(edi_files):
    with open("cached_guide_segments.json", "r") as f:
        guide_data = json.load(f)

    extracted_claims = []
    for edi_file in edi_files:
        content = edi_file.read().decode('utf-8', errors='ignore')
        segments = content.split('~')

        current_claim = {}
        for segment in segments:
            seg_id = segment.split('*')[0]
            guide_segment = next((item for item in guide_data if item['segment_id'] == seg_id), None)

            if guide_segment:
                extracted = extract_data_from_segment(segment, guide_segment)
                current_claim.update(extracted)

            if seg_id == 'CLM' and current_claim:
                extracted_claims.append(current_claim.copy())
                current_claim = {}

    return extracted_claims

def extract_data_from_segment(segment_line, guide_segment):
    prompt = f"""
    You are an EDI 837i Data Extractor.

    EDI Segment Line:
    {segment_line}

    Companion Guide Mapping:
    {guide_segment}

    Extract only the data elements from the segment based on the mapping. Return as JSON key-value pairs.
    """
    response = client.chat.completions.create(
        model="gpt-35-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    try:
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.warning(f"Failed to extract data from segment: {segment_line[:30]}... Error: {e}")
        return {}

def export_to_excel(extracted_data):
    df = pd.DataFrame(extracted_data)
    df.to_excel("Extracted_EDI_Data.xlsx", index=False)


















ISA*00*          *00*          *ZZ*SENDERID       *ZZ*RECEIVERID     *240805*1230*^*00501*000000905*0*T*:~
GS*HC*SENDERID*RECEIVERID*20240805*1230*1*X*005010X223A2~
ST*837*0001*005010X223A2~
BHT*0019*00*0123*20240805*1230*CH~
NM1*41*2*SENDER COMPANY*****46*SENDERID~
PER*IC*CONTACT NAME*TE*1234567890~
NM1*40*2*RECEIVER COMPANY*****46*RECEIVERID~
HL*1**20*1~
NM1*85*2*PROVIDER NAME*****XX*1234567893~
N3*123 PROVIDER ADDRESS~
N4*CITY*ST*12345~
CLM*123456789*500***11:B:1*Y*A*Y*I~
DTP*434*RD8*20240801-20240801~
NM1*QC*1*DOE*JOHN****MI*987654321~
N3*456 PATIENT STREET~
N4*PATIENTCITY*ST*98765~
CLM*987654321*300***11:B:1*Y*A*Y*I~
DTP*434*RD8*20240802-20240802~
NM1*QC*1*SMITH*JANE****MI*123123123~
N3*789 PATIENT AVE~
N4*ANOTHERCITY*ST*67890~
SE*20*0001~
GE*1*1~
IEA*1*000000905~









ISA*00*          *00*          *ZZ*PAYERSENDER    *ZZ*PAYEERECEIVER  *240805*1245*^*00501*000000789*0*T*:~
GS*RA*PAYERSENDER*PAYEERECEIVER*20240805*1245*1*X*005010X218~
ST*820*0001*005010X218~
BPR*C*1500*C*ACH*CTX*01*123456789*DA*987654321*20240805~
TRN*1*1234567890*9876543210~
REF*EV*9876543210~
DTM*097*20240805~
N1*PR*PAYER NAME~
N1*PE*PAYEE NAME~
ENT*1*2N*987654321~
RMR*IV*123456789*PO*500~
DTM*003*20240730~
ENT*2*2N*123123123~
RMR*IV*987654321*PO*1000~
DTM*003*20240731~
SE*15*0001~
GE*1*1~
IEA*1*000000789~




@st.cache_data
def build_faiss_index(guide_files):
    chunks = []
    for guide_file in guide_files:
        try:
            content = guide_file.read().decode('utf-8')
        except UnicodeDecodeError:
            guide_file.seek(0)
            content = guide_file.read().decode('latin1')  # fallback to latin1 (windows-1252)
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                chunks.append(line.strip())
    embeddings = OpenAIEmbeddings(openai_api_key=AZURE_OPENAI_KEY)
    vectorstore = FAISS.from_texts(chunks, embeddings)
    return vectorstore







import streamlit as st
import pandas as pd
import os
import faiss
import tempfile
from langchain.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from openai import AzureOpenAI

# Azure OpenAI Config
AZURE_OPENAI_KEY = "YOUR_AZURE_OPENAI_KEY"
AZURE_OPENAI_ENDPOINT = "YOUR_AZURE_OPENAI_ENDPOINT"

# Initialize OpenAI Client
client = AzureOpenAI(api_key=AZURE_OPENAI_KEY, azure_endpoint=AZURE_OPENAI_ENDPOINT)

st.title("EDI Data Extractor - RAG Powered")

# Sidebar - Upload Companion Guides
st.sidebar.header("Upload Companion Guides")
uploaded_guides = st.sidebar.file_uploader("Upload Guide Files (TXT/PDF)", type=["txt", "pdf"], accept_multiple_files=True)

# Main - Upload EDI Files
uploaded_edi_files = st.file_uploader("Upload EDI Files", type=["txt"], accept_multiple_files=True)

# File Type Selection for EDI Files
file_types = {}
if uploaded_edi_files:
    for edi_file in uploaded_edi_files:
        file_type = st.selectbox(f"Select EDI Type for {edi_file.name}", ["837i", "820"], key=edi_file.name)
        file_types[edi_file.name] = file_type

# Load Companion Guides into FAISS
@st.cache_data
def build_faiss_index(guide_files):
    chunks = []
    for guide_file in guide_files:
        content = guide_file.read().decode('utf-8')
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                chunks.append(line.strip())
    embeddings = OpenAIEmbeddings(openai_api_key=AZURE_OPENAI_KEY)
    vectorstore = FAISS.from_texts(chunks, embeddings)
    return vectorstore

if uploaded_guides:
    st.sidebar.success("Companion Guides Uploaded")
    vectorstore = build_faiss_index(uploaded_guides)

def preprocess_edi(content):
    edi_cleaned = content.replace('\n', '').replace('\r', '')
    return [seg.strip() for seg in edi_cleaned.split('~') if seg.strip()]

def retrieve_context(segment_line):
    docs = vectorstore.similarity_search(segment_line, k=1)
    return docs[0].page_content if docs else ""

def extract_data(segment_line, context):
    prompt = f"""
    You are an EDI Data Extractor.
    Companion Guide Context:
    {context}

    Given EDI Segment Line:
    {segment_line}

    Extract data elements as key-value JSON.
    """
    response = client.chat.completions.create(
        model="gpt-35-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# Process Button
if st.button("Process EDI Files") and uploaded_edi_files and uploaded_guides:
    all_claims_data = []

    for edi_file in uploaded_edi_files:
        content = edi_file.read().decode('utf-8')
        segments = preprocess_edi(content)

        current_claim = {}
        for segment in segments:
            context = retrieve_context(segment)
            extracted_json = extract_data(segment, context)
            try:
                segment_data = eval(extracted_json)
                if 'Claim_ID' in segment_data:  # Detecting new claim start
                    if current_claim:
                        all_claims_data.append(current_claim)
                    current_claim = {}
                current_claim.update(segment_data)
            except:
                continue
        if current_claim:
            all_claims_data.append(current_claim)

    # Export to Excel
    df = pd.DataFrame(all_claims_data)
    tmp_download_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    df.to_excel(tmp_download_file.name, index=False)

    st.success("Extraction Complete!")
    st.download_button(label="Download Extracted Data", data=open(tmp_download_file.name, 'rb'), file_name='extracted_claims.xlsx')


















def record_text_input():
    global typed_text, last_clicked_control
    if not typed_text.strip():
        return
    if last_clicked_control:
        action = {
            "event": "text_input",
            "text": typed_text,
            "time": time.time(),
            "control": last_clicked_control  # We use last clicked control!
        }
        recorded_actions.append(action)
        print(f"Typed Text: {typed_text} in {last_clicked_control['name']}")
    else:
        print("No control clicked before typing. Text recorded without control info.")
    typed_text = ""









from pywinauto.uia_defines import IUIA

def record_text_input():
    global typed_text
    if not typed_text.strip():
        return
    try:
        focused_elem = IUIA().get_focused_element()
        info = focused_elem.element_info
        action = {
            "event": "text_input",
            "text": typed_text,
            "time": time.time(),
            "control": {
                "name": info.name,
                "control_type": info.control_type,
                "automation_id": info.automation_id,
                "rectangle": str(info.rectangle)
            }
        }
        recorded_actions.append(action)
        print(f"Typed Text: {typed_text} in {info.name}")
    except Exception as e:
        print(f"Failed to get focused control: {e}")
    typed_text = ""
















def record_text_input():
    global typed_text
    if not typed_text.strip():
        return
    try:
        focused_ctrl = Desktop(backend="uia").get_focus().element_info
        action = {
            "event": "text_input",
            "text": typed_text,
            "time": time.time(),
            "control": {
                "name": focused_ctrl.name,
                "control_type": focused_ctrl.control_type,
                "automation_id": focused_ctrl.automation_id,
                "rectangle": str(focused_ctrl.rectangle)
            }
        }
        recorded_actions.append(action)
        print(f"Typed Text: {typed_text} in {focused_ctrl.name}")
    except Exception as e:
        print(f"Failed to get focused control: {e}")
    typed_text = ""









import mouse
import keyboard
import json
import time
import threading
from pywinauto import Desktop
import win32gui

recorded_actions = []
typed_text = ""
recording = True  # Flag to stop recording on ESC

print("Recording started... Perform actions and press ESC to stop.")

# Get control from current mouse position
def control_from_current_position():
    try:
        x, y = mouse.get_position()
        elem = Desktop(backend="uia").from_point(x, y)
        info = elem.element_info
        return {
            "name": info.name,
            "control_type": info.control_type,
            "automation_id": info.automation_id,
            "rectangle": str(info.rectangle),
            "coords": (x, y)
        }
    except Exception as e:
        print(f"Failed to get control at ({x},{y}): {e}")
        return None

# Mouse Event Listener
def mouse_listener():
    while recording:
        # Mouse is event-driven, no need for polling
        pass  # Hook is handled outside (non-blocking)
    print("Mouse listener stopped.")

# Mouse Hook Callback
def on_mouse_event(event):
    if hasattr(event, 'event_type') and hasattr(event, 'button'):
        if event.event_type == 'down' and event.button in ['left', 'right']:
            time.sleep(0.2)
            ctrl = control_from_current_position()
            if ctrl:
                recorded_actions.append({
                    "event": "mouse_click",
                    "button": event.button,
                    "time": time.time(),
                    "control": ctrl
                })
                print(f"Clicked on {ctrl['name']} at {ctrl['coords']}")

# Keyboard Hook Callback
def on_key_event(event):
    global typed_text, recording
    if event.name == 'esc' and event.event_type == 'down':
        # Stop Recording
        print("ESC pressed. Stopping recording.")
        recording = False
        if typed_text:
            record_text_input()
        # Unhook listeners
        mouse.unhook(on_mouse_event)
        keyboard.unhook_all()
        # Save actions
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        print("Actions saved to recorded_user_actions.json")
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    if not typed_text.strip():
        return
    foreground_hwnd = win32gui.GetForegroundWindow()
    active_ctrl = Desktop(backend="uia").window(handle=foreground_hwnd).element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": active_ctrl.name,
            "control_type": active_ctrl.control_type,
            "automation_id": active_ctrl.automation_id,
            "rectangle": str(active_ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {active_ctrl.name}")
    typed_text = ""

# Start Mouse Hook (non-blocking)
mouse.hook(on_mouse_event)

# Start Keyboard Hook (non-blocking)
keyboard.hook(on_key_event)

# Keep main thread alive while recording
while recording:
    time.sleep(0.5)

print("Recorder stopped.")





















import mouse
import keyboard
import json
import time
from pywinauto import Desktop
import win32gui

recorded_actions = []

print("Recording actions... Open your app and perform actions. Press ESC to stop.")

# Get control from current mouse position
def control_from_current_position():
    try:
        x, y = mouse.get_position()
        elem = Desktop(backend="uia").from_point(x, y)
        info = elem.element_info
        return {
            "name": info.name,
            "control_type": info.control_type,
            "automation_id": info.automation_id,
            "rectangle": str(info.rectangle),
            "coords": (x, y)
        }
    except Exception as e:
        print(f"Failed to get control at ({x},{y}): {e}")
        return None

# Mouse Event Hook
def on_mouse_event(event):
    if hasattr(event, 'event_type') and hasattr(event, 'button'):
        if event.event_type == 'down' and event.button in ['left', 'right']:
            time.sleep(0.2)  # Small UI stabilization delay
            ctrl = control_from_current_position()
            if ctrl:
                recorded_actions.append({
                    "event": "mouse_click",
                    "button": event.button,
                    "time": time.time(),
                    "control": ctrl
                })
                print(f"Clicked on {ctrl['name']} at {ctrl['coords']}")

# Keyboard Hook
typed_text = ""
def on_key(event):
    global typed_text
    if event.name == 'esc':
        if typed_text:
            record_text_input()
        print("Recording stopped.")
        mouse.unhook_all()
        keyboard.unhook_all()
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        exit()
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    foreground_hwnd = win32gui.GetForegroundWindow()
    active_ctrl = Desktop(backend="uia").window(handle=foreground_hwnd).element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": active_ctrl.name,
            "control_type": active_ctrl.control_type,
            "automation_id": active_ctrl.automation_id,
            "rectangle": str(active_ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {active_ctrl.name}")
    typed_text = ""

mouse.hook(on_mouse_event)
keyboard.hook(on_key)

keyboard.wait('esc')















import mouse
import keyboard
import json
import time
from pywinauto import Desktop
import win32gui

recorded_actions = []

print("Recording actions... Open your app and perform actions. Press ESC to stop.")

# Resolve control dynamically at click position
def control_from_point(x, y):
    try:
        elem = Desktop(backend="uia").from_point(x, y)
        info = elem.element_info
        return {
            "name": info.name,
            "control_type": info.control_type,
            "automation_id": info.automation_id,
            "rectangle": str(info.rectangle)
        }
    except:
        return None

# Mouse Event Hook
def on_mouse_event(event):
    if hasattr(event, 'event_type') and hasattr(event, 'button'):
        if event.event_type == 'down' and event.button in ['left', 'right']:
            time.sleep(0.2)  # Small delay to let UI settle
            ctrl = control_from_point(event.x, event.y)
            if ctrl:
                recorded_actions.append({
                    "event": "mouse_click",
                    "button": event.button,
                    "time": time.time(),
                    "control": ctrl
                })
                print(f"Clicked on {ctrl['name']} at ({event.x},{event.y})")

# Keyboard Hook
typed_text = ""
def on_key(event):
    global typed_text
    if event.name == 'esc':
        if typed_text:
            record_text_input()
        print("Recording stopped.")
        mouse.unhook_all()
        keyboard.unhook_all()
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        exit()
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    foreground_hwnd = win32gui.GetForegroundWindow()
    active_ctrl = Desktop(backend="uia").window(handle=foreground_hwnd).element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": active_ctrl.name,
            "control_type": active_ctrl.control_type,
            "automation_id": active_ctrl.automation_id,
            "rectangle": str(active_ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {active_ctrl.name}")
    typed_text = ""

mouse.hook(on_mouse_event)
keyboard.hook(on_key)

keyboard.wait('esc')













import mouse
import keyboard
import json
import time
from pywinauto import Desktop
from pywinauto.controls.uiawrapper import UIAWrapper

recorded_actions = []
control_map = []

print("Loading UI Tree...")

def dump_control_tree(element: UIAWrapper, depth=0):
    node = {
        "name": element.element_info.name,
        "control_type": element.element_info.control_type,
        "automation_id": element.element_info.automation_id,
        "rectangle": element.element_info.rectangle,
        "depth": depth
    }
    control_map.append(node)
    try:
        for child in element.children():
            dump_control_tree(child, depth + 1)
    except:
        pass

# Load Active App Tree
active_app = Desktop(backend="uia").get_active()
dump_control_tree(active_app)

print(f"Loaded {len(control_map)} controls.")

# Find Control by Mouse Position
def find_control_by_position(x, y):
    for ctrl in control_map:
        rect = ctrl["rectangle"]
        if rect.left <= x <= rect.right and rect.top <= y <= rect.bottom:
            return ctrl
    return None

# Mouse Event Hook
def on_mouse_event(event):
    if hasattr(event, 'event_type') and hasattr(event, 'button'):
        if event.event_type == 'down' and event.button in ['left', 'right']:
            time.sleep(0.2)  # Let system stabilize UI state
            ctrl = find_control_by_position(event.x, event.y)
            if ctrl:
                recorded_actions.append({
                    "event": "mouse_click",
                    "button": event.button,
                    "time": time.time(),
                    "control": {
                        "name": ctrl['name'],
                        "control_type": ctrl['control_type'],
                        "automation_id": ctrl['automation_id'],
                        "rectangle": str(ctrl['rectangle'])
                    }
                })
                print(f"Clicked on {ctrl['name']} at ({event.x},{event.y})")

# Keyboard Hook
typed_text = ""
def on_key(event):
    global typed_text
    if event.name == 'esc':
        if typed_text:
            record_text_input()
        print("Recording stopped.")
        mouse.unhook_all()
        keyboard.unhook_all()
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        exit()
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    active_ctrl = Desktop(backend="uia").get_active().element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": active_ctrl.name,
            "control_type": active_ctrl.control_type,
            "automation_id": active_ctrl.automation_id,
            "rectangle": str(active_ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {active_ctrl.name}")
    typed_text = ""

mouse.hook(on_mouse_event)
keyboard.hook(on_key)

keyboard.wait('esc')












import mouse
import keyboard
import json
import time
from pywinauto import Desktop

recorded_actions = []

print("Recording actions... Perform actions and press ESC to stop.")

def control_from_point(x, y):
    try:
        elem = Desktop(backend="uia").from_point(x, y)
        info = elem.element_info
        return {
            "name": info.name,
            "control_type": info.control_type,
            "automation_id": info.automation_id,
            "rectangle": str(info.rectangle)
        }
    except:
        return None

# Mouse Event Hook
def on_mouse_event(event):
    if hasattr(event, 'event_type') and hasattr(event, 'button'):
        if event.event_type == 'down' and event.button in ['left', 'right']:
            if hasattr(event, 'x') and hasattr(event, 'y'):
                ctrl = control_from_point(event.x, event.y)
                if ctrl:
                    recorded_actions.append({
                        "event": "mouse_click",
                        "button": event.button,
                        "time": time.time(),
                        "control": ctrl
                    })
                    print(f"Clicked on {ctrl['name']} at ({event.x},{event.y})")

# Keyboard Hook
typed_text = ""
def on_key(event):
    global typed_text
    if event.name == 'esc':
        if typed_text:
            record_text_input()
        print("Recording stopped.")
        mouse.unhook_all()
        keyboard.unhook_all()
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        exit()
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    ctrl = Desktop(backend="uia").get_active().element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": ctrl.name,
            "control_type": ctrl.control_type,
            "automation_id": ctrl.automation_id,
            "rectangle": str(ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {ctrl.name}")
    typed_text = ""

mouse.hook(on_mouse_event)
keyboard.hook(on_key)

keyboard.wait('esc')









..
import mouse
import keyboard
import json
import time
from pywinauto import Desktop

recorded_actions = []

print("Recording actions... Perform actions on app and press ESC to stop.")

def control_from_point(x, y):
    try:
        elem = Desktop(backend="uia").from_point(x, y)
        info = elem.element_info
        return {
            "name": info.name,
            "control_type": info.control_type,
            "automation_id": info.automation_id,
            "rectangle": str(info.rectangle)
        }
    except:
        return None

# Mouse Hook for click events
def on_mouse_event(event):
    if event.event_type == 'down' and event.button in ['left', 'right']:
        ctrl = control_from_point(event.x, event.y)
        if ctrl:
            recorded_actions.append({
                "event": "mouse_click",
                "button": event.button,
                "time": time.time(),
                "control": ctrl
            })
            print(f"Clicked on {ctrl['name']} at ({event.x},{event.y})")

# Keyboard Listener to capture text inputs
typed_text = ""
def on_key(event):
    global typed_text
    if event.name == 'esc':
        if typed_text:
            record_text_input()
        print("Recording stopped.")
        mouse.unhook_all()
        keyboard.unhook_all()
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        exit()
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    ctrl = Desktop(backend="uia").get_active().element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": ctrl.name,
            "control_type": ctrl.control_type,
            "automation_id": ctrl.automation_id,
            "rectangle": str(ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {ctrl.name}")
    typed_text = ""

mouse.hook(on_mouse_event)
keyboard.hook(on_key)

keyboard.wait('esc')










import mouse
import keyboard
import json
import time
from pywinauto import Desktop

recorded_actions = []

print("Recording actions... Click on app and Press ESC to stop.")

def control_from_point(x, y):
    try:
        elem = Desktop(backend="uia").from_point(x, y)
        info = elem.element_info
        return {
            "name": info.name,
            "control_type": info.control_type,
            "automation_id": info.automation_id,
            "rectangle": str(info.rectangle)
        }
    except:
        return None

# Mouse Click Listener using from_point
def on_click(event):
    ctrl = control_from_point(event.x, event.y)
    if ctrl:
        recorded_actions.append({
            "event": "mouse_click",
            "button": event.button,
            "time": time.time(),
            "control": ctrl
        })
        print(f"Clicked on {ctrl['name']} at ({event.x},{event.y})")

# Keyboard Listener to capture text inputs
typed_text = ""
def on_key(event):
    global typed_text
    if event.name == 'esc':
        if typed_text:
            record_text_input()
        print("Recording stopped.")
        mouse.unhook_all()
        keyboard.unhook_all()
        with open("recorded_user_actions.json", "w") as f:
            json.dump(recorded_actions, f, indent=4)
        exit()
    elif event.event_type == 'down':
        if event.name == 'enter':
            record_text_input()
        else:
            typed_text += event.name if len(event.name) == 1 else ''

def record_text_input():
    global typed_text
    ctrl = Desktop(backend="uia").get_active().element_info
    action = {
        "event": "text_input",
        "text": typed_text,
        "time": time.time(),
        "control": {
            "name": ctrl.name,
            "control_type": ctrl.control_type,
            "automation_id": ctrl.automation_id,
            "rectangle": str(ctrl.rectangle)
        }
    }
    recorded_actions.append(action)
    print(f"Typed Text: {typed_text} in {ctrl.name}")
    typed_text = ""

mouse.on_click(on_click)
keyboard.hook(on_key)

keyboard.wait('esc')









import streamlit as st
import subprocess
import pyautogui
import keyboard
import time
import json
from pywinauto import Desktop
import requests
import os

# ========== CONFIG ==========
AZURE_OPENAI_ENDPOINT = 'https://<YOUR-ENDPOINT>.openai.azure.com/'
AZURE_OPENAI_KEY = '<YOUR-API-KEY>'
DEPLOYMENT_NAME = '<YOUR-DEPLOYMENT-NAME>'  # GPT-4 deployment name in Azure
# ============================

# List of apps you want to automate
apps = {
    "Notepad": r"C:\Windows\System32\notepad.exe",
    "Calculator": r"C:\Windows\System32\calc.exe"
}

st.title("AI-Powered RPA Recorder (Streamlit + Azure OpenAI)")

# Streamlit UI Elements
selected_app = st.selectbox("Select Application", list(apps.keys()))

if st.button("Start Recording"):
    subprocess.Popen(apps[selected_app])
    st.success(f"Recording started for {selected_app}. Perform your actions. Press 'Esc' key to stop.")
    recorded_actions = []

    st.info("Recording in Progress... (Press 'Esc' to Stop Recording)")

    while True:
        # Stop recording if ESC is pressed
        if keyboard.is_pressed('esc'):
            st.success("Recording Stopped.")
            break
        
        # Capture Mouse Clicks
        if pyautogui.mouseDown():
            x, y = pyautogui.position()
            element_info = {}
            try:
                elem = Desktop(backend="uia").get_active()
                element_info = {
                    'name': elem.element_info.name,
                    'control_type': elem.element_info.control_type,
                    'automation_id': elem.element_info.automation_id,
                    'rectangle': str(elem.element_info.rectangle),
                }
            except Exception:
                pass

            action = {
                'type': 'click',
                'position': (x, y),
                'element': element_info,
                'timestamp': time.time()
            }
            recorded_actions.append(action)
            time.sleep(0.5)

        # Capture Keypresses
        keys = keyboard.get_hotkey_name()
        if keys and keys != 'esc':
            element_info = {}
            try:
                elem = Desktop(backend="uia").get_active()
                element_info = {
                    'name': elem.element_info.name,
                    'control_type': elem.element_info.control_type,
                    'automation_id': elem.element_info.automation_id,
                    'rectangle': str(elem.element_info.rectangle),
                }
            except Exception:
                pass

            action = {
                'type': 'keypress',
                'keys': keys,
                'element': element_info,
                'timestamp': time.time()
            }
            recorded_actions.append(action)
            time.sleep(0.5)

    # Save recorded actions
    with open("recorded_actions.json", "w") as f:
        json.dump(recorded_actions, f, indent=4)

    st.success("Actions Recorded Successfully!")

# Generate Script using Azure OpenAI API
if st.button("Generate Script with Azure OpenAI"):
    if not os.path.exists("recorded_actions.json"):
        st.error("No recorded actions found. Please record first.")
    else:
        with open("recorded_actions.json", "r") as f:
            recorded_data = json.load(f)

        prompt = f"""
You are an expert RPA Automation Script Generator using Python's pywinauto library.
Based on the following recorded user actions, generate a Python script to automate them:
{json.dumps(recorded_data, indent=2)}

- Launch the application using pywinauto Application(backend="uia").
- For clicks, try to identify controls by title, control_type, automation_id.
- If identification fails, fallback to click_input(coords=(x, y)).
- For keypress actions, use type_keys.
- Add appropriate waits for control readiness.
- Generate clean executable Python code.
"""

        headers = {
            'Content-Type': 'application/json',
            'api-key': AZURE_OPENAI_KEY
        }

        payload = {
            "messages": [
                {"role": "system", "content": "You are a Python automation script writer."},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 1500,
            "temperature": 0.2
        }

        response = requests.post(
            f"{AZURE_OPENAI_ENDPOINT}openai/deployments/{DEPLOYMENT_NAME}/chat/completions?api-version=2023-05-15",
            headers=headers,
            json=payload
        )

        if response.status_code == 200:
            ai_response = response.json()
            generated_code = ai_response['choices'][0]['message']['content']
            st.code(generated_code, language="python")

            with open("generated_script.py", "w") as f:
                f.write(generated_code)

            st.download_button("Download Script", generated_code, file_name="generated_script.py")
        else:
            st.error(f"Failed to get response from Azure OpenAI: {response.text}")

# Replay Script
if st.button("Replay Generated Script"):
    if not os.path.exists("generated_script.py"):
        st.error("No generated script found. Please generate it first.")
    else:
        subprocess.run(["python", "generated_script.py"])
        st.success("Replay Completed.")


















import streamlit as st
import subprocess
import pyautogui
import keyboard
import time
import json
from pywinauto import Desktop
import requests
import os

# ========== CONFIG ==========
AZURE_OPENAI_ENDPOINT = 'https://<YOUR-ENDPOINT>.openai.azure.com/'
AZURE_OPENAI_KEY = '<YOUR-API-KEY>'
DEPLOYMENT_NAME = '<YOUR-DEPLOYMENT-NAME>'  # GPT-4 deployment name in Azure
# ============================

# List of apps you want to automate
apps = {
    "Notepad": r"C:\Windows\System32\notepad.exe",
    "Calculator": r"C:\Windows\System32\calc.exe"
}

st.title("AI-Powered RPA Recorder (Streamlit + Azure OpenAI)")

# Streamlit UI Elements
selected_app = st.selectbox("Select Application", list(apps.keys()))

if st.button("Start Recording"):
    subprocess.Popen(apps[selected_app])
    st.success(f"Recording started for {selected_app}. Perform your actions. Press 'Esc' key to stop.")
    recorded_actions = []

    st.info("Recording in Progress... (Press 'Esc' to Stop Recording)")

    while True:
        if keyboard.is_pressed('esc'):
            st.success("Recording Stopped.")
            break
        
        x, y = pyautogui.position()
        if pyautogui.mouseDown():
            element_info = {}
            try:
                elem = Desktop(backend="uia").get_active()
                element_info = {
                    'name': elem.element_info.name,
                    'control_type': elem.element_info.control_type,
                    'automation_id': elem.element_info.automation_id,
                    'rectangle': str(elem.element_info.rectangle),
                }
            except Exception:
                pass

            action = {
                'type': 'click',
                'position': (x, y),
                'element': element_info,
                'timestamp': time.time()
            }
            recorded_actions.append(action)
            time.sleep(0.5)

        if keyboard.is_pressed():
            keys = keyboard.get_hotkey_name()
            element_info = {}
            try:
                elem = Desktop(backend="uia").get_active()
                element_info = {
                    'name': elem.element_info.name,
                    'control_type': elem.element_info.control_type,
                    'automation_id': elem.element_info.automation_id,
                    'rectangle': str(elem.element_info.rectangle),
                }
            except Exception:
                pass

            action = {
                'type': 'keypress',
                'keys': keys,
                'element': element_info,
                'timestamp': time.time()
            }
            recorded_actions.append(action)
            time.sleep(0.5)

    # Save recorded actions
    with open("recorded_actions.json", "w") as f:
        json.dump(recorded_actions, f, indent=4)

    st.success("Actions Recorded Successfully!")

# Generate Script using Azure OpenAI API
if st.button("Generate Script with Azure OpenAI"):
    if not os.path.exists("recorded_actions.json"):
        st.error("No recorded actions found. Please record first.")
    else:
        with open("recorded_actions.json", "r") as f:
            recorded_data = json.load(f)

        prompt = f"""
You are an RPA Script Generator using Python's pywinauto library.
Based on the following recorded user actions, generate a Python script that automates them:
{json.dumps(recorded_data, indent=2)}

Use pywinauto Application(backend="uia") to launch apps.
For each click, locate the control using title, control_type, automation_id.
If precise element details are unavailable, fallback to click_input(coords=(x, y)).
Handle keypresses using type_keys.
Ensure waits and proper flow is maintained.
Generate clean executable Python code.
"""

        headers = {
            'Content-Type': 'application/json',
            'api-key': AZURE_OPENAI_KEY
        }

        payload = {
            "messages": [
                {"role": "system", "content": "You are a Python automation script writer."},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 1500,
            "temperature": 0.2
        }

        response = requests.post(
            f"{AZURE_OPENAI_ENDPOINT}openai/deployments/{DEPLOYMENT_NAME}/chat/completions?api-version=2023-05-15",
            headers=headers,
            json=payload
        )

        if response.status_code == 200:
            ai_response = response.json()
            generated_code = ai_response['choices'][0]['message']['content']
            st.code(generated_code, language="python")

            with open("generated_script.py", "w") as f:
                f.write(generated_code)

            st.download_button("Download Script", generated_code, file_name="generated_script.py")
        else:
            st.error(f"Failed to get response from Azure OpenAI: {response.text}")

# Replay Script
if st.button("Replay Generated Script"):
    if not os.path.exists("generated_script.py"):
        st.error("No generated script found. Please generate it first.")
    else:
        subprocess.run(["python", "generated_script.py"])
        st.success("Replay Completed.")














| **S.No** | **Achievement Category**                      | **Details**                                                                                                                                                         |
| -------- | --------------------------------------------- | ------------------------------------------------------------------------------------------------------------------------------------------------------------------- |
| 1        | Multi-PDF Ingestion & Processing              | Successfully designed a workflow to ingest multiple PDFs and extract relevant content chunk-wise using a combination of text and image (OCR) layers.                |
| 2        | Hybrid Chunking (Text + OCR)                  | Implemented fallback OCR extraction using `pdf2image` and `pytesseract` to handle image-based PDFs in absence of text layer.                                        |
| 3        | Vector Store Migration to FAISS               | Transitioned from Qdrant DB (which couldn't be installed on VDI) to FAISS for local vector storage with persistent index management across multiple sessions.       |
| 4        | Metadata-based File Isolation in FAISS        | Maintained file-level segregation in FAISS using metadata (filename tags) to avoid mixing chunks from different PDFs during similarity searches.                    |
| 5        | Retrieval-Augmented Generation (RAG) Pipeline | Successfully built an end-to-end RAG flow: Extract relevant content from contract → Summarize → Generate test cases → Refine them iteratively using LLM.            |
| 6        | Azure OpenAI Integration                      | Integrated with Azure OpenAI endpoint for GPT-4o completions, switching from Groq API to ensure enterprise compliance.                                              |
| 7        | Streamlit UI for Non-Technical Users          | Developed a Streamlit-based tool that allows users to upload PDFs, trigger chunk extraction, and generate refined test scenarios in a user-friendly interface.      |
| 8        | FAISS Persistent Index Save/Load              | Enabled saving/loading of FAISS indexes to disk (local JSON and binary files), ensuring persistence across sessions.                                                |
| 9        | Chunk Querying with Filename Filtering        | Enhanced search mechanism to retrieve only relevant chunks based on filename and query similarity, simulating Qdrant's filter feature.                              |
| 10       | Test Case Refinement & Expansion              | Developed logic to post-process LLM outputs to ensure sufficient depth of clinical journey test cases, dynamically adding service lines based on scenarios.         |
| 11       | DataFrame-driven Post Processing              | Leveraged Pandas for refining, merging, and validating the generated test cases, ensuring alignment with contract-provided codes and SDOH billing data.             |
| 12       | Flexible Deployment without Admin Access      | All code changes were made to work within VDI constraints (no environment variable edits, no Docker, no Qdrant), using local binary paths and in-memory operations. |





import faiss
import numpy as np
import pickle
from langchain.docstore.in_memory import InMemoryDocstore
from langchain.vectorstores.faiss import FAISS
from langchain.schema import Document

faiss_index = None
docstore = None
index_to_docstore_id = None
current_id = 0  # Global unique ID across all docs

def store_chunks_in_faiss(chunks, filename, embedding_model):
    global faiss_index, docstore, index_to_docstore_id, current_id

    texts = [chunk.page_content for chunk in chunks]
    metadatas = [{"filename": filename, "text": chunk.page_content} for chunk in chunks]
    vectors = embedding_model.embed_documents(texts)

    dim = len(vectors[0])
    if faiss_index is None:
        index = faiss.IndexIDMap2(faiss.IndexFlatL2(dim))
        faiss_index = index
        docstore = {}
        index_to_docstore_id = {}

    ids = []
    for i, (vec, meta, text) in enumerate(zip(vectors, metadatas, texts)):
        doc_id = str(current_id)
        faiss_index.add_with_ids(np.array([vec]).astype('float32'), np.array([current_id]))
        docstore[doc_id] = Document(page_content=text, metadata=meta)
        index_to_docstore_id[current_id] = doc_id
        ids.append(current_id)
        current_id += 1

    print(f"Added {len(ids)} chunks from {filename} to FAISS.")


def query_relevant_chunks(query, filename, embedding_model, threshold=0.1, top_k=30):
    query_vector = embedding_model.embed_query(query)

    scores, indices = faiss_index.search(np.array([query_vector]).astype('float32'), top_k)

    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx == -1 or score > threshold:
            continue  # Filter low scores
        doc_id = index_to_docstore_id.get(idx)
        if doc_id:
            doc = docstore[doc_id]
            if doc.metadata.get("filename") == filename:
                results.append(doc.page_content)
    return results

def save_faiss_index(path="faiss_data/multi_pdf"):
    faiss.write_index(faiss_index, f"{path}.index")
    data = {
        "docstore": docstore,
        "index_to_docstore_id": index_to_docstore_id,
        "current_id": current_id
    }
    with open(f"{path}.pkl", "wb") as f:
        pickle.dump(data, f)
    print(f"FAISS index saved at {path}.index and {path}.pkl")

def load_faiss_index(path="faiss_data/multi_pdf"):
    global faiss_index, docstore, index_to_docstore_id, current_id
    faiss_index = faiss.read_index(f"{path}.index")
    with open(f"{path}.pkl", "rb") as f:
        data = pickle.load(f)
    docstore = data["docstore"]
    index_to_docstore_id = data["index_to_docstore_id"]
    current_id = data["current_id"]
    print(f"FAISS index loaded with {len(index_to_docstore_id)} entries.")
    










def query_relevant_chunks(query, filename, threshold=0.1, top_k=30):
    if faiss_index is None:
        print("FAISS index is not built.")
        return []

    # Embed and Normalize Query Vector
    query_vector = embedding_model.embed_query(query)
    query_vector = query_vector / np.linalg.norm(query_vector)

    scores, indices = faiss_index.index.search(np.array([query_vector]).astype("float32"), top_k)

    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx == -1:
            continue  # No result
        doc_id = faiss_index.index_to_docstore_id.get(idx)
        if doc_id is None:
            continue
        doc = faiss_index.docstore.search(doc_id)
        if doc and doc.metadata.get("filename") == filename:
            print(f"[Match] Filename: {doc.metadata.get('filename')} | Score: {score}")
            if score >= threshold:
                results.append(doc.page_content)

    if not results:
        print("No relevant chunks found with current threshold/filename filter.")

    return results







import faiss
import numpy as np
from langchain.vectorstores import FAISS
from langchain.docstore.in_memory import InMemoryDocstore
from langchain.schema import Document

faiss_index = None  # Global FAISS Index

def store_chunks_in_faiss(chunks, filename):
    global faiss_index

    texts = [chunk.page_content for chunk in chunks]
    metadatas = [{"filename": filename, "text": chunk.page_content} for chunk in chunks]

    vectors = embedding_model.embed_documents(texts)
    # Normalize vectors for Cosine Similarity
    vectors = vectors / np.linalg.norm(vectors, axis=1, keepdims=True)

    dim = len(vectors[0])
    index = faiss.IndexFlatIP(dim)  # Inner Product for Cosine Similarity
    index.add(np.array(vectors).astype("float32"))

    documents = [Document(page_content=text, metadata=meta) for text, meta in zip(texts, metadatas)]
    docstore = InMemoryDocstore({str(i): doc for i, doc in enumerate(documents)})
    index_to_docstore_id = {i: str(i) for i in range(len(documents))}

    faiss_index = FAISS(
        embedding_function=embedding_model.embed_query,
        index=index,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
    )

    print(f"FAISS index built with {len(vectors)} vectors for file: {filename}")







import os
import json
import faiss
import numpy as np
from langchain.docstore.in_memory import InMemoryDocstore
from langchain.vectorstores import FAISS
from langchain.schema import Document

# Global FAISS index
faiss_index = None

# Embedding model (Assumed to be initialized elsewhere)
embedding_model = None  # You should initialize your HuggingFaceBgeEmbeddings here

# ---------- Store Chunks in FAISS and Persist ----------
def store_chunks_in_faiss(chunks, filename):
    global faiss_index

    texts = [chunk.page_content for chunk in chunks]
    metadatas = [{"filename": filename, "text": chunk.page_content} for chunk in chunks]

    vectors = embedding_model.embed_documents(texts)

    documents = [Document(page_content=text, metadata=meta) for text, meta in zip(texts, metadatas)]
    dim = len(vectors[0])

    index = faiss.IndexFlatL2(dim)
    index.add(np.array(vectors).astype("float32"))

    os.makedirs("faiss_indices", exist_ok=True)
    faiss.write_index(index, f"faiss_indices/{filename}.index")

    # Persist docstore metadata
    with open(f"faiss_indices/{filename}_docstore.json", "w", encoding="utf-8") as f:
        json.dump([
            {"id": i, "text": doc.page_content, "metadata": doc.metadata}
            for i, doc in enumerate(documents)
        ], f, indent=2)

    # In-memory index for immediate use
    docstore = InMemoryDocstore({str(i): doc for i, doc in enumerate(documents)})
    index_to_docstore_id = {i: str(i) for i in range(len(documents))}

    faiss_index = FAISS(
        embedding_function=embedding_model.embed_query,
        index=index,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
    )

# ---------- Load FAISS Index & Docstore ----------
def load_faiss_index(filename):
    global faiss_index

    index = faiss.read_index(f"faiss_indices/{filename}.index")

    with open(f"faiss_indices/{filename}_docstore.json", "r", encoding="utf-8") as f:
        data = json.load(f)

    docstore = InMemoryDocstore({str(d["id"]): Document(page_content=d["text"], metadata=d["metadata"]) for d in data})
    index_to_docstore_id = {d["id"]: str(d["id"]) for d in data}

    faiss_index = FAISS(
        embedding_function=embedding_model.embed_query,
        index=index,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
    )

# ---------- Query Relevant Chunks ----------
def query_relevant_chunks(query, filename, threshold=0.75, top_k=30):
    global faiss_index

    query_vector = embedding_model.embed_query(query)

    scores, indices = faiss_index.index.search(
        np.array([query_vector]).astype("float32"), top_k
    )

    results = []

    for score, idx in zip(scores[0], indices[0]):
        if idx == -1:
            continue

        doc_id = faiss_index.index_to_docstore_id.get(idx)
        if doc_id is None:
            continue

        doc = faiss_index.docstore.search(doc_id)
        if doc and doc.metadata.get("filename") == filename and score >= threshold:
            results.append(doc.page_content)

    return results

# ---------- Usage Example ----------
# 1. Store Chunks:
# store_chunks_in_faiss(chunks, filename="contract1")

# 2. Load FAISS index when app restarts:
# load_faiss_index(filename="contract1")

# 3. Query:
# results = query_relevant_chunks(query="HIV SNP services", filename="contract1")
# print(results)





















from pywinauto import Application, Desktop
import time

# Start Calculator
Application(backend="uia").start("calc.exe")
time.sleep(2)  # Allow time for full load

# Connect to Calculator window
calc_win = Desktop(backend="uia").window(title_re=".*Calculator.*")
calc_win.wait("visible", timeout=10)
calc_win.set_focus()

# Perform 2 + 3 =
calc_win.child_window(title="Two", control_type="Button").click_input()
calc_win.child_window(title="Plus", control_type="Button").click_input()
calc_win.child_window(title="Three", control_type="Button").click_input()
calc_win.child_window(title="Equals", control_type="Button").click_input()

# Extract and clean result
result_text = calc_win.child_window(auto_id="CalculatorResults", control_type="Text").window_text()
result = result_text.replace("Display is", "").strip()

print("Result is:", result)



from pywinauto import Desktop
windows = Desktop(backend="uia").windows()
for w in windows:
    print(w.window_text())
    


from pywinauto.application import Application
from pywinauto.findwindows import find_windows

app = Application(backend="uia").start("calc.exe")
time.sleep(2)

windows = find_windows(process=app.process, backend="uia")
print("Handles found:", windows)

from pywinauto import Desktop
import time

# Launch Calculator manually first or via script
app = Application(backend="uia").start("calc.exe")
time.sleep(2)

# Use Desktop to list top-level windows
win = Desktop(backend="uia").window(title_re=".*Calculator.*")
win.wait("visible", timeout=10)
win.set_focus()

win.child_window(title="Two", control_type="Button").click_input()
win.child_window(title="Plus", control_type="Button").click_input()
win.child_window(title="Three", control_type="Button").click_input()
win.child_window(title="Equals", control_type="Button").click_input()

result = win.child_window(auto_id="CalculatorResults", control_type="Text").window_text()
print("Result is:", result.replace("Display is", "").strip())











from pywinauto.application import Application
from pywinauto.findwindows import ElementNotFoundError
import time

# Start Calculator
app = Application(backend="uia").start("calc.exe")
time.sleep(1)  # Initial wait

# Print all windows in the app (to debug titles)
windows = app.windows()
print(f"Found {len(windows)} windows:")
for w in windows:
    print(w.window_text())

# Try to attach to the first visible window (fallback logic)
main_win = None
for _ in range(10):  # Try for 5 seconds max
    try:
        windows = app.windows()
        for w in windows:
            if w.is_visible() and "calc" in w.window_text().lower():
                main_win = w
                break
        if main_win:
            break
        time.sleep(0.5)
    except Exception as e:
        print("Retrying due to:", e)
        time.sleep(0.5)

if not main_win:
    raise TimeoutError("Calculator window not found or visible.")

# Focus and operate
main_win.set_focus()
main_win.child_window(title="Two", control_type="Button").click_input()
main_win.child_window(title="Plus", control_type="Button").click_input()
main_win.child_window(title="Three", control_type="Button").click_input()
main_win.child_window(title="Equals", control_type="Button").click_input()

# Get and clean the result
result = main_win.child_window(auto_id="CalculatorResults", control_type="Text").window_text()
print("Result is:", result.replace("Display is", "").strip())












from pywinauto.application import Application
from pywinauto.findwindows import ElementNotFoundError
import time

# Start the Calculator app
app = Application(backend="uia").start("calc.exe")
time.sleep(1)

# Retry to find the window for a few seconds
for _ in range(10):
    try:
        calc = app.window(title_re=".*Calculator.*")
        calc.wait("visible", timeout=2)
        break
    except ElementNotFoundError:
        time.sleep(0.5)
else:
    raise Exception("Calculator window not found after waiting.")

# Bring to front
calc.set_focus()

# Do operation: 2 + 3 =
calc.child_window(title="Two", control_type="Button").click_input()
calc.child_window(title="Plus", control_type="Button").click_input()
calc.child_window(title="Three", control_type="Button").click_input()
calc.child_window(title="Equals", control_type="Button").click_input()

# Get result
result = calc.child_window(auto_id="CalculatorResults", control_type="Text").window_text()
print("Result is:", result)














from pywinauto.application import Application
import time

# Start the Calculator app
app = Application(backend="uia").start('calc.exe')
time.sleep(1)  # wait for it to open

# Connect to the calculator window
calc = app.window(title_re='Calculator')

# Make sure it's in focus
calc.set_focus()

# Perform 2 + 3 =
calc.child_window(title="Two", control_type="Button").click_input()
calc.child_window(title="Plus", control_type="Button").click_input()
calc.child_window(title="Three", control_type="Button").click_input()
calc.child_window(title="Equals", control_type="Button").click_input()

# Wait for result to update
time.sleep(0.5)

# Get the result
result = calc.child_window(auto_id="CalculatorResults", control_type="Text").window_text()

print("Result is:", result)

# Optional: Close Calculator
# app.kill()



from langchain.vectorstores.faiss import FAISS
from langchain.docstore.document import Document
from langchain.docstore.in_memory import InMemoryDocstore
from langchain.embeddings.base import Embeddings
import numpy as np
import faiss

def store_chunks_in_faiss(chunks, filename):
    global faiss_index  # Optional: if you want to reuse or update it later

    texts = [chunk.page_content for chunk in chunks]
    metadatas = [{"filename": filename, "text": chunk.page_content} for chunk in chunks]

    # 1. Embed the documents manually (just like you do in Qdrant)
    vectors = embedding_model.embed_documents(texts)

    # 2. Create Document objects (LangChain uses these with metadata)
    documents = [Document(page_content=text, metadata=meta) for text, meta in zip(texts, metadatas)]

    # 3. Create FAISS index with same vector size
    dim = len(vectors[0])
    index = faiss.IndexFlatL2(dim)

    # 4. Add vectors to FAISS
    index.add(np.array(vectors).astype("float32"))

    # 5. Create LangChain-compatible FAISS object
    docstore = InMemoryDocstore({str(i): doc for i, doc in enumerate(documents)})
    index_to_docstore_id = {i: str(i) for i in range(len(documents))}

    faiss_index = FAISS(
        embedding_function=embedding_model.embed_query,
        index=index,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
    )


def query_relevant_chunks(query, filename, threshold=0.75, top_k=30):
    # Embed the query
    query_vector = embedding_model.embed_query(query)

    # Search FAISS index
    scores, indices = faiss_index.index.search(
        np.array([query_vector]).astype("float32"), top_k
    )

    # Filter results by score threshold and filename match
    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx == -1:
            continue  # No result

        doc_id = faiss_index.index_to_docstore_id.get(idx)
        if doc_id is None:
            continue

        doc = faiss_index.docstore.search(doc_id)
        if doc and doc.metadata.get("filename") == filename and score >= threshold:
            results.append(doc.page_content)

    return results











max_rows_per_file = 10000

# Convert columns to a single header string
header = "".join(col for col in all_columns)

# Total number of rows
total_rows = len(all_data_rows)

# Calculate how many output files are needed
total_files = (total_rows + max_rows_per_file - 1) // max_rows_per_file

for file_num in range(total_files):
    start_index = file_num * max_rows_per_file
    end_index = min(start_index + max_rows_per_file, total_rows)
    chunk_rows = all_data_rows[start_index:end_index]
    
    filename = f"output_claims25-1_latest_{file_num + 1}.txt"
    
    with open(filename, "w") as file:
        file.write(header + "\n")  # write header in every file
        
        for row in chunk_rows:
            s = "|".join(row_val for row_val in row)
            file.write(s + "\n")

















from langchain.document_loaders import PyPDFLoader
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pdf2image import convert_from_path
import pytesseract

def fallback_ocr_loader(pdf_path):
    images = convert_from_path(pdf_path)
    docs = []
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img)
            page_text = f"\n[OCR Page {i+1}]\n{text.strip()}"
            docs.append(Document(page_content=page_text))
        except Exception as e:
            print(f"OCR failed on page {i+1}: {e}")
    return docs

def merge_text_and_ocr(text_docs, ocr_docs):
    merged_docs = []
    total_pages = max(len(text_docs), len(ocr_docs))
    
    for i in range(total_pages):
        text_content = text_docs[i].page_content.strip() if i < len(text_docs) else ""
        ocr_content = ocr_docs[i].page_content.strip() if i < len(ocr_docs) else ""
        
        combined = f"{text_content}\n\n{ocr_content}".strip()
        merged_docs.append(Document(page_content=combined))
    return merged_docs

def load_and_chunk_pdf(pdf_path):
    try:
        loader = PyPDFLoader(pdf_path)
        text_docs = loader.load()
    except:
        text_docs = []

    if not text_docs or not text_docs[0].page_content.strip():
        text_docs = []

    ocr_docs = fallback_ocr_loader(pdf_path)

    # Merge both sources (text + OCR)
    merged_docs = merge_text_and_ocr(text_docs, ocr_docs)

    splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=150)
    chunks = splitter.split_documents(merged_docs)

    return chunks











import streamlit as st
import os
import tempfile
import json
import pandas as pd
import re

from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import HumanMessage
from langchain.chat_models import AzureChatOpenAI
from langchain_groq import ChatGroq
from langchain.document_loaders import PyPDFLoader
from pdf2image import convert_from_path
import pytesseract

# Configuration
QDRANT_HOST = "localhost"
COLLECTION_NAME = "hiv_snp_chunks"
embedding_model = HuggingFaceBgeEmbeddings(model_name="BAAI/bge-large-en", model_kwargs={"device": "cpu"})
qdrant = QdrantClient(QDRANT_HOST)

# OCR fallback
def fallback_ocr_loader(pdf_path):
    images = convert_from_path(pdf_path)
    all_text = ""
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img)
            all_text += f"\nPage {i+1}\n{text}"
        except Exception as e:
            print(f"OCR failed on page {i+1}: {e}")
    return all_text

# Load and chunk PDF
def load_and_chunk_pdf(pdf_path):
    try:
        loader = PyPDFLoader(pdf_path)
        docs = loader.load()
    except:
        docs = [fallback_ocr_loader(pdf_path)]
    if not docs or not docs[0].page_content.strip():
        docs = [fallback_ocr_loader(pdf_path)]
    splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=150)
    return splitter.split_documents(docs)

# Store chunks in Qdrant with file name tag
def store_chunks_in_qdrant(chunks, filename):
    vectors = embedding_model.embed_documents([chunk.page_content for chunk in chunks])
    points = [
        PointStruct(
            id=i,
            vector=vec,
            payload={"text": chunk.page_content, "filename": filename}
        )
        for i, (vec, chunk) in enumerate(zip(vectors, chunks))
    ]
    qdrant.recreate_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=len(vectors[0]), distance=Distance.COSINE)
    )
    qdrant.upload_points(collection_name=COLLECTION_NAME, points=points)

# Query Qdrant for relevant chunks
def query_relevant_chunks(query, filename, threshold=0.75, top_k=30):
    query_vector = embedding_model.embed_query(query)
    results = qdrant.search(
        collection_name=COLLECTION_NAME,
        query_vector=query_vector,
        limit=top_k,
        query_filter=Filter(
            must=[FieldCondition(key="filename", match=MatchValue(value=filename))]
        ),
        score_threshold=threshold
    )
    return [res.payload["text"] for res in results]

# Choose LLM client
def get_llm_client(source, groq_key=None, azure_key=None, azure_url=None, azure_deployment=None):
    if source == "Groq":
        return ChatGroq(api_key=groq_key, model="meta-llama/llama-4-scout-17b-16e-instruct")
    else:
        return AzureChatOpenAI(
            azure_endpoint=azure_url,
            openai_api_key=azure_key,
            deployment_name=azure_deployment,
            model_name="gpt-4o",
            api_version="2025-01-01-preview"
        )

# Streamlit App
st.title("📄 HIV SNP Test Case Generator")

with st.sidebar:
    st.header("⚙️ LLM Settings")
    llm_source = st.selectbox("LLM Provider", ["Groq", "Azure OpenAI"])
    groq_key = st.text_input("Groq API Key", type="password")
    azure_key = st.text_input("Azure API Key", type="password")
    azure_url = st.text_input("Azure Endpoint")
    azure_deployment = st.text_input("Azure Deployment (e.g. gpt-4o)")

uploaded_files = st.file_uploader("Upload PDF files", type="pdf", accept_multiple_files=True)

if uploaded_files:
    for uploaded in uploaded_files:
        with st.spinner(f"Processing {uploaded.name}..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(uploaded.read())
                chunks = load_and_chunk_pdf(tmp.name)
                store_chunks_in_qdrant(chunks, uploaded.name)
            st.success(f"{uploaded.name} indexed successfully with {len(chunks)} chunks.")

    query = st.text_area("Enter your query", value="HIV SNP services in compensation/payment sections")

    if st.button("Generate Summary & Test Cases"):
        for uploaded in uploaded_files:
            st.subheader(f"📄 {uploaded.name}")
            relevant_chunks = query_relevant_chunks(query, uploaded.name)
            if not relevant_chunks:
                st.warning("No relevant content found.")
                continue

            combined_text = "\n--\n".join(relevant_chunks)
            llm = get_llm_client(
                llm_source, groq_key=groq_key, azure_key=azure_key,
                azure_url=azure_url, azure_deployment=azure_deployment
            )

            # Step 1: Summarize
            prompt_summary = f"""You are reviewing extracted paragraphs from a provider contract.

Text:
"""{combined_text}"""

Summarize HIV SNP compensation terms, services, codes, units, rates, and limits."""
            summary = llm([HumanMessage(content=prompt_summary)]).content.strip()
            st.info(summary)

            # Step 2: Generate test cases
            prompt_testcases = f"""You are a healthcare QA engineer. Based on the contract summary:

{summary}

Generate at least 6 test case scenarios in JSON list format with fields:
- Summary
- Test Scenario
- Line Number
- Requirement
- Service Type
- Revenue Code
- Diagnosis Code
- Units
- POS
- Bill Amount
- Expected Output

Respond with only a valid JSON array."""
            raw = llm([HumanMessage(content=prompt_testcases)]).content.strip()
            match = re.search(r'\[.*\]', raw, re.DOTALL)
            if match:
                try:
                    json_data = json.loads(match.group(0))
                    df = pd.DataFrame(json_data)
                    st.dataframe(df)
                    output_file = f"out/testcases_{uploaded.name}.xlsx"
                    df.to_excel(output_file, index=False)
                    with open(output_file, "rb") as f:
                        st.download_button("⬇️ Download Excel", f, file_name=output_file)
                except json.JSONDecodeError as e:
                    st.error(f"JSON Parse Error: {e}")
                    st.code(raw[:500])
            else:
                st.error("No valid JSON found in LLM response.")
